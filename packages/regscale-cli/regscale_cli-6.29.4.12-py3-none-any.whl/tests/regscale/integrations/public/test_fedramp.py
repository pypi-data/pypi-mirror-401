#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import json
from pathlib import Path
from unittest.mock import MagicMock, mock_open, patch
from xml.etree.ElementTree import Element

import pytest
from docx import Document
from lxml import etree

from regscale.integrations.public.fedramp.fedramp_cis_crm import gen_key
from regscale.integrations.public.fedramp.fedramp_common import (
    check_profile,
    create_component_mapping,
    create_port,
    create_ssp_components,
    get_base_contact,
    get_profile,
    get_tables,
    get_text_between_headers,
    get_xpath_data_detailed,
    post_links,
    post_regscale_object,
    validate_date_str,
)
from regscale.integrations.public.fedramp.fedramp_five import extract_parts
from regscale.integrations.public.fedramp.import_workbook import _update_ip_addresses, determine_ip_address_version
from regscale.models.regscale_models.component import Component
from regscale.models.regscale_models.control_implementation import (
    ControlImplementation,
    ControlImplementationOrigin,
    ControlImplementationStatus,
)
from regscale.models.regscale_models.ports_protocol import PortsProtocol
from tests import CLITestFixture
from tests.mocks.response import create_mock_response


class TestFedRamp(CLITestFixture):
    mock_ssp_id = 123
    existing_links = []
    content_type = "application/json"
    accept = content_type

    @staticmethod
    @pytest.fixture
    def element():
        # This creates a new Element object for each test
        return Element("Control")

    @staticmethod
    @pytest.fixture
    def implementation():
        # This provides a fresh instance of ControlImplementation for each test
        return ControlImplementation(
            **{
                "id": 3367,
                "uuid": "28485bbd-ea6a-400f-9211-8be5e792d05d",
                "isPublic": True,
                "inheritable": True,
                "controlOwner": "Sumka, James",
                "controlOwnerId": "758e40b9-8f47-498f-9581-bb1f2f63c52f",
                "systemRole": "Service Provider Management",
                "policy": "",
                "implementation": "My implemen",
                "cloudImplementation": "",
                "status": "Fully Implemented",
                "stepsToImplement": "",
                "lastAssessmentResult": "",
                "controlName": "15.3",
                "controlTitle": "Classify Service Providers",
                "controlID": 10797,
                "controlSource": "Baseline",
                "exclusionJustification": "",
                "practiceLevel": "",
                "processLevel": "",
                "cyberFunction": "",
                "implementationType": "",
                "implementationMethod": "",
                "qdWellDesigned": "",
                "qdProcedures": "",
                "qdSegregation": "",
                "qdFlowdown": "",
                "qdAutomated": "",
                "qdOverall": "",
                "qiResources": "",
                "qiMaturity": "",
                "qiReporting": "",
                "qiVendorCompliance": "",
                "qiIssues": "",
                "qiOverall": "",
                "bBaseline": False,
                "bInherited": False,
                "bOverlay": False,
                "bTailored": False,
                "bStatusImplemented": False,
                "bStatusPartiallyImplemented": False,
                "bStatusPlanned": False,
                "bStatusAlternative": False,
                "bStatusNotApplicable": False,
                "bServiceProviderCorporate": False,
                "bServiceProviderSystemSpecific": False,
                "bServiceProviderHybrid": False,
                "bConfiguredByCustomer": False,
                "bProvidedByCustomer": False,
                "bShared": False,
                "bInheritedFedrampAuthorization": False,
                "responsibility": "Customer",
                "assessmentFrequency": 365,
                "parentId": 22,
                "parentModule": "securityplans",
                "createdById": "758e40b9-8f47-498f-9581-bb1f2f63c52f",
                "dateCreated": "2024-01-08T17:50:07.9305152",
                "lastUpdatedById": "758e40b9-8f47-498f-9581-bb1f2f63c52f",
                "dateLastUpdated": "2024-02-07T13:54:45.5246858",
                "weight": 0,
                "parameters": [],
                "tests": [],
                "objectives": [],
                "family": "Service Provider Management",
            }
        )

    @staticmethod
    @pytest.fixture
    def example_table():
        """Create an example XML table"""
        xml_string = """
        <w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex" xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex" xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex" xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex" xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex" xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex" xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex" xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink" xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:oel="http://schemas.microsoft.com/office/2019/extlst" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash" xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"><w:tblPr><w:tblW w:w="5000" w:type="pct"/><w:jc w:val="center"/><w:tblBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="969996" w:themeColor="text1" w:themeTint="80"/><w:left w:val="single" w:sz="4" w:space="0" w:color="969996" w:themeColor="text1" w:themeTint="80"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="969996" w:themeColor="text1" w:themeTint="80"/><w:right w:val="single" w:sz="4" w:space="0" w:color="969996" w:themeColor="text1" w:themeTint="80"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="969996" w:themeColor="text1" w:themeTint="80"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="969996" w:themeColor="text1" w:themeTint="80"/></w:tblBorders><w:tblCellMar><w:top w:w="72" w:type="dxa"/><w:left w:w="72" w:type="dxa"/><w:bottom w:w="72" w:type="dxa"/><w:right w:w="72" w:type="dxa"/></w:tblCellMar><w:tblLook w:val="0000" w:firstRow="0" w:lastRow="0" w:firstColumn="0" w:lastColumn="0" w:noHBand="0" w:noVBand="0"/></w:tblPr><w:tblGrid><w:gridCol w:w="3541"/><w:gridCol w:w="9409"/></w:tblGrid><w:tr w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w14:paraId="25F8059F" w14:textId="77777777" w:rsidTr="00551FB0"><w:trPr><w:cantSplit/><w:trHeight w:val="20"/><w:tblHeader/><w:jc w:val="center"/></w:trPr><w:tc><w:tcPr><w:tcW w:w="1367" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="1D396B" w:themeFill="accent5"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar><w:vAlign w:val="center"/></w:tcPr><w:p w14:paraId="37EBFE1E" w14:textId="77777777" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00CC17A4" w:rsidP="0094101B"><w:pPr><w:pStyle w:val="GSATableHeading"/><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:szCs w:val="20"/></w:rPr></w:pPr><w:r w:rsidRPr="00551FB0"><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:szCs w:val="20"/></w:rPr><w:t>Security Objective</w:t></w:r></w:p></w:tc><w:tc><w:tcPr><w:tcW w:w="3633" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="1D396B" w:themeFill="accent5"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar><w:vAlign w:val="center"/></w:tcPr><w:p w14:paraId="73808EED" w14:textId="77777777" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00CC17A4" w:rsidP="0094101B"><w:pPr><w:pStyle w:val="GSATableHeading"/><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:szCs w:val="20"/></w:rPr></w:pPr><w:r w:rsidRPr="00551FB0"><w:rPr><w:rFonts w:asciiTheme="majorHAnsi" w:hAnsiTheme="majorHAnsi"/><w:szCs w:val="20"/></w:rPr><w:t>Low, Moderate or High</w:t></w:r></w:p></w:tc></w:tr><w:tr w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w14:paraId="1208D909" w14:textId="77777777" w:rsidTr="00551FB0"><w:trPr><w:cantSplit/><w:trHeight w:val="20"/><w:tblHeader/><w:jc w:val="center"/></w:trPr><w:tc><w:tcPr><w:tcW w:w="1367" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="auto"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar></w:tcPr><w:p w14:paraId="36FC716A" w14:textId="77777777" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00CC17A4" w:rsidP="00551FB0"><w:pPr><w:pStyle w:val="GSATableHeading"/><w:keepNext w:val="0"/><w:keepLines w:val="0"/><w:rPr><w:szCs w:val="20"/></w:rPr></w:pPr><w:r w:rsidRPr="00551FB0"><w:rPr><w:szCs w:val="20"/></w:rPr><w:t>Confidentiality</w:t></w:r></w:p></w:tc><w:sdt><w:sdtPr><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr><w:alias w:val="Low Moderate High"/><w:tag w:val="lowmoderatehigh"/><w:id w:val="-1240482717"/><w:placeholder><w:docPart w:val="BBFD3EFCFAB74D189DFADB8014C4B0A2"/></w:placeholder><w:dropDownList><w:listItem w:value="Choose an item."/><w:listItem w:displayText="Low (L)" w:value="Low (L)"/><w:listItem w:displayText="Moderate (M)" w:value="Moderate (M)"/><w:listItem w:displayText="High (H)" w:value="High (H)"/></w:dropDownList></w:sdtPr><w:sdtEndPr/><w:sdtContent><w:tc><w:tcPr><w:tcW w:w="3633" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="auto"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar></w:tcPr><w:p w14:paraId="0D97C262" w14:textId="43014E66" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00D87316" w:rsidP="00551FB0"><w:pPr><w:pStyle w:val="GSATableText"/><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr></w:pPr><w:r><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr><w:t>Moderate (M)</w:t></w:r></w:p></w:tc></w:sdtContent></w:sdt></w:tr><w:tr w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w14:paraId="1274421F" w14:textId="77777777" w:rsidTr="00551FB0"><w:trPr><w:cantSplit/><w:trHeight w:val="20"/><w:tblHeader/><w:jc w:val="center"/></w:trPr><w:tc><w:tcPr><w:tcW w:w="1367" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="auto"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar></w:tcPr><w:p w14:paraId="787AE0EB" w14:textId="77777777" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00CC17A4" w:rsidP="00551FB0"><w:pPr><w:pStyle w:val="GSATableHeading"/><w:keepNext w:val="0"/><w:keepLines w:val="0"/><w:rPr><w:szCs w:val="20"/></w:rPr></w:pPr><w:r w:rsidRPr="00551FB0"><w:rPr><w:szCs w:val="20"/></w:rPr><w:t>Integrity</w:t></w:r></w:p></w:tc><w:sdt><w:sdtPr><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr><w:alias w:val="Low Moderate High"/><w:tag w:val="lowmoderatehigh"/><w:id w:val="-577524295"/><w:placeholder><w:docPart w:val="0BF26A562B574649976781A6EEC52519"/></w:placeholder><w:dropDownList><w:listItem w:value="Choose an item."/><w:listItem w:displayText="Low (L)" w:value="Low (L)"/><w:listItem w:displayText="Moderate (M)" w:value="Moderate (M)"/><w:listItem w:displayText="High (H)" w:value="High (H)"/></w:dropDownList></w:sdtPr><w:sdtEndPr/><w:sdtContent><w:tc><w:tcPr><w:tcW w:w="3633" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="auto"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar></w:tcPr><w:p w14:paraId="1D138F7C" w14:textId="6E8E85A3" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00D87316" w:rsidP="00551FB0"><w:pPr><w:pStyle w:val="GSATableText"/><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr></w:pPr><w:r><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr><w:t>Moderate (M)</w:t></w:r></w:p></w:tc></w:sdtContent></w:sdt></w:tr><w:tr w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w14:paraId="7AA0A70F" w14:textId="77777777" w:rsidTr="00551FB0"><w:trPr><w:cantSplit/><w:trHeight w:val="20"/><w:tblHeader/><w:jc w:val="center"/></w:trPr><w:tc><w:tcPr><w:tcW w:w="1367" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="auto"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar></w:tcPr><w:p w14:paraId="111949F9" w14:textId="77777777" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00CC17A4" w:rsidP="00551FB0"><w:pPr><w:pStyle w:val="GSATableHeading"/><w:keepNext w:val="0"/><w:keepLines w:val="0"/><w:rPr><w:szCs w:val="20"/></w:rPr></w:pPr><w:r w:rsidRPr="00551FB0"><w:rPr><w:szCs w:val="20"/></w:rPr><w:t>Availability</w:t></w:r></w:p></w:tc><w:sdt><w:sdtPr><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr><w:alias w:val="Low Moderate High"/><w:tag w:val="lowmoderatehigh"/><w:id w:val="829017778"/><w:placeholder><w:docPart w:val="4D57AA5107874C0EAE6342AA135AAA7E"/></w:placeholder><w:dropDownList><w:listItem w:value="Choose an item."/><w:listItem w:displayText="Low (L)" w:value="Low (L)"/><w:listItem w:displayText="Moderate (M)" w:value="Moderate (M)"/><w:listItem w:displayText="High (H)" w:value="High (H)"/></w:dropDownList></w:sdtPr><w:sdtEndPr/><w:sdtContent><w:tc><w:tcPr><w:tcW w:w="3633" w:type="pct"/><w:shd w:val="clear" w:color="auto" w:fill="auto"/><w:tcMar><w:top w:w="0" w:type="dxa"/><w:left w:w="101" w:type="dxa"/><w:bottom w:w="115" w:type="dxa"/><w:right w:w="101" w:type="dxa"/></w:tcMar></w:tcPr><w:p w14:paraId="12379D77" w14:textId="4BC04DBB" w:rsidR="00CC17A4" w:rsidRPr="00551FB0" w:rsidRDefault="00D87316" w:rsidP="00551FB0"><w:pPr><w:pStyle w:val="GSATableText"/><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr></w:pPr><w:r><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr><w:t>Moderate (M)</w:t></w:r></w:p></w:tc></w:sdtContent></w:sdt></w:tr></w:tbl>
        """
        # Create a new document
        root = etree.fromstring(xml_string)
        table = Document().add_table(rows=1, cols=1)
        table._tbl = root
        # Return the document as a binary string
        return table

    @staticmethod
    @pytest.fixture
    def document():
        """Create a docx document"""
        doc = Document()
        doc.add_heading("Header 1", level=1)
        doc.add_paragraph("Some text")
        doc.add_paragraph("More text")
        doc.add_heading("Header 2", level=1)
        doc.add_paragraph("Some more text")
        doc.add_paragraph("Even more text")
        return doc

    @pytest.fixture
    def test_document(self):
        """Load test document"""
        doc = Document(f"{self.get_tests_dir('tests')}/test_data/RegScale FedRAMP Mock Authorization Package.docx")
        return doc

    @staticmethod
    def test_gen_key():
        """Test the gen_key function with various control ID patterns"""

        # Test cases for basic control IDs without enhancements
        assert gen_key("AC-1") == "AC-1"
        assert gen_key("SI-2") == "SI-2"
        assert gen_key("AU-3") == "AU-3"

        # Test cases for control IDs with numeric enhancements
        assert gen_key("AC-1(1)") == "AC-1(1)"
        assert gen_key("SI-2(5)") == "SI-2(5)"
        assert gen_key("AU-3(10)") == "AU-3(10)"

        # Test cases for control IDs with letter enhancements (should be stripped)
        assert gen_key("AC-1(a)") == "AC-1"
        assert gen_key("SI-2(b)") == "SI-2"
        assert gen_key("AU-3(z)") == "AU-3"

        # Test cases for control IDs with both numeric and letter enhancements
        assert gen_key("AC-1(1)(a)") == "AC-1(1)"
        assert gen_key("SI-2(5)(b)") == "SI-2(5)"
        assert gen_key("AU-3(10)(z)") == "AU-3(10)"

        # Test cases for edge cases
        assert gen_key("AC-1(1)(A)") == "AC-1(1)"  # Uppercase letter
        assert gen_key("SI-2(5)(B)") == "SI-2(5)"  # Uppercase letter
        assert gen_key("AU-3(10)(Z)") == "AU-3(10)"  # Uppercase letter

        # Test cases for malformed or non-matching patterns
        assert gen_key("AC1") == "AC1"  # No dash
        assert gen_key("AC-1a") == "AC-1a"  # Letter without parentheses
        assert gen_key("AC-1(1a)") == "AC-1(1a)"  # Mixed in parentheses
        assert gen_key("") == ""  # Empty string
        assert gen_key("AC-1(b)") == "AC-1"  # Multiple numeric enhancements (keeps first)

        # Test cases for different control families
        assert gen_key("CM-1") == "CM-1"
        assert gen_key("IA-1") == "IA-1"
        assert gen_key("MP-1") == "MP-1"
        assert gen_key("PE-1") == "PE-1"
        assert gen_key("PL-1") == "PL-1"
        assert gen_key("PS-1") == "PS-1"
        assert gen_key("RA-1") == "RA-1"
        assert gen_key("SA-1") == "SA-1"
        assert gen_key("SC-1") == "SC-1"
        assert gen_key("SR-1") == "SR-1"

    @staticmethod
    @pytest.mark.skip("Not implemented")
    def test_validate_date_str():
        # Test valid date
        assert validate_date_str("05/05/2023")
        # Test invalid date
        assert validate_date_str("May 5, 2023") is False
        # Test exception handling
        mock_datetime = MagicMock()
        mock_datetime.strptime.side_effect = ValueError()
        with patch("datetime.datetime", mock_datetime):
            assert validate_date_str("05/05/2023")
            assert validate_date_str("marches 5 2021") is False

    @pytest.mark.skip("Not implemented")
    def test_check_profile(self):
        # Create mock response data
        api = self.api
        data = [
            {
                "id": 576,
                "uuid": "47bdf9eb-e885-40e6-b1f6-e7b906f231bf",
                "name": "CDS_Access",
                "confidentiality": "High",
                "integrity": "High",
                "availability": "Moderate",
                "category": "High",
                "profileOwner": "Eaton, Bryan",
                "profileOwnerId": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
                "createdById": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
                "dateCreated": "2023-03-29T20:59:41.4556531",
                "lastUpdatedById": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
                "dateLastUpdated": "2023-03-29T20:59:41.4557173",
                "isPublic": True,
            }
        ]

        mock_response = create_mock_response(json_data=data)
        # Mock the API object and its `get` method to return the mock response
        api.get = MagicMock(return_value=mock_response)
        # Define the input parameters
        config = self.config
        title = "CDS_Access"

        # Call the `check_profile` function with the input parameters
        result = check_profile(api, config, title)

        # Assert that the function returns the correct profile ID
        assert result == 576

    @pytest.mark.skip("Not implemented")
    def test_create_component_mapping(self, caplog):
        # Create mock response data
        api = self.api
        data = {
            "id": 614,
            "uuid": "4176063c-ae0d-4727-a6df-e582b01d6d27",
            "securityPlanId": 279,
            "componentId": 676,
            "createdById": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
            "dateCreated": "2023-05-04T17:30:17.2216596+00:00",
            "lastUpdatedById": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
            "isPublic": True,
            "dateLastUpdated": "2023-05-04T17:30:17.2216599+00:00",
            "tenantsId": 1,
        }
        mock_response = create_mock_response(json_data=data)

        # Mock the API object and its `post` method to return the mock response
        api.post = MagicMock(return_value=mock_response)

        # Define the input parameters
        config = self.config
        ssp_id = 279
        cmp_id = 676
        headers = {"Content-Type": self.content_type}

        # Call the `create_component_mapping` function with the input parameters
        create_component_mapping(api, config, ssp_id, cmp_id)

        # Check that the `post` method was called with the correct URL and data
        expected_data = {
            "securityPlanId": ssp_id,
            "componentId": cmp_id,
            "isPublic": True,
            "createdById": config["userId"],
            "lastUpdatedById": config["userId"],
        }
        api.post.assert_called_once_with(
            url=f"{config['domain']}/api/componentmapping",
            headers=headers,
            data=json.dumps(expected_data),
        )

        # Check that a warning message is logged if the status code is not 200
        mock_response.status_code = 400
        create_component_mapping(api, config, ssp_id, cmp_id)
        assert "Unable to post Mapping Response" in caplog.records[-1].getMessage()

    @pytest.mark.skip("Not implemented")
    def test_create_port(self):
        api = self.api
        data = {
            "parentId": 1016,
            "parentModule": "components",
            "startPort": 80,
            "endPort": 80,
            "protocol": "http",
            "service": "",
            "purpose": "service",
            "used_by": "",
            "createdById": None,
        }
        mock_response = create_mock_response(json_data=data)

        # Mock the API object and its `post` method to return the mock response
        api.post = MagicMock(return_value=mock_response)

        # Define the input parameters
        config = self.config
        headers = {
            "authorization": config["token"],
            "accept": self.accept,
            "Content-Type": self.content_type,
        }
        ports = PortsProtocol(**data)
        create_port(api, config, ports)

        api.post.assert_called_once_with(
            url=f"{config['domain']}/api/portsProtocols",
            headers=headers,
            data=json.dumps(ports.__dict__),
        )

        assert api.post.call_count == 1

    @pytest.mark.skip("Not implemented")
    def test_create_ssp_components(self):
        # Create mock response data
        api = self.api
        data = {
            "id": 280,
            "title": "This System",
            "description": "The entire system as depicted in the system authorization boundary\\n\\nEmail is employed",
            "componentType": "compliance artifact",
            "componentOwnerId": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
            "purpose": None,
            "securityPlansId": 280,
            "cmmcAssetType": None,
            "createdBy": None,
            "createdById": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
            "dateCreated": None,
            "lastUpdatedBy": None,
            "lastUpdatedById": "62b632b9-20c0-45c8-87d9-edd34942c3ae",
            "dateLastUpdated": None,
            "status": "Active",
            "uuid": None,
            "componentOwner": None,
            "cmmcExclusion": False,
            "isPublic": True,
        }
        mock_response = create_mock_response(json_data=data)
        api.post = MagicMock(return_value=mock_response)

        config = self.config
        components = [
            {
                "title": "Component 1",
                "type": "Hardware",
                "description": "This is a hardware component",
                "status": {"state": "operational"},
                "protocols": [
                    {
                        "name": "HTTP",
                        "port-ranges": [{"start": "80", "end": "80"}],
                    }
                ],
            },
            {
                "title": "Component 2",
                "type": "Software",
                "description": "This is a software component",
                "status": {"state": "operational"},
            },
        ]
        ssp_id = 280

        # Call the method
        create_ssp_components(api, config, components, ssp_id)

        # Check that the API was called with the expected parameters
        assert api.post.call_count == 5

        # Check the parameters of the first API call (creating Component 1)
        expected_comp1 = Component(
            title="Component 1",
            securityPlansId=ssp_id,
            componentType="Hardware",
            lastUpdatedById=config["userId"],
            createdById=config["userId"],
            cmmcExclusion=False,
            componentOwnerId=config["userId"],
            description="This is a hardware component",
            status="Active",
        ).dict()
        expected_headers = {
            "authorization": config["token"],
            "accept": self.accept,
            "Content-Type": self.content_type,
        }
        api.post.assert_any_call(
            url=config["domain"] + "/api/components",
            headers=expected_headers,
            data=json.dumps(expected_comp1),
        )

        # Check the parameters of the second API call (creating Component 2)
        expected_comp2 = Component(
            title="Component 2",
            securityPlansId=ssp_id,
            componentType="Software",
            lastUpdatedById=config["userId"],
            createdById=config["userId"],
            cmmcExclusion=False,
            componentOwnerId=config["userId"],
            description="This is a software component",
            status="Active",
        ).dict()
        api.post.assert_any_call(
            url=config["domain"] + "/api/components",
            headers=expected_headers,
            data=json.dumps(expected_comp2),
        )

    @staticmethod
    @pytest.mark.skip("Not implemented")
    def test_get_tables(test_document):
        # Create a mock document object with two tables

        # Call the function
        tables = get_tables(test_document)

        # Check the results
        assert len(tables) == 700

    @pytest.mark.skip("Not implemented")
    def test_post_regscale_object(self):
        config = self.config
        api = self.api
        obj = {"id": 1, "name": "test_obj"}
        endpoint = "controlimplementation"
        data = {
            "id": 1385,
            "lastUpdatedById": "62b632b9-20c0-45c8-8",
            "isPublic": False,
            "url": "https://www.whitehouse.com",
            "title": "Safeguarding Against whatevertion (PII)",
            "parentID": 282,
            "parentModule": "securityplans",
        }
        mock_response = create_mock_response(json_data=data)
        api.post = MagicMock(return_value=mock_response)
        post_regscale_object(api, config, obj, endpoint)

        # Verify that api.post() was called with the correct arguments
        api.post.assert_called_once_with(
            config["domain"] + f"/api/{endpoint}",
            headers={
                "authorization": config["token"],
                "accept": self.accept,
                "Content-Type": self.content_type,
            },
            data='{"id": 1, "name": "test_obj"}',
        )

    @staticmethod
    @pytest.mark.skip("Not implemented")
    def test_get_xpath_data_detailed(example_table):
        key = "Security Objective"
        ident = "Confidentiality"
        xpath = "//w:r/w:t"
        count_array = [2, 3, 4]
        tables = [example_table]
        result = get_xpath_data_detailed(tables, key, ident, xpath, count_array)
        assert result["type"] == "Security Objective"
        assert result["nist_ident"] == "Confidentiality"
        assert result["confidentiality"] is None
        assert result["integrity"] is None
        assert result["availability"] is None

    @pytest.mark.skip("Not implemented")
    def test_post_links(self):
        api = self.api
        config = self.config
        document = MagicMock()
        file_path: Path = Path("test_file_path.docx")
        regscale_ssp = {"id": 407}
        data = [{"title": "example link", "link": "www.google.com"}]
        mock_response = create_mock_response(json_data=data)
        api.get = MagicMock(return_value=mock_response)
        post_links(config, api, document, file_path, regscale_ssp, post_embeds=False)

        # Make assertions here
        api.get.assert_called_once_with(
            config["domain"] + "/api/links/getAllByParent/407/securityplans",
        )

    @staticmethod
    @pytest.mark.skip("Not implemented")
    def test_get_text(test_document):
        document = test_document
        full_text = []
        [full_text.append(para.text) for para in document.paragraphs]
        start_header = "Information System Components and Boundaries"
        end_header = "Types of Users"
        expected_output = "RegScale leverages Microsoft Azure commercial services to provide a Software and SaaS (software as a Service) Figure 9-1 "
        assert get_text_between_headers(full_text, start_header, end_header) == expected_output

    @staticmethod
    @pytest.mark.skip("Not implemented")
    def test_base_contact(test_document):
        expected_output = {
            "name": "",
            "title": "",
            "company": "",
            "address": "",
            "phone": "",
            "email": "",
        }
        assert get_base_contact(test_document) == expected_output

    @staticmethod
    def test_implementation_status_fully_implemented(element, implementation):
        element.set("name", "implementation-status")
        element.set("value", "implemented")
        ControlImplementation._update_implementation_status(element, implementation)
        assert implementation.status == ControlImplementationStatus.FullyImplemented
        assert implementation.bStatusImplemented is True

    @staticmethod
    def test_implementation_status_invalid(element, implementation):
        element.set("name", "implementation-status")
        element.set("value", "unknown-status")
        ControlImplementation._update_implementation_status(element, implementation)
        assert implementation.status is None

    @staticmethod
    def test_control_origination_provider(element, implementation):
        element.set("name", "control-origination")
        element.set("value", "sp-corporate")
        ControlImplementation._update_implementation_status(element, implementation)
        assert implementation.responsibility == ControlImplementationOrigin.Provider
        assert implementation.bServiceProviderCorporate is True

    @staticmethod
    def test_control_origination_invalid(element, implementation):
        element.set("name", "control-origination")
        element.set("value", "invalid-origin")
        ControlImplementation._update_implementation_status(element, implementation)
        assert implementation.responsibility is None

    @staticmethod
    def test_extract_parts():
        # test case for the extract_parts function
        # test case 1
        # when the input is a string
        # the function should return a Dict of the individual parts
        test_dict = {
            "T-1": "<p>Part a:</p><p>The FedRAMP Compliance Director is responsible for developing,documenting, and disseminating system-level security policies andprocedures, as well as ensuring that the proper NIST SP 800-53, Revision5, high controls are implemented to protect the confidentiality,integrity, and availability of the CC2CGS information system and itsdata. All revisions to policies and procedures are reviewed and approvedby the FedRAMP Compliance Director. Upon final approval, policies andprocedures are published to Amazon S3 and made available to employees onan as-needed basis.</p><p>The access control policy is comprised of the following sections:</p><ul><li><p>Introduction</p></li><li><p>Purpose</p></li><li><p>Scope</p></li><li><p>Roles and Responsibilities</p></li><li><p>Management Commitment</p></li><li><p>Authority</p></li><li><p>Document Control</p></li><li><p>Compliance</p></li><li><p>Policy Requirements</p></li></ul><p>The access control policy is consistent with:</p><ul><li><p>Applicable laws</p></li><li><p>Executive orders</p></li><li><p>Directives</p></li><li><p>Regulations</p></li><li><p>Policies</p></li><li><p>Standards</p></li><li><p>Guidelines</p></li></ul><p>The system-level policy provides the security framework upon whichall access control efforts will be based. The provisions of the policyand procedures pertain to all Test Case employees, contingent workers,contractors, business partners, and third parties accessing all CC2CGScomponents that process, transmit, and store government information.This includes information held on behalf of government customers,partners, and other third parties, to protect the confidentiality,integrity, and availability of Test Case confidential information andsystems. Operational systems and systems under development must meetaccess control requirements commensurate with the sensitivity of theinformation they handle and the potential impact of unauthorizeddisclosure of information and data.</p><p>Part b:</p><p>Test Case has designated the FedRAMP Compliance Director to managethe development, documentation, and dissemination of the access controlpolicy and procedures.</p><p>Part c:</p><p>The FedRAMP Compliance Director reviews and updates the accesscontrol policy and procedures at least annually. Policy and procedurerevisions may also be carried out when required due to significantchanges in the environment.</p>",
            "T-2": "<p>Part a:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>Test Case has identified the following types of information systemaccounts that support its organizational mission and business functionsrelated to the CC2CGS federal system:</p><ul><li><p>Microsoft Entra ID/ Microsoft Entra ID accounts – Test Case usesMicrosoft Entra ID as a Security Assertion Markup Language (SAML) andMFA provider to centrally control access to the CC2CGS federalproduction environment. All Test Case team members who administer andmanage the production environment are provisioned Microsoft Entra IDaccounts, which serve as an SSO mechanism for backend access via AmazonWorkspaces to the production environment. Access to the CC2CGS requirespersonnel to have Microsoft Entra ID credentials, as well as FIPS 5CYubiKey Federal MFA. Authorization and access are determined during theaccount provisioning process. Personnel are given access based on theirroles and responsibilities in accordance with separation of duties andleast privileged principles.</p></li><li><p>AWS root/ IAM accounts – The AWS root account for the CC2CGSenvironment is managed by the SRE Team and is not used for day-to-dayoperations. Instead, IAM accounts are used by authorized personnel andare separated using groups and roles based on their jobfunctions.</p></li><li><p>Security tools and applications – The SRE Team access third-partytools such as Trend Micro Deep Security Manager (DSM) and SplunkEnterprise with Enterprise Security using a secured HTTPS web connectionand a combination of username and password. These accounts are managedby the SRE Team.</p></li><li><p>Initial customer account: After a new customer signs a contract,SRE Team creates the initial domain administrator account and emails thecustomer administrator with a link to the account. The customer domainadministrator is responsible for setting up the rest of theirorganization's accounts.</p></li></ul><p><strong><u>Federal Customer Responsibility</u></strong></p><p>Once the initial administrator account is set up by Test Case,customers are responsible for establishing conditions for rolemembership for their CC2CGS accounts.</p><p>Part b:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>Test Case assigns account managers for all internal informationsystem accounts. The SRE Team and Customer Support Team own andadminister the Microsoft Entra ID, AWS IAM accounts, and security tools.The SRE Team is considered the owner of the AWS root account. The SRETeam and Customer Support Team set up the initial Customer Administratoraccount.</p><p><strong>Federal Customer Responsibility</strong></p><p>Customers are responsible for designating an account owner to beprovided the Customer Administrator account for their access to CC2CGS.After the initial Customer Administrator account is created byTest Case, the customer administrator is responsible for creating theCC2CGS accounts for the individual users within theirorganization.</p><p>Part c:</p><p><strong>Test Case Responsibility</strong></p><p>Access requests are submitted via Jira Service Management. All accessto CC2CGS (both the Application and the Management plane) is assignedusing role-based access control (RBAC). A user's role determines theirassigned permissions, privileges, and access within the CC2CGSenvironment. Users must complete prerequisites based on their role, suchas signed access agreements, completion of applicable trainingscompleting incident response training, or signing the rules of behavior(RoB), prior to being granted access to CC2CGS</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>Customers are responsible for establishing conditions for rolemembership for their CC2CGS user accounts.</p><p>Part d:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>Test Case has developed an access request process for approval andcreation of all internal CC2CGS accounts. Each request must be based ona business need, and access is limited to the access required for theindividual to perform their identified role within the system. Allrequests for account creation, modification, and disabling areinitiated, approved, and tracked in a Jira Service Management ticket.All users are assigned an individual account to access theenvironment.</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>Once Test Case creates the initial Customer Administrator account inCC2CGS, the Customer Administrator is responsible for creating andmanaging CC2CGS accounts for the individual users within theirorganization and for developing a process to identify users who requireaccess to the system and their associated levels ofpermissions.</p><p>Part e:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>All new Test Case employees are provided with basic access to email,corporate information technology (IT) tools, and other Test Caseresources, as part of the initial onboarding process. For personnel whorequire access to the CC2CGS production environment, SRE Team andCustomer Support Team must submit a Jira Service Management ticket witha workflow is used to obtain authorization from the employee'ssupervisor and the account owner.</p><p>Federal Customers are onboarded by SRE Team and Customer Support Teampersonnel only after the customer has signed a contract. The SRE Teamand the Customer Support Team create the initial Customer Administratoraccount with a randomly generated password and email the CustomerAdministrator with a link to their account. The Customer Administratormust change the initial password upon first login. Accounts are approvedand signed off by the SRE Team and Customer Support Team.</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>After the initial Customer Administrator account is created byTest Case, the Customer Administrator is responsible for creating theCC2CGS accounts for the individual users within their organization,vendors, and partners and for ensuring that all access isapproved.</p><p>Part f:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>For Test Case personnel who require access to the CC2CGS productionenvironment, a standard access request process is used, in which a JiraService Management ticket with a workflow is used to obtainauthorization from the employee's supervisor and the account owner.Ticket information specifies the user, identifier, role(s) requested,and approver. Each request must be based on a business need, and accessis limited to the access required for that individual to perform theiridentified role within the system.</p><p>When an existing employee is terminated or leaves the company, thereporting manager notifies HR, and HR begins the off-boarding process inthe corporate HR management system. A notification is sent to IT Teamwho then disables all accounts on the date indicated in the off-boardingticket.</p><p>Federal Customers are onboarded by the SRE Team and Customer SupportTeam only after the customer has signed a contract. The SRE Team andCustomer Support Team create an initial customer administrator accountin CC2CGS and email the user ID and initial password to the customer.Accounts are created in accordance with the CC2CGS access control policyand procedure.</p><p>If an update needs to be made to an account, a request is submittedto the SRE Team and Customer Support Team through Jira ServiceManagement. The SRE Team and Customer Support Team review the requestand update the account if approved. If there is an update in the user'sMicrosoft Entra ID, the SRE Team and Customer Support Team ensures theupdate is reflected within the account as necessary.</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>After the initial Customer Administrator account is created byTest Case, the Customer Administrator is responsible for creating theCC2CGS accounts for the individual users and for modifying, disabling,and removing those accounts in accordance with their organizationalpolicies and procedures.</p><p>Part g:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>Logging and monitoring of account activity in the CC2CGS environmentis performed using Splunk Enterprise Security. Alerts are set up inSplunk Enterprise Security for specific unauthorized actions, andnotifications are sent to on-call personnel if threats are detected.</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>After the initial Customer Administrator account is created byTest Case, the Customer Administrator is responsible for creating theCC2CGS accounts for the individual users and for monitoring the useraccounts they create within CC2CGS.</p><p>Part h:</p><p><strong>Test Case Responsibility</strong></p><p>When an existing employee is terminated, transferred, or leavesTest Case, the reporting manager notifies HR within 24 hours, and HRbegins the off-boarding process. A notification is also sent to the SRETeam and the Customer Support Team, who then disables all accountswithin eight (8) hours in the off-boarding ticket. All access isdisabled on the employee's eight (8) hours as a part of Test Case or ifthey transfer to another role within Test Case. Additionally,notifications are sent to SRE Team, Customer Support Team, FedRAMPCompliance Director when an account or role is no longer required withinCC2CGS.</p><p><strong>Federal Customer Responsibility</strong></p><p>Customer account owners are responsible for implementing processes toensure that they are notified when user access is no longer required,when a user is transferred or terminated, or when an individual'sneed-to-know changes.</p><p>Part i:</p><p><strong>Test Case Responsibility</strong></p><p>All CC2CGS Management plane personnel are granted a Microsoft EntraID account and a FIPS 5C YubiKey Federal MFA account. All accountrequests are submitted through the Test Case access request form andtracked via the Jira Service Management ticketing system. Account accessis based on user roles and responsibilities and must be approved by theSRE Team and the Customer Support Team.</p><p><strong>Federal Customer Responsibility</strong></p><p>Customers are responsible for authorizing access to their CC2CGSaccounts based on valid access authorization, intended system usage, andother attributes defined by the customer organization.</p><p>Part j:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>Test Case assigns account managers for all internal informationsystem accounts. The SRE Team and Customer Support Team own andadminister the Microsoft Entra ID, AWS IAM accounts, and security tools.The SRE Team and Customer Support Team set up the initial CustomerAdministrator account. The FedRAMP Compliance Director performs accountreviews at least monthly for privileged access and every six (6) monthsfor non-privileged access.</p><p><strong>Federal Customer Responsibility</strong></p><p>Customers are responsible for designating an account owner to beprovided with the Customer Administrator account for their access to theCC2CGS system. After the initial Customer Administrator account iscreated by Test Case, the customer administrator is responsible forcreating the CC2CGS accounts for the individual users within theirorganization and performs account reviews at least monthly forprivileged access and every six (6) months for non-privilegedaccess.</p><p>Part k:</p><p><strong>Test Case Responsibility</strong></p><p>Shared credentials are not used within CC2CGS. The SRE Team isconsidered the owner of the AWS root account and Test Case prohibitsthe user of shared or group accounts through the access control policy.All users must have unique credentials to access the CC2CGS productionenvironment.</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>Customers are responsible for determining whether shared or groupaccounts will be used and for determining a process for reissuing sharedor group account credentials when individuals are removed from thegroup, if applicable.</p><p>Part l:</p><p><strong><u>Test Case Responsibility</u></strong></p><p>Test Case integrates the personnel termination and transferprocesses with the account management process for CC2CGS. Forterminations and transfers, HR begins the process to have access removedor updated by SRE Team and Customer Support Team by submitting a JiraService Management ticket.</p><p><strong><u>Federal Customer Responsibility</u></strong></p><p>Customers are responsible for aligning account management processeswith personnel termination and transfer processes.</p>",
            "T-3": "test one part or default value for one part",
        }
        test_1 = test_dict.get("T-1")
        test_2 = test_dict.get("T-2")
        test_3 = test_dict.get("T-3")
        dict_test_1 = extract_parts(test_1)
        dict_test_2 = extract_parts(test_2)
        dict_test_3 = extract_parts(test_3)
        # test case 1
        expected_t1_keys = ["Part a", "Part b", "Part c"]
        assert list(dict_test_1.keys()) == expected_t1_keys

        # test case 2
        # assert each key in dict_test_2 is equal to the expected value
        expected_t2_keys = [
            "Part a",
            "Part b",
            "Part c",
            "Part d",
            "Part e",
            "Part f",
            "Part g",
            "Part h",
            "Part i",
            "Part j",
            "Part k",
            "Part l",
        ]
        assert list(dict_test_2.keys()) == expected_t2_keys

        # test case 3
        assert dict_test_3 == {"default": "test one part or default value for one part"}

    def test_parse_ips(self):
        """
        Test parsing v6 and v4 IP addresses
        """
        test_ips: list[dict] = [
            {
                "ip": "23.45.55.55",
                "asset_field": "ipAddress",
                "valid_ip": True,
            },
            {
                "ip": "10.25.255.191",
                "asset_field": "ipAddress",
                "valid_ip": True,
            },
            {
                "ip": "23.45.265.254",
                "asset_field": "ipAddress",
                "valid_ip": False,
            },
            {
                "ip": "45,34,54,56",
                "asset_field": "ipAddress",
                "valid_ip": False,
            },
            {
                "ip": "0.42.34.76",
                "asset_field": "ipAddress",
                "valid_ip": False,
            },
            {
                "ip": "2",
                "asset_field": "ipAddress",
                "valid_ip": False,
            },
            {
                "ip": "1.2.3",
                "asset_field": "ipAddress",
                "valid_ip": False,
            },
            {
                "ip": "255.256.267.300",
                "asset_field": "ipAddress",
                "valid_ip": False,
            },
            {
                "ip": "2001:db8:3333:4444:5555:6666:7777:8888",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "1050:0000:0000:0000:0005:0600:300c:326b",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "1050:0:0:0:5:600:300c:326b",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:db8::",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:db8::1234:5678",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "::1234:5678",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:db8:a0b:12f0::::0:1",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "::ffff:192.0.2.128",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:0db8:85a3:0000:0000:8a2e:0370:zzzz",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "2001:db8:85a3::8a2e:370g:7334",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "2001::85a3::8a2e:370:7334",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "12345::85a3:8a2e:0370:7334",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "::ffff:192.0.2.256",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "2001:0db8:85a3::8a2e::7334",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "::",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:db8:85a3:8d3:1319:8a2e:370:7348",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:db8:85a3:8d3:1319:8a2e:370g:7348",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "abcd:ef01:2345:6789:abcd:ef01:2345:6789",
                "asset_field": "iPv6Address",
                "valid_ip": True,
            },
            {
                "ip": "2001:db8:1234:5678:abcd:ef01:2345:12345",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
            {
                "ip": "1200::AB00:1234::2552:7777:1313",
                "asset_field": "iPv6Address",
                "valid_ip": False,
            },
        ]

        for test in test_ips:
            self.logger.info(f"Testing IP: {test['ip']}")
            asset_field = determine_ip_address_version(test["ip"])
            self.logger.info(f"Asset Field: {asset_field}")
            if test["valid_ip"]:
                assert asset_field == test["asset_field"]
            else:
                assert asset_field is None

    @staticmethod
    def test_update_ip_addresses_ipv4():
        asset = {}
        inventory = {"IPV4_OR_IPV6_ADDRESS": "192.168.1.1"}
        mapping = MagicMock()
        mapping.get_value.return_value = "192.168.1.1"

        _update_ip_addresses(asset, inventory, mapping)

        assert asset["ipAddress"] == "192.168.1.1"
        assert asset["iPv6Address"] is None

    @staticmethod
    def test_update_ip_addresses_ipv6():
        asset = {}
        inventory = {"IPV4_OR_IPV6_ADDRESS": "2001:db8:3333:4444:5555:6666:7777:8888"}
        mapping = MagicMock()
        mapping.get_value.return_value = "2001:db8:3333:4444:5555:6666:7777:8888"

        _update_ip_addresses(asset, inventory, mapping)

        assert asset["ipAddress"] is None
        assert asset["iPv6Address"] == "2001:db8:3333:4444:5555:6666:7777:8888"

    @staticmethod
    def test_update_ip_addresses_multiple_addresses():
        asset = {}
        inventory = {"IPV4_OR_IPV6_ADDRESS": "192.168.1.1, 2001:db8:3333:4444:5555:6666:7777:8888"}
        mapping = MagicMock()
        mapping.get_value.return_value = "192.168.1.1, 2001:db8:3333:4444:5555:6666:7777:8888"

        _update_ip_addresses(asset, inventory, mapping)

        assert asset["ipAddress"] == "192.168.1.1"
        assert asset["iPv6Address"] == "2001:db8:3333:4444:5555:6666:7777:8888"

    @staticmethod
    def test_update_ip_addresses_multiple_addresses_with_semicolons():
        asset = {}
        inventory = {"IPV4_OR_IPV6_ADDRESS": "192.168.1.1; 2001:db8:3333:4444:5555:6666:7777:8888"}
        mapping = MagicMock()
        mapping.get_value.return_value = "192.168.1.1; 2001:db8:3333:4444:5555:6666:7777:8888"

        _update_ip_addresses(asset, inventory, mapping)

        assert asset["ipAddress"] == "192.168.1.1"
        assert asset["iPv6Address"] == "2001:db8:3333:4444:5555:6666:7777:8888"

    @staticmethod
    def test_update_ip_addresses_multiple_addresses_with_commas_and_semicolons():
        asset = {}
        inventory = {
            "IPV4_OR_IPV6_ADDRESS": "192.168.1.1, 2001:db8:3333:4444:5555:6666:7777:8888; 192.168.1.2; 2001:db8:3333:4444:5555:6666:7777:8888"
        }
        mapping = MagicMock()
        mapping.get_value.return_value = (
            "192.168.1.1, 2001:db8:3333:4444:5555:6666:7777:8888; 192.168.1.2; 2001:db8:3333:4444:5555:6666:7777:8888"
        )

        _update_ip_addresses(asset, inventory, mapping)

        assert asset["ipAddress"] == "192.168.1.1, 192.168.1.2"
        assert asset["iPv6Address"] == "2001:db8:3333:4444:5555:6666:7777:8888, 2001:db8:3333:4444:5555:6666:7777:8888"

    @staticmethod
    def test_multiple_ip_addresses_with_newlines():
        asset = {}
        inventory = {"IPV4_OR_IPV6_ADDRESS": "192.168.1.1\n2001:db8:3333:4444:5555:6666:7777:8888"}
        mapping = MagicMock()
        mapping.get_value.return_value = "192.168.1.1\n2001:db8:3333:4444:5555:6666:7777:8888"

        _update_ip_addresses(asset, inventory, mapping)

        assert asset["ipAddress"] == "192.168.1.1"
        assert asset["iPv6Address"] == "2001:db8:3333:4444:5555:6666:7777:8888"

    @staticmethod
    def test_update_ip_addresses_empty():
        asset = {}
        inventory = {"IPV4_OR_IPV6_ADDRESS": ""}
        mapping = MagicMock()
        mapping.get_value.return_value = ""

        _update_ip_addresses(asset, inventory, mapping)

        assert asset.get("ipAddress") is None
        assert asset.get("iPv6Address") is None

    @staticmethod
    def test_ssp_parser_vertical_table_single_header():
        """Test vertical table detection with single header (original behavior)."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        xml_content = b"""
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tbl>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>System Owner Information</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Name</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>John Doe</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Title</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>System Owner</w:t></w:r></w:p></w:tc>
                </w:tr>
            </w:tbl>
        </w:document>
        """

        parser = SSPDocParser.__new__(SSPDocParser)
        parser.doc = MagicMock()
        parser.doc.element.body = []  # No preceding text needed for this test

        tables = parser.parse_xml_for_tables(xml_content)

        assert len(tables) == 1
        table_data = tables[0]["table_data"]
        assert len(table_data) == 2
        assert table_data[0] == {"Name": "John Doe"}
        assert table_data[1] == {"Title": "System Owner"}

    @staticmethod
    def test_ssp_parser_vertical_table_two_headers_empty_second():
        """Test vertical table detection with two headers where second is empty (new fix)."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        xml_content = b"""
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tbl>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>System Information</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t></w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>CSP Name:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Public Safety Platform (PSP)</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>CSO Name:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Mark43, Inc.</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Service Model:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Software-as-a-Service (SaaS)</w:t></w:r></w:p></w:tc>
                </w:tr>
            </w:tbl>
        </w:document>
        """

        parser = SSPDocParser.__new__(SSPDocParser)
        parser.doc = MagicMock()
        parser.doc.element.body = []  # No preceding text needed for this test

        tables = parser.parse_xml_for_tables(xml_content)

        assert len(tables) == 1
        table_data = tables[0]["table_data"]
        assert len(table_data) == 3
        assert table_data[0] == {"CSP Name:": "Public Safety Platform (PSP)"}
        assert table_data[1] == {"CSO Name:": "Mark43, Inc."}
        assert table_data[2] == {"Service Model:": "Software-as-a-Service (SaaS)"}

    @staticmethod
    def test_ssp_parser_vertical_table_two_headers_whitespace_second():
        """Test vertical table detection with two headers where second is only whitespace."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        xml_content = b"""
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tbl>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>System Information</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>   </w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>FedRAMP Package ID:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>FR2235965777</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>FIPS PUB 199 Level:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>High</w:t></w:r></w:p></w:tc>
                </w:tr>
            </w:tbl>
        </w:document>
        """

        parser = SSPDocParser.__new__(SSPDocParser)
        parser.doc = MagicMock()
        parser.doc.element.body = []  # No preceding text needed for this test

        tables = parser.parse_xml_for_tables(xml_content)

        assert len(tables) == 1
        table_data = tables[0]["table_data"]
        assert len(table_data) == 2
        assert table_data[0] == {"FedRAMP Package ID:": "FR2235965777"}
        assert table_data[1] == {"FIPS PUB 199 Level:": "High"}

    @staticmethod
    def test_ssp_parser_horizontal_table_not_detected_as_vertical():
        """Test that horizontal tables with meaningful headers are not detected as vertical."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        xml_content = b"""
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tbl>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Date</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Description</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Version</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>05/14/2024</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Initial version</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>1.0</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>11/01/2024</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Updated version</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>2.0</w:t></w:r></w:p></w:tc>
                </w:tr>
            </w:tbl>
        </w:document>
        """

        parser = SSPDocParser.__new__(SSPDocParser)
        parser.doc = MagicMock()
        parser.doc.element.body = []  # No preceding text needed for this test

        tables = parser.parse_xml_for_tables(xml_content)

        assert len(tables) == 1
        table_data = tables[0]["table_data"]
        assert len(table_data) == 2
        # Should be horizontal format with named columns
        assert table_data[0] == {"Date": "05/14/2024", "Description": "Initial version", "Version": "1.0"}
        assert table_data[1] == {"Date": "11/01/2024", "Description": "Updated version", "Version": "2.0"}

    @staticmethod
    def test_ssp_parser_two_column_table_non_vertical():
        """Test that 2-column tables with meaningful second header are not detected as vertical."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        xml_content = b"""
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tbl>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Security Objective</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Level</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Confidentiality</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>High</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Integrity</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>High</w:t></w:r></w:p></w:tc>
                </w:tr>
            </w:tbl>
        </w:document>
        """

        parser = SSPDocParser.__new__(SSPDocParser)
        parser.doc = MagicMock()
        parser.doc.element.body = []  # No preceding text needed for this test

        tables = parser.parse_xml_for_tables(xml_content)

        assert len(tables) == 1
        table_data = tables[0]["table_data"]
        assert len(table_data) == 2
        # Should be horizontal format since second header is not empty
        assert table_data[0] == {"Security Objective": "Confidentiality", "Level": "High"}
        assert table_data[1] == {"Security Objective": "Integrity", "Level": "High"}

    @staticmethod
    def test_ssp_parser_vertical_table_first_header_not_in_list():
        """Test that tables with first header not in vertical_tables list are not detected as vertical."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        xml_content = b"""
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tbl>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Random Table Header</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t></w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Field One:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Value One</w:t></w:r></w:p></w:tc>
                </w:tr>
                <w:tr>
                    <w:tc><w:p><w:r><w:t>Field Two:</w:t></w:r></w:p></w:tc>
                    <w:tc><w:p><w:r><w:t>Value Two</w:t></w:r></w:p></w:tc>
                </w:tr>
            </w:tbl>
        </w:document>
        """

        parser = SSPDocParser.__new__(SSPDocParser)
        parser.doc = MagicMock()
        parser.doc.element.body = []  # No preceding text needed for this test

        tables = parser.parse_xml_for_tables(xml_content)

        assert len(tables) == 1
        table_data = tables[0]["table_data"]
        assert len(table_data) == 2
        # Should be horizontal format since "Random Table Header" is not in vertical_tables list
        assert table_data[0] == {"Random Table Header": "Field One:", "": "Value One"}
        assert table_data[1] == {"Random Table Header": "Field Two:", "": "Value Two"}

    @staticmethod
    def test_ssp_parser_all_vertical_table_types():
        """Test all table types that should be detected as vertical."""
        from regscale.integrations.public.fedramp.docx_parser import SSPDocParser
        from unittest.mock import MagicMock

        vertical_table_headers = [
            "Identification of Organization that Prepared this Document",
            "Identification of Cloud Service Provider",
            "System Owner Information",
            "System Information",
            "System Component Information",
            "ISSO (or Equivalent) Point of Contact",
        ]

        for header in vertical_table_headers:
            # Test with single header
            xml_content = f"""
            <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:tbl>
                    <w:tr>
                        <w:tc><w:p><w:r><w:t>{header}</w:t></w:r></w:p></w:tc>
                    </w:tr>
                    <w:tr>
                        <w:tc><w:p><w:r><w:t>Field:</w:t></w:r></w:p></w:tc>
                        <w:tc><w:p><w:r><w:t>Value</w:t></w:r></w:p></w:tc>
                    </w:tr>
                </w:tbl>
            </w:document>
            """.encode()

            parser = SSPDocParser.__new__(SSPDocParser)
            parser.doc = MagicMock()
            parser.doc.element.body = []  # No preceding text needed for this test

            tables = parser.parse_xml_for_tables(xml_content)

            assert len(tables) == 1
            table_data = tables[0]["table_data"]
            assert len(table_data) == 1
            assert table_data[0] == {"Field:": "Value"}

            # Test with two headers (empty second)
            xml_content_two_headers = f"""
            <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:tbl>
                    <w:tr>
                        <w:tc><w:p><w:r><w:t>{header}</w:t></w:r></w:p></w:tc>
                        <w:tc><w:p><w:r><w:t></w:t></w:r></w:p></w:tc>
                    </w:tr>
                    <w:tr>
                        <w:tc><w:p><w:r><w:t>Field:</w:t></w:r></w:p></w:tc>
                        <w:tc><w:p><w:r><w:t>Value</w:t></w:r></w:p></w:tc>
                    </w:tr>
                </w:tbl>
            </w:document>
            """.encode()

            parser2 = SSPDocParser.__new__(SSPDocParser)
            parser2.doc = MagicMock()
            parser2.doc.element.body = []  # No preceding text needed for this test

            tables2 = parser2.parse_xml_for_tables(xml_content_two_headers)

            assert len(tables2) == 1
            table_data2 = tables2[0]["table_data"]
            assert len(table_data2) == 1
            assert table_data2[0] == {"Field:": "Value"}
