#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""standard python imports"""

import dataclasses
import json
import os
import re
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime
from io import StringIO
from pathlib import Path
from tempfile import gettempdir
from threading import Lock
from typing import Any, Dict, List, Optional, Tuple, Type, Union
from urllib.parse import urljoin

import click
import requests
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.table import Table
from lxml import etree
from pydantic import BaseModel
from ssp import SSP

from regscale.core.app.api import Api
from regscale.core.app.application import Application
from regscale.core.app.utils.app_utils import (
    capitalize_words,
    check_file_path,
    download_file,
    error_and_exit,
    get_current_datetime,
)
from regscale.integrations.public.fedramp.parameter_utils import format_parameter_name
from regscale.integrations.public.fedramp.ssp_logger import SSPLogger
from regscale.models.regscale_models import (
    Component,
    ControlImplementation,
    ControlParameter,
    File,
    InterConnection,
    LeveragedAuthorization,
    Parameter,
    PortsProtocol,
    Privacy,
    ProfileMapping,
    Requirement,
    SecurityControl,
    SecurityPlan,
    SystemRole,
)
from regscale.models.regscale_models.control_implementation import ControlImplementationStatus

ssp_logger = SSPLogger()
logger = ssp_logger

namespaces = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}
DATE_FORMAT = "%Y-%m-%d %H:%M:%S"
NEW_LINE_OUTPUT = "\n------------------------------\n"
SYSTEM_TYPE = "Major Application"
SYSTEM_STATUS = "System Status"
SERVICE_ARCHS = "Service Provider Architecture Layers"
DEPLOY_MODEL = "Service Provider Cloud Deployment Model"
END_MARKER = "!!!"
SSP_URL_SUFFIX = "/api/securityplans/getList"
XPATH_TAG = "//w:r/w:t"
TABLE_TAG = "//w:tbl/w:tr"
ORGANIZATION_TAG = "Organization Name"
CONTROL_ID = "Control ID"
TBD = "To be determinned"
IMPACT_LEVEL = "Impact Level"
YES = "Yes"
CSP = "CSP"

# Checkbox characters used in FedRAMP documents to indicate selections
# Includes various checkbox symbols from different fonts and rendering systems
CHECKBOX_CHARS = ["☒", "☑", "☑️", "✓", "✔", "✔️", "✅", "⬜", "▣", "■", "□", "⊠", "⊗", "×"]


def decode_access_level(key: str) -> str:
    """
    Decodes the access level from the FedRAMP document

    :param str key: Key used to decode the access level
    :return: Access level as a string
    :rtype: str
    """
    access_levels = {
        "P": "Privileged",
        "NP": "Non-Privileged",
        "NLA": "No Logical Access",
    }
    return access_levels.get(key, "Non-Privileged")


def create_responsible_roles(app: Application, table_data: list, ssp_id: int) -> None:
    """
    [BETA] Inserts the actual the Responsible Roles into the Security Plan.

    :param Application app: Application object
    :param list table_data: list of dicts
    :param int ssp_id: RegScale SSP ID
    :rtype: None
    """
    na_text = ControlImplementationStatus.NA
    roles = [table for table in table_data if "Role" in table.keys() and "Internal or External" in table.keys()]
    logger.info(
        event_msg=f"Found {len(roles)} Responsible Roles",
        record_type="role",
        model_layer="system-roles",
    )
    user_id = app.config.get("userId")
    with ThreadPoolExecutor(max_workers=20) as executor:
        futures = []
        for role in roles:
            try:
                access_level = decode_access_level(
                    role.get(
                        "Privileged (P), Non-Privileged (NP), or No Logical Access (NLA)",
                        "Unknown",
                    )
                )
                future = executor.submit(
                    SystemRole.get_or_create,
                    app=app,
                    role_name=role.get("Role"),
                    ssp_id=ssp_id,
                    roleType=role.get("Internal or External", "Internal"),
                    accessLevel=access_level,
                    sensitivityLevel=role.get("Sensitivity Level", na_text),
                    assignedUserId=user_id,
                    privilegeDescription=role.get("Authorized Privileges", na_text),
                    securityPlanId=ssp_id,
                    functions=role.get("Functions Performed", na_text),
                    createdById=user_id,
                    logger=logger,
                )
                futures.append(future)
            except Exception as e:
                logger.error(
                    f"Failed to create Responsible Roles with error: {str(e)}",
                    record_type="role",
                    model_layer="system-roles",
                )
        for future in as_completed(futures):
            try:
                future.result()
            except Exception as e:
                logger.error(
                    f"Failed to create Responsible Roles with error: {str(e)}",
                    record_type="role",
                    model_layer="system-roles",
                )

    logger.info(
        "Successfully Created Responsible Roles",
        record_type="role",
        model_layer="system-roles",
    )


def assign_role_to_control(
    control: Any,
    system_role: dict,
    ctrl_roles: dict,
    ctrl_roles_lock: Lock,
) -> None:
    """
    Assign control_roles using the system_role specified

    :param Any control: The control to process
    :param dict system_role: The system role to assign to control
    :param dict ctrl_roles: A dict of control roles
    :param Lock ctrl_roles_lock: A lock to protect the shared resource
    :rtype: None
    """
    friendly_control_id = get_friendly_control_id(control)
    with ctrl_roles_lock:  # Acquire the lock to modify shared resource
        if friendly_control_id in ctrl_roles:
            ctrl_roles[friendly_control_id].append(system_role["id"])
        else:
            ctrl_roles[friendly_control_id] = [system_role["id"]]


def process_role(
    value: Any,
    control: Any,
    unique_values: set,
    system_roles: List,
    app: Application,
    ssp_id: int,
    ctrl_roles: dict,
    ctrl_roles_lock: Lock,
) -> None:
    """
    Process the Responsible Role

    :param Any value: The value to process
    :param Any control: The control to process
    :param set unique_values: A set of unique values
    :param List system_roles: A list of system roles
    :param Application app: The application object
    :param int ssp_id: The SSP ID
    :param dict ctrl_roles: A dict of control roles
    :param Lock ctrl_roles_lock: A lock to protect the shared resource
    :rtype: None
    """
    if type(value) is str and value.startswith("Responsible Role:"):
        role = value.split(":", 1)[1].strip()

        # Handle case whwere there are multiples comma delimited
        myrolelist = role.split(",")

        for role in myrolelist:
            role = role.strip(":")

            if role.lower() not in unique_values:
                unique_values.add(role.lower())
                system_roles.append(role.strip())

            system_role = SystemRole.get_or_create(
                app=app,
                role_name=role.strip(),
                ssp_id=ssp_id,
                roleType="Internal",
                accessLevel="Privileged",
                sensitivityLevel=ControlImplementationStatus.NA,
                assignedUserId=app.config.get("userId"),
                privilegeDescription=role,
                securityPlanId=ssp_id,
                functions=role,
                createdById=app.config.get("userId"),
                logger=logger,
            )

            if isinstance(system_role, SystemRole):
                system_role = system_role.dict()

            if control:
                assign_role_to_control(
                    control=control, system_role=system_role, ctrl_roles=ctrl_roles, ctrl_roles_lock=ctrl_roles_lock
                )


def post_responsible_roles(app: Application, table_data: list, ssp_id: int) -> dict:
    """
    [BETA] Insert the Responsible Roles into the Security Plan

    :param Application app: Application object
    :param list table_data: list of dicts
    :param int ssp_id: RegScale SSP ID
    :return: dict of the control to role mappings
    :rtype: dict
    """
    data = [table for table in table_data if "Control Summary Information" in table.keys()]
    system_roles = list()

    unique_values = set(system_roles)
    ctrl_roles = dict()
    ctrl_roles_lock = Lock()  # Create a lock to protect shared resource

    for obj in data:
        try:
            control = list(obj.keys())[0] if isinstance(obj, dict) and obj.keys() else None
            for value in obj.values():
                process_role(
                    value,
                    control,
                    unique_values,
                    system_roles,
                    app,
                    ssp_id,
                    ctrl_roles,
                    ctrl_roles_lock,
                )
        except Exception as e:
            logger.error(
                f"Failed to parse Responsible Roles with error: {str(e)}",
                record_type="role",
                model_layer="system-roles",
            )

    return ctrl_roles


def process_fedramp_oscal_ssp(file_path: click.Path, submission_date: date, expiration_date: date) -> None:
    """
    OSCAL FedRAMP to RegScale SSP

    :param click.Path file_path: A click file path object to the oscal file
    :param date submission_date: The Submission date YYYY-MM-DD
    :param date expiration_date: The Expiration date YYYY-MM-DD
    :rtype: None
    """
    app = Application()
    api = Api()
    config = app.config
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            ssp_dict = json.load(file)
    except FileNotFoundError:
        error_and_exit(f"File not found!\n{file_path}")
    except json.JSONDecodeError as jex:
        logger.error("JSONDecodeError, something is wrong with the file: %s\n%s", file_path, jex)

    # Create SSP
    create_ssp(api, config, ssp_dict, submission_date, expiration_date)


def check_profile(api: Api, config: dict, title: str) -> list:
    """
    Check if the profile exists in RegScale

    :param Api api: The api instance
    :param dict config: The application configuration
    :param str title: The title of the profile in question
    :raises: ValueError if the provided title doesn't exist in RegScale
    :return: List of filtered profiles
    :rtype: list
    """
    profiles = []
    profiles_response = api.get(config["domain"] + "/api/profiles/getList")
    if profiles_response.ok:
        profiles = profiles_response.json()
    if filtered := [dat for dat in profiles if dat["name"] == title]:
        return filtered[0]["id"]
    else:
        raise ValueError(
            f"The profile {title} does not exist in RegScale, \
                please create it and re-run this task."
        )


def create_port(api: Api, config: dict, dat: PortsProtocol) -> None:
    """Create a port and protocol for a component

    :param Api api: An api instance
    :param dict config: Configuration
    :param PortsProtocol dat: Port and protocol data
    :rtype: None
    """

    existing_ports = api.get(
        url=config["domain"] + f"/api/portsProtocols/getAllByParent/{dat.parentId}/components",
    ).json()
    if dat not in [PortsProtocol.from_dict(port, True) for port in existing_ports]:
        # Check if obj exists
        port_res = api.post(
            url=config["domain"] + "/api/portsProtocols",
            data=json.dumps(dat.__dict__),
        )
        if port_res.status_code == 200:
            logger.info("Port and Protocol for component %i added!", dat.parentId)
        else:
            logger.warning(
                "Unable to post Port and Protocol: %s.",
                json.dumps(dat),
            )


def create_ssp_components(api: Api, config: dict, components: list[dict], ssp_id: int) -> None:
    """
    Creates SSP Components

    :param Api api: The API instance
    :param dict config: The application's configuration
    :param list[dict] components: The components
    :param int ssp_id: The ID of the SSP in RegScale
    :rtype: None
    """
    component_types = [
        "hardware",
        "software",
        "policy",
        "service",
        "process",
        "procedure",
        "compliance artifact",
    ]
    ports = set()
    for component in components:
        comp_type = component["type"] if component["type"].lower() in component_types else "compliance artifact"
        status = "Inactive/Retired"
        if component["status"]["state"] == "operational":
            status = "Active"

        comp = Component(
            title=component["title"],
            securityPlansId=ssp_id,
            componentType=comp_type,
            lastUpdatedById=config["userId"],
            createdById=config["userId"],
            cmmcExclusion=False,
            componentOwnerId=config["userId"],
            description=component["description"],
            status=status,
        )

        # save component
        cmp_id = None
        url = urljoin(config["domain"], "/api/components")
        cmp_response = api.post(
            url=url,
            json=comp.dict(),
        )
        if cmp_response.ok:
            cmp = cmp_response.json()
            cmp_id = cmp["id"]
            logger.info(
                "Successfully posted new component# %i as %s for ssp# %i.",
                cmp_id,
                cmp["title"],
                ssp_id,
            )
        if cmp_id and "protocols" in component.keys():
            for protocol in component["protocols"]:
                ports_protocols = PortsProtocol(
                    service="",
                    usedBy="",
                    parentId=cmp_id,
                    purpose=component["type"],
                    startPort=int(protocol["port-ranges"][0]["start"]),
                    endPort=int(protocol["port-ranges"][0]["end"]),
                    protocol=protocol["name"],
                    parentModule="components",
                    lastUpdatedById=config["userId"],
                    createdById=config["userId"],
                )
                ports.add(ports_protocols)

        if component["type"].lower() == "interconnection" and cmp_id:
            # Create ports and protocols object
            ports_protocols = PortsProtocol(
                service="",
                usedBy="",
                parentId=cmp_id,
                purpose=component["type"],
                startPort=0,
                endPort=0,
                protocol="",
                parentModule="components",
                lastUpdatedById=config["userId"],
                createdById=config["userId"],
            )
            ports_protocols.parentId = cmp_id
            ports_protocols.purpose = component["type"]
            # loop through properties to find port number
            if "props" in component.keys():
                for prop in component["props"]:
                    if prop["name"] == "information":
                        ports_protocols.purpose = prop["value"]
                    if prop["name"] == "port":
                        ports_protocols.startPort = int(prop["value"])
                        ports_protocols.endPort = int(prop["value"])
            ports.add(ports_protocols)
        create_component_mapping(api, config, ssp_id, cmp_id)
    if ports:
        for dat in ports:
            create_port(api, config, dat)


def create_component_mapping(api: Api, config: dict, ssp_id: int, cmp_id: int) -> None:
    """
    Create Component Mapping

    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param int ssp_id: The SSP ID.
    :param int cmp_id: The component ID.
    :rtype: None
    """
    mapping = {
        "securityPlanId": ssp_id,
        "componentId": cmp_id,
        "isPublic": True,
        "createdById": config["userId"],
        "lastUpdatedById": config["userId"],
    }
    mapping_response = api.post(
        url=config["domain"] + "/api/componentmapping",
        data=mapping,
    )
    if mapping_response.status_code != 200:
        logger.warning("Unable to post Mapping Response: %s.", json.dumps(mapping))


def create_ssp_stakeholders(api: Api, config: dict, ssp_id: int, ssp_dict: dict) -> None:
    """
    Create Stakeholders in RegScale

    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param int ssp_id: The SSP ID.
    :param dict ssp_dict: An SSP Dictionary.
    :rtype: None
    """
    parties = ssp_dict["system-security-plan"]["metadata"]["parties"]
    filtered_parties = list(filter(lambda x: x["type"] == "person", parties))
    for party in filtered_parties:
        title = [dat["value"] for dat in party["props"] if dat["name"] == "job-title"]
        phone = [dat["number"] for dat in party["telephone-numbers"]]
        email = list(party["email-addresses"])
        addresses = list(party["addresses"]) if "addresses" in party.keys() else None
        stakeholder = {
            "name": party["name"],
            "title": title[0] if title else "",
            "phone": phone[0] if phone else "",
            "email": email[0] if email else "",
            "address": (
                addresses[0]["addr-lines"][0]
                + " "
                + addresses[0]["city"]
                + " "
                + addresses[0]["state"]
                + ", "
                + addresses[0]["postal-code"]
                if addresses
                else ""
            ),
            "otherID": party["uuid"],
            "notes": email[0] if email else "",
            "parentId": ssp_id,
            "parentModule": "securityplans",
        }
        post_stakeholder(api, config, stakeholder)


def post_stakeholder(api: Api, config: dict, stakeholder: dict) -> Optional[list]:
    """Post Stakeholders to RegScale

    :param Api api: API instance
    :param dict config: An application configuration
    :param dict stakeholder: A stakeholder dictionary
    :return: A list of stakeholders, if any
    :rtype: Optional[list]
    """
    response = api.post(
        url=urljoin(config["domain"], "/api/stakeholders"),
        json=stakeholder,
    )
    if response.ok:
        logger.info(
            f"Created Stakeholder {response.json()} ",
            record_type="stakeholder",
            model_layer="stakeholder",
        )
        return response.json()
    else:
        logger.warning(
            f"Unable to create stakeholder: {stakeholder}",
            record_type="stakeholder",
            model_layer="stakeholder",
        )
        return None


def create_ssp_control_implementations(
    api: Api,
    config: dict,
    ssp_id: int,
    controls: dict,
    ssp_dict: dict,
) -> None:
    """
    Create the control implementations from the oscal SSP object

    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param int ssp_id: The SSP ID.
    :param dict controls: A dict of existing controls in RegScale.
    :param dict ssp_dict: An SSP Dictionary.
    :rtype: None
    """
    if not controls:
        return
    control_implementations = ssp_dict["system-security-plan"]["control-implementation"]["implemented-requirements"]

    for implementation in control_implementations:
        status = ControlImplementationStatus.NotImplemented

        for prop in implementation["props"]:
            if prop["name"] == "implementation-status":
                status = capitalize_words(prop["value"].replace("-", " "))
                if prop["value"].lower() == "implemented":
                    status = ControlImplementationStatus.FullyImplemented
                if prop["value"].lower() == "partial":
                    status = ControlImplementationStatus.PartiallyImplemented

        control_id = [
            control["controlID"]
            for control in controls
            if control["controlId"].lower() == implementation["control-id"].lower()
        ][0]
        imp = ControlImplementation(
            parentId=ssp_id,
            parentModule="securityplans",
            controlID=control_id,
            controlOwnerId=config["userId"],
            lastUpdatedById=config["userId"],
            createdById=config["userId"],
            status=status,
        )
        # Post Implementation
        post_regscale_object(api=api, config=config, obj=imp)


def post_regscale_object(
    api: Api, config: dict, obj: Any, endpoint: str = "controlimplementation"
) -> requests.Response:
    """
    Post RegScale control implementation

    :param Api api: API instance
    :param dict config: Application config
    :param Any obj: data object
    :param str endpoint: Endpoint to use in RegScale, defaults to "controlimplementation"
    :raises: TypeError if obj is not a dataclass, BaseModel, or dict
    :return: Response from API call to RegScale
    :rtype: requests.Response
    """
    response = None
    if dataclasses.is_dataclass(obj):
        dat = dataclasses.asdict(obj)
    elif isinstance(obj, BaseModel):
        dat = obj.dict()
    elif isinstance(obj, dict):
        dat = obj
    else:
        raise TypeError("Object must be a dataclass, BaseModel, or dict to post to RegScale.")
    try:
        response = api.post(config["domain"] + f"/api/{endpoint}", json=dat)
    except Exception as ex:
        logger.error("Unable to Post %s: %s to RegScale.\n%s", endpoint, dat, ex)

    return response


def create_ssp(api: Api, config: dict, ssp_dict: dict, submission_date: date, expiration_date: date) -> int:
    """
    Create a basic SSP in RegScale

    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param dict ssp_dict: An SSP Dictionary.
    :param date submission_date: The Submission date YYYY-MM-DD
    :param date expiration_date: The Expiration date YYYY-MM-DD
    :return: A newly created RegScale security plan id.
    :rtype: int
    """
    existing_ssps = []
    metadata = ssp_dict["system-security-plan"]["metadata"]
    system = ssp_dict["system-security-plan"]["system-characteristics"]
    fedramp_profile = get_profile(ssp_dict["system-security-plan"]["import-profile"]["href"])["profile"]
    profile_id = check_profile(api, config, fedramp_profile["metadata"]["title"])
    components = ssp_dict["system-security-plan"]["system-implementation"]["components"]
    ssp_payload = {
        "uuid": ssp_dict["system-security-plan"]["uuid"],
        "systemName": system.get("system-name", None),  # Required
        "planInformationSystemSecurityOfficerId": config["userId"],
        "planAuthorizingOfficialId": config["userId"],
        "systemOwnerId": config["userId"],
        "otherIdentifier": system["system-ids"][0]["id"],
        "confidentiality": capitalize_words(
            system["system-information"]["information-types"][0]["confidentiality-impact"]["selected"].split("-")[2]
        ),  # Required
        "integrity": capitalize_words(
            system["system-information"]["information-types"][0]["integrity-impact"]["selected"].split("-")[2]
        ),  # Required
        "availability": capitalize_words(
            system["system-information"]["information-types"][0]["availability-impact"]["selected"].split("-")[2]
        ),  # Required
        "status": capitalize_words(system["status"].get("state", "operational")),  # Required
        "description": system.get("description", None),
        "dateSubmitted": submission_date.strftime(DATE_FORMAT),
        "approvalDate": (submission_date + relativedelta(years=1)).strftime(DATE_FORMAT),  # User must be changed
        "expirationDate": expiration_date.strftime(DATE_FORMAT),
        "systemType": SYSTEM_TYPE,  # User must change
        "purpose": metadata.get("", None),
        "conditionsOfApproval": metadata.get("", None),
        "environment": metadata.get("", None),
        "lawsAndRegulations": metadata.get("", None),
        "authorizationBoundary": metadata.get("", None),
        "networkArchitecture": metadata.get("", None),
        "dataFlow": metadata.get("", None),
        "overallCategorization": capitalize_words(system["security-sensitivity-level"].split("-")[2]),
        "maturityTier": metadata.get("", None),
        "createdById": config["userId"],
        "hva": False,
        "practiceLevel": metadata.get("", None),
        "processLevel": metadata.get("", None),
        "cmmcLevel": metadata.get("", None),
        "cmmcStatus": metadata.get("", None),
        "isPublic": True,
        "executiveSummary": metadata.get("", None),
        "recommendations": metadata.get("", None),
        "importProfile": metadata.get("version", "fedramp1.1.0-oscal1.0.0"),
        "parentId": profile_id,
        "parentModule": "profiles",
    }
    logger.warning("Unknown System Type, defaulting to %s.", ssp_payload["systemType"])
    logger.warning("Unknown HVA status, defaulting to %r.", ssp_payload["hva"])

    existing_ssp_response = api.get(url=urljoin(config["domain"], SSP_URL_SUFFIX))
    if existing_ssp_response.ok:
        existing_ssps = existing_ssp_response.json()

    if system["system-name"] in {ssp["systemName"] for ssp in existing_ssps}:
        dat = {ssp["id"] for ssp in existing_ssps if ssp["systemName"] == system["system-name"]}
        click.confirm(
            f"This SSP Title already exists in the system, \
                SSP: {dat.pop() if len(dat) < 2 else dat}.  Would you still like to continue?",
            abort=True,
        )

    response = api.post(url=urljoin(config["domain"], "/api/securityplans"), json=ssp_payload)
    if response.ok:
        logger.info("SSP Created with an id of %i!", response.json()["id"])
        ssp_id = response.json()["id"]
    controls_response = api.get(urljoin(config["domain"], f"/api/profilemapping/getByProfile/{profile_id}"))
    controls = controls_response.json() if controls_response.ok else []
    create_ssp_components(api, config, components, ssp_id)
    create_ssp_control_implementations(api, config, ssp_id, controls, ssp_dict)
    create_ssp_stakeholders(api, config, ssp_id, ssp_dict)
    # update_ssp_contacts(api, config, ssp_id, ssp_dict)

    return ssp_id


def get_profile(url: str) -> dict:
    """
    Downloads the FedRAMP profile

    :param str url: A profile URL.
    :return: A dictionary with the profile json data.
    :rtype: dict
    """
    dl_path = download_file(url)
    with open(dl_path, encoding="utf-8") as json_file:
        data = json.load(json_file)
    return data


def get_tables(document: Any) -> list:
    """
    Return all document tables

    :param Any document: document object
    :return: List of all document tables
    :rtype: list
    """
    tables = list(document.tables)
    for t_table in document.tables:
        for row in t_table.rows:
            for cell in row.cells:
                tables.extend(iter(cell.tables))
    return tables


def get_xpath_privacy_detailed(tables: Table, key: str, xpath: str, count_array: list[int] = [0, 2, 4, 6]) -> dict:
    """
    Use Xpath to pull data from XML tables

    :param Table tables: XML tables
    :param str key: specific key in XML table
    :param str xpath: xpath of the element
    :param list count_array: array of numbers, default is [0, 2, 4, 6]
    :return: Dictionary of specific items found
    :rtype: dict
    """
    result = {"piishare": None, "piipublic": None, "piaperformed": None, "piasorn": None}
    for t_var in tables:
        if key in t_var._element.xml:
            tree = etree.parse(StringIO(t_var._element.xml))
            tags = tree.xpath(xpath, namespaces=namespaces)
            for p_var in tags:
                t_tags = p_var.xpath(XPATH_TAG, namespaces=namespaces)
                for idx, t_var in enumerate(t_tags):
                    if idx in count_array:
                        result[list(result.keys())[count_array.index(idx)]] = t_var.text

    return result


def get_xpath_data_detailed(tables: Table, key: str, ident: str, xpath: str, count_array: list = None) -> dict:
    """
    Use Xpath to pull data from XML tables

    :param Table tables: XML tables
    :param str key: specific key in XML table
    :param str ident:
    :param str xpath: xpath of the element
    :param list count_array: array of numbers, default is [2, 3, 4]
    :return: Dictionary of items found
    :rtype: dict
    """
    if count_array is None:
        count_array = [2, 3, 4]
    tables = iter(tables)
    confidentiality = None
    integrity = None
    availability = None
    for t_var in tables:
        if key in t_var._element.xml:
            f = StringIO(t_var._element.xml)
            tree = etree.parse(f)
            tags = tree.xpath(xpath, namespaces=namespaces)
            for p_var in tags:
                t_tags = p_var.xpath(XPATH_TAG, namespaces=namespaces)
                count = 0
                for t_var in t_tags:
                    if t_var.text == ident or count > 0:
                        count += 1
                        if count == count_array[0]:
                            confidentiality = t_var.text
                        if count == count_array[1]:
                            integrity = t_var.text
                        if count == count_array[2]:
                            availability = t_var.text
    return {
        "type": key,
        "nist_ident": ident,
        "confidentiality": confidentiality,
        "integrity": integrity,
        "availability": availability,
    }


def extract_between_strings(text: str, start_marker: str, end_marker: str) -> Optional[str]:
    """
    Extract sub string between start marker and end marker strings

    :param str text: source string
    :param str start_marker: start string to look for
    :param str end_marker: end string to look for
    :return: extracted sub string or None
    :rtype: Optional[str]
    """
    start_index = text.find(start_marker)
    if start_index == -1:
        return None  # Start marker not found
    start_index += len(start_marker)
    end_index = text.find(end_marker, start_index)
    if end_index == -1:
        return None  # End marker not found
    return text[start_index:end_index]


def get_xpath_sysinfo_detailed(tables: Table, key: str, xpath: str, count_array: list = None) -> dict:
    """
    Use Xpath to pull data from XML tables

    :param Table tables: XML tables
    :param str key: specific key in XML table
    :param str xpath: xpath of the element
    :param list count_array: array of numbers, default is [0,2, 3, 4]
    :return: Dictionary of specific items found
    :rtype: dict
    """
    if count_array is None:
        count_array = [0, 2, 4, 6]
    tables = iter(tables)
    uniqueident = None
    systemname = None
    keycount = 0
    for t_var in tables:
        # there are multiple occurences of the key in the document
        # we only want the first one.
        if keycount > 0:
            break
        if key in t_var._element.xml:
            f = StringIO(t_var._element.xml)
            tree = etree.parse(f)
            tags = tree.xpath(xpath, namespaces=namespaces)
            for p_var in tags:
                t_tags = p_var.xpath(XPATH_TAG, namespaces=namespaces)
                count = 0
                for t_var in t_tags:
                    if count == count_array[0]:
                        uniqueident = t_var.text.strip()
                    if count == count_array[1]:
                        systemname = t_var.text.strip()
                    count += 1
            keycount += 1

    return {
        "uniqueidentifier": uniqueident,
        "systemname": systemname,
    }


def get_xpath_prepdata_detailed(tables: Table, key: str, ident: str, xpath: str) -> dict:
    """
    Use Xpath to pull data from XML tables

    :param Table tables: XML tables
    :param str key: specific key in XML table
    :param str ident: document identifier tag
    :param str xpath: xpath of the element
    :return: Dictionary of items found
    :rtype: dict
    """
    tables = iter(tables)
    orgname = ""
    street = ""
    office = ""
    citystate = ""
    for t_var in tables:
        if key in t_var._element.xml:
            f = StringIO(t_var._element.xml)
            tree = etree.parse(f)
            tags = tree.xpath(xpath, namespaces=namespaces)
            for p_var in tags:
                t_tags = p_var.xpath(XPATH_TAG, namespaces=namespaces)
                preptext = ""
                for t_var in t_tags:
                    preptext += t_var.text
            preptext += END_MARKER
            orgname = extract_between_strings(preptext, ORGANIZATION_TAG, "Street Address")
            street = extract_between_strings(preptext, "Street Address", "Suite/Room/Building")
            office = extract_between_strings(preptext, "Suite/Room/Building", "City, State Zip")
            citystate = extract_between_strings(preptext, "City, State Zip", END_MARKER)

    return {
        "type": key,
        "nist_ident": ident,
        "orgname": orgname,
        "office": office,
        "street": street,
        "citystate": citystate,
    }


def get_contact_info(tables: list, key: str, xpath: str) -> dict:
    """
    Use Xpath to pull data from XML tables

    :param list tables: XML tables
    :param str key: key to look for
    :param str xpath: xpath of the element
    :return: Dictionary of sorted data
    :rtype: dict
    """
    idents = [
        "Name",
        "Title",
        "Company / Organization",
        "Address",
        "Phone Number",
        "Email Address",
    ]
    dat = {}

    def loop_and_update(element_list: list) -> dict:
        """
        Loop through the element list and update the data dictionary

        :param list element_list: List of elements
        :return: Updated dictionary
        :rtype: dict
        """
        value = ""
        if idents:
            count = 0
            field = idents.pop(0)

            while count < len(element_list) - 1:
                if element_list[count].text == field:
                    value = "".join([value, element_list[count + 1].text])
                    dat[field] = value
                    value = ""
                    count += 1
                try:
                    if element_list[count + 1].text in idents:
                        field = idents.pop(0)
                    else:
                        if field in dat:
                            dat[field] = "".join([dat[field], element_list[count + 1].text])
                        count += 1
                except IndexError:
                    logger.debug("Unable to continue, index error on row: %i.", count)

    tables = iter(tables)

    tag_data = []
    for _, t_enum in enumerate(tables):
        if key in t_enum._element.xml:
            f_var = StringIO(t_enum._element.xml)
            tree = etree.parse(f_var)
            tags = tree.xpath(xpath, namespaces=namespaces)
            for tag in tags:
                p_tags = tag.xpath("//w:p", namespaces=namespaces)
                for p_var in p_tags:
                    t_tags = p_var.xpath(XPATH_TAG, namespaces=namespaces)
                    for tags in t_tags:
                        tag_data.append(tags)
    loop_and_update(tag_data)

    return dat


def extract_title_from_row(row_text: list) -> Optional[str]:
    """
    Extracts the title from the first row's data.

    :param list row_text: List of cell text from the row
    :return: Title or None if not found
    :rtype: Optional[str]
    """
    title_data = [x for x in row_text if x]  # Ignore empty values
    if title_data:
        title_parts = title_data.pop().split("\n")
        return title_parts[1] if len(title_parts) > 1 else title_parts[0]
    return None


def process_contact_info(table, contact_mapping: Dict[int, str], result: Dict[str, str]) -> Dict[str, str]:
    """
    Processes contact information from the table based on the mapping of cell indices to fields.

    :param table: The table from the Word document
    :param contact_mapping: Dictionary mapping cell indices to contact information fields
    :param result: Dictionary to store the contact information
    :return: Updated result dictionary with contact information
    :rtype: dict
    """
    for cell_index, cell in enumerate(table._cells):
        field = contact_mapping.get(cell_index)
        if field:
            result[field] = cell.text.strip()
    return result


def get_base_contact(document: Document, key: Optional[str] = "Point of Contact") -> Dict[str, str]:
    """
    Gets contact information from a document.

    :param Document document: SSP document
    :param Optional[str] key: key to parse for, defaults to 'Point of Contact'
    :return: dictionary with contact information
    :rtype: dict
    """
    result = {}
    title = None

    # Define a mapping of cell indices to contact information fields
    contact_mapping = {3: "name", 5: "title", 7: "company", 9: "address", 11: "phone", 13: "email"}

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]

            if row_index == 0 and not title:
                # Extract the title from the first row
                title = extract_title_from_row(row_text)
                if title:
                    result["title"] = title
                continue

            if row_text and key == row_text[0]:
                # Process the contact information
                result = process_contact_info(table, contact_mapping, result)
                return result  # Early return after contact info is processed

    return result


def post_interconnects(app: Application, table_data: list, regscale_ssp: dict) -> None:
    """
    Interconnects map to SSP in RegScale

    :param Application app: Application object
    :param list table_data: List of tables
    :param dict regscale_ssp: SecurityPlan object
    :rtype: None
    """
    api = Api()
    user_id = app.config["userId"]
    key = "SP* IP Address and Interface"
    existing_interconnects = []
    dat = [table for table in table_data if key in table.keys()]
    existing_interconnect_response = api.get(
        app.config["domain"] + f"/api/interconnections/getAllByParent/{regscale_ssp['id']}/securityplans"
    )
    if not existing_interconnect_response.raise_for_status() and (
        existing_interconnect_response.headers.get("content-type") == "application/json; charset=utf-8"
    ):
        existing_interconnects = existing_interconnect_response.json()

    logger.info(
        f"Found {len(existing_interconnects)} existing interconnects",
        record_type="interconnects",
        model_layer="interconnects",
    )
    for interconnect in dat:
        interconnection = InterConnection(
            name=interconnect[key],
            aOId=user_id,
            interconnectOwnerId=user_id,
            dateCreated=get_current_datetime(),
            dateLastUpdated=get_current_datetime(),
            lastUpdatedById=user_id,
            createdById=user_id,
            description=(
                interconnect["Information Being Transmitted"]
                if "Information Being Transmitted" in interconnect.keys()
                else ""
            ),
            parentId=regscale_ssp["id"],
            parentModule="securityplans",
            agreementDate=get_current_datetime(),
            expirationDate=(datetime.now() + relativedelta(years=3)).strftime(DATE_FORMAT),
            status="Approved",
            organization=regscale_ssp["systemName"],
            categorization=regscale_ssp["overallCategorization"],
            connectionType="Web Service or API",
            authorizationType="Interconnect Security Agreement (ISA)",
        )
        if interconnection.name + interconnection.description not in {
            inter["name"] + inter["description"] for inter in existing_interconnects
        }:
            post_regscale_object(
                api=api,
                config=app.config,
                obj=interconnection.dict(),
                endpoint="interconnections",
            )


def create_privacy_data(app: Application, privacy_data: dict, ssp_id: int) -> None:
    """
    Post Privacy settings for SSP

    :param Application app: Application object
    :param list privacy_data: list of tables
    :param int ssp_id: RegScale SSP ID
    :rtype: None
    """
    use_default = YES.lower() == privacy_data.get("piishare", "").lower()

    privacy = Privacy(
        id=0,
        piiCollection=privacy_data["piishare"],
        piiPublicCollection=privacy_data["piipublic"],
        piaConducted=privacy_data["piaperformed"],
        sornExists=privacy_data["piasorn"],
        sornId=None,
        ombControlId=None,
        infoCollected="Collection info not supplied" if use_default else None,
        justification="Justification not supplied" if use_default else None,
        businessUse="Business use case not found" if use_default else None,
        pointOfContactId=app.config["userId"],
        privacyOfficerId=app.config["userId"],
        informationSharing="System sharing info not found" if use_default else None,
        consent="Consent info not found" if use_default else None,
        security="Security Information not found" if use_default else None,
        privacyActSystem=YES if YES in privacy_data else "No",
        recordsSchedule=None,
        securityPlanId=ssp_id,
        status="Not Applicable",
        dateApproved=None,
        notes="Imported from NIST 800.r3 SSP document.",
    )

    if not Privacy.get_all_by_parent(ssp_id):
        if new_privacy := privacy.create():
            logger.info(
                f"Privacy #{new_privacy.id} created.",
                record_type="privacy",
                model_layer="privacy",
            )
    else:
        logger.info(
            "Privacy settings already exist, skipping...",
            record_type="privacy",
            model_layer="privacy",
        )


def post_ports(app: Application, table_data: list, ssp_id: int) -> None:
    """
    Ports map to interconnects

    :param Application app: Application object
    :param list table_data: list of tables
    :param int ssp_id: RegScale SSP ID
    :rtype: None
    """
    dat = [table for table in table_data if "Protocols" in table.keys()]
    existing_ports = PortsProtocol.get_all_by_parent(parent_id=ssp_id, parent_module="securityplans")

    logger.info(
        f"Found {len(existing_ports)} existing ports",
        record_type="ports-protocols",
        model_layer="ports-protocols",
    )
    process_port_protocols(protocols=dat, ssp_id=ssp_id, app_config=app.config)


def process_port_protocols(protocols: List[dict], ssp_id: int, app_config: Dict) -> None:
    """
    Process port protocols
    :param List[dict] protocols: List of protocols
    :param int ssp_id: SSP ID
    :param dict app_config: Application configuration
    :rtype: None
    """
    key = "Ports (TCP/UDP)*"

    for protocol in protocols:
        try:
            port_field, port_protocol, purpose = extract_protocol_details(protocol, key)
            port_entries, start = parse_port_field(port_field)

            if start:
                start, port_entries = update_start_port(start, port_entries)

            create_ports_protocol_records(protocol, port_entries, start, ssp_id, port_protocol, purpose, app_config)

            if not port_entries and port_field.strip().isdigit():
                single_port = int(port_field.strip())
                create_single_port_record(protocol, single_port, ssp_id, port_protocol, purpose, app_config)

        except Exception as e:
            logger.info(f"Error processing port: {port_field} with error: {e}")


def extract_protocol_details(protocol: dict, key: str) -> tuple:
    """
    Extracts and cleans up the port field, protocol, and purpose.

    :param dict protocol: Protocol data
    :param str key: Key to extract the port field
    :return: Tuple of (port_field, port_protocol, purpose)
    """
    port_field = protocol.get(key, "")
    unwanted_substrings = ["(TCP)", "(UDP)", "(UDP/TCP)", "TCP", "UDP", "TCP/UDP", "//n", "/n", "\n", "/"]

    # Use a loop to remove all unwanted substrings
    for substr in unwanted_substrings:
        port_field = port_field.replace(substr, "")

    port_protocol = protocol.get("Protocols") or parse_port_or_protocol(key, protocol, str)
    purpose = protocol.get("Purpose", "Unknown") if protocol.get("Purpose") != "" else "Unknown"
    return port_field, port_protocol, purpose


def parse_port_field(port_field: str) -> tuple:
    """
    Parses the port field into start port and entries.

    :param str port_field: The port field value
    :return: Tuple of (port_entries, start)
    """
    if "-" in port_field:
        start, end = int(port_field.split("-")[0]), port_field.split("-")[1]
        port_entries = end.split(",")
    else:
        start = None
        port_entries = port_field.split(",")
    return port_entries, start


def create_ports_protocol_records(
    protocol: dict, port_entries: List[str], start: int, ssp_id: int, port_protocol: str, purpose: str, app_config: Dict
) -> None:
    """
    Creates and updates PortsProtocol records for each port entry.

    :param dict protocol: Protocol data
    :param List[str] port_entries: List of port entries
    :param int start: Start port
    :param int ssp_id: SSP ID
    :param str port_protocol: Protocol type
    :param str purpose: Purpose of the protocol
    :param dict app_config: Application configuration
    :rtype: None
    """
    for port_entry in port_entries:
        port_entry = port_entry.strip()
        if start:
            create_ports_protocol(protocol, start, port_entry, ssp_id, port_protocol, purpose, app_config)
        else:
            single_port = int(port_entry)
            create_ports_protocol(protocol, single_port, single_port, ssp_id, port_protocol, purpose, app_config)


def create_single_port_record(
    protocol: dict, single_port: int, ssp_id: int, port_protocol: str, purpose: str, app_config: Dict
) -> None:
    """
    Creates a PortsProtocol record for a single port.

    :param dict protocol: Protocol data
    :param int single_port: Single port value
    :param int ssp_id: SSP ID
    :param str port_protocol: Protocol type
    :param str purpose: Purpose of the protocol
    :param dict app_config: Application configuration
    :rtype: None
    """
    create_ports_protocol(protocol, single_port, single_port, ssp_id, port_protocol, purpose, app_config)


def create_ports_protocol(
    protocol: dict, start_port: int, end_port: int, ssp_id: int, port_protocol: str, purpose: str, app_config: Dict
) -> None:
    """
    Creates or updates a PortsProtocol object.

    :param dict protocol: Protocol data
    :param int start_port: Start port
    :param int end_port: End port
    :param int ssp_id: SSP ID
    :param str port_protocol: Protocol type
    :param str purpose: Purpose of the protocol
    :param dict app_config: Application configuration
    :rtype: None
    """
    PortsProtocol(
        service=protocol.get("Services", "Unknown"),
        usedBy=protocol.get("Used By", "Unknown"),
        parentId=ssp_id,
        purpose=purpose,
        startPort=start_port,
        endPort=end_port,
        protocol=port_protocol.strip().replace("(", "").replace(")", ""),
        parentModule="securityplans",
        lastUpdatedById=app_config.get("userId"),
        createdById=app_config.get("userId"),
    ).create_or_update()


def update_start_port(start: str, port_entries: List[str]) -> tuple[str, list[str]]:
    """
    Update the start port if necessary
    :param start:
    :param port_entries:
    :return: Updated start, port entries
    :rtype: tuple[str, list[str]]
    """
    # Convert start to an integer for comparison
    start = int(start)

    # Convert port_entries to integers
    port_entries_int = [int(port) for port in port_entries]

    # Check if there's a smaller number in port_entries
    if port_entries_int and min(port_entries_int) < start:
        # Get the smallest number from port_entries
        new_start = min(port_entries_int)

        # Log the change
        if new_start != start:
            logger.info(f"Start port changed from {start} to {new_start}")

        # Remove the new start from port_entries and append the old start
        port_entries_int.remove(new_start)
        port_entries_int.append(start)

        # Update start
        start = new_start

    # Return updated start and updated port_entries
    return str(start), [str(port) for port in port_entries_int]


def parse_port_or_protocol(key: str, protocol: dict, return_type: Union[Type[str], Type[int]]) -> Union[str, int]:
    """
    Parse port number from protocol

    :param str key: Key to parse
    :param dict protocol: Protocol dictionary
    :param Union[Type[str], Type[int]] return_type: Data type to return
    :return: Port number or protocol
    :rtype: Union[str, int]
    """
    try:
        if return_type == int:
            port_to_be = protocol.get(key, [])
            if isinstance(port_to_be, str) == str:
                port_to_be.replace(" ", "").replace("/n", "").replace("//n", "").replace("\n", "")
            if port := "".join(c for c in protocol.get(key, []) if c.isdigit()):
                return int(port)
            if port := "".join(c for c in protocol.get("Protocols", []) if c.isdigit()):
                return int(port)
        elif return_type == str:
            return "".join(c for c in protocol[key] if not c.isdigit()) or "".join(
                c for c in protocol["Protocols"] if c.isdigit()
            )
    except ValueError:
        return 0


def get_current_implementations(app: Application, regscale_id: int) -> list[dict]:
    """Pull current implementations for a given regscale id

    :param Application app: Application instance
    :param int regscale_id: RegScale ID
    :return: List of dictionaries
    :rtype: list[dict]
    """
    current_imps = []
    api = Api()
    try:
        current_imps_response = api.get(
            url=app.config["domain"] + f"/api/controlImplementation/getAllByPlan/{regscale_id}",
            params=("skip_check", True),
        )
        if not current_imps_response.raise_for_status():
            current_imps = current_imps_response.json()
    except requests.HTTPError:  # This endpoint returns 404 when empty.
        current_imps = []
    return current_imps


def get_friendly_control_id(control_number: str) -> str:
    """Get friendly control id from control number

    :param str control_number: Control number
    :return: Friendly control id
    :rtype: str
    """
    # exp = r"^.*?\([^\d]*(\d+)[^\d]*\).*$"
    # the above regex allows for a denial of service attack
    # the below regex should mitigate that
    exp = r"\((\d+)\)"
    return (
        f"{control_number[:match.regs[1][0] - 1].strip()}.{match.groups()[0]}".lower()
        if (match := re.search(exp, control_number))
        else control_number.lower()
    )


def post_implementations(
    app: Application,
    ssp_obj: SSP,
    regscale_ssp: dict,
    mapping: List[ProfileMapping],
    ctrl_roles: dict,
    save_data: bool = False,
    load_missing: bool = False,
) -> List:
    """
    Post implementations to RegScale

    :param Application app: Application object
    :param SSP ssp_obj: SecurityPlan object (python-docx)
    :param dict regscale_ssp: RegScale ssp
    :param List[ProfileMapping] mapping: mapping
    :param dict ctrl_roles: Control roles
    :param bool save_data: Whether to save data to a file
    :param bool load_missing: Whether to load missing controls
    :return: List of new implementations
    :rtype: List
    """
    api = Api()
    current_imps = ControlImplementation.get_all_by_parent(
        parent_id=regscale_ssp.get("id"), parent_module="securityplans"
    )  # get_current_implementations(app, regscale_ssp["id"])
    for imp in current_imps:
        imp.control = SecurityControl.get_object(object_id=imp.controlID)
    for map in mapping:
        map.control = SecurityControl.get_object(object_id=map.controlID)

    missing_controls = check_control_list_length(ssp_obj, mapping)
    log_controls_info(ssp_obj, current_imps)
    (
        mapped_controls_log,
        unmapped_controls_log,
        implemented_controls_log,
        new_implementations,
        _,
    ) = process_controls(
        app,
        api,
        ssp_obj,
        regscale_ssp,
        mapping,
        ctrl_roles,
        current_imps,
        missing_controls,
    )

    if load_missing:
        new_imps = load_non_matched_profile_controls(app, regscale_ssp=regscale_ssp, mapping=mapping)
        new_implementations.extend(new_imps)

    if save_data:
        save_control_logs(mapped_controls_log, unmapped_controls_log, implemented_controls_log)
    return new_implementations


def get_responsibility_and_status(fedramp_control: Any) -> Tuple[str, str]:
    """
    Get responsibility and implementation status

    :param Any fedramp_control: FedrampControl object
    :return: Tuple of responsibility and implementation status
    :rtype: Tuple[str, str]
    """
    responsibility = None
    if fedramp_control.control_origination:
        if "Shared".lower() in fedramp_control.control_origination[0].lower():
            responsibility = "Shared"
        elif "Customer".lower() in fedramp_control.control_origination[0].lower():
            responsibility = "Customer"
        elif "Provider".lower() in fedramp_control.control_origination[0].lower():
            responsibility = "Provider"
        else:
            responsibility = fedramp_control.control_origination[0]

    if fedramp_control.implementation_status and fedramp_control.implementation_status[0] in [
        "Alternative Implementation",
        "Implemented",
    ]:
        implementation_status = ControlImplementationStatus.FullyImplemented
    elif ControlImplementationStatus.PartiallyImplemented in fedramp_control.implementation_status:
        implementation_status = ControlImplementationStatus.PartiallyImplemented
    else:
        implementation_status = (
            fedramp_control.implementation_status[0] if fedramp_control.implementation_status else "Not Implemented"
        )

    return responsibility, implementation_status


def find_control_in_mapping(mapping: List[ProfileMapping], friendly_control_id: str) -> str:
    """
    Find control id in mapping

    :param List[ProfileMapping] mapping: Mapping
    :param str friendly_control_id: Friendly control id
    :return: ControlId
    :rtype: str
    """
    control_id = None
    if control := [control for control in mapping if control.control.controlId.lower() == friendly_control_id]:
        control_id = control[0].control.controlId
    return control_id


def get_implementation_text(fedramp_control: Any) -> str:
    """
    Get implementation text

    :param Any fedramp_control: FedrampControl object
    :return: Implementation text
    :rtype: str
    """
    if len(fedramp_control.parts) > 1:
        return "<br>".join(fedramp_control.part(x).text for x in fedramp_control.parts)
    else:
        try:
            return fedramp_control.part(None).text
        except IndexError:
            return ""


def handle_control_implementation(
    app: Application,
    regscale_ssp: Dict,
    control_id: str,
    responsibility: str,
    implementation_status: str,
    implementation_text: str,
    ctrl_roles: Dict,
    friendly_control_id: str,
    new_implementations: list,
    update_implmentations: list,
    current_implementations_dict: Dict[str, ControlImplementation],
    mapping: Dict[str, ProfileMapping],
) -> ControlImplementation:
    """
    Handle control implementation

    :param Application app: Application instance
    :param Dict regscale_ssp: RegScale ssp
    :param str control_id: Control id
    :param str responsibility: Responsibility
    :param str implementation_status: Implementation status
    :param str implementation_text: Implementation text
    :param Dict ctrl_roles: Control roles
    :param str friendly_control_id: Friendly control id
    :param list new_implementations: List of new implementations
    :param list update_implmentations: List of updated implementations
    :param Dict[str, ControlImplementation] current_implementations_dict: Dictionary of current implementations
    :param List[ProfileMapping] mapping: Mapping
    :return: Response object
    :rtype: Optional[Response]
    """
    implementation = None
    if control_id in current_implementations_dict.keys():
        logger.info(
            f"Updating Implementation: {control_id}",
            "control-implementation",
            "control",
        )
        implementation = current_implementations_dict.get(control_id)
        implementation.status = implementation_status
        implementation.responsibility = responsibility
        implementation.implementation = implementation_text
        implementation.lastUpdatedById = app.config.get("userId")
        implementation.systemRoleId = (
            ctrl_roles.get(friendly_control_id)[0]
            if isinstance(ctrl_roles, dict)
            and friendly_control_id in ctrl_roles.keys()
            and ctrl_roles.get(friendly_control_id)[0]
            else None
        )
        update_implmentations.append(implementation)
    else:
        mapping_control = mapping.get(control_id)
        logger.info(
            f"Creating Implementation: {control_id}",
            record_type="control",
            model_layer="control-implementation",
        )
        implementation = ControlImplementation(
            parentId=regscale_ssp["id"],
            parentModule="securityplans",
            controlOwnerId=app.config["userId"],
            status=implementation_status,
            controlID=mapping_control.controlID,
            responsibility=responsibility,
            implementation=implementation_text,
            systemRoleId=(
                ctrl_roles.get(friendly_control_id)[0]
                if isinstance(ctrl_roles, dict)
                and friendly_control_id in ctrl_roles.keys()
                and ctrl_roles.get(friendly_control_id)[0]
                else None
            ),
        )
        implementation = implementation.create()
        new_implementations.append(implementation)
    return implementation


def handle_requirements(
    app: Application,
    api: Api,
    fedramp_control: Any,
    mapping: List[ProfileMapping],
    friendly_control_id: str,
    implementation_status: str,
    regscale_ssp: Dict,
) -> None:
    """
    Handle requirements

    :param Application app: Application instance
    :param Api api: API instance
    :param Any fedramp_control: FedrampControl object
    :param List[ProfileMapping] mapping: Mapping
    :param str friendly_control_id: Friendly control id
    :param str implementation_status: Implementation status
    :param Dict regscale_ssp: RegScale ssp
    :rtype: None
    """
    parent_security_control_id = [
        control["controlID"] for control in mapping if control["controlId"] == friendly_control_id.split()[0]
    ][0]
    current_imps = get_current_implementations(app=app, regscale_id=regscale_ssp["id"])
    parent_security_control = [imp for imp in current_imps if imp["controlID"] == parent_security_control_id][0]

    for part in fedramp_control.parts:
        implementation_text = fedramp_control.part(part).text
        title = f"{friendly_control_id.split()[0]} - Req. {part}"
        requirement = Requirement(
            id=0,
            description=implementation_text.split("\n")[0],
            implementation=implementation_text,
            title=title,
            lastUpdatedById=app.config["userId"],
            status=implementation_status,
            controlID=parent_security_control_id,
            parentId=parent_security_control["id"],
            parentModule="controls",
            requirementOwnerId=app.config["userId"],
            createdById=app.config["userId"],
        )

        existing_requirement = api.get(
            url=app.config["domain"] + f"/api/requirements/getByParent/{parent_security_control['id']}/controls"
        ).json()

        if title not in {req["title"] for req in existing_requirement}:
            logger.info("Posting Requirement: %s", title)
            post_regscale_object(
                api=api,
                config=app.config,
                obj=requirement,
                endpoint="requirements",
            )
        else:
            logger.info("Requirement %s already exists, skipping...", title)


# format_parameter_name has been moved to parameter_utils.py
# It now uses OSCAL format (_odp.) instead of legacy format (_prm_)
# for consistency with NIST 800-53 Rev5 and FedRAMP Rev5 OSCAL documents


def handle_parameters(fedramp_control: Any, control_imp: ControlImplementation) -> None:
    """
    Handle parameters

    :param Any fedramp_control: FedrampControl object
    :param ControlImplementation control_imp: ControlImplementation object
    :rtype: None
    """
    pnum = 0
    existing_params = Parameter.get_all_by_parent(parent_id=control_imp.id)
    existing_param_names_dict = {param.name: param for param in existing_params}
    base_control_params = ControlParameter.get_by_control(control_id=control_imp.controlID)
    base_control_params_dict = {param.parameterId: param for param in base_control_params}

    for parameter in fedramp_control.parameters:
        pnum = pnum + 1
        try:
            param_dict = get_parameter_value(parameter)
            pname = format_parameter_name(str(fedramp_control), pnum)
            control_param_name = pname
            base_control_param = base_control_params_dict.get(control_param_name)
            if base_control_param:
                if not existing_params or control_param_name not in existing_param_names_dict:
                    Parameter(
                        controlImplementationId=control_imp.id,
                        name=control_param_name,
                        value=param_dict.get("value"),
                        parentParameterId=base_control_param.id,
                    ).create()
                else:
                    existing_param = existing_param_names_dict.get(control_param_name)
                    if existing_param.name == control_param_name:
                        existing_param.value = param_dict.get("value")
                        existing_param.parentParameterId = base_control_param.id
                        existing_param.save()
        except Exception as e:
            logger.warning("Unable to map parameter %s to RegScale: %s", parameter, e)


def process_controls(
    app: Application,
    api: Api,
    ssp_obj: dict,
    regscale_ssp: dict,
    mapping: List[ProfileMapping],
    ctrl_roles: dict,
    current_imps: List[ControlImplementation],
    missing_controls: Optional[list],
) -> Tuple[List[str], List[str], List[str], List, List]:
    """
    Process controls

    :param Application app: Application instance
    :param Api api: API instance
    :param dict ssp_obj: SSP object
    :param dict regscale_ssp: RegScale ssp
    :param dict mapping: Mapping
    :param dict ctrl_roles: Control roles
    :param list current_imps: List of current implementations
    :param Optional[list] missing_controls: List of missing controls in the selected profile
    :return Tuple[List[str], List[str], List[str], List]:
        List of mapped controls
        List of unmapped controls
        List of implemented controls
        List new implementations
        List of updated implementations
    """
    mapped_controls_log = []
    unmapped_controls_log = []
    implemented_controls_log = []
    new_implementations = []
    update_implmentations = []
    has_requirements = False
    if not missing_controls:
        missing_controls = []

    logger.info(f"Processing {len(ssp_obj.control_list)} controls..")
    mapping_dict = {control.control.controlId: control for control in mapping}
    for fedramp_control in ssp_obj.control_list:
        responsibility, implementation_status = get_responsibility_and_status(fedramp_control)
        friendly_control_id = get_friendly_control_id(fedramp_control.number)
        control_id = find_control_in_mapping(mapping, friendly_control_id)
        logger.info(f"Processing Control: {friendly_control_id.upper()}")
        if not control_id or friendly_control_id in missing_controls:
            unmapped_controls_log.append(friendly_control_id.upper())
            continue

        implementation_text = get_implementation_text(fedramp_control)
        current_implementations_dict = {c.control.controlId: c for c in current_imps}
        control_imp = handle_control_implementation(
            app,
            regscale_ssp,
            control_id,
            responsibility,
            implementation_status,
            implementation_text,
            ctrl_roles,
            friendly_control_id,
            new_implementations,
            update_implmentations,
            current_implementations_dict,
            mapping_dict,
        )

        if "Req" in fedramp_control.number:
            handle_requirements(
                app,
                api,
                fedramp_control,
                mapping,
                friendly_control_id,
                implementation_status,
                regscale_ssp,
            )
            has_requirements = True

        if control_imp:
            handle_parameters(fedramp_control, control_imp)

        if has_requirements:
            has_requirements = False  # Reset for the next control

    ControlImplementation.batch_update(items=update_implmentations)
    return (
        mapped_controls_log,
        unmapped_controls_log,
        implemented_controls_log,
        new_implementations,
        update_implmentations,
    )


def check_control_list_length(ssp_obj: SSP, mapping: dict) -> Optional[list]:
    """
    Check control list length

    :param SSP ssp_obj: SSP object
    :param dict mapping: Mapping
    :return: List of missing controls, if found
    :rtype: Optional[list]
    """
    profile_control_ids = [get_friendly_control_id(item.control.controlId) for item in mapping]
    parsed_control_ids = [get_friendly_control_id(control.number) for control in ssp_obj.control_list]
    if len(ssp_obj.control_list) > len(mapping):
        missing_controls = set(profile_control_ids) - set(parsed_control_ids)
        logger.error(
            f"There are more controls in the source document ({len(ssp_obj.control_list)}) than in the base profile ({len(mapping)})!",
            record_type="implementations",
            model_layer="implementations",
        )
        logger.error(
            f"Extra controls found in source document and missing from base profile: {', '.join(missing_controls)}",
            record_type="implementations",
            model_layer="implementations",
        )
        return missing_controls


def log_controls_info(ssp_obj: SSP, current_imps: List[dict]) -> None:
    """
    Log controls info

    :param SSP ssp_obj: SSP object
    :param List[dict] current_imps: List of current implementations
    :return None:
    :rtype None:
    """
    logger.info(
        f"Attempting to post {len(ssp_obj.control_list)} controls from this FedRAMP SSP Document to RegScale!",
        record_type="control",
        model_layer="control-implementation",
    )
    if len(current_imps) > 0:
        logger.info(
            f"This RegScale Security plan already has {len(current_imps)} implementations..",
            record_type="control",
            model_layer="control-implementation",
        )


def save_control_logs(
    mapped_controls_log: List[str],
    unmapped_controls_log: List[str],
    implemented_controls_log: List[str],
) -> None:
    """
    Save control logs

    :param List[str] mapped_controls_log: List of mapped controls
    :param List[str] unmapped_controls_log: List of unmapped controls
    :param List[str] implemented_controls_log: List of implemented controls
    :return None:
    :rtype None:
    """
    check_file_path("./artifacts", output=False)
    with open("./artifacts/control_implementation.log", "w") as f:
        f.write("|*** Unmapped Controls ***|\n")
        f.write(NEW_LINE_OUTPUT)
        f.write("\n".join(unmapped_controls_log))
        f.write(NEW_LINE_OUTPUT)
        f.write("|*** Mapped Controls ***|\n")
        f.write(NEW_LINE_OUTPUT)
        f.write("\n".join(mapped_controls_log))
        f.write(NEW_LINE_OUTPUT)
        f.write("|*** Already Implemented Controls ***|\n")
        f.write(NEW_LINE_OUTPUT)
        f.write("\n".join(implemented_controls_log))
        f.write(NEW_LINE_OUTPUT)


def get_parameter_value(param: str) -> Dict:
    """
    Get the value of a Parameter

    :param str param: Parameter as a string
    :return: Dictionary of parameter name and value
    :rtype: Dict
    """
    param_dict = dict()
    if ":" in param:
        param_dict["name"] = param.split(":")[0]
        param_dict["value"] = param.split(":")[1] if len(param.split(":")) > 1 else param
    else:
        param_dict["name"] = param
        param_dict["value"] = param
    return param_dict


def load_non_matched_profile_controls(app: Application, regscale_ssp: dict, mapping: List[ProfileMapping]) -> List:
    """Load controls from a given profile mapping that are not matched by the document

    :param Application app: Application instance
    :param dict regscale_ssp: RegScale SSP as a dictionary
    :param List[ProfileMapping] mapping: Profile mapping
    :return: List of newly created implementations
    :rtype: List
    """
    api = Api()
    current_imps = get_current_implementations(app, regscale_ssp["id"])
    if ssp := [
        ssp
        for ssp in api.get(url=urljoin(app.config["domain"], SSP_URL_SUFFIX)).json()
        if ssp["title"] == regscale_ssp["systemName"]
    ]:
        created_imps = []
        ssp_id = ssp[0]["id"]
        existing_controls = {imp["controlID"] for imp in current_imps}
        controls_to_add = [control for control in mapping if control.controlID not in existing_controls]
        logger.info(
            f"Adding {len(controls_to_add)} additional controls from profile",
            record_type="control",
            model_layer="control-implementation",
        )
        existing_control_ids = {imp["controlID"] for imp in current_imps}
        for control in controls_to_add:
            if isinstance(control, dict):
                control_id = control["controlID"]
            elif isinstance(control, ProfileMapping):
                control_id = control.controlID
            else:
                continue
            if control_id not in existing_control_ids:
                implementation = ControlImplementation(
                    parentId=ssp_id,
                    parentModule="securityplans",
                    controlOwnerId=app.config["userId"],
                    status="Not Implemented",
                    controlID=control_id,
                    responsibility=None,
                    implementation=None,
                ).dict()
                logger.info(
                    f"Posting implementation: {control_id}.",
                    record_type="control",
                    model_layer="control-implementation",
                )
                created_imps.append(post_regscale_object(api, app.config, implementation))
        return created_imps


def post_attachments(api: Api, link: str, regscale_ssp: dict) -> None:
    """
    Download and post Attachments to RegScale

    :param Api api: API object
    :param str link: link to download file onary of RegScale SSP
    :param dict regscale_ssp: RegScale SSP
    :rtype: None
    """
    try:
        dl_path = download_file(link["link"])
        logger.info(
            f"Posting linked image to RegScale.. {link}",
            record_type="attachments",
            model_layer="attachments",
        )
        File.upload_file_to_regscale(
            file_name=(dl_path.absolute()),
            parent_id=regscale_ssp["id"],
            parent_module="securityplans",
            api=api,
        )

    except Exception as ex:
        logger.warning(
            f"Unable to download file: {link}\n{ex}",
            record_type="attachments",
            model_layer="attachments",
        )


def posted_embedded_attachments(api: Api, file_path: Path, regscale_ssp: dict) -> None:
    """
    Find and post embedded picture files to RegScale

    :param Api api: API object
    :param Path file_path: file_path
    :param dict regscale_ssp: RegScale SSP
    :return None:
    """
    # Minimum file size for diagrams (50KB) - lowered from 200KB to capture more diagrams
    min_file_size = 50000
    filename = file_path
    logger.info(f"Extracting embedded images from {filename}")
    with zipfile.ZipFile(filename, mode="r") as archive:
        file_dump_path = gettempdir() + os.sep + "imagedump"
        extracted_count = 0
        for file in archive.filelist:
            if file.filename.startswith("word/media/") and file.file_size > min_file_size:
                archive.extract(file, path=file_dump_path)
                extracted_count += 1
                logger.debug(f"Extracted: {file.filename} ({file.file_size} bytes)")
        logger.info(f"Extracted {extracted_count} images from document")
        # Create directories in case they do not exist.
        media_path = file_dump_path + os.sep + "word" + os.sep + "media"
        if not os.path.exists(media_path):
            os.makedirs(media_path)
        uploaded_count = 0
        for filename in os.listdir(file_dump_path + os.sep + "word" + os.sep + "media"):
            full_file_path = os.path.join(file_dump_path + os.sep + "word" + os.sep + "media", filename)
            if os.path.isfile(full_file_path):
                logger.info(
                    f"Posting embedded image to RegScale... {full_file_path}",
                    record_type="attachments",
                    model_layer="attachments",
                )
                try:
                    File.upload_file_to_regscale(
                        file_name=full_file_path,
                        parent_id=regscale_ssp["id"],
                        parent_module="securityplans",
                        api=api,
                    )
                    uploaded_count += 1
                except Exception as e:
                    logger.warning(
                        f"Unable to upload image -- continuing {e}",
                        record_type="attachments",
                        model_layer="attachments",
                    )
        logger.info(
            f"Successfully uploaded {uploaded_count} embedded images to RegScale",
            record_type="attachments",
            model_layer="attachments",
        )


def post_links(
    config: Dict, api: Api, document: Document, file_path: Path, regscale_ssp: Dict, post_embeds: bool = True
) -> None:
    """
    Use XPath to pull data from XML tables and post links to RegScale.

    :param dict config: Application config
    :param Api api: Api object
    :param Document document: Word Document
    :param Path file_path: File path
    :param dict regscale_ssp: RegScale SSP
    :param bool post_embeds: Whether to post embedded items to RegScale
    :rtype: None
    """
    # Post embedded attachments if needed
    if post_embeds:
        posted_embedded_attachments(api, file_path, regscale_ssp)

    # Extract and post attachments from the document tables
    attachments = extract_attachments_from_document(document)

    # Fetch existing links from the API
    existing_links = api.get(f"{config['domain']}/api/links/getAllByParent/{regscale_ssp['id']}/securityplans").json()
    logger.info(f"Found {len(existing_links)} existing links", record_type="links", model_layer="links")

    # Post new links to RegScale
    post_new_links(api, config, regscale_ssp, attachments, existing_links)


def extract_attachments_from_document(document: Document) -> List[Dict[str, str]]:
    """
    Extracts attachments (links) from a document by parsing its tables.

    :param Document document: Word Document to parse
    :return: List of dictionaries with attachment info
    """
    attachments = []
    titles = []

    for table in document.tables:
        if table._cells and "Identification Number" in table.cell(0, 0).text.strip():
            titles = extract_titles_from_table(table)

            for link in table._element.xpath(".//w:hyperlink"):
                inner_run = link.xpath("w:r", namespaces=link.nsmap)[0]
                title = titles.pop()["title"] if titles else inner_run.text

                # Extract the relationship ID and target URL
                r_id = link.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                link_url = document._part.rels[r_id]._target
                attachments.append({"title": title, "link": link_url})

    return attachments


def extract_titles_from_table(table) -> List[Dict[str, str]]:
    """
    Extracts titles from a table by parsing each row's text.

    :param table: The table from the Word document to parse.
    :return: List of dictionaries with title and link data
    """
    titles = []
    previous_text = None
    current_data = {}

    for _, element in enumerate(table._element.xpath(".//w:r/w:t")):
        current_text = element.text.strip()

        if current_text:
            if previous_text and previous_text.lower() == "link":
                current_data["id"] = current_text
            elif validate_date_str(current_text):
                current_data["title"] = f"{current_data.get('title', '')} {previous_text}".strip()
                current_data["date"] = current_text
            elif "date" in current_data:
                current_data["link"] = current_text
                titles.append(current_data.copy())
                current_data = {}

        previous_text = current_text if len(current_data) != 4 else "link"

    titles.reverse()
    return titles


def post_new_links(
    api: Api, config: Dict, regscale_ssp: Dict, attachments: List[Dict[str, str]], existing_links: List[Dict]
) -> None:
    """
    Posts new links to RegScale if they do not already exist.

    :param Api api: Api object
    :param dict config: Application config
    :param dict regscale_ssp: RegScale SSP
    :param List[Dict[str, str]] attachments: List of extracted links
    :param List[Dict] existing_links: List of existing links from RegScale
    :rtype: None
    """
    existing_link_urls = {link["url"] for link in existing_links}

    for attachment in attachments:
        link_data = {
            "id": 0,
            "url": attachment["link"],
            "title": attachment["title"],
            "parentID": regscale_ssp["id"],
            "parentModule": "securityplans",
        }

        if attachment["link"] not in existing_link_urls:
            post_regscale_object(api, config, link_data, endpoint="links")
            post_attachments(api=api, link=attachment, regscale_ssp=regscale_ssp)
        else:
            logger.info(
                f"{attachment['link']} already exists in Security Plan, skipping...",
                record_type="links",
                model_layer="links",
            )


def validate_date_str(date_text: str, fmt: str = "%m/%d/%Y") -> bool:
    """
    Validate provided text is a date in mm/dd/yyyy format

    :param str date_text: Date as a string
    :param str fmt: date format of the date_text, defaults to %m/%d/%Y
    :return: Whether provided text can be converted as a date
    :rtype: bool
    """
    try:
        datetime.strptime(date_text, fmt)
    except ValueError:
        return False
    return True


def get_text_between_headers(full_text: list, start_header: str, end_header: str) -> str:
    """
    Parses text from a document

    :param list full_text: Full text of the document
    :param str start_header: Starting header
    :param str end_header: Where the text ends
    :return: String parsed from document
    :rtype: str
    """
    try:
        start_index = next(i for i, j in enumerate(full_text) if j.lower().strip() == start_header.lower())
        end_index = next(i for i, j in enumerate(full_text) if j.lower().strip() == end_header.lower())
        logger.debug(f"{start_index=}, {end_index=}")
        found_text = " ".join(full_text[start_index + 1 : end_index])
        logger.debug(f"{found_text=}")
        return found_text
    except StopIteration:
        # If the start header is not found, return an empty string or raise an error
        return ""


def replace_content_control(element, namespaces=None):
    """Replaces xml w:sdt tags with the elements found under w:sdtContent
    :param element: lxml.etree.Element
    :param namespaces: dict
    """

    kwargs = {} if not namespaces else {"namespaces": namespaces}  # xpath args can vary for docx objects vs lxml.etree
    sdt_elements = element.xpath(".//w:sdt", **kwargs)  # Get all w:sdt elements
    for e in sdt_elements:
        inner_elements = e.xpath(".//w:sdt", namespaces=e.nsmap)
        for inner_element in inner_elements:
            replace_content_control(inner_element, namespaces=e.nsmap)
        content_elements = e.xpath(".//w:sdtContent/*", namespaces=e.nsmap)  # Get all elements under sdtContent
        for c in content_elements:
            e.addprevious(c)  # Add content outside of sdt tag
        parent = e.getparent()
        if parent is not None:
            parent.remove(e)  # Remove sdt tag


def gather_stakeholders(tables: list, regscale_ssp: dict, document: Document):
    """Gather Stakeholders

    :param list tables: A list of tables from the XML document.
    :param dict regscale_ssp: A dict of RegScale SSP data.
    :param SSP ssp: Object of docx SSP data.
    """
    app = Application()
    api = Api()
    pocs = collect_points_of_contact(tables, document)
    existing_stakeholders = fetch_existing_stakeholders(api, app.config, regscale_ssp)
    logger.info(
        f"Found {len(existing_stakeholders)} existing stakeholders",
        record_type="stakeholders",
        model_layer="stakeholders",
    )
    filtered_pocs = filter_valid_pocs(pocs)

    insert_new_stakeholders(api, app.config, filtered_pocs, existing_stakeholders, regscale_ssp)


def collect_points_of_contact(tables: list, document: Document) -> List:
    """Collect various points of contact

    :param list tables: A list of tables from the XML document.
    :param SSP ssp: SSP Object of docx SSP data.
    :return: A list of points of contact
    :rtype: List
    """
    pocs = list()
    pocs.append(extract_management_poc(tables, pocs))
    pocs.extend(extract_information_poc(tables, document))
    pocs.append(extract_csp_poc(tables, pocs))
    return pocs


def fetch_existing_stakeholders(api: Api, config: dict, regscale_ssp: dict) -> list:
    """Fetch existing stakeholders from the API

    :param Api api: An API instance
    :param dict config: A configuration dictionary
    :param dict regscale_ssp: A dict of RegScale SSP data.
    :return: A list of existing stakeholders
    :rtype: list
    """
    url = urljoin(
        config.get("domain"),
        f"/api/stakeholders/getAllByParent/{str(regscale_ssp.get('id'))}/securityplans",
    )
    response = api.get(url=url, headers={"Content-Type": "application/json"})
    if response.ok:
        logger.info(
            f"Found {len(response.json())} existing stakeholders",
            record_type="stakeholders",
            model_layer="stakeholders",
        )
        return response.json()
    return []


def filter_valid_pocs(pocs: list) -> list:
    """Filter out invalid (non-dict) points of contact

    :param list pocs: A list of points of contact
    :return: A list of valid points of contact
    :rtype: list
    """
    return [poc for poc in pocs if isinstance(poc, dict)]


def insert_new_stakeholders(
    api: Api, config: dict, pocs: list, existing_stakeholders: list, regscale_ssp: dict
) -> None:
    """Insert new stakeholders into the system

    :param Api api: An API instance
    :param dict config: A configuration dictionary
    :param list pocs: A list of points of contact
    :param list existing_stakeholders: A list of existing stakeholders
    :param dict regscale_ssp: A dict of RegScale SSP data.
    :rtype: None
    """
    pocs_inserted = []
    for poc in pocs:
        stakeholder = create_stakeholder_dict(poc, regscale_ssp)
        if should_insert_stakeholder(poc, pocs_inserted, existing_stakeholders):
            post_stakeholder(api, config, stakeholder)
            pocs_inserted.append(stakeholder.get("name").strip())


def extract_management_poc(tables: list, pocs: list) -> Dict:
    """Extract Management POC

    :param list tables: A list of tables from the XML document.
    :param list pocs: A dict of points of contact
    :return: A dict of the Management POC
    :rtype: Dict
    """
    return _extracted_from_gather_stakeholders(
        tables,
        "Owner Information",
        pocs,
        "Information System Management Point of Contact",
    )


def extract_information_poc(tables: list, document: Document) -> List:
    """Extract Information POC

    :param list tables: A list of tables from the XML document.
    :param SSP ssp: SSP Object of docx SSP data.
    :return: A List of the Information POC
    :rtype: List
    """
    return [
        get_contact_info(
            tables,
            key="Information System Technical Point of Contact",
            xpath=TABLE_TAG,
        ),
        get_base_contact(document),
    ]


def extract_csp_poc(tables: List, pocs: List) -> Dict:
    """Extract CSP POC

    :param List tables: A list of tables from the XML document.
    :param List pocs: A list of points of contact
    :return: A dict of the CSP POC
    :rtype: Dict
    """
    return _extracted_from_gather_stakeholders(
        tables,
        "AO Point of Contact",
        pocs,
        "CSP Name Internal ISSO (or Equivalent) Point of Contact",
    )


def create_stakeholder_dict(poc: dict, regscale_ssp: dict) -> dict:
    """Create a stakeholder dictionary

    :param dict poc: A point of contact
    :param dict regscale_ssp: A dict of RegScale SSP data.
    :return: A stakeholder dictionary
    :rtype: dict
    """
    poc = {k.lower(): v for k, v in poc.items()}  # Make this case-insensitive.
    email = name = title = phone = ""
    if "email" in (keys := [key.lower() for key in poc]):
        email = poc[(_ := _check_if_string_in_list_of_string("email", keys)).lower()]
    if "name" in (keys := [key.lower() for key in poc]):
        name = poc[(_ := _check_if_string_in_list_of_string("name", keys)).lower()]
    if "title" in (keys := [key.lower() for key in poc]):
        name = poc[(_ := _check_if_string_in_list_of_string("title", keys)).lower()]
    if "phone" in (keys := [key.lower() for key in poc]):
        name = poc[(_ := _check_if_string_in_list_of_string("phone", keys)).lower()]

    return {
        "name": name,
        "title": title,
        "phone": phone,
        "email": email,
        "address": poc.get("address", "") if "address" in poc else "",
        "parentId": regscale_ssp["id"],
        "parentModule": "securityplans",
    }


def should_insert_stakeholder(poc: Dict, pocs_inserted: List, existing_stakeholders: List) -> bool:
    """Check if the stakeholder should be inserted

    :param Dict poc: A point of contact
    :param List pocs_inserted: A list of pocs already inserted
    :param List existing_stakeholders: A list of existing stakeholders
    :return: True if the stakeholder should be inserted
    :rtype: bool
    """
    return "name" in poc.keys() and (
        poc["name"].strip() not in pocs_inserted
        and poc["name"].strip() not in {guy["name"] for guy in existing_stakeholders if "name" in guy.keys()}
    )


def _check_if_string_in_list_of_string(string: str, list_of_strings: list) -> str:
    """
    Check if a string is in a list of strings

    :param str string: The string to check
    :param list list_of_strings: The list of strings to check
    :return: the found string
    :rtype: str
    """
    dat = any(string in s for s in list_of_strings)
    return {s for s in list_of_strings if string in s}.pop() if dat else ""


def _extracted_from_gather_stakeholders(tables: Any, key: str, pocs: list, key2: str) -> dict:
    """
    Extract system owner and add to pocs

    :param Any tables: Tables from the XML document
    :param str key: Key used to find system owner
    :param list pocs: List of points of contacts from a SSP
    :param str key2: Second key to find system owner details
    :return: Dictionary of system owner details
    :rtype: dict
    """
    system_owner = get_contact_info(tables, key=key, xpath=TABLE_TAG)
    pocs.append(system_owner)
    return get_contact_info(tables, key=key2, xpath=TABLE_TAG)


def post_leveraged_authorizations(table_data: list, ssp_id: int) -> None:
    """
    Function to post leveraged authorizations

    :param list table_data: Data used to post to RegScale
    :param int ssp_id: SSP ID # in RegScale
    :rtype: None
    """
    date_key = "Date Granted"
    app = Application()
    key = "Leveraged Information System Name"
    data = [table for table in table_data if key in table.keys()]
    for la in data:
        if not la.get(key, None):
            continue
        if la.get(date_key) is None or not validate_date_str(la.get(date_key)):
            logger.warning(
                f"Using today's date because of bad Date Granted for {la.get(key)}: {la.get(date_key)}",
                record_type="leveraged-authorizations",
                model_layer="leveraged-authorizations",
            )
            la[date_key] = get_current_datetime()
        try:
            LeveragedAuthorization(
                title=la.get(key, " "),
                servicesUsed=la.get("Leveraged Service Provider Owner"),
                fedrampId=la.get("FedRAMP ID", "Fxxxxxxxxxx"),
                authorizationType=la.get("Authorization Type", "Joint Authorization Board (JAB)"),
                dateAuthorized=la.get("Date Granted"),
                natureOfAgreement=la.get("Nature of Agreement", "Other"),
                dataTypes=la.get("Data Types", TBD),
                authenticationType=la.get("Authentication Type", TBD),
                authorizedUserTypes=la.get("Authorizad User Types", TBD),
                impactLevel=la.get(IMPACT_LEVEL, "High"),
                createdById=app.config.get("userId"),
                securityPlanId=ssp_id,
                ownerId=app.config.get("userId"),
                lastUpdatedById=app.config.get("userId"),
            ).create()
            logger.info(
                f"Leveraged Authorizations for {la.get(key)} created in RegScale.",
                record_type="leveraged-authorizations",
                model_layer="leveraged-authorizations",
            )
        except Exception as e:
            logger.error(
                f"Error creating leveraged authorizations: {e}",
                record_type="leveraged-authorizations",
                model_layer="leveraged-authorizations",
            )


def find_profile_by_name(profile_name: str) -> Dict:
    """Find profile by name

    :param str profile_name: Name of the profile
    :raises ValueError: If the profile is not found
    :return: Dictionary of the profile
    :rtype: Dict
    """
    app = Application()
    api = Api()
    logger.info(
        f"Using the {profile_name} profile to import controls.",
        record_type="profile",
        model_layer="profile",
    )
    profile_response = api.get(url=urljoin(app.config["domain"], "/api/profiles/getList"))

    if profile_response.ok:
        profiles = profile_response.json()
        logger.info(
            f"Found {len(profiles)} profiles in RegScale.",
            record_type="profile",
            model_layer="profile",
        )
    else:
        profiles = []
        logger.error(
            "Unable to get profiles from RegScale.",
            record_type="profile",
            model_layer="profile",
        )
        profile_response.raise_for_status()
    profile = None
    try:
        for profile_obj in profiles:
            if profile_obj["name"] == profile_name:
                profile = profile_obj
        if profile is None:
            raise ValueError(f"Unable to find profile: {profile_name}")
    except Exception as ex:
        logger.error(
            f"Unable to continue, {profile_name} is not found!\n{ex}",
            record_type="profile",
            model_layer="profile",
        )
    return profile


def get_profile_info_by_id(profile_id: int) -> Dict:
    """
    Get a profile by the profile_id from the Regscale Api

    :param int profile_id: The profile_id to get
    :return: Profile
    :rtype: Dict
    """
    profile = None
    try:
        app = Application()
        api = Api()
        profile_response = api.get(urljoin(app.config["domain"], f"/api/profiles/{profile_id}"))
        if profile_response.ok:
            profile = profile_response.json()
        else:
            logger.error(
                "Unable to get profile from RegScale.",
                record_type="profile",
                model_layer="profile",
            )
    except (IndexError, AttributeError) as ex:
        logger.error(
            f"Error Profile, {profile_id} is not found!\n{ex}",
            record_type="profile",
            model_layer="profile",
        )
    return profile


def process_fedramp_docx_by_profile_id(
    file_path: Union[click.Path, str],
    profile_id: int,
    save_data: bool = False,
    load_missing: bool = False,
) -> Any:
    """
    Process a FedRAMP docx by the profile_id from the Regscale Api

    :param Union[click.Path, str] file_path: The file path to the FedRAMP docx
    :param int profile_id: The profile_id to process
    :param bool save_data: Whether to save the data
    :param bool load_missing: Whether to load missing controls
    :return: RegScale SSP
    :rtype: Any
    """
    profile = get_profile_info_by_id(profile_id)
    new_implementations, regscale_ssp = process_fedramp_docx(
        fedramp_file_path=file_path,
        base_fedramp_profile=profile["name"],
        base_fedramp_profile_id=profile["id"],
        save_data=save_data,
        add_missing=load_missing,
        profile=profile,
    )
    # implementation_results
    logger.write_events()
    return (
        "artifacts/import-results.csv",
        {
            "ssp_title": regscale_ssp.get("systemName", "New SSP Default Name"),
            "ssp_id": regscale_ssp.get("id"),
        },
        new_implementations,
    )


def get_profile_mapping(profile_id: int) -> Optional[list]:
    """
    Get a profile mapping by the profile_id from the Regscale Api

    :param int profile_id: The profile_id to get
    :return: Profile Mapping, if found
    :rtype: Optional[list]

    """
    app = Application()
    api = Api()
    profile_mapping = None
    try:
        profile_mapping_resp = api.get(
            urljoin(
                app.config["domain"],
                f"/api/profileMapping/getByProfile/{profile_id}",
            )
        )
        if profile_mapping_resp.ok:
            profile_mapping = profile_mapping_resp.json()
        else:
            logger.error(
                "Unable to get profile mapping from RegScale.",
                record_type="profile-mapping",
                model_layer="profile-mapping",
            )
    except Exception as e:
        logger.error(
            f"Failed to get profile-mappings by profile id with error: {str(e)}",
            record_type="profile-mapping",
            model_layer="profile-mapping",
        )
    return profile_mapping


def parse_ssp_docx_tables(tables: Any) -> Tuple[str, str, str, str, str, str, list]:
    """
    Parse the SSP docx tables

    :param tables: List of tables from the SSP docx
    :return Tuple[str, str, str, str, str, list]: System Status, System Type, Title, Cloud Model, Cloud Sevice, Version, Table Data
    :rtype: Tuple[str, str, str, list]
    """
    count = 0
    title = None
    version = None
    system_status = "Other"
    system_type = SYSTEM_TYPE
    cloud_model = None
    cloud_service = None
    table_data = []
    for table in tables:
        for i, row in enumerate(table.rows):
            checked = False
            rem = row._element
            check_boxes = rem.xpath(".//w14:checked")
            text = (cell.text.strip() for cell in row.cells)
            if check_boxes:
                for checks in check_boxes:
                    if checks.items()[0][1] == "1":
                        count = count + 1
                        checked = True
            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                if not title:
                    dat = [x.strip() for x in keys if x].pop().split("\n")
                    title = dat[1] if len(dat) > 1 else dat[0]
                    if len(dat) > 1:
                        version = dat[2] if len(dat) > 1 else dat[1]
                    if version:
                        version = version.replace("Version", "")
                    else:
                        version = "1.0"
                continue
            row_data = dict(zip(keys, text))
            if checked:
                if SYSTEM_STATUS in row_data:
                    system_status = row_data[SYSTEM_STATUS]
                if SERVICE_ARCHS in row_data:
                    cloud_service = system_type = row_data[SERVICE_ARCHS]
                if DEPLOY_MODEL in row_data:
                    cloud_model = row_data[DEPLOY_MODEL]
            row_data["checked"] = checked
            row_data["element"] = rem
            table_data.append(row_data)
    return (
        system_status,
        system_type,
        title,
        cloud_model,
        cloud_service,
        version,
        table_data,
    )


def process_fedramp_docx(  # noqa: C901
    fedramp_file_path: click.Path,
    base_fedramp_profile: str,
    base_fedramp_profile_id: int,
    save_data: bool = False,
    add_missing: bool = False,
    profile: Dict = None,
) -> Tuple[List, Document]:
    """
    Convert a FedRAMP file to a RegScale SSP

    :param click.Path fedramp_file_path: The click file path object
    :param str base_fedramp_profile: base fedramp profile
    :param int base_fedramp_profile_id: base fedramp profile id
    :param bool save_data: Whether to save the data
    :param bool add_missing: Whether to add missing controls
    :param Dict profile: The profile to use
    :return: Tuple of new implementations count and RegScale SSP
    :rtype: Tuple[List, SSP]
    """
    # If list of controls is more than profile mapping, make sure i get them from somewhere? Get base catalog from profile.
    load_missing = add_missing
    app = Application()
    api = Api()
    ssp = SSP(fedramp_file_path)
    document = Document(fedramp_file_path)
    for p in document.paragraphs:
        replace_content_control(p._element)

    full_text = [para.text for para in document.paragraphs]
    system_text_lookup = "System Function or Purpose"
    description_lookup_str = "General System Description"

    description = get_text_between_headers(
        full_text,
        start_header=description_lookup_str,
        end_header=system_text_lookup,
    )

    environment = get_text_between_headers(
        full_text,
        start_header="SYSTEM ENVIRONMENT AND INVENTORY",
        end_header="Data Flow",
    )

    purpose = get_text_between_headers(
        full_text,
        start_header=system_text_lookup,
        end_header="System Description:",
    )
    if not purpose or purpose == "":
        purpose = get_text_between_headers(
            full_text,
            start_header=system_text_lookup,
            end_header="Information System Components and Boundaries",
        )

    confidentiality = "Low"
    integrity = "Low"
    availability = "Low"
    tables = get_tables(document)
    security_objective = get_xpath_data_detailed(
        tables,
        key="Security Objective",
        ident="Confidentiality",
        xpath=TABLE_TAG,
        count_array=[2, 4, 6],
    )

    availability = (
        security_objective["availability"].split(" ")[0] if "availability" in security_objective else availability
    )
    confidentiality = (
        security_objective["confidentiality"].split(" ")[0]
        if "confidentiality" in security_objective
        else confidentiality
    )
    integrity = security_objective["integrity"].split(" ")[0] if "integrity" in security_objective else integrity

    (
        system_status,
        system_type,
        title,
        cloud_model,
        cloud_service,
        version,
        table_data,
    ) = parse_ssp_docx_tables(document.tables)

    mdeploypublic = True if "multiple organizations " in cloud_model else False
    mdeploypriv = True if "specific organization/agency" in cloud_model else False
    mdeploygov = True if "organizations/agencies" in cloud_model else False
    mdeployhybrid = True if "shared across all clients/agencies" in cloud_model else False

    msaas = True if SYSTEM_TYPE in cloud_service else False
    mpaas = True if SYSTEM_TYPE in cloud_service and not msaas else False
    miaas = True if "General Support System" in cloud_service else False
    mother = True if "Explain:" in cloud_service else False

    privacydata = get_xpath_privacy_detailed(
        tables,
        key="Does the ISA collect, maintain, or share PII in any identifiable form?",
        xpath=TABLE_TAG,
        count_array=[0, 2, 4, 6],
    )

    sysinfo = get_xpath_sysinfo_detailed(
        tables,
        key="Unique Identifier",
        xpath=TABLE_TAG,
        count_array=[3, 5],
    )
    if sysinfo["systemname"]:
        title = sysinfo["systemname"].strip() if "systemname" in sysinfo else title
    if sysinfo["uniqueidentifier"]:
        uniqueidentifier = sysinfo["uniqueidentifier"].strip() if "uniqueidentifier" in sysinfo else None
    else:
        uniqueidentifier = None

    prepdata = get_xpath_prepdata_detailed(
        tables,
        key="Identification of Organization that Prepared this Document",
        ident=ORGANIZATION_TAG,
        xpath=TABLE_TAG,
    )
    preporgname = prepdata["orgname"] if "orgname" in prepdata else None
    prepaddress = prepdata["street"] if "street" in prepdata else None
    prepoffice = prepdata["office"] if "office" in prepdata else None
    prepcitystate = prepdata["citystate"] if "citystate" in prepdata else None
    cspdata = get_xpath_prepdata_detailed(
        tables,
        key="Identification of Cloud Service Provider",
        ident=ORGANIZATION_TAG,
        xpath=TABLE_TAG,
    )
    csporgname = cspdata["orgname"] if "orgname" in cspdata else None
    cspaddress = cspdata["street"] if "street" in cspdata else None
    cspoffice = cspdata["office"] if "office" in cspdata else None
    cspcitystate = cspdata["citystate"] if "citystate" in cspdata else None
    status = "Operational" if "in production" in system_status else "Other"
    # Links are posted to links mapped to ssp
    # post_links(app, table_data, ssp_id)
    # Parts will go in implementation fields.
    if base_fedramp_profile_id:
        profile = get_profile_info_by_id(profile_id=base_fedramp_profile_id)
    if not profile:
        profile = find_profile_by_name(base_fedramp_profile) or {}
    profile_mapping = ProfileMapping.get_by_profile(profile_id=profile.get("id"))
    if len(profile_mapping) == 0:
        error_and_exit(f"Unable to continue, please load {base_fedramp_profile} with controls!")

    logger.info(
        f"Utilizing profile: {profile.get('name')}",
        record_type="profile",
        model_layer="profile",
    )
    args = {
        "profile": profile,
        "title": title,
        "otheridentifier": uniqueidentifier,
        # get_profile_mapping(profile["id"]))
        "version": version,
        "confidentiality": confidentiality,
        "integrity": integrity,
        "availability": availability,
        "status": status,
        "system_type": system_type,
        # "ssp": ssp,
        "revision": revision(document),
        "description": description,
        "environment": environment,
        "purpose": purpose,
        "modeiaas": miaas,
        "modepaas": mpaas,
        "modeother": mother,
        "modesaas": msaas,
        "deploypubic": mdeploypublic,
        "deployprivate": mdeploypriv,
        "deploygov": mdeploygov,
        "deployhybrid": mdeployhybrid,
        "preporgname": preporgname,
        "prepaddress": prepaddress,
        "prepoffice": prepoffice,
        "prepcitystate": prepcitystate,
        "csporgname": csporgname,
        "cspaddress": cspaddress,
        "cspoffice": cspoffice,
        "cspcitystate": cspcitystate,
    }
    regscale_ssp = create_initial_ssp(args)

    try:
        logger.info("Parsing and creating Privacy Data", record_type="privacy", model_layer="privacy")
        create_privacy_data(app=app, privacy_data=privacydata, ssp_id=regscale_ssp.get("id"))
        logger.info(
            "Successfully Created Privacy data.",
            record_type="privacy",
            model_layer="privacy",
        )
    except Exception as e:
        logger.error(
            f"Unable to create privacy record: {e}",
            record_type="privacy",
            model_layer="privacy",
        )

    try:
        logger.info(
            "Parsing and creating System Information", record_type="responsible-roles", model_layer="responsible-roles"
        )
        create_responsible_roles(app, table_data, ssp_id=regscale_ssp["id"])
        ctrl_roles = post_responsible_roles(app, table_data, ssp_id=regscale_ssp["id"])
    except Exception as e:
        logger.error(
            f"Unable to gather responsible roles: {e}",
            record_type="responsible-roles",
            model_layer="responsible-roles",
        )

    try:
        logger.info("Parsing and creating Stakeholders", record_type="stakeholder", model_layer="stakeholder")
        gather_stakeholders(tables, regscale_ssp, document)
    except Exception as e:
        logger.error(
            f"Unable to gather stakeholders: {e}",
            record_type="stakeholder",
            model_layer="stakeholder",
        )
    try:
        logger.info("Parsing and creating Interconnects", record_type="interconnect", model_layer="interconnects")
        post_interconnects(app, table_data, regscale_ssp)
    except Exception as e:
        logger.error(
            f"Unable to gather interconnects: {e}",
            record_type="interconnect",
            model_layer="interconnects",
        )
    try:
        logger.info(
            "Parsing and creating Ports and Protocols", record_type="ports-protocols", model_layer="ports-protocols"
        )
        tables_dict = tables_to_dict(document)
        ports_table_data = [row for t in tables_dict for row in t if "Ports (TCP/UDP)*" in row]
        post_ports(app, ports_table_data, ssp_id=regscale_ssp["id"])
    except Exception as e:
        logger.error(
            f"Unable to gather ports: {e}",
            record_type="ports-protocols",
            model_layer="ports-protocols",
        )
    try:
        logger.info("Parsing and creating Links", record_type="links", model_layer="links")
        post_links(
            config=app.config, api=api, document=document, file_path=fedramp_file_path, regscale_ssp=regscale_ssp
        )
    except Exception as e:
        logger.error(
            f"Unable to gather links: {e}",
            record_type="links",
            model_layer="links",
        )
    try:
        logger.info(
            "Parsing and creating implementations", record_type="implementations", model_layer="implementations"
        )
        new_implementations = post_implementations(
            app=app,
            ssp_obj=ssp,
            regscale_ssp=regscale_ssp,
            mapping=profile_mapping,
            ctrl_roles=ctrl_roles,
            save_data=save_data,
            load_missing=load_missing,
        )
    except Exception as e:
        logger.debug(e, exc_info=True)
        logger.error(
            f"Unable to gather implementations: {e}",
            record_type="implementations",
            model_layer="implementations",
        )
        new_implementations = []
    try:
        logger.info(
            "Parsing and creating Leveraged Authorizations",
            record_type="leveraged-authorizations",
            model_layer="leveraged-authorizations",
        )
        post_leveraged_authorizations(table_data, ssp_id=regscale_ssp.get("id"))
    except Exception as e:
        logger.error(
            f"Unable to gather leveraged authorizations: {e}",
            record_type="leveraged-authorizations",
            model_layer="leveraged-authorizations",
        )
    return new_implementations, regscale_ssp


def tables_to_dict(document: Document) -> List[List[Dict]]:
    """
    Convert tables in a document to a list of dictionaries
    :param Document document: document to convert
    :return: List of table and its rows as dictionaries
    :rtype: List[List[Dict]]
    """
    tables_as_dicts = []

    for table in document.tables:
        replace_content_control(table._element)  # remove content controls
        # Assuming first row is the header
        keys = [cell.text.strip() for cell in table.rows[0].cells]
        table_dict = []

        # Iterate over the rest of the rows
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            # Create a dict by zipping keys and values together
            row_dict = dict(zip(keys, values))
            table_dict.append(row_dict)

        tables_as_dicts.append(table_dict)

    return tables_as_dicts


def revision(doc: Document) -> str:
    """The SSP revision."""
    result = None
    regex = r"Version\s#*([\d.]+)"
    try:
        fed_ramp_revision_string = doc.tables[0].cell(0, 0).text  # First cell of the first table.
        result = re.search(regex, fed_ramp_revision_string).group(1)
    except (AttributeError, IndexError):
        print("Warning, unable to return revision information.")
    return result


def _determine_overall_categorization(
    confidentiality: str, integrity: str, availability: str, default: str = "Moderate"
) -> str:
    """
    Determine the overall FIPS 199 categorization based on C/I/A values.
    The overall categorization is the highest of the three values.

    :param str confidentiality: Confidentiality level (Low, Moderate, High)
    :param str integrity: Integrity level (Low, Moderate, High)
    :param str availability: Availability level (Low, Moderate, High)
    :param str default: Default value if unable to determine
    :return: The overall categorization
    :rtype: str
    """
    level_order = {"low": 1, "moderate": 2, "high": 3}
    levels = [
        level_order.get(confidentiality.lower(), 0),
        level_order.get(integrity.lower(), 0),
        level_order.get(availability.lower(), 0),
    ]
    max_level = max(levels)

    if max_level == 3:
        return "High"
    elif max_level == 2:
        return "Moderate"
    elif max_level == 1:
        return "Low"
    return default


def create_initial_ssp(args: Dict) -> Any:
    """
    Create an initial SSP

    :param Dict args: Arguments to create the initial SSP
    :return: SSP
    :rtype: Any
    """
    app = Application()
    api = Api()
    today_dt = date.today()
    expiration_date = date(today_dt.year + 3, today_dt.month, today_dt.day)
    default = "Moderate"
    profile = args.get("profile")
    title = args.get("title")
    version = args.get("version", "")
    otheridentifier = args.get("otheridentifier", "")
    confidentiality = capitalize_words(args.get("confidentiality", default))
    integrity = capitalize_words(args.get("integrity", default))
    availability = capitalize_words(args.get("availability", default))
    status = args.get("status", "Operational")
    system_type = args.get("system_type", SYSTEM_TYPE)
    revision = args.get("revision", "1.0")
    description = args.get("description", "Unable to determine System Description")
    environment = args.get("environment", "")
    purpose = args.get("purpose", "")
    modeliaas = args.get("modeiaas", False)
    modelother = args.get("modeother", False)
    modelpaas = args.get("modepaas", False)
    modelsaas = args.get("modesaas", False)
    deploygov = args.get("deploygov", False)
    deployhybrid = args.get("deployhybrid", False)
    deployprivate = args.get("deployprivate", False)
    deploypublic = args.get("deploypublic", False)
    deployother = args.get("deployother", False)
    preporgname = args.get("preporgname", "")
    prepaddress = args.get("prepaddress", "")
    prepoffice = args.get("prepoffice", "")
    prepcitystate = args.get("prepcitystate", "")
    csporgname = args.get("csporgname", "")
    cspaddress = args.get("cspaddress", "")
    cspoffice = args.get("cspoffice", "")
    cspcitystate = args.get("cspcitystate", "")

    # Additional fields for FedRAMP compliance
    authorization_boundary = args.get("authorization_boundary", "")
    network_architecture = args.get("network_architecture", "")
    data_flow = args.get("data_flow", "")
    fedramp_id = args.get("fedramp_id", "")
    fedramp_authorization_type = args.get("fedramp_authorization_type", "")
    fedramp_authorization_status = args.get("fedramp_authorization_status", "")
    laws_and_regulations = args.get("laws_and_regulations", "")

    # Determine overall categorization based on confidentiality, integrity, availability
    overall_cat = _determine_overall_categorization(confidentiality, integrity, availability, default)

    regscale_ssp = SecurityPlan(
        dateSubmitted=get_current_datetime(),
        expirationDate=expiration_date.strftime(DATE_FORMAT),
        approvalDate=expiration_date.strftime(DATE_FORMAT),
        parentId=profile["id"],
        parentModule="profiles",
        systemName=title or "Unable to determine System Name",
        otherIdentifier=otheridentifier,
        confidentiality=confidentiality,
        integrity=integrity,
        availability=availability,
        status=status,
        bDeployGov=deploygov,
        bDeployHybrid=deployhybrid,
        bDeployPrivate=deployprivate,
        bDeployPublic=deploypublic,
        bDeployOther=deployother,
        bModelIaaS=modeliaas,
        bModelOther=modelother,
        bModelPaaS=modelpaas,
        bModelSaaS=modelsaas,
        createdById=app.config["userId"],
        lastUpdatedById=app.config["userId"],
        systemOwnerId=app.config["userId"],
        planAuthorizingOfficialId=app.config["userId"],
        planInformationSystemSecurityOfficerId=app.config["userId"],
        systemType=system_type,
        overallCategorization=overall_cat,
        description=description,
        purpose=purpose,
        environment=environment,
        authorizationBoundary=authorization_boundary,
        networkArchitecture=network_architecture,
        dataFlow=data_flow,
        lawsAndRegulations=laws_and_regulations,
        fedrampId=fedramp_id,
        fedrampAuthorizationType=fedramp_authorization_type,
        fedrampAuthorizationStatus=fedramp_authorization_status or ("Authorized" if status == "Operational" else ""),
        fedrampAuthorizationLevel=overall_cat,
        executiveSummary=f"Revision: {revision}",
        version=version,
        prepOrgName=preporgname,
        prepAddress=prepaddress,
        prepOffice=prepoffice,
        prepCityState=prepcitystate,
        cspOrgName=csporgname,
        cspAddress=cspaddress,
        cspOffice=cspoffice,
        cspCityState=cspcitystate,
    )

    if regscale_ssp.status != "Operational":
        regscale_ssp.explanationForNonOperational = "Unable to determine status from SSP during FedRAMP .docx import."
    existing_security_plans_reponse = api.get(
        url=urljoin(app.config["domain"], SSP_URL_SUFFIX),
    )
    existing_security_plans = []
    if not existing_security_plans_reponse.ok:
        logger.info("No Security Plans found")
    else:
        existing_security_plans = existing_security_plans_reponse.json()
    if len(existing_security_plans) >= 1:
        existing_ssps_dict = {
            ssp.get("title").lower(): ssp for ssp in existing_security_plans if ssp and "title" in ssp
        }
        if regscale_ssp.systemName.lower() not in existing_ssps_dict:
            regscale_ssp = regscale_ssp.create()
        else:
            existing_ssp = existing_ssps_dict.get(regscale_ssp.systemName.lower())
            regscale_ssp.id = existing_ssp.get("id")
            regscale_ssp = regscale_ssp.save()
    else:
        regscale_ssp = regscale_ssp.create()
    return regscale_ssp.dict()
