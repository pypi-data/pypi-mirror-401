"""This module contains the SSPDocParser class,
which is used to parse an SSP .docx file and return a list of dictionaries representing the table data.
"""

import logging
import re
import zipfile
from collections import defaultdict
from typing import Any, Dict, List, Optional

import docx
import lxml.etree as etree

logger = logging.getLogger("regscale")

# Precompiled pattern for detecting generic image captions like "Image 1", "Figure 2"
GENERIC_CAPTION_PATTERN = re.compile(
    r"^(?:figure|image|diagram|illustration|exhibit)\s*\d+[:\s]*$",
    re.IGNORECASE,
)

# Pattern to detect broken Word cross-reference errors
BROKEN_REFERENCE_PATTERN = re.compile(
    r"error[!\s]*reference\s*source\s*not\s*found",
    re.IGNORECASE,
)


def _is_broken_reference_caption(caption: str) -> bool:
    """
    Check if a caption contains a broken Word cross-reference error.

    Word displays "Error! Reference source not found" when a cross-reference
    link is broken (e.g., "See Figure X below" where the reference is invalid).

    :param str caption: The caption text to check.
    :return: True if the caption contains a broken reference error, False otherwise.
    :rtype: bool
    """
    if not caption:
        return False
    return bool(BROKEN_REFERENCE_PATTERN.search(caption))


# Standard caption pattern prefixes
CAPTION_PATTERNS = ("figure", "image", "diagram", "illustration", "exhibit")

# WordprocessingML text element namespace
WML_TEXT_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"


def _extract_paragraph_text(elem: Any) -> str:
    """
    Extract text content from a Word paragraph element.

    :param elem: XML element containing the paragraph.
    :return: Concatenated text from all text nodes.
    :rtype: str
    """
    return "".join(node.text for node in elem.iter(WML_TEXT_NS) if node.text)


def _is_valid_caption(para_text: str) -> bool:
    """
    Check if paragraph text looks like a valid caption.

    :param para_text: The paragraph text to check.
    :return: True if text starts with a caption pattern.
    :rtype: bool
    """
    para_lower = para_text.strip().lower()
    return any(para_lower.startswith(pattern) for pattern in CAPTION_PATTERNS)


def _enhance_generic_caption(caption: str, context: str) -> str:
    """
    Enhance a generic caption with section context if applicable.

    :param caption: The caption text.
    :param context: Section heading context.
    :return: Enhanced caption or original.
    :rtype: str
    """
    if _is_generic_image_caption(caption) and context:
        return f"{caption} - {context}"
    return caption


def _is_generic_image_caption(caption: str) -> bool:
    """
    Check if a caption is generic (just "Image X" or "Figure X" without description).

    :param str caption: The caption text to check.
    :return: True if the caption is generic, False otherwise.
    :rtype: bool
    """
    if not caption:
        return False
    # Check if caption matches generic pattern (e.g., "Image 1", "Figure 2-1")
    # A generic caption has no descriptive text after the number
    cleaned = caption.strip()
    # If it matches patterns like "Figure 1" or "Image 2:" without more content
    if GENERIC_CAPTION_PATTERN.match(cleaned):
        return True
    # Also check for very short captions that are just "Figure X-Y" without description
    # Use bounded quantifiers to prevent ReDoS attacks
    if len(cleaned.split()) <= 3 and re.match(
        r"^(?:figure|image)\s{1,10}[\d\-\.]+\s{0,5}:?\s{0,5}$", cleaned, re.IGNORECASE
    ):
        return True
    return False


class SSPDocParser:
    """
    Parses an SSP .docx file and returns a list of dictionaries representing the table data.
    """

    docx_path: str
    xml_content: bytes
    tables: List
    text: Dict

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = docx.Document(self.docx_path)

    def parse(self) -> List[Dict]:
        """
        Parses the .docx file and returns a list of dictionaries representing the table data.
        :return: A list of dictionaries, each representing a table's data.
        :rtype: List[Dict]
        """
        self.tables = self.parse_xml_for_tables(self.docx_to_xml(self.docx_path))
        self.text = self.group_paragraphs_by_headings()
        return self.tables

    # Namespaces for DOCX XML parsing
    DOCX_NAMESPACES = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }

    @staticmethod
    def _extract_heading_name(heading_text: str) -> Optional[str]:
        """Extract clean heading name by removing section numbers. Skip if contains broken references."""
        if not heading_text:
            return None
        # Remove section numbers like "9.1 " from the beginning
        cleaned = re.sub(r"^[\d\.]+\s*", "", heading_text.strip())
        if not cleaned:
            return None
        # Skip entirely if heading contains broken Word cross-reference errors
        if _is_broken_reference_caption(cleaned):
            logger.debug("Skipping heading with broken reference: %s", cleaned[:50])
            return None
        return cleaned

    @staticmethod
    def _find_image_relationship_id(elem: Any, namespaces: Dict) -> Optional[str]:
        """Find the relationship ID for an image element."""
        xpath_patterns = [
            ".//a:blip/@r:embed",
            ".//a:blip/@r:link",
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            "/@{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            "/@{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
        ]
        for xpath in xpath_patterns:
            try:
                results = elem.xpath(xpath, namespaces=namespaces)
                if results:
                    return results[0]
            except Exception:
                continue

        # Fallback: search for any embed or link attribute in blip elements
        try:
            for blip in elem.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed:
                    return embed
                link = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link")
                if link:
                    return link
        except Exception:
            pass

        return None

    @staticmethod
    def _find_figure_caption(elem: Any, current_heading: str = "") -> Optional[str]:
        """
        Look for a Figure or Image caption in the paragraphs following an image.

        Handles various caption formats:
        - "Figure X-X: Description"
        - "Figure X: Description"
        - "Image X: Description"
        - Generic "Image 1", "Image 2" (enhanced with section context)

        :param elem: The XML element containing the image.
        :param current_heading: The current section heading for context.
        :return: The caption text if found, None otherwise.
        """
        next_elem = elem.getnext()
        check_count = 0

        while next_elem is not None and check_count < 5:
            caption = SSPDocParser._extract_caption_from_element(next_elem, current_heading)
            if caption:
                return caption
            if next_elem.tag.endswith("p"):
                check_count += 1
            next_elem = next_elem.getnext()
        return None

    @staticmethod
    def _extract_caption_from_element(elem: Any, context: str = "") -> Optional[str]:
        """
        Extract caption text from an element if it looks like a valid caption.

        :param elem: XML element to check.
        :param context: Section heading for context enhancement.
        :return: Caption text or None.
        """
        if not elem.tag.endswith("p"):
            return None

        para_text = _extract_paragraph_text(elem)
        if not para_text or not _is_valid_caption(para_text):
            return None

        caption = para_text.strip()
        if _is_broken_reference_caption(caption):
            logger.debug("Skipping broken reference caption: %s", caption[:50])
            return None

        return _enhance_generic_caption(caption, context)

    @staticmethod
    def _find_caption_before_image(elem: Any) -> Optional[str]:
        """Look for a caption in paragraphs before the image (some templates put captions above)."""
        prev_elem = elem.getprevious()
        check_count = 0

        while prev_elem is not None and check_count < 3:
            # No context for before-image captions - enhancement happens in caller
            caption = SSPDocParser._extract_caption_from_element(prev_elem, "")
            if caption:
                return caption
            if prev_elem.tag.endswith("p"):
                check_count += 1
            prev_elem = prev_elem.getprevious()
        return None

    # Image marker strings to detect embedded images
    IMAGE_MARKERS = ("graphicData", "blip", "drawing", "pict", "imagedata")

    def get_figure_captions(self) -> Dict:
        """
        Fetches the figure captions from the .docx file.

        Returns a dictionary mapping relationship IDs to caption text.
        Looks for Figure captions following images, checks for captions
        before images, and falls back to the current section heading
        if no caption is found.

        :return: A dictionary mapping image relationship IDs to their captions.
        :rtype: Dict
        """
        captions: Dict[str, str] = {}
        heading_image_count: Dict[str, int] = defaultdict(int)
        stats = {"found": 0, "with_rid": 0}

        # Process paragraph images
        current_heading = self._process_paragraph_images(captions, heading_image_count, stats)

        # Process table images
        self._process_table_images(captions, heading_image_count, current_heading)

        logger.info(
            "Image caption scan: %d images detected, %d with relationship IDs, %d caption mappings",
            stats["found"],
            stats["with_rid"],
            len(captions),
        )
        return captions

    def _process_paragraph_images(
        self, captions: Dict[str, str], heading_counts: Dict[str, int], stats: Dict[str, int]
    ) -> str:
        """
        Process images in document paragraphs.

        :param captions: Dictionary to populate with caption mappings.
        :param heading_counts: Counter for images per heading.
        :param stats: Statistics dictionary for tracking.
        :return: The last heading found.
        """
        current_heading = "Document_Image"

        for paragraph in self.doc.paragraphs:
            current_heading = self._update_heading(paragraph, current_heading)
            elem = paragraph._element
            if not self._element_has_image(elem):
                continue

            stats["found"] += 1
            r_id = self._get_relationship_id(elem)
            if not r_id:
                logger.debug("Found image but could not extract relationship ID")
                continue

            stats["with_rid"] += 1
            caption = self._resolve_caption(elem, current_heading, heading_counts)
            captions[r_id] = caption
            logger.debug("Mapped relationship %s to: %s", r_id, caption)

        return current_heading

    def _process_table_images(
        self, captions: Dict[str, str], heading_counts: Dict[str, int], current_heading: str
    ) -> None:
        """
        Process images in document tables.

        :param captions: Dictionary to populate with caption mappings.
        :param heading_counts: Counter for images per heading.
        :param current_heading: Current section heading for context.
        """
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_cell_images(cell, captions, heading_counts, current_heading)

    def _process_cell_images(
        self, cell: Any, captions: Dict[str, str], heading_counts: Dict[str, int], heading: str
    ) -> None:
        """Process images within a table cell."""
        for paragraph in cell.paragraphs:
            elem = paragraph._element
            if not self._element_has_image(elem):
                continue

            r_id = self._get_relationship_id(elem)
            if not r_id or r_id in captions:
                continue

            caption = self._resolve_table_caption(elem, cell, heading, heading_counts)
            captions[r_id] = caption
            logger.debug("Mapped table image %s to: %s", r_id, caption)

    def _update_heading(self, paragraph: Any, current: str) -> str:
        """Update current heading if paragraph is a heading."""
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            cleaned = self._extract_heading_name(paragraph.text)
            if cleaned:
                return cleaned
        return current

    def _element_has_image(self, elem: Any) -> bool:
        """Check if element contains an image."""
        elem_xml = elem.xml if hasattr(elem, "xml") else str(etree.tostring(elem))
        return any(marker in elem_xml for marker in self.IMAGE_MARKERS)

    def _get_relationship_id(self, elem: Any) -> Optional[str]:
        """Get relationship ID for an image element using multiple methods."""
        r_id = self._find_image_relationship_id(elem, self.DOCX_NAMESPACES)
        if not r_id:
            r_id = self._find_image_relationship_id_extended(elem)
        return r_id

    def _resolve_caption(self, elem: Any, heading: str, heading_counts: Dict[str, int]) -> str:
        """Resolve caption for a paragraph image."""
        caption = self._find_figure_caption(elem, heading)
        if not caption:
            caption = self._find_caption_before_image(elem)
            caption = _enhance_generic_caption(caption, heading) if caption else None

        caption = self._validate_caption(caption)
        return caption if caption else self._generate_fallback_caption(heading, heading_counts)

    def _resolve_table_caption(self, elem: Any, cell: Any, heading: str, heading_counts: Dict[str, int]) -> str:
        """Resolve caption for a table image."""
        caption = self._find_figure_caption(elem, heading) or self._find_caption_before_image(elem)
        if not caption:
            caption = self._get_cell_text_caption(cell)

        caption = self._validate_caption(caption)
        return caption if caption else self._generate_fallback_caption(heading, heading_counts)

    @staticmethod
    def _get_cell_text_caption(cell: Any) -> Optional[str]:
        """Try to get caption from cell text."""
        cell_text = cell.text.strip() if cell.text else ""
        if cell_text and len(cell_text) < 200 and not _is_broken_reference_caption(cell_text):
            return cell_text
        return None

    @staticmethod
    def _validate_caption(caption: Optional[str]) -> Optional[str]:
        """Validate caption and return None if it contains broken references."""
        if caption and _is_broken_reference_caption(caption):
            logger.debug("Skipping caption with broken reference: %s", caption[:50])
            return None
        return caption

    @staticmethod
    def _generate_fallback_caption(heading: str, heading_counts: Dict[str, int]) -> str:
        """Generate a fallback caption based on heading."""
        heading_counts[heading] += 1
        count = heading_counts[heading]
        return f"{heading}_{count}" if count > 1 else heading

    @staticmethod
    def _find_image_relationship_id_extended(elem: Any) -> Optional[str]:
        """Find image relationship ID using extended XPath patterns for various embedding methods."""
        # Extended patterns for VML images, inline shapes, etc.
        extended_patterns = [
            ".//{urn:schemas-microsoft-com:vml}imagedata/@{urn:schemas-microsoft-com:office:office}relid",
            ".//{urn:schemas-microsoft-com:vml}imagedata/@r:id",
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing//"
            "{http://schemas.openxmlformats.org/drawingml/2006/main}blip/@"
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
        ]
        for xpath in extended_patterns:
            try:
                results = elem.xpath(xpath)
                if results:
                    return results[0]
            except Exception:
                continue
        return None

    def group_paragraphs_by_headings(self) -> Dict:
        """
        Groups the paragraphs in the .docx file by their headings.
        :return: A dictionary mapping headings to their paragraphs.
        :rtype: Dict
        """
        grouped_text = defaultdict(list)
        current_heading = None

        for para in self.doc.paragraphs:
            # Check if the paragraph is a heading based on its style
            if para.style.name.startswith("Heading"):
                # A new heading is encountered; update the current heading variable
                current_heading = para.text
            elif current_heading:
                # Add the paragraph text under the current heading key
                grouped_text[current_heading].append(para.text)
            # Note: Paragraphs before the first heading won't be included

        return dict(grouped_text)

    @staticmethod
    def docx_to_xml(docx_path: str) -> bytes:
        """
        Converts a .docx file to XML.
        :param str docx_path: The path to the .docx file.
        :return: The XML content of the .docx file.
        :rtype: bytes
        """
        with zipfile.ZipFile(docx_path, "r") as docx_zip:
            xml_content = docx_zip.read("word/document.xml")
        return xml_content

    @staticmethod
    def parse_checkbox_string(text: str) -> List[Dict]:
        """
        Parses a string like "Implementation Status (check all that apply): ☒ Implemented ..." and returns a list of dictionaries.

        :param str text: The text to parse.
        :return: A list of dictionaries, each representing a status option and its status.
        :rtype: List[Dict]
        """

        # Regular expression pattern to match status options change .*? to [^☒☐]* to match any character except ☒ and ☐
        pattern = r"([☒☐]\s*([^☒☐]*))(?=[☒☐]|$)"

        # Find all non-overlapping matches with the pattern
        matches = re.findall(pattern, text)

        # Process matches and create dictionaries
        dict_list = []
        for match in matches:
            status_indicator, option_text = match
            dict_list.append({option_text.strip(): status_indicator.startswith("☒")})

        return dict_list

    def parse_xml_for_tables(self, xml_content: bytes) -> List[Dict]:
        """
        Parses the XML content for tables and returns a list of dictionaries representing the table data.

        :param bytes xml_content: The XML content to parse.
        :return: A list of dictionaries, each representing a table's data.
        :rtype: List[Dict]
        """
        # Parse the XML content
        tree = etree.ElementTree(etree.fromstring(xml_content))
        root = tree.getroot()

        # Define namespace map to handle XML namespaces
        # Extended namespace support for page break detection
        namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",  # noqa
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",  # noqa
            "w15": "http://schemas.microsoft.com/office/word/2012/wordml",  # noqa
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",  # noqa
        }
        return self.parse_tables(root, namespaces)

    @staticmethod
    def extract_table_headers(table: Any, namespaces: Any) -> List:
        """
        Extracts headers from a table.

        :param Any table: The table element to extract headers from.
        :param Any namespaces: The XML namespaces.
        :return: A list of headers.
        :rtype: List
        """
        headers = []
        for cell in table.findall(".//w:tr[1]//w:tc", namespaces):
            cell_text = "".join(node.text for node in cell.findall(".//w:t", namespaces) if node.text)
            headers.append(cell_text.strip() if cell_text else "")
        return headers

    @staticmethod
    def has_page_break(para: Any, namespaces: Dict) -> bool:
        """
        Check if a paragraph contains a page break marker.

        :param Any para: The paragraph element to check.
        :param Dict namespaces: The XML namespaces.
        :return: True if a page break is found, False otherwise.
        :rtype: bool
        """
        # Check for lastRenderedPageBreak (automatic page break marker)
        page_breaks = para.findall(".//w:lastRenderedPageBreak", namespaces)
        # Check for manual page breaks (w:br with type='page')
        manual_breaks = para.findall(".//w:br", namespaces)
        manual_page_breaks = [br for br in manual_breaks if br.get("{%s}type" % namespaces.get("w", "")) == "page"]
        return bool(page_breaks or manual_page_breaks)

    @staticmethod
    def _process_run_element(child: Any, namespaces: Dict) -> Optional[str]:
        """
        Process a single element within a run and return its text representation.

        :param Any child: The child element to process.
        :param Dict namespaces: The XML namespaces.
        :return: The text representation or None if no text.
        :rtype: Optional[str]
        """
        tag = etree.QName(child).localname

        if tag == "t" and child.text:
            return child.text
        elif tag == "br":
            br_type = child.get("{%s}type" % namespaces.get("w", ""))
            if br_type == "page":
                logger.debug("Manual page break found within cell content")
            return "\n"
        elif tag == "lastRenderedPageBreak":
            logger.debug("Automatic page break marker found, continuing")
            return None
        elif tag == "tab":
            return "\t"
        return None

    @staticmethod
    def _extract_paragraph_text(para: Any, namespaces: Dict) -> str:
        """
        Extract text from a paragraph element, processing all runs.

        :param Any para: The paragraph element to process.
        :param Dict namespaces: The XML namespaces.
        :return: The text content from the paragraph.
        :rtype: str
        """
        para_parts = []
        for run in para.findall(".//w:r", namespaces):
            for child in run:
                text = SSPDocParser._process_run_element(child, namespaces)
                if text is not None:
                    para_parts.append(text)
        return "".join(para_parts)

    @staticmethod
    def fetch_cell_text(cell: Any, namespaces: Dict) -> str:
        """
        Fetches the text from a table cell, handling page breaks and line breaks.

        This method properly handles content that spans page breaks by:
        - Detecting w:lastRenderedPageBreak elements and continuing to read
        - Converting w:br elements to appropriate separators
        - Preserving paragraph structure

        :param Any cell: The cell element to fetch text from.
        :param Dict namespaces: The XML namespaces.
        :return: The text from the cell.
        :rtype: str
        """
        cell_parts = []
        page_break_count = 0

        for para in cell.findall(".//w:p", namespaces):
            if SSPDocParser.has_page_break(para, namespaces):
                page_break_count += 1
                logger.debug("Page break detected in cell, continuing to read content")

            para_text = SSPDocParser._extract_paragraph_text(para, namespaces)
            if para_text.strip():
                cell_parts.append(para_text.strip())

        if page_break_count > 0:
            logger.debug("Cell contained %d page break(s), all content extracted", page_break_count)

        return " ".join(cell_parts) + " " if cell_parts else ""

    def extract_vertical_row_data(self, table: any, namespaces: any) -> List[Dict]:
        """
        Extracts data from a table organized vertically, with the first cell as the key
        and the second cell as the value.

        :param any table: The table element to extract data from.
        :param any namespaces: The XML namespaces.
        :return: A list of dictionaries representing the table's data.
        :rtype: List[Dict]
        """
        dicts_list = []
        for row in table.findall(".//w:tr", namespaces)[1:]:
            cells = row.findall(".//w:tc", namespaces)
            if len(cells) >= 2:
                vertical_data = {
                    self.fetch_cell_text(cells[0], namespaces)
                    .strip(): self.fetch_cell_text(cells[1], namespaces)
                    .strip()
                }
                dicts_list.append(vertical_data)
        return dicts_list

    def extract_row_data(self, row: any, headers: any, namespaces: any) -> Dict:
        """
        Extracts data from a table row.

        :param any row: The row element to extract data from.
        :param any headers: The headers of the table.
        :param any namespaces: The XML namespaces.
        :return: A dictionary representing the row's data.
        :rtype: Dict
        """
        row_data = {}
        for header, cell in zip(headers, row.findall(".//w:tc", namespaces)):
            cell_text = self.fetch_cell_text(cell, namespaces)
            if "☒" in cell_text or "☐" in cell_text:
                row_data[header] = self.parse_checkbox_string(cell_text)
            else:
                row_data[header] = cell_text.strip() if cell_text else None
        return row_data

    def fetch_preceding_text(self) -> List[str]:
        """
        Fetches the text immediately preceding a table.
        """
        preceding_texts = []
        for element in self.doc.element.body:
            if element.tag.endswith("tbl"):
                para = element.getprevious()
                text = para.text if para is not None and para.tag.endswith("p") else ""
                preceding_texts.append(text)
        return preceding_texts

    def parse_tables(self, root: any, namespaces: any) -> List[Dict]:
        """
        Parses all tables in the XML root.

        :param any root: The XML root element.
        :param any namespaces: The XML namespaces.
        :return: A list of dictionaries, each representing a table's data.
        :rtype: List[Dict]
        """
        vertical_tables = [
            "Identification of Organization that Prepared this Document".lower(),
            "Identification of Cloud Service Provider".lower(),
            "System Owner Information".lower(),
            "System Information".lower(),
            "System Component Information".lower(),
            "ISSO (or Equivalent) Point of Contact".lower(),
        ]
        tables_list = []
        preceding_text = self.fetch_preceding_text()
        for i, table in enumerate(root.findall(".//w:tbl", namespaces)):
            tables_dicts_list = {}
            headers = self.extract_table_headers(table, namespaces)
            table_data = []
            tables_dicts_list["preceding_text"] = preceding_text[i] if i < len(preceding_text) else ""
            # Check if this is a vertical table:
            is_vertical = False
            # - Single header that matches vertical_tables list
            first_option = len(headers) == 1 and headers[0].lower() in vertical_tables
            # - OR two headers where first matches and second is empty/whitespace
            second_option = len(headers) == 2 and headers[0].lower() in vertical_tables and not headers[1].strip()
            if first_option or second_option:
                is_vertical = True

            if is_vertical:
                table_data = self.extract_vertical_row_data(table, namespaces)
                tables_dicts_list["table_data"] = table_data
                # tables_dicts_list.append({headers[0].lower(): table_data})
            else:
                for row in table.findall(".//w:tr", namespaces)[1:]:  # Skip header row
                    row_data = self.extract_row_data(row, headers, namespaces)
                    table_data.append(row_data)
                tables_dicts_list["table_data"] = table_data
            tables_list.append(tables_dicts_list)
        return tables_list
