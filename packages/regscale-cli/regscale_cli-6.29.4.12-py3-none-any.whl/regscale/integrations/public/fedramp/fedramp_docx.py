"""
FedRAMP Docx parsing and import support
"""

from typing import Union, Any, Optional, Dict, Tuple, List

import click
from docx import Document
from docx.document import Document as DocxDocument
from pathlib import Path
from ssp import SSP  # type: ignore

from regscale.core.app.api import Api
from regscale.core.app.application import Application
from regscale.core.app.utils.app_utils import error_and_exit
from regscale.integrations.public.fedramp.fedramp_common import (
    get_profile_info_by_id,
    logger,
    replace_content_control,
    get_text_between_headers,
    get_tables,
    get_xpath_data_detailed,
    TABLE_TAG,
    parse_ssp_docx_tables,
    SYSTEM_TYPE,
    get_xpath_privacy_detailed,
    get_xpath_sysinfo_detailed,
    get_xpath_prepdata_detailed,
    ORGANIZATION_TAG,
    find_profile_by_name,
    revision,
    create_initial_ssp,
    create_privacy_data,
    create_responsible_roles,
    post_responsible_roles,
    gather_stakeholders,
    post_interconnects,
    tables_to_dict,
    post_ports,
    post_links,
    post_implementations,
    post_leveraged_authorizations,
)
from regscale.models import ProfileMapping


def process_fedramp_docx_by_profile_id(
    file_path: Union[click.Path, str],
    profile_id: int,
    save_data: bool = False,
    load_missing: bool = False,
) -> Any:
    """
    Process a FedRAMP docx by the profile_id from the Regscale Api

    :param Union[click.Path, str] file_path: The file path to the FedRAMP docx
    :param int profile_id: The profile_id to process
    :param bool save_data: Whether to save the data
    :param bool load_missing: Whether to load missing controls
    :return: RegScale SSP
    :rtype: Any
    """
    profile = get_profile_info_by_id(profile_id)
    new_implementations, regscale_ssp = process_fedramp_docx(
        fedramp_file_path=file_path,
        base_fedramp_profile=profile["name"],
        base_fedramp_profile_id=profile["id"],
        save_data=save_data,
        add_missing=load_missing,
        profile=profile,
    )
    # implementation_results
    logger.write_events()
    return (
        "artifacts/import-results.csv",
        {
            "ssp_title": regscale_ssp.get("systemName", "New SSP Default Name"),  # type: ignore
            "ssp_id": regscale_ssp.get("id"),  # type: ignore
        },
        new_implementations,
    )


# flake8: noqa: C901
def process_fedramp_docx(
    fedramp_file_path: Union[click.Path, str],
    base_fedramp_profile: str,
    base_fedramp_profile_id: Optional[int],
    save_data: bool = False,
    add_missing: bool = False,
    profile: Optional[Dict[Any, Any]] = None,
) -> Tuple[List, DocxDocument]:
    """
    Convert a FedRAMP file to a RegScale SSP

    :param Union[click.Path, str] fedramp_file_path: The click file path object
    :param str base_fedramp_profile: base fedramp profile
    :param Optional[int] base_fedramp_profile_id: base fedramp profile id
    :param bool save_data: Whether to save the data
    :param bool add_missing: Whether to add missing controls
    :param Optional[Dict[Any, Any]] profile: The profile to use
    :return: Tuple of new implementations count and RegScale SSP
    :rtype: Tuple[List, SSP]
    """
    # If list of controls is more than profile mapping, make
    # sure i get them from somewhere? Get base catalog from profile.
    load_missing = add_missing
    app = Application()
    api = Api()
    ssp = SSP(fedramp_file_path)
    document = Document(str(fedramp_file_path))
    for p in document.paragraphs:
        replace_content_control(p._element)  # noqa

    full_text = [para.text for para in document.paragraphs]
    system_text_lookup = "System Function or Purpose"
    description_lookup_str = "General System Description"

    description = get_text_between_headers(
        full_text,
        start_header=description_lookup_str,
        end_header=system_text_lookup,
    )

    environment = get_text_between_headers(
        full_text,
        start_header="SYSTEM ENVIRONMENT AND INVENTORY",
        end_header="Data Flow",
    )

    purpose = get_text_between_headers(
        full_text,
        start_header=system_text_lookup,
        end_header="System Description:",
    )
    if not purpose or purpose == "":
        purpose = get_text_between_headers(
            full_text,
            start_header=system_text_lookup,
            end_header="Information System Components and Boundaries",
        )

    confidentiality = "Low"
    integrity = "Low"
    availability = "Low"
    tables = get_tables(document)
    security_objective = get_xpath_data_detailed(
        tables,
        key="Security Objective",
        ident="Confidentiality",
        xpath=TABLE_TAG,
        count_array=[2, 4, 6],
    )

    availability = (
        security_objective["availability"].split(" ")[0] if "availability" in security_objective else availability
    )
    confidentiality = (
        security_objective["confidentiality"].split(" ")[0]
        if "confidentiality" in security_objective
        else confidentiality
    )
    integrity = security_objective["integrity"].split(" ")[0] if "integrity" in security_objective else integrity

    (
        system_status,
        system_type,
        title,
        cloud_model,
        cloud_service,
        version,
        table_data,
    ) = parse_ssp_docx_tables(document.tables)

    mdeploypublic = True if "multiple organizations " in cloud_model else False
    mdeploypriv = True if "specific organization/agency" in cloud_model else False
    mdeploygov = True if "organizations/agencies" in cloud_model else False
    mdeployhybrid = True if "shared across all clients/agencies" in cloud_model else False

    msaas = True if SYSTEM_TYPE in cloud_service else False
    mpaas = True if SYSTEM_TYPE in cloud_service and not msaas else False
    miaas = True if "General Support System" in cloud_service else False
    mother = True if "Explain:" in cloud_service else False

    privacydata = get_xpath_privacy_detailed(
        tables,
        key="Does the ISA collect, maintain, or share PII in any identifiable form?",
        xpath=TABLE_TAG,
        count_array=[0, 2, 4, 6],
    )

    sysinfo = get_xpath_sysinfo_detailed(
        tables,
        key="Unique Identifier",
        xpath=TABLE_TAG,
        count_array=[3, 5],
    )
    if sysinfo["systemname"]:
        title = sysinfo["systemname"].strip() if "systemname" in sysinfo else title
    if sysinfo["uniqueidentifier"]:
        uniqueidentifier = sysinfo["uniqueidentifier"].strip() if "uniqueidentifier" in sysinfo else None
    else:
        uniqueidentifier = None

    prepdata = get_xpath_prepdata_detailed(
        tables,
        key="Identification of Organization that Prepared this Document",
        ident=ORGANIZATION_TAG,
        xpath=TABLE_TAG,
    )
    preporgname = prepdata["orgname"] if "orgname" in prepdata else None
    prepaddress = prepdata["street"] if "street" in prepdata else None
    prepoffice = prepdata["office"] if "office" in prepdata else None
    prepcitystate = prepdata["citystate"] if "citystate" in prepdata else None
    cspdata = get_xpath_prepdata_detailed(
        tables,
        key="Identification of Cloud Service Provider",
        ident=ORGANIZATION_TAG,
        xpath=TABLE_TAG,
    )
    csporgname = cspdata["orgname"] if "orgname" in cspdata else None
    cspaddress = cspdata["street"] if "street" in cspdata else None
    cspoffice = cspdata["office"] if "office" in cspdata else None
    cspcitystate = cspdata["citystate"] if "citystate" in cspdata else None
    status = "Operational" if "in production" in system_status else "Other"
    # Links are posted to links mapped to ssp
    # post_links(app, table_data, ssp_id)
    # Parts will go in implementation fields.
    if base_fedramp_profile_id:
        profile = get_profile_info_by_id(profile_id=base_fedramp_profile_id)
    if not profile:
        profile = find_profile_by_name(base_fedramp_profile) or {}
    if not (profile_id := int(profile.get("id", 0))):
        error_and_exit(f"Unable to continue, please load {base_fedramp_profile} with controls!")
    profile_mapping = ProfileMapping.get_by_profile(profile_id=profile_id)
    if len(profile_mapping) == 0:
        error_and_exit(f"Unable to continue, please load {base_fedramp_profile} with controls!")

    logger.info(
        f"Utilizing profile: {profile.get('name')}",
        record_type="profile",
        model_layer="profile",
    )
    args = {
        "profile": profile,
        "title": title,
        "otheridentifier": uniqueidentifier,
        # get_profile_mapping(profile["id"]))
        "version": version,
        "confidentiality": confidentiality,
        "integrity": integrity,
        "availability": availability,
        "status": status,
        "system_type": system_type,
        # "ssp": ssp,
        "revision": revision(document),
        "description": description,
        "environment": environment,
        "purpose": purpose,
        "modeiaas": miaas,
        "modepaas": mpaas,
        "modeother": mother,
        "modesaas": msaas,
        "deploypubic": mdeploypublic,
        "deployprivate": mdeploypriv,
        "deploygov": mdeploygov,
        "deployhybrid": mdeployhybrid,
        "preporgname": preporgname,
        "prepaddress": prepaddress,
        "prepoffice": prepoffice,
        "prepcitystate": prepcitystate,
        "csporgname": csporgname,
        "cspaddress": cspaddress,
        "cspoffice": cspoffice,
        "cspcitystate": cspcitystate,
    }
    regscale_ssp = create_initial_ssp(args)

    try:
        logger.info("Parsing and creating Privacy Data", record_type="privacy", model_layer="privacy")
        create_privacy_data(app=app, privacy_data=privacydata, ssp_id=regscale_ssp.get("id"))
        logger.info(
            "Successfully Created Privacy data.",
            record_type="privacy",
            model_layer="privacy",
        )
    except Exception as e:
        logger.error(
            f"Unable to create privacy record: {e}",
            record_type="privacy",
            model_layer="privacy",
        )

    try:
        logger.info(
            "Parsing and creating System Information", record_type="responsible-roles", model_layer="responsible-roles"
        )
        create_responsible_roles(app, table_data, ssp_id=regscale_ssp["id"])
        ctrl_roles = post_responsible_roles(app, table_data, ssp_id=regscale_ssp["id"])
    except Exception as e:
        logger.error(
            f"Unable to gather responsible roles: {e}",
            record_type="responsible-roles",
            model_layer="responsible-roles",
        )

    try:
        logger.info("Parsing and creating Stakeholders", record_type="stakeholder", model_layer="stakeholder")
        gather_stakeholders(tables, regscale_ssp, document)
    except Exception as e:
        logger.error(
            f"Unable to gather stakeholders: {e}",
            record_type="stakeholder",
            model_layer="stakeholder",
        )
    try:
        logger.info("Parsing and creating Interconnects", record_type="interconnect", model_layer="interconnects")
        post_interconnects(app, table_data, regscale_ssp)
    except Exception as e:
        logger.error(
            f"Unable to gather interconnects: {e}",
            record_type="interconnect",
            model_layer="interconnects",
        )
    try:
        logger.info(
            "Parsing and creating Ports and Protocols", record_type="ports-protocols", model_layer="ports-protocols"
        )
        tables_dict = tables_to_dict(document)
        ports_table_data = [row for t in tables_dict for row in t if "Ports (TCP/UDP)*" in row]
        post_ports(app, ports_table_data, ssp_id=regscale_ssp["id"])
    except Exception as e:
        logger.error(
            f"Unable to gather ports: {e}",
            record_type="ports-protocols",
            model_layer="ports-protocols",
        )
    try:
        logger.info("Parsing and creating Links", record_type="links", model_layer="links")
        post_links(
            config=app.config,
            api=api,
            document=document,
            file_path=Path(str(fedramp_file_path)),
            regscale_ssp=regscale_ssp,
        )
    except Exception as e:
        logger.error(
            f"Unable to gather links: {e}",
            record_type="links",
            model_layer="links",
        )
    try:
        logger.info(
            "Parsing and creating implementations", record_type="implementations", model_layer="implementations"
        )
        new_implementations = post_implementations(
            app=app,
            ssp_obj=ssp,
            regscale_ssp=regscale_ssp,
            mapping=profile_mapping,
            ctrl_roles=ctrl_roles,
            save_data=save_data,
            load_missing=load_missing,
        )
    except Exception as e:
        logger.info(e)
        logger.error(
            f"Unable to gather implementations: {e}",
            record_type="implementations",
            model_layer="implementations",
        )
        new_implementations = []
    try:
        logger.info(
            "Parsing and creating Leveraged Authorizations",
            record_type="leveraged-authorizations",
            model_layer="leveraged-authorizations",
        )
        post_leveraged_authorizations(table_data, ssp_id=regscale_ssp.get("id"))
    except Exception as e:
        logger.error(
            f"Unable to gather leveraged authorizations: {e}",
            record_type="leveraged-authorizations",
            model_layer="leveraged-authorizations",
        )
    return new_implementations, regscale_ssp
