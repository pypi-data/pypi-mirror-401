from copy import deepcopy
from dataclasses import dataclass
from typing import List, Optional, Tuple

import structlog
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from adeu.utils.docx import get_paragraph_prefix, get_run_text, get_visible_runs, iter_document_parts

logger = structlog.get_logger(__name__)


@dataclass
class TextSpan:
    start: int
    end: int
    text: str
    run: Optional[Run]
    paragraph: Paragraph


class DocumentMapper:
    def __init__(self, doc: DocumentObject):
        self.doc = doc
        self.full_text = ""
        self.spans: List[TextSpan] = []
        self._build_map()

    def _build_map(self):
        current_offset = 0
        self.spans = []
        self.full_text = ""

        # We must mirror ingest.py logic exactly:
        # 1. Iterate Parts (Header -> Body -> Footer)
        # 2. In each part: Paragraphs first, then Tables.
        # 3. Tables are formatted: Cells joined by " | ", Rows joined by "\n\n".
        #    Cell paragraphs joined by "\n".

        for part in iter_document_parts(self.doc):
            # 1. Loose Paragraphs
            for p in part.paragraphs:
                # Add Header Prefix (Virtual)
                prefix = get_paragraph_prefix(p)
                if prefix:
                    self._add_virtual_text(prefix, current_offset, p)
                    current_offset += len(prefix)

                current_offset = self._map_paragraph_runs(p, current_offset)

                # Add Paragraph Separator (\n\n)
                self._add_virtual_text("\n\n", current_offset, p)
                current_offset += 2

            # 2. Tables
            for table in part.tables:
                for row in table.rows:
                    row_has_content = False
                    for cell in row.cells:
                        # Map cell paragraphs
                        cell_paras = list(cell.paragraphs)
                        cell_non_empty = any(get_visible_runs(p) for p in cell_paras)

                        if cell_non_empty:
                            if row_has_content:
                                # Add separator from previous cell
                                self._add_virtual_text(" | ", current_offset, cell_paras[0])
                                current_offset += 3

                            row_has_content = True

                            for j, p in enumerate(cell_paras):
                                if j > 0:
                                    # Join paragraphs within cell with \n
                                    self._add_virtual_text("\n", current_offset, p)
                                    current_offset += 1

                                # Add Header Prefix (Virtual) inside table?
                                # ingest.py does it, so we must too.
                                prefix = get_paragraph_prefix(p)
                                if prefix:
                                    self._add_virtual_text(prefix, current_offset, p)
                                    current_offset += len(prefix)

                                current_offset = self._map_paragraph_runs(p, current_offset)

                    if row_has_content:
                        # End of Row
                        self._add_virtual_text("\n\n", current_offset, row.cells[0].paragraphs[0])
                        current_offset += 2

        # Remove trailing newline to match ingest.py
        if self.spans and self.spans[-1].text == "\n\n":
            self.spans.pop()
            self.full_text = self.full_text[:-2]

    def _map_paragraph_runs(self, paragraph: Paragraph, start_offset: int) -> int:
        current = start_offset
        runs = get_visible_runs(paragraph)
        for run in runs:
            text = get_run_text(run)
            text_len = len(text)
            if text_len == 0:
                continue

            span = TextSpan(start=current, end=current + text_len, text=text, run=run, paragraph=paragraph)
            self.spans.append(span)
            self.full_text += text
            current += text_len
        return current

    def _add_virtual_text(self, text: str, offset: int, context_paragraph: Paragraph):
        span = TextSpan(
            start=offset,
            end=offset + len(text),
            text=text,
            run=None,  # Virtual
            paragraph=context_paragraph,
        )
        self.spans.append(span)
        self.full_text += text

    def _replace_smart_quotes(self, text: str) -> str:
        return text.replace("", '"').replace("", '"').replace("‘", "'").replace("’", "'")

    def find_match_index(self, target_text: str) -> int:
        start_idx = self.full_text.find(target_text)
        if start_idx == -1:
            norm_full = self._replace_smart_quotes(self.full_text)
            norm_target = self._replace_smart_quotes(target_text)
            start_idx = norm_full.find(norm_target)
        return start_idx

    def find_target_runs(self, target_text: str) -> List[Run]:
        start_idx = self.find_match_index(target_text)
        if start_idx == -1:
            return []
        return self._resolve_runs_at_range(start_idx, start_idx + len(target_text))

    def find_target_runs_by_index(self, start_index: int, length: int) -> List[Run]:
        end_index = start_index + length
        return self._resolve_runs_at_range(start_index, end_index)

    def _resolve_runs_at_range(self, start_idx: int, end_idx: int) -> List[Run]:
        affected_spans = [s for s in self.spans if s.end > start_idx and s.start < end_idx]
        if not affected_spans:
            return []

        working_runs = [s.run for s in affected_spans if s.run is not None]
        if not working_runs:
            return []

        dom_modified = False

        first_span = affected_spans[0]
        local_start = start_idx - first_span.start
        if local_start > 0 and first_span.run is not None:
            _, right_run = self._split_run_at_index(working_runs[0], local_start)
            working_runs[0] = right_run
            dom_modified = True

        last_span = affected_spans[-1]
        extra_len = last_span.end - end_idx
        if extra_len > 0 and last_span.run is not None:
            last_run = working_runs[-1]
            split_point = len(last_run.text) - extra_len
            left_run, _ = self._split_run_at_index(last_run, split_point)
            working_runs[-1] = left_run
            dom_modified = True

        if dom_modified:
            self._build_map()

        return working_runs

    def get_insertion_anchor(self, index: int) -> Optional[Run]:
        preceding = [s for s in self.spans if s.end == index]
        if preceding:
            if preceding[-1].run:
                return preceding[-1].run
        containing = [s for s in self.spans if s.start < index < s.end]
        if containing:
            span = containing[0]
            if span.run is None:
                pass
            else:
                offset = index - span.start
                left, _ = self._split_run_at_index(span.run, offset)
                return left

        if index == 0 and self.spans:
            for s in self.spans:
                if s.run:
                    return s.run
            return None

        preceding_gap = [s for s in self.spans if s.end < index]
        if preceding_gap:
            for s in reversed(preceding_gap):
                if s.run:
                    return s.run
        return None

    def _split_run_at_index(self, run: Run, split_index: int) -> Tuple[Run, Run]:
        text = run.text
        left_text = text[:split_index]
        right_text = text[split_index:]

        run.text = left_text
        new_r_element = deepcopy(run._element)
        t_list = new_r_element.findall(qn("w:t"))
        for t in t_list:
            new_r_element.remove(t)

        new_t = OxmlElement("w:t")
        new_t.text = right_text
        if right_text.strip() != right_text:
            new_t.set(qn("xml:space"), "preserve")
        new_r_element.append(new_t)
        run._element.addnext(new_r_element)
        new_run = Run(new_r_element, run._parent)
        return run, new_run
