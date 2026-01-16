import io

import structlog
from docx import Document

from adeu.utils.docx import get_paragraph_prefix, get_run_text, get_visible_runs, iter_document_parts

logger = structlog.get_logger(__name__)


def extract_text_from_stream(file_stream: io.BytesIO, filename: str = "document.docx") -> str:
    """
    Extracts text from a file stream using raw run concatenation.

    CRITICAL: This must match DocumentMapper._build_map logic exactly.
    We iterate runs and join them. We do not use para.text.
    """
    try:
        # Ensure stream is at start
        file_stream.seek(0)

        doc = Document(file_stream)
        full_text = []

        for part in iter_document_parts(doc):
            # 1. Paragraphs
            for para in part.paragraphs:
                # Use the visible runs helper to see <w:ins> content
                runs = get_visible_runs(para)
                # Use get_run_text to include tabs and breaks
                p_text = "".join([get_run_text(r) for r in runs])

                # Add Markdown prefix if heading
                prefix = get_paragraph_prefix(para)
                full_text.append(prefix + p_text)

            # 2. Tables
            for table in part.tables:
                for row in table.rows:
                    row_parts = []
                    for cell in row.cells:
                        # Cell paragraphs
                        cell_text_parts = []
                        for p in cell.paragraphs:
                            # Note: We probably don't want headers inside tables usually,
                            # but for consistency we should allow it if styled.
                            prefix = get_paragraph_prefix(p)
                            p_content = "".join([get_run_text(r) for r in get_visible_runs(p)])
                            cell_text_parts.append(prefix + p_content)

                        cell_text = "\n".join(cell_text_parts)
                        if cell_text:
                            row_parts.append(cell_text)

                    if row_parts:
                        full_text.append(" | ".join(row_parts))

        return "\n\n".join(full_text)

    except Exception as e:
        logger.error(f"Text extraction failed: {e}", exc_info=True)
        raise ValueError(f"Could not extract text: {str(e)}") from e
