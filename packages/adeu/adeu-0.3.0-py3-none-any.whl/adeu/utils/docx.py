"""
Low-level utilities for manipulating DOCX XML structures.
Contains normalization logic ported from Open-Xml-PowerTools concepts.
"""

import structlog
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = structlog.get_logger(__name__)


def create_element(name: str):
    return OxmlElement(name)


def create_attribute(element, name: str, value: str):
    element.set(qn(name), value)


def _is_page_instr(instr: str) -> bool:
    if not instr:
        return False
    instr = instr.upper().strip()
    # Check for PAGE or NUMPAGES keyword at start of instruction
    parts = instr.split()
    if not parts:
        return False
    return parts[0] in ("PAGE", "NUMPAGES")


def get_paragraph_prefix(paragraph: Paragraph) -> str:
    """
    Returns the Markdown prefix for a paragraph based on its style.
    e.g. 'Heading 1' -> '# ', 'Heading 2' -> '## '
    """
    # 1. Check Outline Level (Structural Truth)
    # python-docx outline_level: 0=Level 1, ..., 8=Level 9, 9=Body Text
    try:
        lvl = paragraph.paragraph_format.outline_level
        if lvl is not None and 0 <= lvl <= 8:
            return "#" * (lvl + 1) + " "
    except Exception:
        pass

    if not paragraph.style:
        return ""

    style_name = paragraph.style.name
    if not style_name:
        return ""

    # 2. Check Style Name
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading", "").strip())
            return "#" * level + " "
        except ValueError:
            pass

    if style_name == "Title":
        return "# "

    # 3. Heuristic for "Normal" style headers (Lazy Lawyer / Manually formatted)
    # If text is short (<100 chars), All Caps, and Bold -> Likely a Header
    if style_name == "Normal":
        text = paragraph.text.strip()
        if text and len(text) < 100:
            is_all_caps = text.isupper()

            # Check for Bold (Paragraph style or explicit run formatting)
            is_bold = False
            if paragraph.style.font.bold:
                is_bold = True
            else:
                # Check if visible runs are bold
                # This is a loose check; if the first run is bold, we assume intention
                runs = [r for r in paragraph.runs if r.text.strip()]
                if runs and runs[0].bold:
                    is_bold = True

            if is_all_caps and is_bold:
                return "## "

    return ""


def get_visible_runs(paragraph: Paragraph):
    """
    Iterates over runs in a paragraph, including those inside <w:ins> tags.
    Effectively returns the 'Accepted Changes' view of the runs.
    Filters out dynamic page number fields ({PAGE}, {NUMPAGES}).
    """
    runs = []

    # State for complex fields (w:fldChar)
    in_complex_field = False
    current_instr = ""
    hide_result = False

    def process_run_element(r_element):
        nonlocal in_complex_field, current_instr, hide_result

        # 1. Parse Field Characters (begin/separate/end)
        for fchar in r_element.findall(qn("w:fldChar")):
            fld_type = fchar.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_complex_field = True
                current_instr = ""
            elif fld_type == "separate":
                # End of instruction, start of visible result
                if _is_page_instr(current_instr):
                    hide_result = True
            elif fld_type == "end":
                in_complex_field = False
                current_instr = ""
                hide_result = False

        # 2. Accumulate Instruction Text
        if in_complex_field and not hide_result:
            for instr in r_element.findall(qn("w:instrText")):
                if instr.text:
                    current_instr += instr.text

        # 3. Yield Run (if not hidden)
        if not hide_result:
            runs.append(Run(r_element, paragraph))

    # Iterate over all children of the paragraph XML element
    for child in paragraph._element:
        tag = child.tag
        if tag == qn("w:r"):
            # Standard run
            process_run_element(child)
        elif tag == qn("w:ins"):
            # Inserted runs (Track Changes)
            for subchild in child:
                if subchild.tag == qn("w:r"):
                    process_run_element(subchild)
        # w:del is skipped implies we read the "Future" state (Deletions are gone)
        # w:fldSimple is typically skipped here. If supported in future, add logic to check w:instr attribute.

    return runs


def get_run_text(run: Run) -> str:
    """
    Extracts text from a run, converting <w:tab/> to spaces and <w:br/> to newlines.
    Standard run.text ignores these.
    """
    text = ""
    for child in run._element:
        if child.tag == qn("w:t"):
            text += child.text or ""
        elif child.tag == qn("w:tab"):
            text += " "  # Convert tab to space
        elif child.tag == qn("w:br"):
            text += "\n"
        elif child.tag == qn("w:cr"):
            text += "\n"
    return text


def _are_runs_identical(r1: Run, r2: Run) -> bool:
    """
    Compares two runs to see if they have identical formatting properties.
    """
    rPr1 = r1._r.rPr
    rPr2 = r2._r.rPr

    xml1 = rPr1.xml if rPr1 is not None else ""
    xml2 = rPr2.xml if rPr2 is not None else ""

    return xml1 == xml2


def _coalesce_runs_in_paragraph(paragraph: Paragraph):
    """
    Merges adjacent runs with identical formatting.
    This fixes issues where words are split like ["Con", "tract"] due to editing history.
    """
    i = 0
    # Safe iteration while modifying the list
    while i < len(paragraph.runs) - 1:
        current_run = paragraph.runs[i]
        next_run = paragraph.runs[i + 1]

        if _are_runs_identical(current_run, next_run):
            # Merge content
            current_run.text += next_run.text
            # Remove next_run from the XML tree
            paragraph._p.remove(next_run._r)
            # Do NOT increment i; check the *new* next_run against current_run
        else:
            i += 1


def iter_document_parts(doc: DocumentObject):
    """
    Yields document parts in a linear order for processing:
    1. Unique Headers (Primary, First, Even)
    2. Main Body
    3. Unique Footers (Primary, First, Even)

    Handles 'Link to Previous' to avoid duplication.
    """

    def _iter_section_parts(section, part_type_attr):
        # 1. Primary
        part = getattr(section, part_type_attr)
        if not part.is_linked_to_previous:
            yield part

        # 2. First Page
        if section.different_first_page_header_footer:
            first = getattr(section, f"first_page_{part_type_attr}")
            if not first.is_linked_to_previous:
                yield first

        # 3. Even Page
        if doc.settings.odd_and_even_pages_header_footer:
            even = getattr(section, f"even_page_{part_type_attr}")
            if not even.is_linked_to_previous:
                yield even

    # 1. Headers
    for section in doc.sections:
        yield from _iter_section_parts(section, "header")

    # 2. Main Body (The Document object itself acts as the container)
    yield doc

    # 3. Footers
    for section in doc.sections:
        yield from _iter_section_parts(section, "footer")


def normalize_docx(doc: DocumentObject):
    """
    Applies normalization to a DOCX document to make text mapping reliable.
    1. Removes proof errors (spellcheck squiggles).
    2. Coalesces adjacent runs.
    """
    logger.info("Normalizing DOCX structure...")

    # Remove proof errors (spelling/grammar tags) via XPath
    for proof_err in doc.element.xpath("//w:proofErr"):
        proof_err.getparent().remove(proof_err)

    # Coalesce all parts (Headers, Body, Footers)
    for part in iter_document_parts(doc):
        for p in part.paragraphs:
            _coalesce_runs_in_paragraph(p)

        for table in part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _coalesce_runs_in_paragraph(p)
