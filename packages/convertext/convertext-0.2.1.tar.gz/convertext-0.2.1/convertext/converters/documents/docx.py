"""DOCX format converter."""

from pathlib import Path
from typing import Any, Dict, List

import docx

from convertext.converters.base import BaseConverter, Document


class DocxConverter(BaseConverter):
    """DOCX/DOC format converter."""

    @property
    def input_formats(self) -> List[str]:
        return ['docx', 'doc']

    @property
    def output_formats(self) -> List[str]:
        return ['txt', 'html', 'md']

    def can_convert(self, source: str, target: str) -> bool:
        return source in self.input_formats and target in self.output_formats

    def convert(self, source_path: Path, target_path: Path, config: Dict[str, Any]) -> bool:
        """Convert DOCX to target format."""
        doc = self._read_docx(source_path, config)

        target_fmt = target_path.suffix.lstrip('.').lower()
        if target_fmt == 'txt':
            return self._write_txt(doc, target_path)
        elif target_fmt == 'html':
            return self._write_html(doc, target_path)
        elif target_fmt == 'md':
            return self._write_md(doc, target_path)

        return False

    def _read_docx(self, path: Path, config: Dict[str, Any]) -> Document:
        """Read DOCX into intermediate Document with full formatting."""
        doc = Document()
        docx_doc = docx.Document(path)

        core_props = docx_doc.core_properties
        doc.metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
        }

        for para in docx_doc.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 1
                if para.text.strip():
                    doc.add_heading(para.text.strip(), level)

            elif para.style.name.startswith('List'):
                continue

            else:
                if para.text.strip():
                    if len(para.runs) == 1 and not any([
                        para.runs[0].bold,
                        para.runs[0].italic,
                        para.runs[0].underline
                    ]):
                        doc.add_paragraph(para.text.strip())
                    else:
                        for run in para.runs:
                            if run.text.strip():
                                color = None
                                if run.font.color and run.font.color.rgb:
                                    rgb = run.font.color.rgb
                                    color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"

                                doc.add_run(
                                    text=run.text,
                                    bold=run.bold or False,
                                    italic=run.italic or False,
                                    underline=run.underline or False,
                                    color=color,
                                    font_name=run.font.name,
                                    font_size=int(run.font.size.pt) if run.font.size else None
                                )

        for table in docx_doc.tables:
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)

            if rows:
                doc.add_table(rows=rows[1:], headers=rows[0] if len(rows) > 1 else None)

        return doc

    def _write_txt(self, doc: Document, path: Path) -> bool:
        """Write Document to plain text."""
        with open(path, 'w', encoding='utf-8') as f:
            if doc.metadata.get('title'):
                f.write(doc.metadata['title'] + '\n')
                f.write('=' * len(doc.metadata['title']) + '\n\n')

            for block in doc.content:
                if block['type'] in ['text', 'paragraph']:
                    f.write(block['data'] + '\n\n')
                elif block['type'] == 'heading':
                    f.write('\n' + block['data'].upper() + '\n')
                    f.write('-' * len(block['data']) + '\n\n')
        return True

    def _write_html(self, doc: Document, path: Path) -> bool:
        """Write Document to HTML."""
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '<meta charset="utf-8">',
        ]

        if doc.metadata.get('title'):
            html_parts.append(f"<title>{self._escape_html(doc.metadata['title'])}</title>")
        else:
            html_parts.append('<title>Document</title>')

        html_parts.append('</head>')
        html_parts.append('<body>')

        for block in doc.content:
            if block['type'] == 'paragraph':
                html_parts.append(f"<p>{self._escape_html(block['data'])}</p>")
            elif block['type'] == 'heading':
                level = block['level']
                html_parts.append(f"<h{level}>{self._escape_html(block['data'])}</h{level}>")

        html_parts.append('</body>')
        html_parts.append('</html>')

        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))

        return True

    def _write_md(self, doc: Document, path: Path) -> bool:
        """Write Document to Markdown."""
        with open(path, 'w', encoding='utf-8') as f:
            if doc.metadata.get('title'):
                f.write(f"# {doc.metadata['title']}\n\n")
            if doc.metadata.get('author'):
                f.write(f"**Author:** {doc.metadata['author']}\n\n")

            for block in doc.content:
                if block['type'] == 'paragraph':
                    f.write(block['data'] + '\n\n')
                elif block['type'] == 'heading':
                    f.write('#' * block['level'] + ' ' + block['data'] + '\n\n')
        return True

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))
