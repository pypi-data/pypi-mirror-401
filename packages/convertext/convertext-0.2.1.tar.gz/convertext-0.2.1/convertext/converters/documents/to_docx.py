"""DOCX format output converter."""

from pathlib import Path
from typing import Any, Dict, List
from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from convertext.converters.base import BaseConverter, Document
from convertext.converters.utils import hex_to_rgb


class ToDocxConverter(BaseConverter):
    """Convert various formats to DOCX."""

    @property
    def input_formats(self) -> List[str]:
        return ['txt', 'html', 'md', 'pdf', 'odt', 'epub', 'mobi', 'fb2', 'rtf']

    @property
    def output_formats(self) -> List[str]:
        return ['docx']

    def can_convert(self, source: str, target: str) -> bool:
        return source in self.input_formats and target == 'docx'

    def convert(self, source_path: Path, target_path: Path, config: Dict[str, Any]) -> bool:
        """Convert to DOCX."""
        from convertext.core import ConversionEngine
        from convertext.registry import get_registry

        registry = get_registry()
        source_fmt = source_path.suffix.lstrip('.').lower()

        converter = registry.get_converter(source_fmt, 'html')
        if not converter:
            converter = registry.get_converter(source_fmt, 'txt')
        if not converter:
            return False

        from tempfile import NamedTemporaryFile
        with NamedTemporaryFile(mode='w', suffix='.html', delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            if not converter.convert(source_path, tmp_path, config):
                return False

            if tmp_path.suffix == '.html':
                doc = self._read_html(tmp_path, config)
            else:
                doc = self._read_txt(tmp_path, config)

            return self._create_docx(doc, target_path, config, source_path.stem)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()

    def _read_txt(self, path: Path, config: Dict[str, Any]) -> Document:
        """Read plain text into Document."""
        doc = Document()
        encoding = config.get('documents', {}).get('encoding', 'utf-8')

        with open(path, 'r', encoding=encoding) as f:
            content = f.read()
            for para in content.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())

        return doc

    def _read_html(self, path: Path, config: Dict[str, Any]) -> Document:
        """Read HTML into Document."""
        doc = Document()
        encoding = config.get('documents', {}).get('encoding', 'utf-8')

        with open(path, 'r', encoding=encoding) as f:
            content = f.read()

        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')

        title_tag = soup.find('title')
        if title_tag:
            doc.metadata['title'] = title_tag.get_text()

        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'ul', 'ol']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                doc.add_heading(element.get_text().strip(), level)

            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)

            elif element.name == 'table':
                rows = []
                headers = []
                thead = element.find('thead')
                if thead:
                    header_row = thead.find('tr')
                    if header_row:
                        headers = [th.get_text().strip() for th in header_row.find_all(['th', 'td'])]

                tbody = element.find('tbody') or element
                for row in tbody.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if cells:
                        rows.append(cells)

                if rows:
                    doc.add_table(rows=rows, headers=headers if headers else None)

            elif element.name in ['ul', 'ol']:
                items = [li.get_text().strip() for li in element.find_all('li')]
                if items:
                    doc.add_list(items=items, ordered=(element.name == 'ol'))

        return doc

    def _create_docx(self, doc: Document, path: Path, config: Dict[str, Any], default_title: str) -> bool:
        """Create DOCX file using python-docx."""
        docx_doc = DocxDocument()

        docx_doc.core_properties.title = doc.metadata.get('title', default_title)
        docx_doc.core_properties.author = doc.metadata.get('author', '')
        docx_doc.core_properties.subject = doc.metadata.get('subject', '')

        for block in doc.content:
            if block['type'] == 'heading':
                level = min(block['level'], 9)
                docx_doc.add_heading(block['data'], level=level)

            elif block['type'] in ['paragraph', 'text']:
                docx_doc.add_paragraph(block['data'])

            elif block['type'] == 'run':
                p = docx_doc.add_paragraph()
                run = p.add_run(block['text'])

                if block.get('bold'):
                    run.bold = True
                if block.get('italic'):
                    run.italic = True
                if block.get('underline'):
                    run.underline = True

                if block.get('color'):
                    rgb = hex_to_rgb(block['color'])
                    if rgb:
                        run.font.color.rgb = RGBColor(*rgb)

                if block.get('font'):
                    run.font.name = block['font']

                if block.get('size'):
                    run.font.size = Pt(block['size'])

            elif block['type'] == 'table':
                headers = block.get('headers', [])
                rows = block['rows']

                num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
                num_rows = len(rows) + (1 if headers else 0)

                if num_cols > 0 and num_rows > 0:
                    table = docx_doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'

                    row_idx = 0
                    if headers:
                        for col_idx, header in enumerate(headers):
                            table.rows[row_idx].cells[col_idx].text = str(header)
                        row_idx += 1

                    for row_data in rows:
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < num_cols:
                                table.rows[row_idx].cells[col_idx].text = str(cell_data)
                        row_idx += 1

            elif block['type'] == 'list':
                for item in block['items']:
                    style = 'List Number' if block.get('ordered') else 'List Bullet'
                    docx_doc.add_paragraph(item, style=style)

            elif block['type'] == 'image':
                if block['name'] in doc.images:
                    img_data = doc.images[block['name']]['data']
                    try:
                        docx_doc.add_picture(BytesIO(img_data), width=Inches(4))
                    except Exception:
                        pass

            elif block['type'] == 'link':
                p = docx_doc.add_paragraph()
                p.add_run(f"{block['text']} ")
                p.add_run(f"({block['url']})").font.color.rgb = RGBColor(0, 0, 255)

        docx_doc.save(str(path))
        return True
