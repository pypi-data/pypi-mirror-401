"""
Document Generator for RAG Poisoning

Generates actual binary documents (PDF, DOCX) with embedded malicious payloads
for penetration testing RAG systems.

Security Note: These tools are for authorized red team testing only.
"""

import io
from typing import List
from datetime import datetime

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.colors import white
    from reportlab.lib.units import inch

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from PyPDF2 import PdfReader, PdfWriter

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.utils.logging import get_logger

logger = get_logger(__name__)


class DocumentGeneratorError(Exception):
    """Raised when document generation fails."""


class PoisonedDocumentGenerator:
    """
    Generate poisoned documents for RAG system testing.

    Supports multiple injection methods:
    - PDF metadata injection
    - PDF hidden text (white on white)
    - DOCX custom properties
    - DOCX hidden bookmarks
    - DOCX comments (normally not visible)
    """

    def __init__(self):
        """Initialize document generator."""
        self.check_dependencies()

    def check_dependencies(self):
        """Check if required libraries are available."""
        if not REPORTLAB_AVAILABLE:
            logger.warning(
                "reportlab_not_available",
                message="Install reportlab for PDF generation: pip install reportlab",
            )
        if not PYPDF2_AVAILABLE:
            logger.warning(
                "pypdf2_not_available",
                message="Install PyPDF2 for PDF metadata: pip install PyPDF2",
            )
        if not DOCX_AVAILABLE:
            logger.warning(
                "docx_not_available",
                message="Install python-docx for DOCX generation: pip install python-docx",
            )

    def create_poisoned_pdf(
        self,
        visible_content: str,
        hidden_instruction: str,
        filename: str = "document.pdf",
        method: str = "metadata",
        title: str = "Company Document",
        author: str = "Corporate Team",
    ) -> bytes:
        """
        Create PDF with malicious payload.

        Args:
            visible_content: Normal content user sees
            hidden_instruction: Malicious payload to inject
            filename: Suggested filename
            method: Injection method (metadata, white_text, both)
            title: PDF title
            author: PDF author

        Returns:
            PDF file as bytes

        Raises:
            DocumentGeneratorError: If PDF generation fails
        """
        if not REPORTLAB_AVAILABLE:
            raise DocumentGeneratorError("reportlab not installed. Run: pip install reportlab")

        try:
            # Create PDF in memory
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter

            # Set basic metadata (visible)
            c.setTitle(title)
            c.setAuthor(author)
            c.setSubject("Internal Document")

            # Add visible content
            c.setFont("Helvetica", 12)

            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(inch, height - inch, title)

            # Body content
            c.setFont("Helvetica", 11)
            text_object = c.beginText(inch, height - 1.5 * inch)
            text_object.setFont("Helvetica", 11)
            text_object.setLeading(14)

            # Wrap visible content
            lines = self._wrap_text(visible_content, 80)
            for line in lines[:40]:  # Max 40 lines
                text_object.textLine(line)

            c.drawText(text_object)

            # Add hidden text if method includes white_text
            if method in ["white_text", "both"]:
                c.setFillColor(white)
                c.setFont("Helvetica", 1)  # Tiny font

                # Place at bottom of page (invisible)
                hidden_text = c.beginText(inch, 0.5 * inch)
                hidden_text.setFont("Helvetica", 1)
                hidden_text.setFillColor(white)
                hidden_text.textLine(hidden_instruction)
                c.drawText(hidden_text)

            # Finalize PDF
            c.showPage()
            c.save()

            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()

            # Add metadata injection if requested
            if method in ["metadata", "both"] and PYPDF2_AVAILABLE:
                pdf_bytes = self._inject_pdf_metadata(pdf_bytes, hidden_instruction)

            logger.info("pdf_generated", method=method, size_bytes=len(pdf_bytes))

            return pdf_bytes

        except Exception as e:
            logger.error("pdf_generation_failed", error=str(e))
            raise DocumentGeneratorError(f"Failed to generate PDF: {e}")

    def _inject_pdf_metadata(self, pdf_bytes: bytes, payload: str) -> bytes:
        """
        Inject malicious payload into PDF metadata.

        Args:
            pdf_bytes: Original PDF bytes
            payload: Malicious instruction to inject

        Returns:
            Modified PDF bytes with poisoned metadata
        """
        if not PYPDF2_AVAILABLE:
            logger.warning("pypdf2_unavailable_skipping_metadata")
            return pdf_bytes

        try:
            # Read PDF
            input_buffer = io.BytesIO(pdf_bytes)
            reader = PdfReader(input_buffer)
            writer = PdfWriter()

            # Copy pages
            for page in reader.pages:
                writer.add_page(page)

            # Inject malicious metadata
            # Subject field often included by RAG systems
            writer.add_metadata(
                {
                    "/Subject": f"SYSTEM INSTRUCTION: {payload}",
                    "/Keywords": payload,
                    "/Creator": f"[OVERRIDE] {payload}",
                }
            )

            # Write to new buffer
            output_buffer = io.BytesIO()
            writer.write(output_buffer)

            return output_buffer.getvalue()

        except Exception as e:
            logger.error("metadata_injection_failed", error=str(e))
            return pdf_bytes  # Return original on failure

    def create_chunk_boundary_pdf(
        self, payload: str, chunk_size: int = 1024, filename: str = "boundary_attack.pdf"
    ) -> bytes:
        """
        Create PDF designed to exploit chunk boundaries.

        Positions payload at estimated chunk boundary based on
        common chunking strategies (512, 1024, 2048 tokens).

        Args:
            payload: Malicious instruction
            chunk_size: Target chunk size in characters (approx)
            filename: Suggested filename

        Returns:
            PDF bytes
        """
        if not REPORTLAB_AVAILABLE:
            raise DocumentGeneratorError("reportlab not installed")

        # Create filler content to reach chunk boundary
        filler = self._generate_filler_content(chunk_size - len(payload) // 2)

        # Construct document with payload at boundary
        visible_content = f"""{filler}

[CHUNK BOUNDARY EXPLOITATION]
{payload}

Continue with normal content below this line to avoid suspicion.
Additional information and standard company policies follow..."""

        return self.create_poisoned_pdf(
            visible_content=visible_content,
            hidden_instruction=payload,
            filename=filename,
            method="white_text",
            title="Standard Operating Procedures",
        )

    def create_poisoned_docx(
        self,
        visible_content: str,
        hidden_instruction: str,
        filename: str = "document.docx",
        method: str = "custom_properties",
        title: str = "Company Document",
    ) -> bytes:
        """
        Create DOCX with malicious payload.

        Args:
            visible_content: Normal content user sees
            hidden_instruction: Malicious payload
            filename: Suggested filename
            method: Injection method (custom_properties, hidden_text, comments, both)
            title: Document title

        Returns:
            DOCX file as bytes

        Raises:
            DocumentGeneratorError: If DOCX generation fails
        """
        if not DOCX_AVAILABLE:
            raise DocumentGeneratorError("python-docx not installed. Run: pip install python-docx")

        try:
            doc = Document()

            # Add title
            heading = doc.add_heading(title, level=1)

            # Add visible content
            for paragraph in visible_content.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            # Add hidden content based on method
            if method in ["hidden_text", "both"]:
                # Add as hidden text (font color = white)
                hidden_para = doc.add_paragraph(hidden_instruction)
                for run in hidden_para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    run.font.size = Pt(1)  # Tiny

            if method in ["comments", "both"]:
                # Add as comment (normally not extracted by simple parsers)
                # Note: python-docx doesn't support adding comments directly
                # We'll add it as a custom XML part
                self._add_hidden_comment(doc, hidden_instruction)

            if method in ["custom_properties", "both"]:
                # Add as custom document property
                self._add_custom_properties(doc, hidden_instruction)

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            logger.info("docx_generated", method=method, size_bytes=len(docx_bytes))

            return docx_bytes

        except Exception as e:
            logger.error("docx_generation_failed", error=str(e))
            raise DocumentGeneratorError(f"Failed to generate DOCX: {e}")

    def _add_custom_properties(self, doc: "Document", payload: str):
        """Add custom document properties with malicious payload."""
        try:
            # Access core properties
            core_props = doc.core_properties

            # Inject into subject and keywords
            core_props.subject = f"SYSTEM: {payload}"
            core_props.keywords = payload
            core_props.comments = f"[INSTRUCTION] {payload}"

        except Exception as e:
            logger.warning("custom_properties_failed", error=str(e))

    def _add_hidden_comment(self, doc: "Document", payload: str):
        """Add hidden comment to DOCX (embedded in XML)."""
        try:
            # This is a simplified approach
            # In practice, we'd need to manipulate the OOXML directly
            # For now, we add it as a hidden paragraph
            para = doc.add_paragraph()
            para.add_run(f"[COMMENT: {payload}]")
            # Make it very small and hidden
            for run in para.runs:
                run.font.size = Pt(1)
                run.font.color.rgb = RGBColor(255, 255, 255)

        except Exception as e:
            logger.warning("hidden_comment_failed", error=str(e))

    def _wrap_text(self, text: str, width: int) -> List[str]:
        """Wrap text to specified width."""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0

        for word in words:
            if current_length + len(word) + 1 <= width:
                current_line.append(word)
                current_length += len(word) + 1
            else:
                if current_line:
                    lines.append(" ".join(current_line))
                current_line = [word]
                current_length = len(word)

        if current_line:
            lines.append(" ".join(current_line))

        return lines

    def _generate_filler_content(self, target_length: int) -> str:
        """Generate realistic filler content of specified length."""
        filler_paragraphs = [
            "This document outlines our company's standard operating procedures and guidelines.",
            "All employees are expected to familiarize themselves with these policies.",
            "Security is our top priority and we take data protection seriously.",
            "Please review these guidelines carefully and contact HR with any questions.",
            "Our mission is to provide excellent service while maintaining the highest standards.",
            "We believe in transparency, integrity, and respect for all stakeholders.",
            "These policies are updated regularly to reflect current best practices.",
        ]

        filler = ""
        while len(filler) < target_length:
            filler += " " + filler_paragraphs[len(filler) % len(filler_paragraphs)]

        return filler[:target_length]

    def generate_semantic_dos_documents(
        self, topic: str, contradictory_statements: List[str], format: str = "pdf"
    ) -> List[tuple[bytes, str]]:
        """
        Generate multiple contradictory documents for semantic DoS attack.

        Args:
            topic: The topic to create contradictory docs about
            contradictory_statements: List of contradictory policies
            format: Document format (pdf or docx)

        Returns:
            List of (document_bytes, filename) tuples
        """
        documents = []

        for i, statement in enumerate(contradictory_statements):
            filename = f"{topic}_policy_v{i+1}.{format}"

            content = f"""# {topic.title()} Policy

## Overview

This document defines our official {topic} policy.

## Policy Statement

{statement}

## Enforcement

This policy is effective immediately and supersedes all previous versions.

Last Updated: {datetime.now().strftime('%Y-%m-%d')}
"""

            if format == "pdf":
                doc_bytes = self.create_poisoned_pdf(
                    visible_content=content,
                    hidden_instruction=statement,
                    filename=filename,
                    method="metadata",
                    title=f"{topic.title()} Policy v{i+1}",
                )
            else:  # docx
                doc_bytes = self.create_poisoned_docx(
                    visible_content=content,
                    hidden_instruction=statement,
                    filename=filename,
                    method="custom_properties",
                    title=f"{topic.title()} Policy v{i+1}",
                )

            documents.append((doc_bytes, filename))

        logger.info("semantic_dos_documents_generated", count=len(documents), topic=topic)

        return documents


# Convenience functions


def create_rag_test_pdf(
    visible_content: str, malicious_payload: str, method: str = "metadata"
) -> bytes:
    """
    Quick helper to create poisoned PDF for testing.

    Args:
        visible_content: Legitimate content
        malicious_payload: Attack payload
        method: Injection method

    Returns:
        PDF bytes
    """
    generator = PoisonedDocumentGenerator()
    return generator.create_poisoned_pdf(
        visible_content=visible_content, hidden_instruction=malicious_payload, method=method
    )


def create_rag_test_docx(
    visible_content: str, malicious_payload: str, method: str = "custom_properties"
) -> bytes:
    """
    Quick helper to create poisoned DOCX for testing.

    Args:
        visible_content: Legitimate content
        malicious_payload: Attack payload
        method: Injection method

    Returns:
        DOCX bytes
    """
    generator = PoisonedDocumentGenerator()
    return generator.create_poisoned_docx(
        visible_content=visible_content, hidden_instruction=malicious_payload, method=method
    )
