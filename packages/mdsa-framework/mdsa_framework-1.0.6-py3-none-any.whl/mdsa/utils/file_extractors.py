"""
File Extraction Utilities for RAG Document Processing

Supports extracting text from:
- Plain text (.txt)
- PDF documents (.pdf)
- Word documents (.docx)
- Markdown (.md)
- CSV files (.csv)
- Excel spreadsheets (.xlsx, .xls)

Author: MDSA Framework Team
Date: 2025-12-22
"""

import logging
from pathlib import Path
from typing import Optional
import io

logger = logging.getLogger(__name__)


def extract_text_from_file(file_content: bytes, filename: str) -> Optional[str]:
    """
    Extract text content from various file formats.

    Args:
        file_content: Raw file content as bytes
        filename: Original filename (used to determine file type)

    Returns:
        Extracted text content or None if extraction failed
    """
    file_ext = Path(filename).suffix.lower()

    extractors = {
        '.txt': extract_txt,
        '.md': extract_txt,  # Markdown is plain text
        '.pdf': extract_pdf,
        '.docx': extract_docx,
        '.csv': extract_csv,
        '.xlsx': extract_excel,
        '.xls': extract_excel
    }

    extractor = extractors.get(file_ext)
    if not extractor:
        logger.warning(f"[File Extractor] Unsupported file type: {file_ext}")
        return None

    try:
        text = extractor(file_content)
        if text:
            logger.info(f"[File Extractor] Extracted {len(text)} characters from {filename}")
            return text
        else:
            logger.warning(f"[File Extractor] No text extracted from {filename}")
            return None
    except Exception as e:
        logger.error(f"[File Extractor] Error extracting text from {filename}: {e}")
        return None


def extract_txt(file_content: bytes) -> str:
    """Extract text from plain text or markdown files."""
    try:
        # Try UTF-8 first
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        # Fallback to latin-1
        try:
            return file_content.decode('latin-1')
        except Exception as e:
            logger.error(f"[TXT] Failed to decode text file: {e}")
            return ""


def extract_pdf(file_content: bytes) -> str:
    """Extract text from PDF files."""
    try:
        import pypdf
    except ImportError:
        logger.error("[PDF] pypdf not installed. Install with: pip install pypdf")
        return ""

    try:
        pdf_file = io.BytesIO(file_content)
        reader = pypdf.PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"[PDF] Could not extract text from page {page_num}: {e}")

        full_text = '\n\n'.join(text_parts)
        logger.info(f"[PDF] Extracted text from {len(reader.pages)} pages")
        return full_text

    except Exception as e:
        logger.error(f"[PDF] Error extracting PDF: {e}")
        return ""


def extract_docx(file_content: bytes) -> str:
    """Extract text from Word DOCX files."""
    try:
        from docx import Document
    except ImportError:
        logger.error("[DOCX] python-docx not installed. Install with: pip install python-docx")
        return ""

    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = '\n\n'.join(text_parts)
        logger.info(f"[DOCX] Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        return full_text

    except Exception as e:
        logger.error(f"[DOCX] Error extracting Word document: {e}")
        return ""


def extract_csv(file_content: bytes) -> str:
    """Extract text from CSV files."""
    try:
        import csv
    except ImportError:
        logger.error("[CSV] csv module not available")
        return ""

    try:
        # Decode bytes to string
        csv_text = file_content.decode('utf-8')
        csv_file = io.StringIO(csv_text)

        reader = csv.reader(csv_file)
        rows = list(reader)

        if not rows:
            return ""

        # Format as markdown table
        text_parts = []

        # Header
        if rows:
            header = ' | '.join(rows[0])
            text_parts.append(header)
            text_parts.append('-' * len(header))

        # Data rows
        for row in rows[1:]:
            row_text = ' | '.join(str(cell) for cell in row)
            text_parts.append(row_text)

        full_text = '\n'.join(text_parts)
        logger.info(f"[CSV] Extracted {len(rows)} rows")
        return full_text

    except Exception as e:
        logger.error(f"[CSV] Error extracting CSV: {e}")
        return ""


def extract_excel(file_content: bytes) -> str:
    """Extract text from Excel files (.xlsx, .xls)."""
    try:
        import openpyxl
    except ImportError:
        logger.error("[Excel] openpyxl not installed. Install with: pip install openpyxl")
        return ""

    try:
        excel_file = io.BytesIO(file_content)
        workbook = openpyxl.load_workbook(excel_file, data_only=True)

        text_parts = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            # Add sheet name as header
            text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

            # Get all rows
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            # Format as table
            for row in rows:
                # Filter out None values
                row_values = [str(cell) if cell is not None else '' for cell in row]
                if any(row_values):  # Skip empty rows
                    row_text = ' | '.join(row_values)
                    text_parts.append(row_text)

        full_text = '\n'.join(text_parts)
        logger.info(f"[Excel] Extracted {len(workbook.sheetnames)} sheets")
        return full_text

    except Exception as e:
        logger.error(f"[Excel] Error extracting Excel: {e}")
        return ""


# Supported file extensions
SUPPORTED_EXTENSIONS = {
    '.txt': 'Plain Text',
    '.md': 'Markdown',
    '.pdf': 'PDF Document',
    '.docx': 'Word Document',
    '.csv': 'CSV Spreadsheet',
    '.xlsx': 'Excel Spreadsheet',
    '.xls': 'Excel Spreadsheet (Legacy)'
}


def is_supported_file(filename: str) -> bool:
    """Check if a file type is supported for text extraction."""
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS


def get_supported_extensions() -> list:
    """Get list of supported file extensions."""
    return list(SUPPORTED_EXTENSIONS.keys())
