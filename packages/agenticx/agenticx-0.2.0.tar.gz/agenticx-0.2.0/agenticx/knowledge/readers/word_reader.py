"""Word document reader for AgenticX Knowledge Management System"""

import asyncio
import logging
from pathlib import Path
from typing import List, Optional, Union, Any

from ..base import BaseReader
from ..document import Document, DocumentMetadata

logger = logging.getLogger(__name__)


class WordReader(BaseReader):
    """Reader for Word documents (.doc, .docx)"""
    
    def __init__(
        self,
        extract_images: bool = False,
        extract_tables: bool = True,
        **kwargs
    ):
        """
        Initialize WordReader
        
        Args:
            extract_images: Whether to extract images from document
            extract_tables: Whether to extract tables from document
            **kwargs: Additional configuration
        """
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.config = kwargs
        
        # Try to import required libraries
        self._word_library = self._get_word_library()
    
    def _get_word_library(self) -> Optional[str]:
        """Detect available Word processing library"""
        try:
            import python_docx2txt
            return 'docx2txt'
        except ImportError:
            pass
        
        try:
            from docx import Document as DocxDocument
            return 'python-docx'
        except ImportError:
            pass
        
        try:
            import mammoth
            return 'mammoth'
        except ImportError:
            pass
        
        return None
    
    def _has_docx2txt(self) -> bool:
        """Check if docx2txt is available"""
        try:
            import docx2txt
            return True
        except ImportError:
            return False
    
    def _has_antiword(self) -> bool:
        """Check if antiword is available"""
        import subprocess
        try:
            subprocess.run(['antiword', '-h'], capture_output=True, check=False)
            return True
        except FileNotFoundError:
            return False
    
    async def read(self, source: Union[str, Path]) -> List[Document]:
        """Read Word document and return documents
        
        Args:
            source: File path to read
            
        Returns:
            List containing single document with extracted text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ImportError: If no Word library is available
            Exception: If Word reading fails
        """
        
        if not self._word_library:
            raise ImportError(
                "No Word library available. Install one of: "
                "python-docx2txt (pip install docx2txt), "
                "python-docx (pip install python-docx), "
                "mammoth (pip install mammoth)"
            )
        
        file_path = Path(source)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
        
        # Extract text based on file type and available library
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.doc':
            # For .doc files, try antiword first, then docx2txt
            if self._has_antiword():
                content, metadata_dict = await self._read_with_antiword(file_path)
            elif 'docx2txt' in [self._word_library] or self._has_docx2txt():
                content, metadata_dict = await self._read_with_docx2txt(file_path)
            else:
                raise ImportError("antiword or docx2txt is required for .doc files. Install with: brew install antiword or pip install docx2txt")
        elif file_extension == '.docx':
            # For .docx files, prefer python-docx, fallback to others
            if self._word_library == 'python-docx':
                content, metadata_dict = await self._read_with_python_docx(file_path)
            elif self._word_library == 'mammoth':
                content, metadata_dict = await self._read_with_mammoth(file_path)
            elif self._word_library == 'docx2txt' or self._has_docx2txt():
                content, metadata_dict = await self._read_with_docx2txt(file_path)
            else:
                raise ImportError(f"No suitable library for .docx files. Install one of: python-docx, mammoth, docx2txt")
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")
        
        # Create document metadata
        metadata = DocumentMetadata(
            name=file_path.name,
            source=str(file_path),
            source_type="file",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_path.suffix.lower() == '.docx' else "application/msword",
            size=len(content),
            reader_name=self.__class__.__name__,
            custom=metadata_dict
        )
        
        # Create document
        document = Document(
            content=content,
            metadata=metadata
        )
        
        return [document]
    
    async def read_async(self, source: Union[str, Path], **kwargs):
        """Asynchronously read Word file and yield documents
        
        Args:
            source: File path to read
            **kwargs: Additional arguments
            
        Yields:
            Documents one by one
        """
        documents = await self.read(source)
        for document in documents:
            yield document
    
    async def _read_with_docx2txt(self, file_path: Path) -> tuple[str, dict]:
        """Read Word document using docx2txt"""
        
        def _extract():
            import docx2txt
            
            # Extract text
            text = docx2txt.process(str(file_path))
            
            # Basic metadata
            metadata = {
                'word_reader_library': 'docx2txt',
                'file_extension': file_path.suffix.lower()
            }
            
            return text or "", metadata
        
        # Run in thread pool
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract)
    
    async def _read_with_python_docx(self, file_path: Path) -> tuple[str, dict]:
        """Read Word document using python-docx"""
        
        def _extract():
            from docx import Document as DocxDocument
            
            # Only works with .docx files
            if file_path.suffix.lower() != '.docx':
                raise ValueError("python-docx only supports .docx files")
            
            doc = DocxDocument(str(file_path))
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables if requested
            if self.extract_tables:
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        table_text.append(" | ".join(row_text))
                    if table_text:
                        text_parts.append("\n".join(table_text))
            
            # Get document properties
            core_props = doc.core_properties
            metadata = {
                'word_reader_library': 'python-docx',
                'file_extension': file_path.suffix.lower(),
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            return '\n\n'.join(text_parts), metadata
        
        # Run in thread pool
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract)
    
    async def _read_with_mammoth(self, file_path: Path) -> tuple[str, dict]:
        """Read Word document using mammoth"""
        
        def _extract():
            import mammoth
            
            # Only works with .docx files
            if file_path.suffix.lower() != '.docx':
                raise ValueError("mammoth only supports .docx files")
            
            with open(file_path, 'rb') as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
                messages = result.messages
            
            metadata = {
                'word_reader_library': 'mammoth',
                'file_extension': file_path.suffix.lower(),
                'conversion_messages': [str(msg) for msg in messages]
            }
            
            return text or "", metadata
        
        # Run in thread pool
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _extract)
    
    async def _read_with_antiword(self, file_path: Path) -> tuple[str, dict]:
        """Read Word document using antiword"""
        import subprocess
        
        try:
            # Run antiword to extract text
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                check=True
            )
            
            content = result.stdout
            
            # Create metadata
            metadata_dict = {
                'source': str(file_path),
                'file_type': 'doc',
                'extraction_method': 'antiword',
                'file_size': file_path.stat().st_size,
                'character_count': len(content)
            }
            
            return content, metadata_dict
            
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error reading Word document with antiword: {e}")
        except Exception as e:
            raise Exception(f"Error reading Word document with antiword: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions
        
        Returns:
            List of supported extensions
        """
        
        return ['.doc', '.docx']
    
    def __str__(self) -> str:
        return f"WordReader(library={self._word_library}, extract_tables={self.extract_tables})"