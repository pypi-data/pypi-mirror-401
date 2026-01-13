#!/usr/bin/env python3
"""Generate test documents for integration testing."""

from pathlib import Path


def generate_docx(output_path: Path) -> None:
    """Generate sample DOCX file."""
    try:
        from docx import Document
    except ImportError:
        print("Warning: python-docx not installed, skipping DOCX generation")
        return

    doc = Document()
    doc.add_heading("Sample DOCX Document", 0)
    doc.add_paragraph(
        "This is a test paragraph containing the keyword searchable for testing purposes."
    )
    doc.add_paragraph("Machine learning and artificial intelligence are important topics.")
    doc.add_paragraph("Python programming language is widely used in data science.")

    doc.save(str(output_path))
    print(f"Generated: {output_path}")


def generate_html(output_path: Path) -> None:
    """Generate sample HTML file."""
    html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Sample HTML Document</title>
</head>
<body>
    <h1>Sample HTML Document</h1>
    <p>This is a test paragraph containing the keyword <strong>searchable</strong>.</p>
    <p>Machine learning and artificial intelligence are important topics.</p>
    <p>Python programming language is widely used in data science.</p>
    <ul>
        <li>Item 1</li>
        <li>Item 2</li>
        <li>Item 3</li>
    </ul>
</body>
</html>
"""
    output_path.write_text(html)
    print(f"Generated: {output_path}")


def generate_json(output_path: Path) -> None:
    """Generate sample JSON file."""
    import json

    data = {
        "title": "Sample JSON Document",
        "content": "This is searchable content for testing purposes.",
        "keywords": ["machine learning", "artificial intelligence", "Python"],
        "metadata": {
            "author": "Test Author",
            "date": "2025-12-26",
            "version": "1.0",
        },
        "items": [
            {"id": 1, "name": "Item 1", "description": "First test item"},
            {"id": 2, "name": "Item 2", "description": "Second test item"},
        ],
    }

    output_path.write_text(json.dumps(data, indent=2))
    print(f"Generated: {output_path}")


def generate_xml(output_path: Path) -> None:
    """Generate sample XML file."""
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<document>
    <title>Sample XML Document</title>
    <content>This is searchable content for testing purposes.</content>
    <keywords>
        <keyword>machine learning</keyword>
        <keyword>artificial intelligence</keyword>
        <keyword>Python</keyword>
    </keywords>
    <metadata>
        <author>Test Author</author>
        <date>2025-12-26</date>
    </metadata>
    <items>
        <item id="1">
            <name>Item 1</name>
            <description>First test item</description>
        </item>
        <item id="2">
            <name>Item 2</name>
            <description>Second test item</description>
        </item>
    </items>
</document>
"""
    output_path.write_text(xml)
    print(f"Generated: {output_path}")


def generate_csv(output_path: Path) -> None:
    """Generate sample CSV file."""
    csv = """name,value,category,description
Item 1,100,searchable,First test item with machine learning
Item 2,200,analysis,Second test item with artificial intelligence
Item 3,300,programming,Third test item with Python development
"""
    output_path.write_text(csv)
    print(f"Generated: {output_path}")


def generate_markdown(output_path: Path) -> None:
    """Generate sample Markdown file."""
    md = """# Sample Markdown Document

This is a test paragraph containing the keyword **searchable**.

## Topics

- Machine learning
- Artificial intelligence
- Python programming

## Content

Machine learning and artificial intelligence are important topics in modern technology.
Python programming language is widely used in data science and web development.

### Code Example

```python
def hello_world():
    print("Hello, searchable world!")
```

## Conclusion

This document is for testing purposes.
"""
    output_path.write_text(md)
    print(f"Generated: {output_path}")


def main() -> None:
    """Generate all test documents."""
    fixtures_dir = Path(__file__).parent
    output_dir = fixtures_dir / "documents"
    output_dir.mkdir(exist_ok=True)

    print("Generating test documents...")
    print()

    generate_docx(output_dir / "sample.docx")
    generate_html(output_dir / "sample.html")
    generate_json(output_dir / "sample.json")
    generate_xml(output_dir / "sample.xml")
    generate_csv(output_dir / "sample.csv")
    generate_markdown(output_dir / "sample.md")

    print()
    print(f"Test documents generated in {output_dir}")


if __name__ == "__main__":
    main()
