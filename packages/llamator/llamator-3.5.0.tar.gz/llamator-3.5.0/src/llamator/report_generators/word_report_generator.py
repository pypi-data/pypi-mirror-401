# llamator/src/llamator/report_generators/word_report_generator.py
import logging
import os

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from llamator.attack_provider.attack_registry import test_classes


def set_table_border(table, border_color="404040", border_size=4, border_space=0, border_type="single"):
    """
    Applies borders to a Word table. The color is typically dark gray by default.

    Parameters
    ----------
    table : docx.table.Table
        The table to which borders will be applied.
    border_color : str, optional
        A hex code for the border color (default "404040").
    border_size : int, optional
        Border thickness (default 4).
    border_space : int, optional
        Space between the border and the text (default 0).
    border_type : str, optional
        Type of border (default "single").
    """
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_element = OxmlElement(f"w:{edge}")
        edge_element.set(qn("w:val"), border_type)
        edge_element.set(qn("w:sz"), str(border_size))
        edge_element.set(qn("w:space"), str(border_space))
        edge_element.set(qn("w:color"), border_color)
        tblBorders.append(edge_element)

    tblPr.append(tblBorders)


def set_page_background(document, color):
    """
    Sets a background color for the entire Word document.

    Parameters
    ----------
    document : docx.Document
        The target Word document.
    color : str
        Hex code for the background color (e.g., "F0F8FF" = AliceBlue).
    """
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color)
    document.element.insert(0, background)


def get_tests_mapping() -> dict:
    """
    Constructs a dictionary mapping code_name to the full info dict
    for all registered tests.

    Returns
    -------
    dict
        code_name -> info
    """
    mapping = {}
    for cls in test_classes:
        if hasattr(cls, "info"):
            mapping[cls.info["code_name"]] = cls.info
    return mapping


def set_cell_background(cell, fill_color):
    """
    Applies background color to a single table cell.

    Parameters
    ----------
    cell : docx.table._Cell
        The table cell.
    fill_color : str
        The hex color code for the fill (e.g., "FFDAB9").
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def set_table_background(table, fill_color):
    """
    Sets the background for the entire table.

    Parameters
    ----------
    table : docx.table.Table
        The target table.
    fill_color : str
        The hex color code for the background.
    """
    for row in table.rows:
        for cell in row.cells:
            set_cell_background(cell, fill_color)


def create_word_report(
    artifacts_dir: str,
    csv_folder_name: str = "csv_report",
    docx_file_name: str = "attacks_report.docx",
    status_legend: dict = None,
    main_statuses: tuple = ("Broken", "Resilient", "Errors"),
    language: str = "en",
) -> None:
    """
    Generates a Word report summarizing the results from CSV files,
    including background colors, styles, and localized text.

    Parameters
    ----------
    artifacts_dir : str
        Path to the folder containing artifacts.
    csv_folder_name : str, optional
        Subfolder that contains CSV files (default "csv_report").
    docx_file_name : str, optional
        Name of the resulting Word file (default "attacks_report.docx").
    status_legend : dict, optional
        A dictionary mapping statuses to short descriptions. If None, uses defaults.
    main_statuses : tuple, optional
        Main statuses of interest in summary (default ("Broken", "Resilient", "Errors")).
    language : str, optional
        "en" or "ru". Defaults to "en".
    """

    # Localized strings. The Russian text is fully reverted to the original.
    strings = {
        "ru": {
            "framework_title": "LLAMATOR",
            "testing_report": "Отчёт о тестировании",
            "legend_title": "Легенда статусов:",
            "status": "Статус",
            "attempts": "Количество попыток",
            "attack_results": "Результаты атаки: {}",
            "report_created": "Отчёт в Word создан: {}",
            "report_failed": "Не удалось создать отчёт в Word: {}",
            "csv_not_found": "Папка с CSV-файлами не найдена: {}",
            "csv_read_failed": "Не удалось прочитать CSV-файл {}: {}",
            "status_column_missing": "Столбец 'status' не найден в {}. Пропуск файла.",
            "test_not_found": "Тест с `code_name = '{in_code}'` не найден в описаниях. Пропуск файла {file}.",
            "description_missing": "Описание для теста '{test}' не найдено. Используется название CSV файла.",
            "total": "Итого",
        },
        "en": {
            "framework_title": "LLAMATOR",
            "testing_report": "Testing Report",
            "legend_title": "Status Legend:",
            "status": "Status",
            "attempts": "Number of Attempts",
            "attack_results": "Attack Results: {}",
            "report_created": "Word report created: {}",
            "report_failed": "Failed to create Word report: {}",
            "csv_not_found": "CSV folder not found: {}",
            "csv_read_failed": "Failed to read CSV file {}: {}",
            "status_column_missing": "The 'status' column is missing in {}. Skipping file.",
            "test_not_found": "Test with `code_name = '{in_code}'` not found in descriptions. Skipping file {file}.",
            "description_missing": "Description for test '{test}' not found. Using CSV file name.",
            "total": "Total",
        },
    }

    lang = strings.get(language, strings["en"])
    tests_mapping = get_tests_mapping()

    if not tests_mapping:
        logging.error("Failed to load test information from attack classes. Aborting report generation.")
        return

    if status_legend is None:
        # Default status legend
        if language == "ru":
            status_legend = {
                "Broken": "Модель не смогла противостоять атаке.",
                "Resilient": "Модель успешно отразила атаку.",
                "Errors": "Во время выполнения атаки произошла ошибка.",
            }
        elif language == "en":
            status_legend = {
                "Broken": "The model failed to withstand the attack.",
                "Resilient": "The model successfully resisted the attack.",
                "Errors": "An error occurred during the attack attempt.",
            }

    csv_folder_path = os.path.join(artifacts_dir, csv_folder_name)
    output_docx_path = os.path.join(artifacts_dir, docx_file_name)

    document = Document()

    # Adjust margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    styles = document.styles

    # Create or retrieve style for the main title
    if "CenterTitle" not in styles:
        center_title_style = styles.add_style("CenterTitle", WD_STYLE_TYPE.PARAGRAPH)
        center_title_style.font.name = "Helvetica"
        center_title_style.font.size = Pt(20)
        center_title_style.font.bold = True
        center_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_title_style.paragraph_format.space_after = Pt(10)
        center_title_style.font.eastasia = "Helvetica"
        center_title_style.font.ascii = "Helvetica"
    else:
        center_title_style = styles["CenterTitle"]

    if "Heading2Custom" not in styles:
        heading2_style = styles.add_style("Heading2Custom", WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.base_style = styles["Heading 2"]
        heading2_style.font.name = "Helvetica"
        heading2_style.font.size = Pt(14)
        heading2_style.paragraph_format.space_after = Pt(10)
        heading2_style.font.eastasia = "Helvetica"
        heading2_style.font.ascii = "Helvetica"
    else:
        heading2_style = styles["Heading2Custom"]

    if "NormalStyle" not in styles:
        normal_style = styles.add_style("NormalStyle", WD_STYLE_TYPE.PARAGRAPH)
        normal_style.base_style = styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.font.eastasia = "Times New Roman"
        normal_style.font.ascii = "Times New Roman"
    else:
        normal_style = styles["NormalStyle"]

    # Set background color
    set_page_background(document, "F0F8FF")

    # Main header
    framework_title = document.add_paragraph(lang["framework_title"], style="CenterTitle")
    for run in framework_title.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)

    # Horizontal line
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_para = p._element
    p_pPr = p_para.get_or_add_pPr()
    p_pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_pBdr.append(bottom)
    p_pPr.append(p_pBdr)

    # Subheader
    testing_report_title = document.add_paragraph(lang["testing_report"], style="CenterTitle")
    for run in testing_report_title.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)

    document.add_paragraph()

    # Legend title
    legend_title = document.add_paragraph(lang["legend_title"], style="Heading2Custom")
    for run in legend_title.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)

    # Print out the legend
    for status, description in status_legend.items():
        par = document.add_paragraph(style="NormalStyle")
        run_status = par.add_run(f"{status}: ")
        run_status.bold = True
        run_status.font.color.rgb = RGBColor(80, 80, 80)
        run_description = par.add_run(description)
        run_description.font.color.rgb = RGBColor(80, 80, 80)

    document.add_paragraph()

    # Check CSV folder
    if not os.path.isdir(csv_folder_path):
        logging.error(lang["csv_not_found"].format(csv_folder_path))
        return

    csv_files = sorted([f for f in os.listdir(csv_folder_path) if f.endswith(".csv")])

    for idx, csv_file in enumerate(csv_files):
        in_code_name = os.path.splitext(csv_file)[0]
        csv_path = os.path.join(csv_folder_path, csv_file)

        test_info = tests_mapping.get(in_code_name)
        if not test_info:
            logging.warning(lang["test_not_found"].format(in_code=in_code_name, file=csv_file))
            test_name = in_code_name
            test_description = ""
        else:
            test_name = test_info["name"]
            # Attempt to get the localized description
            test_description = test_info["description"].get(language, "")
            if not test_description:
                logging.warning(lang["description_missing"].format(test=test_name))
                test_name = in_code_name
                test_description = ""

        try:
            df = pd.read_csv(csv_path)
        except Exception as e:
            logging.error(lang["csv_read_failed"].format(csv_path, e))
            continue

        if "status" not in df.columns:
            logging.warning(lang["status_column_missing"].format(csv_path))
            continue

        summary_counts = df["status"].value_counts().to_dict()
        summary_complete = {status: summary_counts.get(status, 0) for status in main_statuses}
        total_attempts = sum(summary_complete.values())

        # Heading for the test
        attack_title = document.add_paragraph(test_name, style="Heading2Custom")
        for run in attack_title.runs:
            run.font.color.rgb = RGBColor(80, 80, 80)
        attack_title.paragraph_format.keep_with_next = True
        attack_title.paragraph_format.keep_together = True

        # Description
        if test_description:
            desc_par = document.add_paragraph(test_description, style="NormalStyle")
            for run in desc_par.runs:
                run.font.color.rgb = RGBColor(80, 80, 80)
            desc_par.paragraph_format.keep_with_next = True
            desc_par.paragraph_format.keep_together = True

        # Summary table
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = lang["status"]
        hdr_cells[1].text = lang["attempts"]

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    run.font.name = "Helvetica"
                    run.font.eastasia = "Helvetica"
                    run.font.ascii = "Helvetica"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fill statuses
        for status, count in summary_complete.items():
            row_cells = table.add_row().cells
            row_cells[0].text = status
            row_cells[1].text = str(count)

            for c in row_cells:
                for paragraph in c.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(80, 80, 80)
                        run.font.name = "Times New Roman"
                        run.font.eastasia = "Times New Roman"
                        run.font.ascii = "Times New Roman"

            if status == "Resilient":
                set_cell_background(row_cells[0], "DFFFD6")  # pale green
            elif status == "Broken":
                set_cell_background(row_cells[0], "FFDAB9")  # pale orange
            elif status == "Errors":
                set_cell_background(row_cells[0], "FFFFE0")  # pale yellow

        # Total row
        total_row = table.add_row().cells
        total_label = lang["total"]
        total_row[0].text = total_label
        total_row[1].text = str(total_attempts)

        for c in total_row:
            for paragraph in c.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    run.font.name = "Times New Roman"
                    run.font.eastasia = "Times New Roman"
                    run.font.ascii = "Times New Roman"

        set_cell_background(total_row[0], "d6d6d6")
        set_cell_background(total_row[1], "d6d6d6")

        set_table_border(table, border_color="404040", border_size=4, border_space=0, border_type="single")
        set_table_background(table, "ededed")

        for row in table.rows:
            row.allow_break_across_pages = False

        if idx < len(csv_files) - 1:
            document.add_paragraph()

    # Save doc
    try:
        document.save(output_docx_path)
        # No console prints about the doc creation here
    except Exception as e:
        logging.error(lang["report_failed"].format(e))
