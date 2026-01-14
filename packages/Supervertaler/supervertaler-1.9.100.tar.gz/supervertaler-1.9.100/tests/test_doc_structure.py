"""
Test script to check if doc.paragraphs includes table paragraphs
"""
from docx import Document
import sys

if len(sys.argv) < 2:
    print("Usage: python test_doc_structure.py <docx_file>")
    sys.exit(1)

doc = Document(sys.argv[1])

print(f"\n{'='*80}")
print(f"DOCUMENT STRUCTURE ANALYSIS")
print(f"{'='*80}\n")

print(f"Total paragraphs in doc.paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Count paragraphs in tables
table_para_count = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            table_para_count += len(cell.paragraphs)

print(f"Paragraphs in tables: {table_para_count}")
print(f"Expected regular paragraphs: {len(doc.paragraphs) - table_para_count}")

print(f"\n{'='*80}")
print(f"PARAGRAPH LIST (first 15):")
print(f"{'='*80}\n")

for i, para in enumerate(doc.paragraphs[:15]):
    in_table = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if para in cell.paragraphs:
                    in_table = True
                    break
    
    marker = " [IN TABLE]" if in_table else ""
    print(f"{i:2d}. {para.text[:60]:<60} {marker}")

print(f"\n{'='*80}\n")
