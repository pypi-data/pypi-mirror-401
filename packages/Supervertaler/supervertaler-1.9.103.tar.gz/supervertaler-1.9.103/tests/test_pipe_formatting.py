"""
Test script to verify pipe symbol formatting (bold + red) in CafeTran export.
This creates a small test file with pipes in the translation.
"""

import sys
import os

# Add modules directory to path
modules_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'modules')
if modules_dir not in sys.path:
    sys.path.insert(0, modules_dir)

from cafetran_docx_handler import CafeTranDOCXHandler
from docx.shared import RGBColor

print("\n" + "="*70)
print("Testing Pipe Symbol Formatting (Bold + Red)")
print("="*70 + "\n")

# Load the test file
test_file = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'projects',
    'test_document (CafeTran bilingual docx).docx'
)

print("Step 1: Loading source file...")
handler = CafeTranDOCXHandler()
if not handler.load(test_file):
    print("Failed to load file")
    sys.exit(1)

segments = handler.extract_source_segments()
print(f"✓ Loaded {len(segments)} segments\n")

print("Step 2: Creating test translations WITH pipe symbols...")
# Create translations with pipes at specific locations
test_translations = [
    "Biagio Pagano",
    "Van Wikipedia, de vrije encyclopedie",
    "Over",
    "Biagio Pagano| (geboren 29 januari 1983) is een Italiaanse voetballer",  # Pipe at end
    "Pagano had 250 wedstrijden gespeeld",
    "Persoonlijke informatie",
    "Geboortedatum",
    "29 januari 1983 (leeftijd 42)",
    "Geboorteplaats",
    "Napels|, Italië",  # Pipe after city name
    "Lengte",
    "1,80 m (5 ft 11 in)",
    "Carrière",
    "Pagano begon zijn carrière bij |Atalanta|.",  # Pipes around team name
    "Hij debuteerde in de Serie A op 17 juni 2001 tegen |Juventus FC| en speelde na terugkeer van |Lumezzane| uit de derde divisie in |2002-03 Serie C1|.",  # Multiple pipes
]

# Pad with more translations if needed
while len(test_translations) < len(segments):
    test_translations.append(f"Vertaling {len(test_translations) + 1}")

print(f"✓ Created {len(test_translations)} test translations\n")

# Show translations with pipes
print("Sample translations with pipes:")
for i, trans in enumerate(test_translations[:15]):
    if '|' in trans:
        print(f"  Row {i+1}: {trans}")

print("\nStep 3: Updating target segments...")
handler.update_target_segments(test_translations[:len(segments)])
print("✓ Updated\n")

print("Step 4: Saving with formatted pipes...")
output_path = test_file.replace('.docx', '_PIPE_FORMATTING_TEST.docx')
if not handler.save(output_path):
    print("Failed to save")
    sys.exit(1)

print(f"✓ Saved to: {output_path}\n")

print("Step 5: Verifying pipe formatting...")
from docx import Document

doc = Document(output_path)
table = doc.tables[0]

# Check specific rows with pipes
test_cases = [
    (4, "Biagio Pagano|"),
    (10, "Napels|, Italië"),
    (14, "|Atalanta|"),
    (15, "|Juventus FC| ... |Lumezzane| ... |2002-03 Serie C1|"),
]

all_correct = True
for row_idx, expected in test_cases:
    cell = table.rows[row_idx].cells[2]
    para = cell.paragraphs[0] if cell.paragraphs else None
    
    if not para or not para.runs:
        print(f"  ✗ Row {row_idx}: No runs found")
        all_correct = False
        continue
    
    pipe_runs = [run for run in para.runs if run.text == '|']
    
    if not pipe_runs:
        print(f"  ⚠ Row {row_idx}: No pipe symbols found")
        continue
    
    # Check if all pipes are bold and red
    pipes_correct = True
    for run in pipe_runs:
        is_bold = run.bold == True
        color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        is_red = color == RGBColor(255, 0, 0) if color else False
        
        if not (is_bold and is_red):
            pipes_correct = False
            break
    
    if pipes_correct:
        print(f"  ✓ Row {row_idx}: {len(pipe_runs)} pipe(s) - all BOLD + RED")
    else:
        print(f"  ✗ Row {row_idx}: {len(pipe_runs)} pipe(s) - formatting INCORRECT")
        all_correct = False

print("\n" + "="*70)
if all_correct:
    print("✓ All pipe symbols are correctly formatted as BOLD + RED!")
else:
    print("✗ Some pipes are missing formatting")
print("="*70)

print(f"\nOpen this file to verify visually: {output_path}")
