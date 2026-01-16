import base64
import os

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocxUtil:

    ########################   表格内容填充   ############################
    @staticmethod
    def insert_table_content(document_object, text_list=None, picture_list=None, width=None, height=None, style=None):
        """
        在表格插入数据
        text_list示例： {"table": 2, "row": 12, "cell": 2, "result": "本地url地址", style=None}
        picture_list示例：
        {"table": 2, "row": 12, "cell": 2, "result": "本地url地址", width=None, height=Cm(1.5)}
        {"table": 2, "row": 12, "cell": 2, "result": ["本地url地址", "url2"], width=None, height=Cm(1.5)}
        """
        DocxUtil.insert_table_text(document_object, text_list, style)
        DocxUtil.insert_table_picture(document_object, picture_list, width, height)

    @staticmethod
    def insert_table_text(document_object, text_list=None, style=None):
        """
        在表格插入文本类型数据
        text_list示例： {"table": 2, "row": 12, "cell": 2, "paragraph": 0, "result": "本地url地址", style=None}
        table 必填
        row 必填
        cell 必填
        paragraph 默认 0
        """
        if not text_list:
            return
        for text in text_list:
            result = text.get("result")
            if not result:
                continue
            cell = document_object.tables[text.get("table")].rows[text.get("row")].cells[text.get("cell")]
            paragraph = cell.paragraphs[text.get("paragraph", 0)]
            paragraph.add_run(result, style=text.get("style", style))

    @staticmethod
    def insert_table_picture(document_object, picture_list=None, width=None, height=None):
        """
        在表格插入图片类型数据
        picture_list示例：
        {"table": 2, "row": 12, "cell": 2, "paragraph": 0, "result": "本地url地址", width=None, height=Cm(1.5)}
        {"table": 2, "row": 12, "cell": 2, "paragraph": 0, "result": ["本地url地址", "url2"], width=None, height=Cm(1.5)}
        """
        if not picture_list:
            return
        for picture in picture_list:
            result = picture.get("result")
            if not result:
                continue
            if type(result) == list:
                img_list = result
            elif type(result) == str:
                img_list = [result]
            else:
                return
            cell = document_object.tables[picture.get("table")].rows[picture.get("row")].cells[picture.get("cell")]
            run = cell.paragraphs[picture.get("paragraph", 0)].add_run()
            for img in img_list:
                if not os.path.exists(img):
                    continue
                run.add_picture(img, width=picture.get("width") or width, height=picture.get("height") or height)

    @staticmethod
    def insert_base64_image_to_table(document_object, base64_img_list, width=None, height=None):
        """
        把 base64 格式的字符串插入docx中
        base64_img_list示例：
        {"table": 2, "row": 12, "cell": 2, "paragraph": 0, "result": "base64_str", width=None, height=Cm(1.5)}
        {"table": 2, "row": 12, "cell": 2, "paragraph": 0, "result": ["base64_str", "base64_str"], width=None, height=Cm(1.5)}
        """
        for base64_img in base64_img_list:
            # 处理参数result
            result = base64_img.get("result")
            if not result:
                continue
            if type(result) == list:
                result_list = result
            elif type(result) == str:
                result_list = [result]
            else:
                continue

            # 解码 Base64 字符串
            base64_list = []
            for base64_str in result_list:
                if base64_str.startswith("data:image/png;base64,"):
                    base64_str = base64_str[len("data:image/png;base64,"):]
                    base64_list.append(base64_str)
            if not base64_list:
                continue
            for base64_str in base64_list:
                image_data = base64.b64decode(base64_str)
                # 临时图片文件名
                temp_image_path = "statics/temp_image.jpg"
                try:
                    # 保存为临时图片文件
                    with open(temp_image_path, 'wb') as f:
                        f.write(image_data)

                    DocxUtil.insert_table_picture(
                        document_object,
                        [
                            {
                                "table": base64_img.get("table"),
                                "row": base64_img.get("row"),
                                "cell": base64_img.get("cell"),
                                "paragraph": base64_img.get("paragraph", 0),
                                "result": temp_image_path,
                                "width": base64_img.get("width"),
                                "height": base64_img.get("height"),
                            }
                        ],
                        width=width,
                        height=height
                    )
                except Exception as e:
                    print(f"插入图片时出错: {e}")
                finally:
                    # 删除临时图片文件
                    if os.path.exists(temp_image_path):
                        os.remove(temp_image_path)

    ########################   文档文字替换   ############################
    @staticmethod
    def replace_text(document_object, table_replace_list=None, para_replace_list=None):
        """
        在指定位置替换文本
        para_replace_list 示例：{"paragraph": 2, "run": 12, "cell": 2, "mark": "$code", "result": "替换结果值"}
        table_replace_list 示例：{"table": 2, "row": 12, "cell": 2, "mark": "$code", "result": "替换结果值"}
        """
        if para_replace_list:
            for para_replace in para_replace_list:
                run = document_object.paragraphs[para_replace.get("paragraph")].runs[para_replace.get("run")]
                run.text = run.text.replace(para_replace.get("mark"), para_replace.get("result") or "")
        if table_replace_list:
            for table_replace in table_replace_list:
                cell = document_object.tables[table_replace.get("table")][table_replace.get("row")][
                    table_replace.get("cell")]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(table_replace.get("mark"), table_replace.get("result") or "")

        return

    @staticmethod
    def replace_range_text(document_object, section=None, section_dict=None, paragraph=None, para_dict=None, table=None,
                     table_dict=None):
        """
        在文档内的字段，
        :param document_object: Document 对象，代表要处理的 Word 文档
        :param section 指定页面，不指定便利所有
        :param section_dict 替换内容
        :param paragraph: 指定段落，默认所有段落
        :param para_dict: 例如 { "$name": "名称" }
        :param table: 指定表，默认所有表
        :param table_dict: 例如 { "$name": "名称" }
        """
        DocxUtil.replace_section(document_object, section_dict, section)
        DocxUtil.replace_paragraph(document_object, para_dict, paragraph)
        DocxUtil.replace_table(document_object, table_dict, table)

        return document_object

    @staticmethod
    def replace_section(document_object, section_dict, section=None):
        """
        页眉替换文本
        :param document_object:
        :param section_dict: 例如 { "$code": "编码" }
        :param section: 指定页眉，默认所有页眉
        :return:
        """
        if not section_dict:
            return
        section_list = [section] if section else document_object.sections
        for section in section_list:
            header = section.header
            for header_para in header.paragraphs:
                for i in range(len(header_para.runs)):
                    for key, value in section_dict.items():
                        header_para.runs[i].text = header_para.runs[i].text.replace(key, value or "")

    @staticmethod
    def replace_paragraph(document_object, para_dict, paragraph=None):
        """
        段落替换文本
        :param document_object:
        :param para_dict: 例如 { "$name": "名称" }
        :param paragraph: 指定段落，默认所有段落
        :return:
        """
        if not para_dict:
            return
        paragraph_list = [paragraph] if paragraph else document_object.paragraphs
        for para in paragraph_list:
            for i in range(len(para.runs)):
                for key, value in para_dict.items():
                    if key in para.runs[i].text:
                        para.runs[i].text = para.runs[i].text.replace(key, value or "")

    @staticmethod
    def replace_table(document_object, table_dict, table=None):
        """
        表格替换文本
        :param document_object:
        :param table_dict: 例如 { "$name": "名称" }
        :param table: 指定表，默认所有表
        :return:
        """
        if not table_dict:
            return
        table_list = [table] if table else document_object.tables
        for table in table_list:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in table_dict.items():
                        if key in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(key, value or "")

    ########################   表格行复制   ############################
    @staticmethod
    def copy_row(target_table, source_row, insert_index=None, row_num=1):
        """
        target_table: 指定操作表
        source_row： 指定复制的行
        insert_index： 复制后的行插入位置
        row_num： 复制行的数量
        完整复制行包括合并属性
        """
        while row_num > 0:
            row_num -= 1
            # 在表格末尾添加新行
            new_row = target_table.add_row()

            # 复制单元格内容和格式
            delete_cell = 0
            for src_idx, src_cell in enumerate(source_row.cells):
                new_cell = new_row.cells[src_idx]
                new_cell = DocxUtil.copy_cell_properties(new_cell, src_cell)

                # 统计合并单元格的占用数量
                grid_span = src_cell._tc.tcPr.find(qn("w:gridSpan"))
                span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
                if span > 1:
                    # 处理合并单元格的列跨度
                    DocxUtil.copy_merged_cell_properties(src_cell, new_cell)
                    delete_cell += span - 1
            while delete_cell > 1:
                delete_cell -= 1
                if len(new_row.cells) > 0:
                    new_row._element.remove(new_row.cells[-1]._element)

            # 复制行高
            source_tr = source_row._tr
            new_tr = new_row._tr
            trPr = source_tr.find(qn('w:trPr'))
            if trPr is not None:
                new_trPr = new_tr.get_or_add_trPr()
                for child in trPr.iterchildren():
                    if child.tag.endswith('}trHeight'):
                        new_trHeight = OxmlElement('w:trHeight')
                        for attr, value in child.attrib.items():
                            new_trHeight.set(attr, value)
                        new_trPr.append(new_trHeight)

            # 移动行到指定位置
            if insert_index and insert_index < len(target_table.rows):
                target_table._tbl.insert(insert_index + 1, new_row._tr)
            else:
                target_table._tbl.append(new_row._tr)

    @staticmethod
    def copy_cell_properties(target_cell, source_cell):
        """
        target_cell: 指定单元格
        source_cell： 指定复制的单元格
        复制cell的属性
        """
        # 清空目标单元格内容,复制文本内容
        target_cell.text = ""
        target_cell.text = source_cell.text

        # 根据目标段落填充数量
        add_para_num = len(source_cell.paragraphs) - len(target_cell.paragraphs)
        while add_para_num > 0:
            add_para_num -= 1
            target_cell.add_paragraph()

        # 复制段落样式
        for para_idx, para in enumerate(source_cell.paragraphs):
            new_para = target_cell.paragraphs[para_idx]
            new_para.style = para.style

            # 复制段落对齐方式
            new_para.alignment = para.alignment

            # 复制行间距
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            new_para.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
            new_para.paragraph_format.page_break_before = para.paragraph_format.page_break_before
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.widow_control = para.paragraph_format.widow_control

            # 复制文本格式
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.style = run.style
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
                new_run.font.size = run.font.size
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb

        # 复制单元格垂直对齐方式
        target_cell.vertical_alignment = source_cell.vertical_alignment
        # 复制单元格边框
        source_tc = source_cell._tc
        target_tc = target_cell._tc
        source_tcPr = source_tc.tcPr
        target_tcPr = target_tc.get_or_add_tcPr()
        for border in ['top', 'left', 'bottom', 'right']:
            source_border = source_tcPr.xpath(f'.//w:{border}')
            if source_border:
                new_border = OxmlElement(f'w:{border}')
                for key, value in source_border[0].attrib.items():
                    new_border.set(key, value)
                target_tcPr.append(new_border)

        # 复制单元格背景色
        shading = source_tcPr.xpath('.//w:shd')
        if shading:
            new_shading = OxmlElement('w:shd')
            for key, value in shading[0].attrib.items():
                new_shading.set(key, value)
            target_tcPr.append(new_shading)

        return target_cell

    @staticmethod
    def copy_merged_cell_properties(target_cell, source_cell):
        """复制合并单元格属性（gridSpan和vMerge）"""
        # 获取源单元格属性元素
        source_tc = source_cell._tc
        source_tcPr = source_tc.tcPr

        # 准备目标单元格属性元素
        target_tc = target_cell._tc
        target_tcPr = target_tc.get_or_add_tcPr()

        # 处理gridSpan（跨列）
        grid_span = source_tcPr.find(qn("w:gridSpan"))
        if grid_span is not None:
            # 删除目标单元格原有gridSpan
            existing = target_tcPr.find(qn("w:gridSpan"))
            if existing is not None:
                target_tcPr.remove(existing)
            # 复制新的gridSpan
            new_gs = OxmlElement("w:gridSpan")
            new_gs.set(qn("w:val"), grid_span.get(qn("w:val")))
            target_tcPr.append(new_gs)

        # 处理vMerge（跨行）
        v_merge = source_tcPr.find(qn("w:vMerge"))
        if v_merge is not None:
            # 删除目标单元格原有vMerge
            existing = target_tcPr.find(qn("w:vMerge"))
            if existing is not None:
                target_tcPr.remove(existing)
            # 复制新的vMerge
            new_vm = OxmlElement("w:vMerge")
            if v_merge.get(qn("w:val")) is not None:
                new_vm.set(qn("w:val"), v_merge.get(qn("w:val")))
            target_tcPr.append(new_vm)

    ########################   其他操作   ############################
    @staticmethod
    def remove_paragraph(target_cell, para_index):
        """
        删除单元格的段落
        """
        para = target_cell.paragraphs[para_index]
        p = para._element
        p.getparent().remove(p)
        para._element = None

        return

