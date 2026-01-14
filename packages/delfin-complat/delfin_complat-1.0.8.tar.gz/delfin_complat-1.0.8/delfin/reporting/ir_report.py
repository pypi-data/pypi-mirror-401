"""Generate IR spectrum reports in Word format."""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-GUI backend
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

from delfin.common.logging import get_logger
from delfin.ir_spectrum import IRMode, calculate_spectrum_range

logger = get_logger(__name__)


def lorentzian_broadening(
    modes: List[IRMode],
    wavenumber_range: tuple[float, float],
    num_points: int = 2000,
    fwhm: float = 10.0
) -> tuple[np.ndarray, np.ndarray]:
    """
    Apply Lorentzian broadening to IR spectrum.

    Args:
        modes: List of IR modes
        wavenumber_range: Range of wavenumbers (min, max) in cm^-1
        num_points: Number of points in broadened spectrum
        fwhm: Full width at half maximum in cm^-1

    Returns:
        Tuple of (wavenumbers, transmittance_percent)
    """
    wavenumbers = np.linspace(wavenumber_range[0], wavenumber_range[1], num_points)
    absorbance = np.zeros(num_points)

    gamma = fwhm / 2.0  # Half width at half maximum

    for mode in modes:
        if mode.intensity_km_mol > 0.1:  # Skip very weak modes
            # Lorentzian line shape
            lorentzian = (gamma**2) / ((wavenumbers - mode.frequency_cm1)**2 + gamma**2)
            # Scale by intensity
            absorbance += mode.intensity_km_mol * lorentzian

    # Normalize absorbance
    if absorbance.max() > 0:
        absorbance = absorbance / absorbance.max()

    # Convert to transmittance (Beer-Lambert: T = 10^(-A))
    # For visualization, use A as arbitrary absorbance units
    transmittance = 100.0 * (1.0 - absorbance)

    return wavenumbers, transmittance


def create_ir_spectrum_plot(
    modes: List[IRMode],
    output_png: Path,
    wavenumber_range: tuple[float, float] | None = None,
    fwhm: float = 10.0,
    dpi: int = 300,
    title: str = 'IR Spectrum'
) -> None:
    """Create IR spectrum plot with stick spectrum and broadened curve.

    Args:
        modes: List of IR modes
        output_png: Path to save PNG image
        wavenumber_range: Wavenumber range (min, max) in cm^-1 (auto if None)
        fwhm: Full width at half maximum for Lorentzian broadening (cm^-1)
        dpi: Resolution of output image
        title: Plot title
    """
    if not MATPLOTLIB_AVAILABLE:
        logger.warning("matplotlib not available, skipping IR spectrum plot")
        return

    if not modes:
        logger.warning("No IR modes to plot")
        return

    fig, ax = plt.subplots(figsize=(10, 6))

    # Auto-determine wavenumber range if not provided
    if wavenumber_range is None:
        wavenumber_range = calculate_spectrum_range(modes)

    # Filter modes with significant intensity
    significant = [m for m in modes if m.intensity_km_mol > 0.1]

    if not significant:
        logger.warning("No significant IR modes found for plotting")
        plt.close(fig)
        return

    # Create stick spectrum (inverted for transmittance representation)
    for mode in significant:
        # Calculate relative stick height (0-100% transmittance scale)
        # Higher intensity = lower transmittance
        intensity_normalized = min(mode.intensity_km_mol / 500.0, 1.0)
        transmittance = 100.0 * (1.0 - intensity_normalized)
        ax.plot([mode.frequency_cm1, mode.frequency_cm1],
                [100, transmittance],
                'r-', linewidth=1.5, alpha=0.6)

    # Add broadened spectrum
    wavenumbers, transmittance = lorentzian_broadening(
        modes,
        wavenumber_range=wavenumber_range,
        num_points=2000,
        fwhm=fwhm
    )
    ax.plot(wavenumbers, transmittance, 'b-', linewidth=2, label='IR Spectrum')

    # Formatting - IR spectra conventionally show high wavenumbers on the left
    ax.set_xlabel('Wavenumber (cm$^{-1}$)', fontsize=12, fontweight='bold')
    ax.set_ylabel('Transmittance (%)', fontsize=12, fontweight='bold')
    ax.set_title(title, fontsize=14, fontweight='bold')
    ax.set_xlim(wavenumber_range[1], wavenumber_range[0])  # Reversed for IR convention
    ax.set_ylim(0, 105)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.legend(loc='lower right')

    # Add minor gridlines
    ax.minorticks_on()
    ax.grid(which='minor', alpha=0.15, linestyle=':')

    plt.tight_layout()
    fig.savefig(output_png, dpi=dpi, bbox_inches='tight')
    plt.close(fig)
    logger.info(f"IR spectrum plot saved to {output_png}")


def set_cell_border(cell, **kwargs):
    """
    Set cell borders in Word table.

    Args:
        cell: Table cell object
        **kwargs: Border properties (top, bottom, left, right, insideH, insideV)
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            edge_el.set(qn('w:space'), str(edge_data.get('space', 0)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)


def generate_ir_report(
    output_file: Path,
    modes: List[IRMode],
    spectrum_png: Path,
    source_file: Optional[Path] = None,
    molecule_name: Optional[str] = None
) -> None:
    """
    Generate Word report for IR spectrum with table and plot.

    Args:
        output_file: Output Word document path
        modes: List of IR modes
        spectrum_png: Path to spectrum plot PNG
        source_file: Optional source .out file path
        molecule_name: Optional molecule name for title
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available, cannot generate Word report")
        return

    if not modes:
        logger.warning("No IR modes to include in report")
        return

    doc = Document()

    # Title
    title_text = f"IR Spectrum Report"
    if molecule_name:
        title_text += f" - {molecule_name}"
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Source file info
    if source_file:
        info = doc.add_paragraph()
        info.add_run("Source file: ").bold = True
        info.add_run(str(source_file))
        info.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Number of modes
    summary = doc.add_paragraph()
    summary.add_run(f"Total vibrational modes: ").bold = True
    summary.add_run(f"{len(modes)}")

    # Add spectrum plot
    if spectrum_png.exists():
        doc.add_heading("IR Spectrum", level=2)
        doc.add_picture(str(spectrum_png), width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        logger.warning(f"Spectrum plot not found: {spectrum_png}")

    # Add table with IR data
    doc.add_heading("Vibrational Frequencies and Intensities", level=2)

    # Filter modes with significant intensity for table
    significant_modes = [m for m in modes if m.intensity_km_mol > 1.0]

    if not significant_modes:
        doc.add_paragraph("No significant IR-active modes found (intensity > 1.0 km/mol).")
    else:
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Mode', 'Frequency\n(cm⁻¹)', 'Intensity\n(km/mol)', 'ε\n(L/(mol·cm))', 'T²\n(a.u.)']

        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            # Format header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for mode in significant_modes:
            row_cells = table.add_row().cells
            row_cells[0].text = str(mode.mode_number)
            row_cells[1].text = f"{mode.frequency_cm1:.2f}"
            row_cells[2].text = f"{mode.intensity_km_mol:.2f}"
            row_cells[3].text = f"{mode.epsilon:.6f}"
            row_cells[4].text = f"{mode.t_squared:.6f}"

            # Center align all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Add note about filtered data
        note = doc.add_paragraph()
        note.add_run(f"Note: ").bold = True
        note.add_run(f"Table shows {len(significant_modes)} modes with intensity > 1.0 km/mol. ")
        note.add_run(f"Total modes in calculation: {len(modes)}.")
        note_format = note.paragraph_format
        note_format.space_before = Pt(6)

    # Add full data table in appendix (optional)
    if len(modes) > len(significant_modes):
        doc.add_page_break()
        doc.add_heading("Appendix: Complete Vibrational Mode Data", level=2)

        # Create full table
        full_table = doc.add_table(rows=1, cols=8)
        full_table.style = 'Light Grid Accent 1'

        # Header
        header_cells = full_table.rows[0].cells
        headers = ['Mode', 'Freq\n(cm⁻¹)', 'Int\n(km/mol)', 'ε', 'T²', 'TX', 'TY', 'TZ']

        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # All modes
        for mode in modes:
            row_cells = full_table.add_row().cells
            row_cells[0].text = str(mode.mode_number)
            row_cells[1].text = f"{mode.frequency_cm1:.2f}"
            row_cells[2].text = f"{mode.intensity_km_mol:.2f}"
            row_cells[3].text = f"{mode.epsilon:.4e}"
            row_cells[4].text = f"{mode.t_squared:.4e}"
            row_cells[5].text = f"{mode.tx:.4f}"
            row_cells[6].text = f"{mode.ty:.4f}"
            row_cells[7].text = f"{mode.tz:.4f}"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    # Save document
    doc.save(str(output_file))
    logger.info(f"IR report saved to {output_file}")
