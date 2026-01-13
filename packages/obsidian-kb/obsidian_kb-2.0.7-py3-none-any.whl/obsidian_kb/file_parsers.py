"""Модуль для извлечения текста из различных форматов файлов."""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: Path) -> str:
    """Извлечение текста из PDF файла.

    Args:
        file_path: Путь к PDF файлу

    Returns:
        Извлечённый текст

    Raises:
        Exception: Если не удалось прочитать файл
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "pypdf не установлен. Установите его: pip install pypdf"
        ) from None

    try:
        reader = PdfReader(file_path)
        text_parts: list[str] = []

        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Ошибка при извлечении текста со страницы {page.page_number + 1}: {e}")
                continue

        text = "\n\n".join(text_parts)
        if not text.strip():
            logger.warning(f"Не удалось извлечь текст из PDF: {file_path}")
        return text

    except Exception as e:
        logger.error(f"Ошибка при чтении PDF файла {file_path}: {e}")
        raise


def extract_text_from_docx(file_path: Path) -> str:
    """Извлечение текста из DOCX файла.

    Args:
        file_path: Путь к DOCX файлу

    Returns:
        Извлечённый текст

    Raises:
        Exception: Если не удалось прочитать файл
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx не установлен. Установите его: pip install python-docx"
        ) from None

    try:
        doc = Document(file_path)
        text_parts: list[str] = []

        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        text = "\n\n".join(text_parts)
        if not text.strip():
            logger.warning(f"Не удалось извлечь текст из DOCX: {file_path}")
        return text

    except Exception as e:
        logger.error(f"Ошибка при чтении DOCX файла {file_path}: {e}")
        raise


def extract_text_from_file(file_path: Path) -> Optional[str]:
    """Извлечение текста из файла в зависимости от его расширения.

    Args:
        file_path: Путь к файлу

    Returns:
        Извлечённый текст или None, если формат не поддерживается
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        try:
            return extract_text_from_pdf(file_path)
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из PDF {file_path}: {e}")
            return None

    elif suffix == ".docx":
        try:
            return extract_text_from_docx(file_path)
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOCX {file_path}: {e}")
            return None

    elif suffix == ".md":
        # Markdown файлы — читаем как текст
        try:
            return file_path.read_text(encoding="utf-8")
        except Exception as e:
            logger.error(f"Ошибка при чтении Markdown файла {file_path}: {e}")
            return None

    else:
        logger.debug(f"Неподдерживаемый формат файла: {suffix}")
        return None


def is_supported_file(file_path: Path) -> bool:
    """Проверка, поддерживается ли формат файла.

    Args:
        file_path: Путь к файлу

    Returns:
        True, если формат поддерживается
    """
    supported_extensions = {".md", ".pdf", ".docx"}
    return file_path.suffix.lower() in supported_extensions

