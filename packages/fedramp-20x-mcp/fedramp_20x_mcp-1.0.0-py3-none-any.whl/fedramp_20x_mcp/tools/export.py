"""
FedRAMP 20x MCP Server - Export Tools

This module contains tool implementation functions for export.
"""
import json
import logging
import os
from pathlib import Path
from datetime import datetime
from typing import Any, Optional

from ..data_loader import FedRAMPDataLoader

logger = logging.getLogger(__name__)

async def export_to_excel(
    export_type: str,
    output_path: Optional[str] = None,
    data_loader: Optional[FedRAMPDataLoader] = None
) -> str:
    """
    Export FedRAMP 20x data to an Excel file.
    
    Args:
        export_type: Type of data to export. Options:
            - "ksi" - All 72 Key Security Indicators
            - "all_requirements" - All 329 requirements across all families
            - "definitions" - All FedRAMP definitions
        output_path: Optional custom output path. If not provided, saves to Downloads folder
        data_loader: FedRAMPDataLoader instance (created if not provided)
        
    Returns:
        Path to the generated Excel file
    """
    # Import openpyxl here to avoid import errors if not installed
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return "Error: openpyxl package is required for Excel export. Install with: pip install openpyxl"
    
    # Create data_loader if not provided
    if data_loader is None:
        data_loader = FedRAMPDataLoader()
    
    await data_loader.load_data()
    
    # Determine output path
    if output_path is None:
        downloads_folder = Path.home() / "Downloads"
        downloads_folder.mkdir(parents=True, exist_ok=True)  # Create if doesn't exist
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"FedRAMP_20x_{export_type}_{timestamp}.xlsx"
        output_path = str(downloads_folder / filename)
    
    # Create workbook
    wb = Workbook()
    if wb.active:
        wb.remove(wb.active)  # Remove default sheet
    
    # Define styles
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_alignment = Alignment(vertical="top", wrap_text=True)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    if export_type == "ksi":
        # Export all KSIs
        ws = wb.create_sheet("Key Security Indicators")
        
        # Headers
        headers = ["KSI ID", "Name", "Category", "Status", "Statement", "Note", "NIST 800-53 Controls", "Reference", "Reference URL", "Impact Levels"]
        ws.append(headers)
        
        # Style headers
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Get all KSIs
        all_ksi = data_loader.list_all_ksi()
        
        for ksi in all_ksi:
            ksi_id = ksi.get('id', '')
            name = ksi.get('name', '')
            category = ksi.get('category', '')
            retired = ksi.get('retired', False)
            status = 'Retired' if retired else 'Active'
            statement = ksi.get('statement', '')
            note = ksi.get('note', '')
            
            # Format controls
            controls = ksi.get('controls', [])
            if controls:
                control_list = [f"{c.get('control_id', '').upper()} - {c.get('title', '')}" for c in controls]
                controls_str = '; '.join(control_list)
            else:
                controls_str = ''
            
            reference = ksi.get('reference', '')
            reference_url = ksi.get('reference_url', '')
            impact = ksi.get('impact', {})
            impact_levels = ', '.join([k.title() for k, v in impact.items() if v]) if impact else ''
            
            row = [ksi_id, name, category, status, statement, note, controls_str, reference, reference_url, impact_levels]
            ws.append(row)
            
            # Style data rows
            for col_num in range(1, len(headers) + 1):
                cell = ws.cell(row=ws.max_row, column=col_num)
                cell.alignment = cell_alignment
                cell.border = border
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 15  # KSI ID
        ws.column_dimensions['B'].width = 40  # Name
        ws.column_dimensions['C'].width = 30  # Category
        ws.column_dimensions['D'].width = 10  # Status
        ws.column_dimensions['E'].width = 60  # Statement
        ws.column_dimensions['F'].width = 40  # Note
        ws.column_dimensions['G'].width = 70  # NIST 800-53 Controls
        ws.column_dimensions['H'].width = 30  # Reference
        ws.column_dimensions['I'].width = 50  # Reference URL
        ws.column_dimensions['J'].width = 20  # Impact Levels
        
        # Freeze header row
        ws.freeze_panes = 'A2'
    
    elif export_type == "all_requirements":
        # Export all requirements
        ws = wb.create_sheet("All Requirements")
        
        headers = ["Requirement ID", "Family", "Term/Name", "Description", "Document"]
        ws.append(headers)
        
        # Style headers
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Get all requirements
        if not data_loader._data_cache:
            return "Error: Data not loaded. Please try again."
        all_reqs = data_loader._data_cache["requirements"]
        
        for req_id, req in sorted(all_reqs.items()):
            family = req_id.split('-')[0] if '-' in req_id else ''
            term = req.get('term', req.get('name', ''))
            description = req.get('description', req.get('definition', ''))
            document = req.get('document_name', '')
            
            row = [req_id, family, term, description, document]
            ws.append(row)
            
            # Style data rows
            for col_num in range(1, len(headers) + 1):
                cell = ws.cell(row=ws.max_row, column=col_num)
                cell.alignment = cell_alignment
                cell.border = border
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 18  # ID
        ws.column_dimensions['B'].width = 12  # Family
        ws.column_dimensions['C'].width = 40  # Term
        ws.column_dimensions['D'].width = 60  # Description
        ws.column_dimensions['E'].width = 30  # Document
        
        ws.freeze_panes = 'A2'
    
    elif export_type == "definitions":
        # Export all definitions
        ws = wb.create_sheet("FedRAMP Definitions")
        
        headers = ["Term", "Definition", "Notes", "References"]
        ws.append(headers)
        
        # Style headers
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Get all definitions
        all_defs = data_loader.list_all_definitions()
        
        for defn in sorted(all_defs, key=lambda x: x.get('term', '')):
            term = defn.get('term', '')
            definition = defn.get('definition', '')
            notes = defn.get('notes', '')
            references = defn.get('references', '')
            
            # Convert lists to strings for Excel compatibility
            if isinstance(notes, list):
                notes = '\n'.join(str(n) for n in notes)
            if isinstance(references, list):
                references = '\n'.join(str(r) for r in references)
            
            row = [term, definition, notes, references]
            ws.append(row)
            
            # Style data rows
            for col_num in range(1, len(headers) + 1):
                cell = ws.cell(row=ws.max_row, column=col_num)
                cell.alignment = cell_alignment
                cell.border = border
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 30  # Term
        ws.column_dimensions['B'].width = 60  # Definition
        ws.column_dimensions['C'].width = 40  # Notes
        ws.column_dimensions['D'].width = 30  # References
        
        ws.freeze_panes = 'A2'
    
    else:
        return f"Error: Unknown export_type '{export_type}'. Valid options: ksi, all_requirements, definitions"
    
    # Save workbook
    wb.save(output_path)
    
    return f"Excel file created successfully at: {output_path}"



async def export_to_csv(
    export_type: str,
    output_path: Optional[str] = None
) -> str:
    """
    Export FedRAMP 20x data to a CSV file.
    
    Args:
        export_type: Type of data to export. Options:
            - "ksi" - All 72 Key Security Indicators
            - "all_requirements" - All 329 requirements across all families
            - "definitions" - All FedRAMP definitions
        output_path: Optional custom output path. If not provided, saves to Downloads folder
        
    Returns:
        Path to the generated CSV file
    """
    await data_loader.load_data()
    
    # Determine output path
    if output_path is None:
        downloads_folder = str(Path.home() / "Downloads")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"FedRAMP_20x_{export_type}_{timestamp}.csv"
        output_path = os.path.join(downloads_folder, filename)
    
    if export_type == "ksi":
        # Export all KSIs
        headers = ["KSI ID", "Name", "Category", "Status", "Statement", "Note", "NIST 800-53 Controls", "Reference", "Reference URL", "Impact Levels"]
        
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(headers)
            
            # Get all KSIs
            all_ksi = data_loader.list_all_ksi()
            
            for ksi in all_ksi:
                ksi_id = ksi.get('id', '')
                name = ksi.get('name', '')
                category = ksi.get('category', '')
                retired = ksi.get('retired', False)
                status = 'Retired' if retired else 'Active'
                statement = ksi.get('statement', '')
                note = ksi.get('note', '')
                
                # Format controls
                controls = ksi.get('controls', [])
                if controls:
                    control_list = [f"{c.get('control_id', '').upper()} - {c.get('title', '')}" for c in controls]
                    controls_str = '; '.join(control_list)
                else:
                    controls_str = ''
                
                reference = ksi.get('reference', '')
                reference_url = ksi.get('reference_url', '')
                impact = ksi.get('impact', {})
                impact_levels = ', '.join([k.title() for k, v in impact.items() if v]) if impact else ''
                
                row = [ksi_id, name, category, status, statement, note, controls_str, reference, reference_url, impact_levels]
                writer.writerow(row)
    
    elif export_type == "all_requirements":
        # Export all requirements
        headers = ["Requirement ID", "Family", "Term/Name", "Description", "Document"]
        
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(headers)
            
            # Get all requirements
            if not data_loader._data_cache:
                return "Error: Data not loaded. Please try again."
            all_reqs = data_loader._data_cache["requirements"]
            
            for req_id, req in sorted(all_reqs.items()):
                family = req_id.split('-')[0] if '-' in req_id else ''
                term = req.get('term', req.get('name', ''))
                description = req.get('description', req.get('definition', ''))
                document = req.get('document_name', '')
                
                row = [req_id, family, term, description, document]
                writer.writerow(row)
    
    elif export_type == "definitions":
        # Export all definitions
        headers = ["Term", "Definition", "Notes", "References"]
        
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(headers)
            
            # Get all definitions
            all_defs = data_loader.list_all_definitions()
            
            for defn in sorted(all_defs, key=lambda x: x.get('term', '')):
                term = defn.get('term', '')
                definition = defn.get('definition', '')
                notes = defn.get('notes', '')
                references = defn.get('references', '')
                
                row = [term, definition, notes, references]
                writer.writerow(row)
    
    else:
        return f"Error: Unknown export_type '{export_type}'. Valid options: ksi, all_requirements, definitions"
    
    return f"CSV file created successfully at: {output_path}"



async def generate_ksi_specification(
    ksi_id: str,
    evidence_collection_strategy: str,
    output_path: Optional[str] = None
) -> str:
    """
    Generate a product specification document for a KSI aligned with FedRAMP 20x requirements.
    
    Args:
        ksi_id: The KSI identifier (e.g., "KSI-AFR-01")
        evidence_collection_strategy: High-level evidence collection strategy description
        output_path: Optional custom output path. If not provided, saves to Downloads folder
        
    Returns:
        Path to the generated Word document
    """
    # Import python-docx here to avoid import errors if not installed
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        return "Error: python-docx package is required for Word document generation. Install with: pip install python-docx"
    
    await data_loader.load_data()
    
    # Get the KSI
    ksi = data_loader.get_ksi(ksi_id.upper())
    if not ksi:
        return f"Error: KSI '{ksi_id}' not found"
    
    # Determine output path
    if output_path is None:
        downloads_folder = str(Path.home() / "Downloads")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = ksi_id.replace('/', '_').replace('\\', '_')
        filename = f"KSI_Spec_{safe_name}_{timestamp}.docx"
        output_path = os.path.join(downloads_folder, filename)
    
    # Create document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Document title
    title = doc.add_heading(f"Product Specification: {ksi.get('name', ksi_id)}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with KSI ID
    subtitle = doc.add_paragraph(f"KSI ID: {ksi_id}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(54, 96, 146)  # FedRAMP blue
    
    # Add metadata table
    doc.add_paragraph()
    metadata_heading = doc.add_heading('Document Information', 2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Populate metadata
    metadata = [
        ('KSI ID', ksi_id),
        ('Category', ksi.get('category', 'N/A')),
        ('Impact Levels', ', '.join([k.title() for k, v in ksi.get('impact', {}).items() if v]) or 'N/A'),
        ('Status', 'Retired' if ksi.get('retired', False) else 'Active'),
        ('Document Date', datetime.now().strftime('%B %d, %Y'))
    ]
    
    for idx, (label, value) in enumerate(metadata):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = str(value)
        # Bold the labels
        table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Overview section
    doc.add_page_break()
    doc.add_heading('1. Overview', 1)
    
    doc.add_heading('1.1 Purpose', 2)
    doc.add_paragraph(
        f"This document provides a comprehensive product specification for implementing "
        f"{ksi.get('name', ksi_id)} in compliance with FedRAMP 20x Key Security Indicators. "
        f"It is designed to guide engineering teams through planning, implementation, and "
        f"evidence collection activities."
    )
    
    doc.add_heading('1.2 Scope', 2)
    impact = ksi.get('impact', {})
    impact_text = "This KSI applies to "
    if impact.get('low') and impact.get('moderate'):
        impact_text += "both Low and Moderate impact systems."
    elif impact.get('low'):
        impact_text += "Low impact systems."
    elif impact.get('moderate'):
        impact_text += "Moderate impact systems."
    else:
        impact_text += "systems as defined by FedRAMP authorization requirements."
    doc.add_paragraph(impact_text)
    
    # Requirement Statement
    doc.add_heading('2. Requirement Statement', 1)
    statement = ksi.get('statement', '')
    if statement:
        doc.add_paragraph(statement)
    else:
        doc.add_paragraph("No statement available for this KSI.")
    
    # Check for retired status
    if ksi.get('retired', False):
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(f"⚠ NOTE: {ksi.get('note', 'This KSI has been retired.')}")
        note_run.font.color.rgb = RGBColor(192, 0, 0)
        note_run.font.bold = True
    
    # Related NIST 800-53 Controls
    doc.add_heading('3. Related NIST 800-53 Controls', 1)
    controls = ksi.get('controls', [])
    if controls:
        doc.add_paragraph(
            "This KSI aligns with the following NIST 800-53 Rev 5 security controls. "
            "Implementation must address these control requirements:"
        )
        for control in controls:
            control_id = control.get('control_id', '').upper()
            control_title = control.get('title', 'N/A')
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{control_id}: ").bold = True
            p.add_run(control_title)
    else:
        doc.add_paragraph("No specific NIST 800-53 controls mapped to this KSI.")
    
    # Reference Documentation
    reference = ksi.get('reference')
    reference_url = ksi.get('reference_url')
    if reference or reference_url:
        doc.add_heading('4. Reference Documentation', 1)
        if reference:
            doc.add_paragraph(f"Primary Reference: {reference}")
        if reference_url:
            doc.add_paragraph(f"Documentation URL: {reference_url}")
    
    # Azure-First Implementation Guidance
    doc.add_heading('5. Azure-First Implementation Guidance', 1)
    
    doc.add_heading('5.1 Recommended Azure Services', 2)
    doc.add_paragraph(
        "Based on FedRAMP 20x requirements and Azure Government compliance capabilities, "
        "consider the following Azure services for implementation:"
    )
    
    # Add Azure-specific recommendations based on category
    category = ksi.get('category', '')
    azure_recommendations = _get_azure_recommendations(category, controls)
    
    for service, description in azure_recommendations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{service}: ").bold = True
        p.add_run(description)
    
    doc.add_heading('5.2 Infrastructure as Code', 2)
    doc.add_paragraph(
        "Implement using Azure Bicep or Terraform with Azure Provider for repeatable, "
        "auditable deployments. Store IaC in version control (Azure Repos or GitHub) "
        "with branch protection and approval workflows."
    )
    
    doc.add_heading('5.3 Automation and Monitoring', 2)
    automation_items = [
        "Use Azure Policy for continuous compliance monitoring and enforcement",
        "Implement Azure Monitor and Log Analytics for centralized logging",
        "Configure Azure Security Center for security posture management",
        "Enable Microsoft Defender for Cloud for threat protection",
        "Use Azure Automation for remediation workflows"
    ]
    for item in automation_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Evidence Collection Strategy
    doc.add_heading('6. Evidence Collection Strategy', 1)
    
    doc.add_heading('6.1 User-Defined Strategy', 2)
    doc.add_paragraph(evidence_collection_strategy)
    
    doc.add_heading('6.2 Recommended Evidence Types', 2)
    evidence_items = [
        "Configuration screenshots from Azure Portal",
        "Azure Policy compliance reports exported as JSON/CSV",
        "Azure Monitor query results and dashboards",
        "IaC templates (Bicep/Terraform) with commit history",
        "Azure DevOps/GitHub Actions pipeline logs",
        "Microsoft Entra ID audit logs for access control",
        "Azure Resource Graph queries demonstrating compliance",
        "Automated test results from security validation pipelines"
    ]
    for item in evidence_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.3 Evidence Collection Schedule', 2)
    doc.add_paragraph(
        "Evidence should be collected regularly to align with FedRAMP "
        "Collaborative Continuous Monitoring (CCM) requirements. Automate evidence "
        "collection where possible using Azure Functions, Logic Apps, or Azure Automation. "
        "Engineering teams should determine the appropriate collection frequency based on "
        "system criticality and organizational requirements."
    )
    
    # Implementation Plan Template
    doc.add_heading('7. Implementation Plan Template', 1)
    
    phases = [
        {
            'phase': 'Phase 1: Requirements Analysis',
            'activities': [
                'Review NIST 800-53 control requirements',
                'Map controls to Azure services and features',
                'Identify gaps in current implementation',
                'Document assumptions and dependencies'
            ]
        },
        {
            'phase': 'Phase 2: Design',
            'activities': [
                'Create Azure architecture diagrams',
                'Design IaC templates (Bicep/Terraform)',
                'Define Azure Policy rules and initiatives',
                'Design monitoring and alerting strategy',
                'Plan evidence collection automation'
            ]
        },
        {
            'phase': 'Phase 3: Implementation',
            'activities': [
                'Deploy Azure infrastructure using IaC',
                'Configure Azure services and policies',
                'Implement monitoring and logging',
                'Set up automated evidence collection',
                'Configure security controls and access policies'
            ]
        },
        {
            'phase': 'Phase 4: Testing and Validation',
            'activities': [
                'Validate Azure Policy compliance',
                'Test evidence collection automation',
                'Perform security scanning and assessment',
                'Validate control implementation against requirements',
                'Document test results'
            ]
        },
        {
            'phase': 'Phase 5: Documentation and Evidence',
            'activities': [
                'Collect and organize all evidence',
                'Document implementation details',
                'Create operational runbooks',
                'Prepare for FedRAMP assessment',
                'Update SSP (System Security Plan) as needed'
            ]
        }
    ]
    
    for phase_info in phases:
        doc.add_heading(phase_info['phase'], 2)
        doc.add_paragraph("Activities:")
        for activity in phase_info['activities']:
            doc.add_paragraph(activity, style='List Bullet')
    
    # Team Roles and Responsibilities
    doc.add_heading('8. Team Roles and Responsibilities', 1)
    
    roles = [
        ('Cloud Architect', 'Design Azure architecture and services selection'),
        ('DevOps Engineer', 'Implement IaC templates and CI/CD pipelines'),
        ('Security Engineer', 'Configure security controls and monitoring'),
        ('Compliance Specialist', 'Ensure FedRAMP requirements are met'),
        ('Product Owner', 'Prioritize requirements and acceptance criteria'),
        ('QA Engineer', 'Validate implementation and test evidence collection')
    ]
    
    for role, responsibility in roles:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{role}: ").bold = True
        p.add_run(responsibility)
    
    # Success Criteria
    doc.add_heading('9. Success Criteria', 1)
    
    success_items = [
        'All NIST 800-53 controls implemented and validated in Azure',
        'Azure Policy reports 100% compliance for defined policies',
        'Evidence collection automation operational and tested',
        'IaC templates reviewed and approved',
        'Security scanning shows no high/critical vulnerabilities',
        'Documentation complete and reviewed',
        'Team training completed on operations and maintenance',
        'Quarterly evidence collection process validated'
    ]
    
    for item in success_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Risks and Mitigation
    doc.add_heading('10. Risks and Mitigation Strategies', 1)
    
    risks = [
        {
            'risk': 'Azure service limitations may not fully satisfy control requirements',
            'mitigation': 'Identify gaps early, use compensating controls, engage Azure support for guidance'
        },
        {
            'risk': 'Evidence collection automation failures',
            'mitigation': 'Implement monitoring and alerting, maintain manual collection procedures as backup'
        },
        {
            'risk': 'Compliance drift over time',
            'mitigation': 'Use Azure Policy in enforcement mode, implement continuous monitoring'
        },
        {
            'risk': 'Resource constraints or timeline delays',
            'mitigation': 'Prioritize critical controls, use phased implementation approach'
        }
    ]
    
    for risk_info in risks:
        doc.add_heading(f"Risk: {risk_info['risk']}", 2)
        doc.add_paragraph(f"Mitigation: {risk_info['mitigation']}")
    
    # Appendix
    doc.add_page_break()
    doc.add_heading('Appendix A: Additional Resources', 1)
    
    resources = [
        'FedRAMP 20x Requirements: https://fedramp.gov/docs/',
        'NIST 800-53 Rev 5: https://csrc.nist.gov/publications/detail/sp/800-53/rev-5/final',
        'Azure Security Documentation: https://learn.microsoft.com/azure/security/',
        'Azure Compliance: https://learn.microsoft.com/azure/compliance/',
        'Azure Government: https://azure.microsoft.com/explore/global-infrastructure/government/',
        'Microsoft Entra ID: https://learn.microsoft.com/entra/',
        'Azure Policy: https://learn.microsoft.com/azure/governance/policy/'
    ]
    
    for resource in resources:
        doc.add_paragraph(resource, style='List Bullet')
    
    # Save document
    doc.save(output_path)
    
    return f"Word document created successfully at: {output_path}"


def _get_azure_recommendations(category: str, controls: list) -> list:
    """Get Azure service recommendations based on KSI category and controls."""
    recommendations = []
    
    # Default recommendations for all KSIs
    recommendations.extend([
        ('Microsoft Entra ID', 'Identity and access management, conditional access, MFA'),
        ('Azure Policy', 'Compliance enforcement and continuous monitoring'),
        ('Azure Monitor', 'Centralized logging, alerting, and diagnostics')
    ])
    
    # Category-specific recommendations
    category_lower = category.lower()
    
    if 'education' in category_lower or 'training' in category_lower:
        recommendations.extend([
            ('Microsoft Viva Learning', 'Security awareness training platform'),
            ('Azure AD Access Reviews', 'Regular access certification and training tracking')
        ])
    
    if 'vulnerability' in category_lower or 'assessment' in category_lower:
        recommendations.extend([
            ('Microsoft Defender for Cloud', 'Vulnerability scanning and security recommendations'),
            ('Azure Security Center', 'Security posture management and compliance tracking')
        ])
    
    if 'incident' in category_lower or 'response' in category_lower:
        recommendations.extend([
            ('Microsoft Sentinel', 'SIEM for incident detection and response'),
            ('Azure Logic Apps', 'Automated incident response workflows')
        ])
    
    if 'data' in category_lower or 'encryption' in category_lower:
        recommendations.extend([
            ('Azure Key Vault', 'Cryptographic key and secret management'),
            ('Azure Storage encryption', 'Data at rest encryption with customer-managed keys')
        ])
    
    if 'network' in category_lower or 'boundary' in category_lower:
        recommendations.extend([
            ('Azure Firewall', 'Network security and filtering'),
            ('Azure DDoS Protection', 'DDoS mitigation'),
            ('Azure Virtual Network', 'Network segmentation and isolation')
        ])
    
    # Control-based recommendations
    control_ids = [c.get('control_id', '').upper() for c in controls]
    
    if any(c.startswith('AU-') for c in control_ids):  # Audit controls
        recommendations.append(('Azure Log Analytics', 'Centralized audit log collection and analysis'))
    
    if any(c.startswith('CM-') for c in control_ids):  # Configuration Management
        recommendations.append(('Azure Automation State Configuration', 'Desired state configuration management'))
    
    if any(c.startswith('SC-') for c in control_ids):  # System and Communications Protection
        recommendations.append(('Azure Front Door', 'Web application firewall and CDN'))
    
    # Remove duplicates while preserving order
    seen = set()
    unique_recommendations = []
    for rec in recommendations:
        if rec[0] not in seen:
            seen.add(rec[0])
            unique_recommendations.append(rec)
    
    return unique_recommendations