"""Document processing utilities for MCP-RAG."""

import logging
import os
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
from dataclasses import dataclass
import tempfile

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Processed document data structure."""
    filename: str
    content: str
    file_type: str
    metadata: Dict[str, Any]
    error: Optional[str] = None


class DocumentProcessor:
    """Document processor for various file formats."""

    def __init__(self):
        self.supported_formats = {
            '.txt': self._process_text,
            '.md': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
        }

    def process_file(self, file_path: Path, filename: str) -> ProcessedDocument:
        """Process a file and extract its content."""
        try:
            file_extension = file_path.suffix.lower()

            if file_extension not in self.supported_formats:
                return ProcessedDocument(
                    filename=filename,
                    content="",
                    file_type=file_extension,
                    metadata={"size": file_path.stat().st_size},
                    error=f"Unsupported file format: {file_extension}"
                )

            processor = self.supported_formats[file_extension]
            content, metadata = processor(file_path)

            return ProcessedDocument(
                filename=filename,
                content=content,
                file_type=file_extension,
                metadata=metadata
            )

        except Exception as e:
            logger.error(f"Failed to process file {filename}: {e}")
            return ProcessedDocument(
                filename=filename,
                content="",
                file_type=file_path.suffix.lower(),
                metadata={"size": file_path.stat().st_size},
                error=str(e)
            )

    def _process_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process text/markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = {
                "size": file_path.stat().st_size,
                "encoding": "utf-8",
                "lines": len(content.splitlines())
            }

            return content, metadata

        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['gbk', 'gb2312', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    metadata = {
                        "size": file_path.stat().st_size,
                        "encoding": encoding,
                        "lines": len(content.splitlines())
                    }
                    return content, metadata
                except UnicodeDecodeError:
                    continue

            raise ValueError("Unable to decode file with supported encodings")

    def _process_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process PDF files."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            content = ""
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"

            metadata = {
                "size": file_path.stat().st_size,
                "pages": len(pdf_reader.pages),
                "lines": len(content.splitlines())
            }

            return content, metadata

    def _process_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX files."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"

        metadata = {
            "size": file_path.stat().st_size,
            "paragraphs": len(doc.paragraphs),
            "lines": len(content.splitlines())
        }

        return content, metadata


# Global processor instance
document_processor = DocumentProcessor()


def get_document_processor() -> DocumentProcessor:
    """Get the global document processor instance."""
    return document_processor