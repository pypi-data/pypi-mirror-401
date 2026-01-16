"""DOCX extractor for Word documents."""

import os
from datetime import datetime
from pathlib import Path

from semplex_cli.types import TextMetadata
from semplex_cli.utils.file_utils import get_file_owner

# Optional dependency - will be imported at runtime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxExtractor:
    """Extract text content from Word documents (.docx)."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install it with: pip install python-docx"
            )

    def extract(self, file_path: Path) -> TextMetadata:
        """
        Extract text content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            TextMetadata with file content and metadata
        """
        file_path = Path(file_path)
        stat = file_path.stat()

        # Read DOCX
        doc = Document(file_path)

        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        content = "\n\n".join(text_parts)

        # Calculate metrics
        lines = content.split("\n")
        words = content.split()

        # Count sections (based on headings)
        section_count = sum(
            1 for para in doc.paragraphs
            if para.style and para.style.name.startswith("Heading")
        )

        return TextMetadata(
            filename=file_path.name,
            filepath=str(file_path.absolute()),
            file_type=".docx",
            file_size=stat.st_size,
            file_owner=get_file_owner(file_path),
            modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            extracted_at=datetime.now().isoformat(),
            content=content,
            char_count=len(content),
            word_count=len(words),
            line_count=len(lines),
            section_count=section_count if section_count > 0 else None,
        )
