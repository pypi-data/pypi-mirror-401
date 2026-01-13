"""Word Document Parser using python-docx."""

from __future__ import annotations

import os
from typing import Optional

from . import ParseOptions, ParseResult, DocumentItem


def parse_docx(file_path: str, options: Optional[ParseOptions] = None) -> ParseResult:
    """
    Parse a Word document, extracting full text content.

    Args:
        file_path: Path to the Word file
        options: Parsing options (not used for DOCX)

    Returns:
        ParseResult with single item containing full document text
    """
    filename = os.path.basename(file_path)

    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document parsing. "
            "Install with: pip install python-docx"
        )

    try:
        doc = Document(file_path)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text: list[str] = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        content = "\n".join(paragraphs)

        items: list[DocumentItem] = []
        if content:
            items.append({
                "number": 1,
                "text": content,
            })
        else:
            print(f"[memvid] No text content found in {filename}")

        return {
            "type": "docx",
            "filename": filename,
            "total_items": 1,
            "items": items,
        }
    except Exception as e:
        raise RuntimeError(
            f'Failed to parse Word file "{filename}": {e}. '
            f"Ensure the file is a valid .docx/.doc file."
        )
