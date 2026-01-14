# -*- coding: utf-8 -*-

"""
本模块功能：处理pdf文件的方便功能
所属工具包：证券投资分析工具SIAT 
SIAT：Security Investment Analysis Tool
创建日期：2025年8月2日
最新修订日期：2025年8月3日
作者：王德宏 (WANG Dehong, Peter)
作者单位：北京外国语大学国际商学院
作者邮件：wdehong2000@163.com
版权所有：王德宏
用途限制：仅限研究与教学使用。
特别声明：作者不对使用本工具进行证券投资导致的任何损益负责！
"""

#==============================================================================

# 首次运行前，请安装依赖：
# !pip install pypdf

#关闭所有警告
import warnings; warnings.filterwarnings('ignore')

#==============================================================================
from pypdf import PdfReader, PdfWriter

def pdf_split(src_path, ranges, dst_template):
    """
    src_path: 原PDF路径
    ranges:   [(start1, end1), (start2, end2), ...]，页码从1开始，闭区间
    dst_template: 输出文件名模板，如 "part_{}.pdf"
    """
    reader = PdfReader(src_path)
    num_pages = len(reader.pages)  # PDF 实际页数

    for idx, (start, end) in enumerate(ranges, 1):
        # 裁剪页码到 [1, num_pages] 区间
        actual_start = max(1, start)
        actual_end = min(end, num_pages)

        # 如果裁剪后没有页可写，则跳过
        if actual_start > actual_end:
            print(f"跳过第 {idx} 段：有效页范围为空（原始 {start}-{end}）")
            continue

        writer = PdfWriter()
        # 页码转为 0-based 索引，循环添加
        for page_num in range(actual_start - 1, actual_end):
            writer.add_page(reader.pages[page_num])

        output_path = dst_template.format(idx)
        with open(output_path, "wb") as f:
            writer.write(f)
        print(f"已生成：{output_path} （页码 {actual_start}-{actual_end}）")
        
    return


if __name__ =="__main__":
    source="校对稿2.pdf"
    
    cs=18      #每章首页偏移
    ce=18-1    #每章尾页偏移
    ranges1 = [(3+cs, 20+ce), 
               (20+cs, 36+ce), 
               (36+cs, 58+cs+1)]             # 第1篇页码范围
        
    pdf_split(source, ranges1, "Chapter_{}.pdf")
    
    
#==============================================================================
#==============================================================================

# 如果还没安装依赖，请先运行：
# !pip install pdf2docx python-docx

import re
import os
from pdf2docx import Converter
from docx import Document

def clean_text(text: str) -> str:
    """
    清理文本：
      - 去除首尾空白
      - 删除标点前后多余空格
      - 中文句子用中文标点，英文/数字用英文标点
      - 合并多余空格
    """
    text = text.strip()
    text = re.sub(r'\s+([,\.!?:;，。！？：；])', r'\1', text)
    text = re.sub(r'([,\.!?:;，。！？：；])\s+', r'\1', text)

    zh2ascii = {'。': '.', '，': ',', '？': '?', '！': '!', '；': ';', '：': ':'}
    for zh, asc in zh2ascii.items():
        text = re.sub(fr'{zh}(?=[A-Za-z0-9])', asc, text)

    ascii2zh = {v: k for k, v in zh2ascii.items()}
    for asc, zh in ascii2zh.items():
        text = re.sub(fr'(?<=[\u4e00-\u9fff]){re.escape(asc)}', zh, text)

    text = re.sub(r'\s+', ' ', text)
    return text

def delete_paragraph(paragraph):
    """彻底删除一个段落（底层 XML 操作）"""
    p = paragraph._p
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def merge_broken_paragraphs(doc: Document):
    """
    合并 PDF 转出来因为硬换行而断开的段落。
    每次合并后重新获取 paragraphs，避免操作已删除的元素。
    """
    pattern = re.compile(r'[。！？…;；:\?!…]$')
    while True:
        paras = doc.paragraphs
        merged = False
        for i in range(len(paras) - 1):
            p, nxt = paras[i], paras[i+1]
            if not p.text.strip():
                continue
            if not pattern.search(p.text.strip()):
                # 在当前段末尾加空格，再把下一段的 run 全部复制过来
                p.add_run(' ')
                for run in list(nxt.runs):
                    nr = p.add_run(run.text)
                    nr.bold = run.bold
                    nr.italic = run.italic
                    nr.underline = run.underline
                    try:
                        nr.font.name = run.font.name
                        nr.font.size = run.font.size
                        nr.font.color.rgb = run.font.color.rgb
                    except Exception:
                        pass
                delete_paragraph(nxt)
                merged = True
                break
        if not merged:
            break

def delete_table(table):
    """彻底删除一个表格（底层 XML 操作）"""
    tbl = table._tbl
    tbl.getparent().remove(tbl)

def merge_split_tables(doc: Document):
    """
    合并跨页被拆成两个表格的情况：
    如果两个相邻表格的表头完全一致，就把后面那个的数据行接到前面。
    """
    i = 0
    while i < len(doc.tables) - 1:
        t1, t2 = doc.tables[i], doc.tables[i+1]
        if len(t1.columns) == len(t2.columns):
            hdr1 = [c.text.strip() for c in t1.rows[0].cells]
            hdr2 = [c.text.strip() for c in t2.rows[0].cells]
            if hdr1 == hdr2:
                for row in t2.rows[1:]:
                    new_row = t1.add_row()
                    for idx, cell in enumerate(row.cells):
                        new_row.cells[idx].text = cell.text
                delete_table(t2)
                continue
        i += 1

def process_docx(doc: Document):
    """对中间生成的 docx 做二次清洗：段落合并、文本清洗、表格拼接"""
    merge_broken_paragraphs(doc)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = clean_text(run.text.replace('\n', ''))
    merge_split_tables(doc)

def pdf2docx(pdf_path: str, docx_path: str):
    """
    效果：能够提取文字、表格和图像，但无法保持原排版格式
    
    主函数：
    1) 用 pdf2docx 转为中间 docx
    2) 用 python-docx 二次清洗段落和表格
    3) 输出最终文件
    """
    base, _ = os.path.splitext(docx_path)
    tmp_docx = base + '_tmp.docx'

    # 第一步：布局级转换
    cv = Converter(pdf_path)
    cv.convert(tmp_docx, start=0, end=None)
    cv.close()

    # 第二步：二次清洗
    doc = Document(tmp_docx)
    process_docx(doc)
    doc.save(docx_path)

    os.remove(tmp_docx)
    print("✅ 已生成：", docx_path)
    
    return
    
    
if __name__ =="__main__":
    source="校对稿2.pdf"
    target="校对稿2x.pdf"
    
    pdf2docx(source, target)
    
#==============================================================================
#==============================================================================
#==============================================================================
