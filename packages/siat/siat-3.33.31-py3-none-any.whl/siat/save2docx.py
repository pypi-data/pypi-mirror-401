# -*- coding: utf-8 -*-
"""
本模块功能：转换ipynb文件为docx，带有的目录，代码行加边框，图像适配页宽。
注意：
    需要安装pandoc并将其路径加入操作系统的PATH。
    可在Anaconda Prompt或macOS Terminal下输入pandoc尝试，若未加入PATH则提示找不到。
尚存问题：
    1. 标题行未居中，且重复生成；
    2. 目录页码不准确，需要手动更新；
    3. 若docx文件已打开出错。
所属工具包：证券投资分析工具SIAT 
SIAT：Security Investment Analysis Tool
创建日期：2025年7月8日
最新修订日期：2025年7月8日
作者：王德宏 (WANG Dehong, Peter)
作者单位：北京外国语大学国际商学院
作者邮件：wdehong2000@163.com
版权所有：王德宏
用途限制：仅限研究与教学使用。
特别声明：作者不对使用本工具进行证券投资导致的任何损益负责！
"""

#==============================================================================
#关闭所有警告
import warnings; warnings.filterwarnings('ignore')

import nbformat
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import markdown as md
import base64
from io import BytesIO
import os
import re

def convert_ipynb_to_docx(ipynb_path):
    """
    将 .ipynb 转为 A3 大小 .docx，修复：
      1. 首页 TOC 可见（去掉分页）
      2. HTML 节点的 font-size, color, text-align, bold/italic/underline 正确继承与渲染
    """
    def parse_style(style_str):
        d = {}
        for part in style_str.split(';'):
            if ':' in part:
                k,v = part.split(':',1)
                d[k.strip()] = v.strip()
        return d

    def apply_style_to_run(run, style):
        """根据 style dict（font-size, color, font-weight, font-style, text-decoration）应用到 run"""
        # 字号
        if 'font-size' in style:
            sz = style['font-size']
            if sz.endswith('px'): val = float(sz[:-2]) * 0.75
            elif sz.endswith('pt'): val = float(sz[:-2])
            else: val = None
            if val: run.font.size = Pt(val)
        # 颜色
        if 'color' in style:
            c = style['color'].lstrip('#')
            run.font.color.rgb = RGBColor(*bytes.fromhex(c))
        # 加粗
        fw = style.get('font-weight','').lower()
        if fw in ('bold','700','800','900'): run.font.bold = True
        # 斜体
        if style.get('font-style','').lower() == 'italic': run.font.italic = True
        # 下划线
        if 'underline' in style.get('text-decoration',''): run.font.underline = True

    def render_inline_html(para, node, parent_style=None):
        """
        递归渲染 node，父级 style 先行应用，再叠加子元素 style
        """
        base_style = parent_style or {}
        for ch in node.children:
            if isinstance(ch, str):
                r = para.add_run(ch)
                apply_style_to_run(r, base_style)
            elif ch.name == 'br':
                para.add_run().add_break()
            elif ch.name in ('b','strong','i','em','u','span','font'):
                # 取出全部样式：父级 + 本级
                own_style = parse_style(ch.get('style',''))
                combined = dict(base_style)
                combined.update(own_style)
                text = ch.get_text()
                r = para.add_run(text)
                # 粗斜下从标签名或 combined 决定
                if ch.name in ('b','strong'): r.font.bold = True
                if ch.name in ('i','em'):     r.font.italic = True
                if ch.name == 'u':            r.font.underline = True
                apply_style_to_run(r, combined)
            else:
                # 其他标签，传递父级 style 递归
                render_inline_html(para, ch, base_style)

    def set_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            node = OxmlElement(f'w:{edge}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), '4')
            node.set(qn('w:color'), 'auto')
            borders.append(node)
        tblPr.append(borders)

    def insert_toc(doc):
        """插入 TOC，占位。去掉分页，避免空白首页。"""
        p = doc.add_paragraph()
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
        p._p.clear_content()
        p._p.append(fld)
        # 不再 add_page_break()

    # —— 主流程 —— #
    nb = nbformat.read(ipynb_path, as_version=4)
    doc = Document()

    # A3 竖排
    sec = doc.sections[0]
    sec.page_width  = Inches(11.69)
    sec.page_height = Inches(16.54)
    avail_w = (sec.page_width - sec.left_margin - sec.right_margin) / Inches(1)

    insert_toc(doc)
    first_h1 = False

    for cell in nb.cells:
        if cell.cell_type == 'markdown':
            html = md.markdown(cell.source, extensions=['extra'])
            soup = BeautifulSoup(html, 'html.parser')

            for elem in soup.contents:
                if not getattr(elem, 'name', None):
                    txt = elem.strip()
                    if txt:
                        doc.add_paragraph(txt)
                    continue

                tag = elem.name.lower()
                m = re.match(r'h([1-6])', tag)
                if m:
                    lvl = int(m.group(1))
                    # H1 开始真正目录
                    if lvl == 1:
                        first_h1 = True
                        para = doc.add_heading(level=1)
                        # 带节点 style
                        parent_style = parse_style(elem.get('style',''))
                        render_inline_html(para, elem, parent_style)
                    else:
                        if not first_h1:
                            p0 = doc.add_paragraph()
                            r0 = p0.add_run(elem.get_text())
                            r0.font.bold = True
                            r0.font.size = Pt({2:20,3:18,4:16,5:14,6:12}[lvl])
                        else:
                            para = doc.add_heading(level=lvl)
                            parent_style = parse_style(elem.get('style',''))
                            render_inline_html(para, elem, parent_style)
                    continue

                if tag in ('p','div'):
                    p = doc.add_paragraph()
                    st = parse_style(elem.get('style',''))
                    if st.get('text-align','') == 'center':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    render_inline_html(p, elem, st)
                    continue

                # 其他
                txt = elem.get_text().strip()
                if txt:
                    doc.add_paragraph(txt)

        elif cell.cell_type == 'code':
            # 代码块
            t0 = doc.add_table(1,1)
            t0.style = 'Table Grid'
            t0.allow_autofit = True
            c0 = t0.cell(0,0).paragraphs[0]
            c0.paragraph_format.left_indent = Cm(0.5)
            for ln in cell.source.splitlines():
                r = c0.add_run(ln + '\n')
                r.font.name = 'Courier New'
                r.font.size = Pt(9)

            for out in cell.get('outputs', []):
                data = out.get('data', {})

                # 图片
                if data.get('image/png'):
                    b64 = data['image/png']
                    img = base64.b64decode(b64)
                    bio = BytesIO(img)
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p_img.add_run().add_picture(bio, width=Inches(avail_w))

                # HTML 输出表格
                elif data.get('text/html'):
                    soup = BeautifulSoup(data['text/html'], 'html.parser')
                    tbl_html = soup.find('table')

                    # 有表格
                    if tbl_html:
                        cap = tbl_html.find('caption')
                        if cap:
                            cpara = doc.add_paragraph(cap.get_text())
                            cpara.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        rows = tbl_html.find_all('tr')
                        nr, nc = len(rows), len(rows[0].find_all(['td','th']))
                        table = doc.add_table(nr, nc)
                        table.style = 'Table Grid'
                        table.allow_autofit = True
                        table.alignment = WD_TABLE_ALIGNMENT.LEFT

                        for i,row in enumerate(rows):
                            cells = row.find_all(['td','th'])
                            for j,cell_html in enumerate(cells):
                                cp = table.cell(i,j).paragraphs[0]
                                cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                # 子节点样式
                                par_style = parse_style(cell_html.get('style',''))
                                render_inline_html(cp, cell_html, par_style)
                                if i == 0:
                                    for run in cp.runs:
                                        run.font.bold = True

                        set_table_borders(table)

                        # 注解
                        for sib in tbl_html.find_next_siblings():
                            txt = sib.get_text().strip()
                            if txt:
                                doc.add_paragraph(txt)

                    # 无表格，退为纯 HTML 渲染
                    else:
                        p_html = doc.add_paragraph()
                        render_inline_html(p_html, soup)

                # 文本输出
                elif out.get('text'):
                    for ln in out['text'].splitlines():
                        doc.add_paragraph(ln)

    # 保存
    base, _ = os.path.splitext(ipynb_path)
    out = base + '.docx'
    doc.save(out)
    #print(f"已生成：{out}\n→ 打开 Word 后，右键 TOC 选择“更新域”即可显示目录")
    
    return out



#==============================================================================
import os
import sys
import psutil

def is_file_opened(file_path: str) -> bool:
    """
    检测文件是否被其他程序打开（跨平台）
    :param file_path: 文件路径
    :return: True-被占用, False-未占用或不存在
    """
    # 检查文件是否存在
    if not os.path.exists(file_path):
        return False
    
    abs_path = os.path.abspath(file_path)  # 转为绝对路径
    
    # 方法1：异常捕获法（快速检测）
    try:
        with open(abs_path, "a") as f:  # 追加模式（不破坏内容）
            pass
        return False  # 成功打开说明未被占用
    except (OSError, PermissionError):
        pass  # 继续尝试其他方法
    
    # 方法2：进程扫描法（精确检测）
    try:
        for proc in psutil.process_iter(['pid', 'name', 'open_files']):
            try:
                open_files = proc.info.get('open_files')
                if open_files and any(f.path == abs_path for f in open_files):
                    return True
            except (psutil.AccessDenied, psutil.NoSuchProcess):
                continue
    except NameError:  # psutil未安装
        pass
    
    # 方法3：文件锁试探法（最终回退）
    try:
        if sys.platform == 'win32':
            import msvcrt
            with open(abs_path, "a") as f:
                msvcrt.locking(f.fileno(), msvcrt.LK_NBLCK, 1)  # 非阻塞锁
        else:
            import fcntl
            with open(abs_path, "a") as f:
                fcntl.flock(f, fcntl.LOCK_EX | fcntl.LOCK_NB)  # 非阻塞独占锁
        return False
    except (OSError, BlockingIOError, ImportError):
        return True  # 所有检测均失败视为占用
    return False

#==============================================================================

def ipynb2docx(ipynb_path):
    """
    将 .ipynb 转为 .docx，特性：
      1. Markdown 首行做文档标题
      2. 在第 2 行插入全文 TOC（1–9 级）
      3. 所有标题左对齐，保留原字号
      4. 仅为“代码段”加边框，不影响输出
      5. 表格均分列宽并居中
      6. 图像放大至可用页宽并居中
      7. 若目标文件已打开，捕获并提示“请先关闭文件”
    """
    base, _ = os.path.splitext(ipynb_path)
    docx_path = base + ".docx"
    
    # 检测docx文件是否已被打开
    if is_file_opened(docx_path):
        print(f"Warning: {docx_path} is occupied by other app")
        print(f"Solution: please close it and try again")
        return
    
    print(f"Converting to docx ...")
    
    #result = convert_ipynb_to_docx(ipynb_path, docx_path=None, page_size=page_size)
    result = convert_ipynb_to_docx(ipynb_path)
    dirpath, filename = os.path.split(result)

    print(f"✅ {filename} is created")
    print(f"✅ In {dirpath}")
    print(f"Note: may need addtional adjustments for better formatting")
    
    return
    

