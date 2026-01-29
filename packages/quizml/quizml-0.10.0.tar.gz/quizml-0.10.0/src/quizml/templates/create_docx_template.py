import os

from docx import Document


def create_template(filename):
    doc = Document()

    # Header
    doc.add_heading("{{ header.title }}", 0)

    p = doc.add_paragraph()
    p.add_run("Duration: ").bold = True
    p.add_run("{{ header.duration }} minutes")

    p = doc.add_paragraph()
    p.add_run("Instructions: ").bold = True
    p.add_run("{{ header.instructions }}")

    doc.add_heading("Questions", level=1)

    # Loop start
    doc.add_paragraph("{% for q in questions %}")

    # Question Title/Number
    doc.add_heading("Question {{ loop.index }}", level=2)

    # Question Text
    doc.add_paragraph("{{ q.question }}")

    # Choices (conditional)
    doc.add_paragraph("{% if q.choices %}")

    # Loop over choices
    doc.add_paragraph("{% for c in q.choices %}")

    # Handle list of dicts (standard quizml format: [{x: 'Wrong'}, {o: 'Right'}])
    # or just simple strings if normalized.
    # We'll use a jinja macro or just simple logic to print values.
    # Since c is a dict like {'x': 'Text'}, we can iterate its values.

    # Simple bullet point
    p = doc.add_paragraph("", style="List Bullet")
    p.add_run("{{ c.values() | list | first }}")

    doc.add_paragraph("{% endfor %}")  # end choices loop
    doc.add_paragraph("{% endif %}")  # end choices check

    # Loop end
    doc.add_paragraph("{% endfor %}")

    # Save
    doc.save(filename)
    print(f"Created {filename}")


if __name__ == "__main__":
    output_path = os.path.join(os.path.dirname(__file__), "prototype.docx")
    create_template(output_path)
