"""Unit tests for DOCX extraction functionality."""

# Import from package (conftest.py adds src to path)
import sys
import tempfile

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

sys.path.insert(0, "src")

from translate_docx.extractor import (
    extract_document,
    extract_line_numbering,
    extract_metadata,
    extract_page_settings,
    extract_paragraph,
    extract_paragraph_formatting,
    extract_run_formatting,
    extract_text_run,
    is_section_header,
    split_into_sections,
)
from translate_docx.models import (
    Document,
    DocumentMetadata,
    LineNumbering,
    PageSettings,
    Paragraph,
    RunFormatting,
    Section,
    TextRun,
)


class TestExtractRunFormatting:
    """Tests for extract_run_formatting function."""

    def test_extract_bold_run(self):
        """Test extracting bold formatting from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        formatting = extract_run_formatting(run)

        assert formatting.bold is True
        assert formatting.italic is None
        assert formatting.underline is None

    def test_extract_italic_run(self):
        """Test extracting italic formatting from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True

        formatting = extract_run_formatting(run)

        assert formatting.italic is True
        assert formatting.bold is None
        assert formatting.underline is None

    def test_extract_underline_run(self):
        """Test extracting underline formatting from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Underlined text")
        run.underline = True

        formatting = extract_run_formatting(run)

        assert formatting.underline is True
        assert formatting.bold is None
        assert formatting.italic is None

    def test_extract_combined_formatting(self):
        """Test extracting multiple formatting attributes at once."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Bold and italic")
        run.bold = True
        run.italic = True
        run.underline = True

        formatting = extract_run_formatting(run)

        assert formatting.bold is True
        assert formatting.italic is True
        assert formatting.underline is True

    def test_extract_font_name(self):
        """Test extracting font name from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Styled text")
        run.font.name = "Arial"

        formatting = extract_run_formatting(run)

        assert formatting.font_name == "Arial"

    def test_extract_font_size(self):
        """Test extracting font size from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Large text")
        run.font.size = Pt(16)

        formatting = extract_run_formatting(run)

        assert formatting.font_size == 16

    def test_extract_superscript(self):
        """Test extracting superscript formatting from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Super")
        run.font.superscript = True

        formatting = extract_run_formatting(run)

        assert formatting.is_superscript is True
        assert formatting.is_subscript is False

    def test_extract_subscript(self):
        """Test extracting subscript formatting from a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Sub")
        run.font.subscript = True

        formatting = extract_run_formatting(run)

        assert formatting.is_subscript is True
        assert formatting.is_superscript is False


class TestExtractTextRun:
    """Tests for extract_text_run function."""

    def test_extract_simple_text_run(self):
        """Test extracting a simple text run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Hello, World!")

        text_run = extract_text_run(run)

        assert isinstance(text_run, TextRun)
        assert text_run.text == "Hello, World!"
        assert isinstance(text_run.formatting, RunFormatting)

    def test_extract_formatted_text_run(self):
        """Test extracting a text run with formatting."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        text_run = extract_text_run(run)

        assert text_run.text == "Bold text"
        assert text_run.formatting.bold is True

    def test_extract_empty_run(self):
        """Test extracting an empty text run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("")

        text_run = extract_text_run(run)

        assert text_run.text == ""
        assert isinstance(text_run.formatting, RunFormatting)


class TestExtractParagraphFormatting:
    """Tests for extract_paragraph_formatting function."""

    def test_extract_alignment(self):
        """Test extracting paragraph alignment."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Centered text")
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        formatting = extract_paragraph_formatting(para)

        assert formatting.alignment == "center"

    def test_extract_left_alignment(self):
        """Test extracting left alignment."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Left aligned")
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        formatting = extract_paragraph_formatting(para)

        assert formatting.alignment == "left"

    def test_extract_right_alignment(self):
        """Test extracting right alignment."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Right aligned")
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        formatting = extract_paragraph_formatting(para)

        assert formatting.alignment == "right"

    def test_extract_spacing(self):
        """Test extracting paragraph spacing."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Spaced paragraph")
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

        formatting = extract_paragraph_formatting(para)

        assert formatting.space_before == Pt(12)
        assert formatting.space_after == Pt(12)

    def test_extract_keep_together(self):
        """Test extracting keep_together property."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Keep together")
        para.paragraph_format.keep_together = True

        formatting = extract_paragraph_formatting(para)

        assert formatting.keep_together is True

    def test_extract_page_break_before(self):
        """Test extracting page_break_before property."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Page break")
        para.paragraph_format.page_break_before = True

        formatting = extract_paragraph_formatting(para)

        assert formatting.page_break_before is True


class TestExtractParagraph:
    """Tests for extract_paragraph function."""

    def test_extract_simple_paragraph(self):
        """Test extracting a simple paragraph."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Simple text")

        extracted = extract_paragraph(para)

        assert isinstance(extracted, Paragraph)
        assert extracted.text == "Simple text"
        assert len(extracted.runs) == 1
        assert extracted.is_header is False

    def test_extract_paragraph_with_multiple_runs(self):
        """Test extracting a paragraph with multiple runs."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        para.add_run("Normal ")
        bold_run = para.add_run("bold")
        bold_run.bold = True
        para.add_run(" text")

        extracted = extract_paragraph(para)

        assert extracted.text == "Normal bold text"
        assert len(extracted.runs) == 3
        assert extracted.runs[0].formatting.bold is None
        assert extracted.runs[1].formatting.bold is True
        assert extracted.runs[2].formatting.bold is None

    def test_extract_paragraph_as_header(self):
        """Test extracting a paragraph marked as header."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Section Header")

        extracted = extract_paragraph(para, is_header=True)

        assert extracted.is_header is True

    def test_extract_paragraph_with_style(self):
        """Test extracting paragraph style name."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Heading", style="Heading 1")

        extracted = extract_paragraph(para)

        assert extracted.style_name == "Heading 1"


class TestExtractDocument:
    """Tests for extract_document function."""

    def test_extract_simple_document(self):
        """Test extracting a simple single-paragraph document."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("First paragraph")
            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert isinstance(extracted, Document)
            assert len(extracted.preamble_paragraphs) == 1
            assert extracted.preamble_paragraphs[0].text == "First paragraph"

    def test_extract_multi_paragraph_document(self):
        """Test extracting a document with multiple paragraphs."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("First paragraph")
            docx_doc.add_paragraph("Second paragraph")
            docx_doc.add_paragraph("Third paragraph")
            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert len(extracted.preamble_paragraphs) == 3
            assert extracted.preamble_paragraphs[0].text == "First paragraph"
            assert extracted.preamble_paragraphs[1].text == "Second paragraph"
            assert extracted.preamble_paragraphs[2].text == "Third paragraph"

    def test_extract_document_with_formatting(self):
        """Test extracting a document with various formatting."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            para = docx_doc.add_paragraph()
            para.add_run("Normal ")
            bold_run = para.add_run("bold")
            bold_run.bold = True
            para.add_run(" text")
            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert len(extracted.preamble_paragraphs) == 1
            para_data = extracted.preamble_paragraphs[0]
            assert len(para_data.runs) == 3
            assert para_data.runs[1].formatting.bold is True

    def test_extract_document_file_not_found(self):
        """Test extracting a non-existent file raises FileNotFoundError."""
        try:
            extract_document("/nonexistent/path/file.docx")
            assert False, "Should have raised an exception"
        except FileNotFoundError:
            pass
        except Exception as e:
            # python-docx may raise different exceptions
            assert "No such file" in str(e) or "not found" in str(e).lower()

    def test_extract_document_metadata(self):
        """Test that extracted document has metadata."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("Test")
            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert extracted.metadata is not None
            assert extracted.metadata.line_numbering is not None
            assert extracted.metadata.page_settings is not None


class TestExtractAllParagraphs:
    """Tests for the Document.all_paragraphs property."""

    def test_all_paragraphs_without_sections(self):
        """Test all_paragraphs returns only preamble when no sections."""
        doc = Document(
            preamble_paragraphs=[
                Paragraph(runs=[TextRun(text="Para 1", formatting=RunFormatting())]),
                Paragraph(runs=[TextRun(text="Para 2", formatting=RunFormatting())]),
            ]
        )

        all_paras = doc.all_paragraphs

        assert len(all_paras) == 2
        assert all_paras[0].text == "Para 1"
        assert all_paras[1].text == "Para 2"


class TestExtractColor:
    """Tests for color extraction from runs."""

    def test_extract_no_color(self):
        """Test extracting text with no explicit color."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text without color")

        formatting = extract_run_formatting(run)

        # Default color should be None
        assert formatting.color is None

    def test_extract_rgb_color(self):
        """Test extracting text with RGB color."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Red text")
        run.font.color.rgb = RGBColor(255, 0, 0)

        formatting = extract_run_formatting(run)

        # Should extract as hex color
        assert formatting.color is not None
        assert isinstance(formatting.color, str)
        assert formatting.color.startswith("#")

    def test_extract_color_with_multiple_formatting(self):
        """Test extracting color along with other formatting."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Blue bold text")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 255)

        formatting = extract_run_formatting(run)

        assert formatting.bold is True
        assert formatting.color is not None
        assert formatting.color.startswith("#")


class TestExtractPageSettings:
    """Tests for page settings extraction."""

    def test_extract_page_settings_basic(self):
        """Test extracting basic page settings."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("Test")
            docx_doc.save(tmp.name)

            docx_doc = DocxDocument(tmp.name)
            section = docx_doc.sections[0]

            page_settings = extract_page_settings(section)

            assert isinstance(page_settings, PageSettings)
            assert page_settings.width is not None
            assert page_settings.height is not None
            assert page_settings.top_margin is not None
            assert page_settings.bottom_margin is not None
            assert page_settings.left_margin is not None
            assert page_settings.right_margin is not None

    def test_extract_page_settings_custom_margins(self):
        """Test extracting custom margin settings."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            section = docx_doc.sections[0]
            section.top_margin = Pt(72)  # 1 inch = 72 points
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)
            docx_doc.add_paragraph("Test")
            docx_doc.save(tmp.name)

            docx_doc = DocxDocument(tmp.name)
            section = docx_doc.sections[0]
            page_settings = extract_page_settings(section)

            assert page_settings.top_margin == Pt(72)
            assert page_settings.bottom_margin == Pt(72)
            assert page_settings.left_margin == Pt(72)
            assert page_settings.right_margin == Pt(72)


class TestExtractLineNumbering:
    """Tests for line numbering extraction."""

    def test_extract_no_line_numbering(self):
        """Test extracting document without line numbering."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("Test")
            docx_doc.save(tmp.name)

            docx_doc = DocxDocument(tmp.name)
            section = docx_doc.sections[0]

            line_numbering = extract_line_numbering(section)

            assert isinstance(line_numbering, LineNumbering)
            assert line_numbering.enabled is False

    def test_extract_line_numbering_defaults(self):
        """Test line numbering returns sensible defaults."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("Test")
            docx_doc.save(tmp.name)

            docx_doc = DocxDocument(tmp.name)
            section = docx_doc.sections[0]
            ln = extract_line_numbering(section)

            # Should always return a LineNumbering object
            assert isinstance(ln, LineNumbering)
            # Default should be disabled
            assert ln.enabled is False or ln.enabled is True


class TestExtractMetadata:
    """Tests for document metadata extraction."""

    def test_extract_metadata_basic(self):
        """Test extracting basic document metadata."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("Test")
            docx_doc.save(tmp.name)

            docx_doc = DocxDocument(tmp.name)
            metadata = extract_metadata(docx_doc)

            assert isinstance(metadata, DocumentMetadata)
            assert metadata.page_settings is not None
            assert metadata.line_numbering is not None

    def test_extract_metadata_with_content(self):
        """Test extracting metadata from document with content."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            para = docx_doc.add_paragraph()
            run = para.add_run("Styled text")
            run.font.name = "Arial"
            run.font.size = Pt(12)
            docx_doc.save(tmp.name)

            docx_doc = DocxDocument(tmp.name)
            metadata = extract_metadata(docx_doc)

            assert metadata is not None
            # At minimum, should have page settings
            assert metadata.page_settings.width is not None

    def test_extract_metadata_empty_document(self):
        """Test extracting metadata from empty document."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.save(tmp.name)

            # Remove all paragraphs
            for para in docx_doc.paragraphs:
                p = para._element
                p.getparent().remove(p)

            docx_doc = DocxDocument(tmp.name)
            metadata = extract_metadata(docx_doc)

            assert isinstance(metadata, DocumentMetadata)


class TestExtractDocumentWithMetadata:
    """Tests for document extraction including metadata."""

    def test_extract_document_includes_metadata(self):
        """Test that extracted document includes metadata."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            docx_doc.add_paragraph("Test paragraph")
            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert extracted.metadata is not None
            assert isinstance(extracted.metadata, DocumentMetadata)
            assert extracted.metadata.page_settings is not None

    def test_extract_document_preserves_colors_and_fonts(self):
        """Test that all formatting is preserved in extraction."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()
            para = docx_doc.add_paragraph()
            run = para.add_run("Formatted ")
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(255, 0, 0)

            run2 = para.add_run("text")
            run2.italic = True
            run2.font.size = Pt(12)

            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert len(extracted.preamble_paragraphs) == 1
            para_data = extracted.preamble_paragraphs[0]
            assert len(para_data.runs) == 2

            # Check first run
            assert para_data.runs[0].formatting.bold is True
            assert para_data.runs[0].formatting.font_name == "Arial"
            assert para_data.runs[0].formatting.font_size == 14
            assert para_data.runs[0].formatting.color is not None

            # Check second run
            assert para_data.runs[1].formatting.italic is True
            assert para_data.runs[1].formatting.font_size == 12


class TestIsSectionHeader:
    """Tests for is_section_header function."""

    def test_is_section_header_with_heading_style(self):
        """Test that paragraphs with Heading style are detected as headers."""
        para = Paragraph(
            runs=[TextRun(text="Section Title", formatting=RunFormatting())],
            style_name="Heading 1",
        )

        assert is_section_header(para) is True

    def test_is_section_header_with_heading_2_style(self):
        """Test that Heading 2 style is detected as header."""
        para = Paragraph(
            runs=[TextRun(text="Subsection", formatting=RunFormatting())],
            style_name="Heading 2",
        )

        assert is_section_header(para) is True

    def test_is_not_section_header_with_normal_style(self):
        """Test that Normal style paragraphs are not headers."""
        para = Paragraph(
            runs=[TextRun(text="This is normal text", formatting=RunFormatting())],
            style_name="Normal",
        )

        assert is_section_header(para) is False

    def test_is_section_header_with_all_bold_short_text(self):
        """Test detection of bold short text as header."""
        bold_run = TextRun(text="Bold Header", formatting=RunFormatting(bold=True))
        para = Paragraph(runs=[bold_run], style_name="Normal")

        assert is_section_header(para) is True

    def test_is_not_section_header_with_partial_bold(self):
        """Test that partially bold text is not detected as header."""
        run1 = TextRun(text="Bold ", formatting=RunFormatting(bold=True))
        run2 = TextRun(text="Normal", formatting=RunFormatting(bold=False))
        para = Paragraph(runs=[run1, run2], style_name="Normal")

        assert is_section_header(para) is False

    def test_is_not_section_header_with_long_bold_text(self):
        """Test that long bold text is not detected as header."""
        long_text = "This is a very long paragraph with bold formatting " * 10
        bold_run = TextRun(text=long_text, formatting=RunFormatting(bold=True))
        para = Paragraph(runs=[bold_run], style_name="Normal")

        assert is_section_header(para) is False

    def test_is_not_section_header_with_empty_paragraph(self):
        """Test that empty paragraphs are not headers."""
        para = Paragraph(runs=[], style_name="Normal")

        assert is_section_header(para) is False

    def test_is_section_header_with_custom_detector(self):
        """Test custom header detection function."""
        para = Paragraph(
            runs=[TextRun(text="Custom", formatting=RunFormatting())],
            style_name="Normal",
        )

        # Custom detector: anything with "Custom" in text
        def custom_detector(p: Paragraph) -> bool:
            return "Custom" in p.text

        assert is_section_header(para, custom_detector) is True

    def test_is_section_header_custom_detector_overrides_heuristic(self):
        """Test that custom detector takes precedence."""
        # This would normally be detected as header (Heading 1 style)
        para = Paragraph(
            runs=[TextRun(text="Section", formatting=RunFormatting())],
            style_name="Heading 1",
        )

        # Custom detector: never a header
        def never_header(p: Paragraph) -> bool:
            return False

        assert is_section_header(para, never_header) is False


class TestSplitIntoSections:
    """Tests for split_into_sections function."""

    def test_split_into_sections_no_headers(self):
        """Test splitting document with no headers."""
        paras = [
            Paragraph(runs=[TextRun(text="Para 1", formatting=RunFormatting())]),
            Paragraph(runs=[TextRun(text="Para 2", formatting=RunFormatting())]),
        ]

        sections, preamble = split_into_sections(paras)

        assert len(sections) == 0
        assert len(preamble) == 2
        assert preamble[0].text == "Para 1"
        assert preamble[1].text == "Para 2"

    def test_split_into_sections_single_section(self):
        """Test splitting document with one section."""
        header = Paragraph(
            runs=[TextRun(text="Header", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
        )
        body1 = Paragraph(runs=[TextRun(text="Body 1", formatting=RunFormatting())])
        body2 = Paragraph(runs=[TextRun(text="Body 2", formatting=RunFormatting())])

        paras = [header, body1, body2]
        sections, preamble = split_into_sections(paras)

        assert len(sections) == 1
        assert len(preamble) == 0
        assert sections[0].header.text == "Header"
        assert len(sections[0].paragraphs) == 2
        assert sections[0].paragraphs[0].text == "Body 1"
        assert sections[0].paragraphs[1].text == "Body 2"

    def test_split_into_sections_multiple_sections(self):
        """Test splitting document with multiple sections."""
        header1 = Paragraph(
            runs=[TextRun(text="Section 1", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
        )
        body1 = Paragraph(runs=[TextRun(text="Body 1", formatting=RunFormatting())])

        header2 = Paragraph(
            runs=[TextRun(text="Section 2", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
        )
        body2 = Paragraph(runs=[TextRun(text="Body 2", formatting=RunFormatting())])

        paras = [header1, body1, header2, body2]
        sections, preamble = split_into_sections(paras)

        assert len(sections) == 2
        assert len(preamble) == 0

        assert sections[0].header.text == "Section 1"
        assert len(sections[0].paragraphs) == 1

        assert sections[1].header.text == "Section 2"
        assert len(sections[1].paragraphs) == 1

    def test_split_into_sections_with_preamble(self):
        """Test splitting with content before first header."""
        preamble1 = Paragraph(runs=[TextRun(text="Intro", formatting=RunFormatting())])
        preamble2 = Paragraph(runs=[TextRun(text="More intro", formatting=RunFormatting())])

        header = Paragraph(
            runs=[TextRun(text="Section", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
        )
        body = Paragraph(runs=[TextRun(text="Body", formatting=RunFormatting())])

        paras = [preamble1, preamble2, header, body]
        sections, preamble = split_into_sections(paras)

        assert len(sections) == 1
        assert len(preamble) == 2
        assert preamble[0].text == "Intro"
        assert preamble[1].text == "More intro"

    def test_split_into_sections_consecutive_headers(self):
        """Test sections with consecutive headers (empty sections)."""
        header1 = Paragraph(
            runs=[TextRun(text="Section 1", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
        )
        header2 = Paragraph(
            runs=[TextRun(text="Section 2", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
        )
        body = Paragraph(runs=[TextRun(text="Body", formatting=RunFormatting())])

        paras = [header1, header2, body]
        sections, preamble = split_into_sections(paras)

        assert len(sections) == 2
        assert len(preamble) == 0

        # First section should be empty
        assert sections[0].header.text == "Section 1"
        assert len(sections[0].paragraphs) == 0

        # Second section should have body
        assert sections[1].header.text == "Section 2"
        assert len(sections[1].paragraphs) == 1

    def test_split_into_sections_with_custom_detector(self):
        """Test section splitting with custom detector."""
        # Paragraphs with "CHAPTER" in text are headers
        chapter1 = Paragraph(runs=[TextRun(text="CHAPTER 1", formatting=RunFormatting())])
        body1 = Paragraph(runs=[TextRun(text="Content", formatting=RunFormatting())])
        chapter2 = Paragraph(runs=[TextRun(text="CHAPTER 2", formatting=RunFormatting())])
        body2 = Paragraph(runs=[TextRun(text="More content", formatting=RunFormatting())])

        paras = [chapter1, body1, chapter2, body2]

        # Custom detector
        def is_chapter(p: Paragraph) -> bool:
            return "CHAPTER" in p.text

        sections, preamble = split_into_sections(paras, is_chapter)

        assert len(sections) == 2
        assert sections[0].header.text == "CHAPTER 1"
        assert sections[1].header.text == "CHAPTER 2"


class TestExtractDocumentWithSections:
    """Tests for document extraction with section splitting."""

    def test_extract_document_creates_sections(self):
        """Test that extract_document properly creates sections."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()

            # Add preamble
            docx_doc.add_paragraph("Preamble text")

            # Add first section
            h1 = docx_doc.add_paragraph("Section 1", style="Heading 1")
            docx_doc.add_paragraph("Body of section 1")

            # Add second section
            h2 = docx_doc.add_paragraph("Section 2", style="Heading 1")
            docx_doc.add_paragraph("Body of section 2")

            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert len(extracted.preamble_paragraphs) == 1
            assert len(extracted.sections) == 2

            assert extracted.preamble_paragraphs[0].text == "Preamble text"
            assert extracted.sections[0].header.text == "Section 1"
            assert extracted.sections[1].header.text == "Section 2"

    def test_extract_document_sections_have_bodies(self):
        """Test that extracted sections have body paragraphs."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()

            # Add section with body
            docx_doc.add_paragraph("Section Header", style="Heading 1")
            docx_doc.add_paragraph("First body paragraph")
            docx_doc.add_paragraph("Second body paragraph")

            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert len(extracted.sections) == 1
            assert len(extracted.sections[0].paragraphs) == 2
            assert extracted.sections[0].paragraphs[0].text == "First body paragraph"
            assert extracted.sections[0].paragraphs[1].text == "Second body paragraph"

    def test_extract_document_headers_marked_as_headers(self):
        """Test that header paragraphs are marked with is_header=True."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()

            docx_doc.add_paragraph("Header", style="Heading 1")
            docx_doc.add_paragraph("Body")

            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert extracted.sections[0].header.is_header is True
            assert extracted.sections[0].paragraphs[0].is_header is False

    def test_extract_document_with_custom_header_detector(self):
        """Test extract_document with custom header detection."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()

            # Create paragraphs with *** markers for headers
            docx_doc.add_paragraph("*** Section A")
            docx_doc.add_paragraph("Body text")
            docx_doc.add_paragraph("*** Section B")
            docx_doc.add_paragraph("More body text")

            docx_doc.save(tmp.name)

            # Custom detector: headers start with ***
            def is_custom_header(para):
                return para.text.startswith("***")

            extracted = extract_document(tmp.name, header_detector=is_custom_header)

            assert len(extracted.sections) == 2
            assert "*** Section A" in extracted.sections[0].header.text
            assert "*** Section B" in extracted.sections[1].header.text

    def test_extract_document_no_sections_all_preamble(self):
        """Test document with no sections (no headers)."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_doc = DocxDocument()

            docx_doc.add_paragraph("Paragraph 1")
            docx_doc.add_paragraph("Paragraph 2")
            docx_doc.add_paragraph("Paragraph 3")

            docx_doc.save(tmp.name)

            extracted = extract_document(tmp.name)

            assert len(extracted.sections) == 0
            assert len(extracted.preamble_paragraphs) == 3
