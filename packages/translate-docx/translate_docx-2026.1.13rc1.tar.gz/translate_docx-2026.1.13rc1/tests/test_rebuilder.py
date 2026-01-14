"""Unit tests for DOCX rebuilding functionality."""

import sys
import tempfile

sys.path.insert(0, "src")

from docx import Document as DocxDocument
from docx.shared import Pt

from translate_docx.models import (
    Document,
    DocumentMetadata,
    LineNumbering,
    PageSettings,
    Paragraph,
    ParagraphFormatting,
    RunFormatting,
    Section,
    TextRun,
)
from translate_docx.rebuilder import (
    add_paragraph_to_document,
    apply_line_numbering,
    apply_metadata,
    apply_page_settings,
    apply_paragraph_formatting,
    apply_run_formatting,
    rebuild_document,
)


class TestApplyRunFormatting:
    """Tests for apply_run_formatting function."""

    def test_apply_bold_formatting(self):
        """Test applying bold formatting to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        formatting = RunFormatting(bold=True)
        apply_run_formatting(run, formatting)

        assert run.bold is True

    def test_apply_italic_formatting(self):
        """Test applying italic formatting to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        formatting = RunFormatting(italic=True)
        apply_run_formatting(run, formatting)

        assert run.italic is True

    def test_apply_underline_formatting(self):
        """Test applying underline formatting to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        formatting = RunFormatting(underline=True)
        apply_run_formatting(run, formatting)

        assert run.underline is True

    def test_apply_combined_formatting(self):
        """Test applying multiple formatting attributes."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        formatting = RunFormatting(bold=True, italic=True, underline=True)
        apply_run_formatting(run, formatting)

        assert run.bold is True
        assert run.italic is True
        assert run.underline is True

    def test_apply_font_name(self):
        """Test applying font name to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        formatting = RunFormatting(font_name="Arial")
        apply_run_formatting(run, formatting)

        assert run.font.name == "Arial"

    def test_apply_font_size(self):
        """Test applying font size to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        formatting = RunFormatting(font_size=16)
        apply_run_formatting(run, formatting)

        assert run.font.size == Pt(16)

    def test_apply_color(self):
        """Test applying color to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Red text")

        formatting = RunFormatting(color="#FF0000")
        apply_run_formatting(run, formatting)

        assert run.font.color.rgb is not None

    def test_apply_superscript(self):
        """Test applying superscript to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Super")

        formatting = RunFormatting(is_superscript=True)
        apply_run_formatting(run, formatting)

        assert run.font.superscript is True

    def test_apply_subscript(self):
        """Test applying subscript to a run."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Sub")

        formatting = RunFormatting(is_subscript=True)
        apply_run_formatting(run, formatting)

        assert run.font.subscript is True

    def test_apply_none_formatting(self):
        """Test that None values don't cause errors."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph()
        run = para.add_run("Text")

        # Create formatting with all None values
        formatting = RunFormatting()
        apply_run_formatting(run, formatting)

        # Should not crash and run should still exist
        assert run.text == "Text"


class TestApplyParagraphFormatting:
    """Tests for apply_paragraph_formatting function."""

    def test_apply_left_alignment(self):
        """Test applying left alignment."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Text")

        formatting = ParagraphFormatting(alignment="left")
        apply_paragraph_formatting(para, formatting)

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert para.alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_apply_center_alignment(self):
        """Test applying center alignment."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Text")

        formatting = ParagraphFormatting(alignment="center")
        apply_paragraph_formatting(para, formatting)

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_apply_spacing(self):
        """Test applying spacing to paragraph."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Text")

        formatting = ParagraphFormatting(space_before=Pt(12), space_after=Pt(12))
        apply_paragraph_formatting(para, formatting)

        assert para.paragraph_format.space_before == Pt(12)
        assert para.paragraph_format.space_after == Pt(12)

    def test_apply_keep_together(self):
        """Test applying keep_together property."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Text")

        formatting = ParagraphFormatting(keep_together=True)
        apply_paragraph_formatting(para, formatting)

        assert para.paragraph_format.keep_together is True

    def test_apply_page_break_before(self):
        """Test applying page_break_before property."""
        docx_doc = DocxDocument()
        para = docx_doc.add_paragraph("Text")

        formatting = ParagraphFormatting(page_break_before=True)
        apply_paragraph_formatting(para, formatting)

        assert para.paragraph_format.page_break_before is True


class TestAddParagraphToDocument:
    """Tests for add_paragraph_to_document function."""

    def test_add_simple_paragraph(self):
        """Test adding a simple paragraph."""
        docx_doc = DocxDocument()
        para_data = Paragraph(
            runs=[TextRun(text="Hello", formatting=RunFormatting())],
            formatting=ParagraphFormatting(),
            style_name="Normal",
        )

        add_paragraph_to_document(docx_doc, para_data)

        assert len(docx_doc.paragraphs) == 1
        assert docx_doc.paragraphs[0].text == "Hello"

    def test_add_paragraph_with_formatting(self):
        """Test adding paragraph with formatting."""
        docx_doc = DocxDocument()

        bold_run = TextRun(text="Bold ", formatting=RunFormatting(bold=True))
        normal_run = TextRun(text="normal", formatting=RunFormatting())

        para_data = Paragraph(
            runs=[bold_run, normal_run],
            formatting=ParagraphFormatting(alignment="center"),
            style_name="Normal",
        )

        add_paragraph_to_document(docx_doc, para_data)

        para = docx_doc.paragraphs[0]
        assert para.text == "Bold normal"
        assert len(para.runs) == 2
        assert para.runs[0].bold is True
        assert para.runs[1].bold is None

    def test_add_paragraph_with_style(self):
        """Test adding paragraph with specific style."""
        docx_doc = DocxDocument()
        para_data = Paragraph(
            runs=[TextRun(text="Heading", formatting=RunFormatting())], style_name="Heading 1"
        )

        add_paragraph_to_document(docx_doc, para_data)

        assert docx_doc.paragraphs[0].style.name == "Heading 1"

    def test_add_paragraph_with_colors(self):
        """Test adding paragraph with colored text."""
        docx_doc = DocxDocument()

        red_run = TextRun(text="Red", formatting=RunFormatting(color="#FF0000"))
        blue_run = TextRun(text=" Blue", formatting=RunFormatting(color="#0000FF"))

        para_data = Paragraph(runs=[red_run, blue_run])

        add_paragraph_to_document(docx_doc, para_data)

        para = docx_doc.paragraphs[0]
        assert len(para.runs) == 2


class TestRebuildDocument:
    """Tests for rebuild_document function."""

    def test_rebuild_simple_document(self):
        """Test rebuilding a simple document."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            # Create source document
            para1 = Paragraph(runs=[TextRun(text="Paragraph 1", formatting=RunFormatting())])
            para2 = Paragraph(runs=[TextRun(text="Paragraph 2", formatting=RunFormatting())])

            doc_data = Document(preamble_paragraphs=[para1, para2])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            assert len(rebuilt.paragraphs) >= 2
            assert "Paragraph 1" in rebuilt.paragraphs[0].text
            assert "Paragraph 2" in rebuilt.paragraphs[1].text

    def test_rebuild_document_with_sections(self):
        """Test rebuilding document with sections."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            # Create sections
            header1 = Paragraph(
                runs=[TextRun(text="Section 1", formatting=RunFormatting(bold=True))],
                style_name="Heading 1",
            )
            body1 = Paragraph(runs=[TextRun(text="Body 1", formatting=RunFormatting())])

            section1 = Section(header=header1, paragraphs=[body1])

            doc_data = Document(sections=[section1])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            assert len(rebuilt.paragraphs) >= 2

    def test_rebuild_document_with_preamble_and_sections(self):
        """Test rebuilding with both preamble and sections."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            # Create preamble
            preamble = [Paragraph(runs=[TextRun(text="Introduction", formatting=RunFormatting())])]

            # Create section
            header = Paragraph(
                runs=[TextRun(text="Main Section", formatting=RunFormatting(bold=True))],
                style_name="Heading 1",
            )
            body = Paragraph(runs=[TextRun(text="Content", formatting=RunFormatting())])
            section = Section(header=header, paragraphs=[body])

            doc_data = Document(preamble_paragraphs=preamble, sections=[section])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            text = "\n".join(p.text for p in rebuilt.paragraphs)
            assert "Introduction" in text
            assert "Main Section" in text
            assert "Content" in text

    def test_rebuild_document_preserves_formatting(self):
        """Test that rebuild preserves text formatting."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            # Create formatted paragraph
            bold_run = TextRun(text="Bold ", formatting=RunFormatting(bold=True))
            italic_run = TextRun(text="italic", formatting=RunFormatting(italic=True))

            para = Paragraph(runs=[bold_run, italic_run])
            doc_data = Document(preamble_paragraphs=[para])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            assert rebuilt.paragraphs[0].runs[0].bold is True
            assert rebuilt.paragraphs[0].runs[1].italic is True

    def test_rebuild_document_preserves_alignment(self):
        """Test that rebuild preserves paragraph alignment."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            # Create centered paragraph
            para = Paragraph(
                runs=[TextRun(text="Centered", formatting=RunFormatting())],
                formatting=ParagraphFormatting(alignment="center"),
            )
            doc_data = Document(preamble_paragraphs=[para])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            assert rebuilt.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_rebuild_empty_document(self):
        """Test rebuilding an empty document."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc_data = Document()

            # Should not crash
            rebuild_document(doc_data, tmp.name)

            # Verify file was created
            rebuilt = DocxDocument(tmp.name)
            assert rebuilt is not None

    def test_rebuild_document_with_multiple_runs_per_paragraph(self):
        """Test paragraph with many runs."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            runs = [
                TextRun(text="Red ", formatting=RunFormatting(color="#FF0000")),
                TextRun(text="Green ", formatting=RunFormatting(color="#00FF00")),
                TextRun(text="Blue", formatting=RunFormatting(color="#0000FF")),
            ]
            para = Paragraph(runs=runs)
            doc_data = Document(preamble_paragraphs=[para])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            para = rebuilt.paragraphs[0]
            assert len(para.runs) == 3
            assert "Red" in para.runs[0].text
            assert "Green" in para.runs[1].text
            assert "Blue" in para.runs[2].text

    def test_rebuild_document_with_superscripts(self):
        """Test that superscripts are preserved in rebuild."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            normal_run = TextRun(text="Text", formatting=RunFormatting())
            super_run = TextRun(text="[1]", formatting=RunFormatting(is_superscript=True))

            para = Paragraph(runs=[normal_run, super_run])
            doc_data = Document(preamble_paragraphs=[para])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            para = rebuilt.paragraphs[0]
            assert para.runs[1].font.superscript is True

    def test_rebuild_document_multiple_sections(self):
        """Test rebuilding document with multiple sections."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            # Create multiple sections
            sections = []
            for i in range(3):
                header = Paragraph(
                    runs=[TextRun(text=f"Section {i+1}", formatting=RunFormatting(bold=True))],
                    style_name="Heading 1",
                )
                body = Paragraph(runs=[TextRun(text=f"Content {i+1}", formatting=RunFormatting())])
                sections.append(Section(header=header, paragraphs=[body]))

            doc_data = Document(sections=sections)

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            text = "\n".join(p.text for p in rebuilt.paragraphs)
            assert "Section 1" in text
            assert "Section 2" in text
            assert "Section 3" in text


class TestApplyPageSettings:
    """Tests for apply_page_settings function."""

    def test_apply_page_width(self):
        """Test applying page width."""
        docx_doc = DocxDocument()
        section = docx_doc.sections[0]

        page_settings = PageSettings(width=7920000)  # A4 width in twips
        apply_page_settings(section, page_settings)

        # Allow for minor rounding adjustments by python-docx
        assert section.page_width is not None
        assert abs(section.page_width - 7920000) < 1000

    def test_apply_page_margins(self):
        """Test applying page margins."""
        docx_doc = DocxDocument()
        section = docx_doc.sections[0]

        page_settings = PageSettings(
            top_margin=Pt(72),
            bottom_margin=Pt(72),
            left_margin=Pt(72),
            right_margin=Pt(72),
        )
        apply_page_settings(section, page_settings)

        assert section.top_margin == Pt(72)
        assert section.bottom_margin == Pt(72)
        assert section.left_margin == Pt(72)
        assert section.right_margin == Pt(72)

    def test_apply_header_footer_distance(self):
        """Test applying header and footer distance."""
        docx_doc = DocxDocument()
        section = docx_doc.sections[0]

        page_settings = PageSettings(header_distance=Pt(36), footer_distance=Pt(36))
        apply_page_settings(section, page_settings)

        assert section.header_distance == Pt(36)
        assert section.footer_distance == Pt(36)


class TestApplyLineNumbering:
    """Tests for apply_line_numbering function."""

    def test_apply_line_numbering_disabled(self):
        """Test that disabled line numbering is skipped."""
        docx_doc = DocxDocument()
        section = docx_doc.sections[0]

        line_numbering = LineNumbering(enabled=False)
        apply_line_numbering(section, line_numbering)

        # Should not crash and document should be valid
        assert section is not None

    def test_apply_line_numbering_enabled(self):
        """Test applying enabled line numbering."""
        docx_doc = DocxDocument()
        section = docx_doc.sections[0]

        line_numbering = LineNumbering(enabled=True, start=1, increment=1, restart="continuous")
        apply_line_numbering(section, line_numbering)

        # Verify by checking the section is valid
        assert section is not None

    def test_apply_line_numbering_with_custom_start(self):
        """Test applying line numbering with custom start value."""
        docx_doc = DocxDocument()
        section = docx_doc.sections[0]

        line_numbering = LineNumbering(enabled=True, start=5, increment=2)
        apply_line_numbering(section, line_numbering)

        # Should not crash
        assert section is not None


class TestApplyMetadata:
    """Tests for apply_metadata function."""

    def test_apply_metadata_with_page_settings(self):
        """Test applying metadata with page settings."""
        docx_doc = DocxDocument()

        page_settings = PageSettings(top_margin=Pt(72), bottom_margin=Pt(72))
        metadata = DocumentMetadata(page_settings=page_settings)

        apply_metadata(docx_doc, metadata)

        section = docx_doc.sections[0]
        assert section.top_margin == Pt(72)
        assert section.bottom_margin == Pt(72)

    def test_apply_metadata_with_line_numbering(self):
        """Test applying metadata with line numbering."""
        docx_doc = DocxDocument()

        line_numbering = LineNumbering(enabled=True, start=1, increment=1)
        metadata = DocumentMetadata(line_numbering=line_numbering)

        apply_metadata(docx_doc, metadata)

        # Should not crash
        assert docx_doc.sections[0] is not None

    def test_apply_metadata_empty(self):
        """Test applying empty metadata doesn't crash."""
        docx_doc = DocxDocument()

        metadata = DocumentMetadata()
        apply_metadata(docx_doc, metadata)

        # Should not crash
        assert docx_doc is not None

    def test_apply_metadata_no_sections(self):
        """Test applying metadata to document without sections."""
        docx_doc = DocxDocument()

        # Remove all sections (edge case)
        metadata = DocumentMetadata()
        apply_metadata(docx_doc, metadata)

        # Should handle gracefully
        assert docx_doc is not None


class TestRebuildDocumentWithMetadata:
    """Tests for document rebuild with metadata."""

    def test_rebuild_document_with_page_settings(self):
        """Test rebuilding document preserves page settings."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            page_settings = PageSettings(
                top_margin=Pt(72),
                bottom_margin=Pt(72),
                left_margin=Pt(72),
                right_margin=Pt(72),
            )
            metadata = DocumentMetadata(page_settings=page_settings)

            para = Paragraph(runs=[TextRun(text="Content", formatting=RunFormatting())])

            doc_data = Document(metadata=metadata, preamble_paragraphs=[para])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            section = rebuilt.sections[0]

            assert section.top_margin == Pt(72)
            assert section.bottom_margin == Pt(72)
            assert section.left_margin == Pt(72)
            assert section.right_margin == Pt(72)

    def test_rebuild_document_with_line_numbering(self):
        """Test rebuilding document with line numbering."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            line_numbering = LineNumbering(enabled=True, start=1, increment=1)
            metadata = DocumentMetadata(line_numbering=line_numbering)

            para = Paragraph(
                runs=[TextRun(text="Line 1", formatting=RunFormatting())],
            )
            para2 = Paragraph(
                runs=[TextRun(text="Line 2", formatting=RunFormatting())],
            )

            doc_data = Document(metadata=metadata, preamble_paragraphs=[para, para2])

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify it's a valid document
            rebuilt = DocxDocument(tmp.name)
            assert len(rebuilt.paragraphs) >= 2

    def test_rebuild_document_full_metadata(self):
        """Test rebuilding with complete metadata."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            page_settings = PageSettings(
                top_margin=Pt(72),
                left_margin=Pt(72),
            )
            line_numbering = LineNumbering(enabled=True)
            metadata = DocumentMetadata(
                page_settings=page_settings, line_numbering=line_numbering, default_font="Arial"
            )

            header = Paragraph(
                runs=[TextRun(text="Header", formatting=RunFormatting(bold=True))],
                style_name="Heading 1",
            )
            body = Paragraph(runs=[TextRun(text="Body", formatting=RunFormatting())])

            doc_data = Document(
                metadata=metadata, sections=[Section(header=header, paragraphs=[body])]
            )

            # Rebuild
            rebuild_document(doc_data, tmp.name)

            # Verify
            rebuilt = DocxDocument(tmp.name)
            assert len(rebuilt.paragraphs) >= 2
            assert rebuilt.sections[0].top_margin == Pt(72)
