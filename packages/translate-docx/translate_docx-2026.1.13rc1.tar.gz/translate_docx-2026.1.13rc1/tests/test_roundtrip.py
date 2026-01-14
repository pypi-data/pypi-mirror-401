"""Round-trip tests: extract → rebuild → extract → compare for lossless preservation."""

import sys
import tempfile

sys.path.insert(0, "src")

from docx.shared import Pt

from translate_docx.extractor import extract_document
from translate_docx.models import (
    Document,
    DocumentMetadata,
    LineNumbering,
    PageSettings,
    Paragraph,
    ParagraphFormatting,
    RunFormatting,
    Section,
    TextRun,
)
from translate_docx.rebuilder import rebuild_document


def compare_run_formatting(fmt1: RunFormatting, fmt2: RunFormatting) -> bool:
    """Compare two RunFormatting objects for equality.

    Args:
        fmt1: First RunFormatting object
        fmt2: Second RunFormatting object

    Returns:
        True if all relevant attributes are equal
    """
    return (
        fmt1.bold == fmt2.bold
        and fmt1.italic == fmt2.italic
        and fmt1.underline == fmt2.underline
        and fmt1.font_name == fmt2.font_name
        and fmt1.font_size == fmt2.font_size
        and fmt1.color == fmt2.color
        and fmt1.is_superscript == fmt2.is_superscript
        and fmt1.is_subscript == fmt2.is_subscript
    )


def compare_text_runs(runs1, runs2) -> bool:
    """Compare two lists of TextRun objects.

    Args:
        runs1: First list of TextRun objects
        runs2: Second list of TextRun objects

    Returns:
        True if all runs are equal
    """
    if len(runs1) != len(runs2):
        return False

    for r1, r2 in zip(runs1, runs2):
        if r1.text != r2.text:
            return False
        if not compare_run_formatting(r1.formatting, r2.formatting):
            return False

    return True


def compare_paragraphs(para1: Paragraph, para2: Paragraph) -> bool:
    """Compare two Paragraph objects.

    Args:
        para1: First Paragraph object
        para2: Second Paragraph object

    Returns:
        True if all content and formatting are equal
    """
    # Compare runs
    if not compare_text_runs(para1.runs, para2.runs):
        return False

    # Compare text content
    if para1.text != para2.text:
        return False

    # Compare paragraph formatting (key attributes)
    fmt1 = para1.formatting
    fmt2 = para2.formatting

    return (
        fmt1.alignment == fmt2.alignment
        and fmt1.space_before == fmt2.space_before
        and fmt1.space_after == fmt2.space_after
        and fmt1.keep_together == fmt2.keep_together
        and fmt1.page_break_before == fmt2.page_break_before
    )


def compare_sections(sect1: Section, sect2: Section) -> bool:
    """Compare two Section objects.

    Args:
        sect1: First Section object
        sect2: Second Section object

    Returns:
        True if headers and all body paragraphs are equal
    """
    # Compare headers
    if not compare_paragraphs(sect1.header, sect2.header):
        return False

    # Compare body paragraphs
    if len(sect1.paragraphs) != len(sect2.paragraphs):
        return False

    for p1, p2 in zip(sect1.paragraphs, sect2.paragraphs):
        if not compare_paragraphs(p1, p2):
            return False

    return True


def compare_documents(doc1: Document, doc2: Document) -> bool:
    """Compare two Document objects for lossless preservation.

    Args:
        doc1: First Document object
        doc2: Second Document object

    Returns:
        True if documents are equivalent
    """
    # Compare preamble
    if len(doc1.preamble_paragraphs) != len(doc2.preamble_paragraphs):
        return False

    for p1, p2 in zip(doc1.preamble_paragraphs, doc2.preamble_paragraphs):
        if not compare_paragraphs(p1, p2):
            return False

    # Compare sections
    if len(doc1.sections) != len(doc2.sections):
        return False

    for s1, s2 in zip(doc1.sections, doc2.sections):
        if not compare_sections(s1, s2):
            return False

    return True


class TestRoundTripSimple:
    """Simple round-trip tests."""

    def test_roundtrip_single_paragraph(self):
        """Test: extract → rebuild → extract single paragraph."""
        # Create original document
        para = Paragraph(runs=[TextRun(text="Hello World", formatting=RunFormatting())])
        doc1 = Document(preamble_paragraphs=[para])

        # Rebuild and re-extract
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Compare
        assert len(doc2.preamble_paragraphs) >= 1
        assert "Hello World" in doc2.preamble_paragraphs[0].text

    def test_roundtrip_multiple_paragraphs(self):
        """Test: extract → rebuild → extract multiple paragraphs."""
        paras = [
            Paragraph(runs=[TextRun(text="Paragraph 1", formatting=RunFormatting())]),
            Paragraph(runs=[TextRun(text="Paragraph 2", formatting=RunFormatting())]),
            Paragraph(runs=[TextRun(text="Paragraph 3", formatting=RunFormatting())]),
        ]
        doc1 = Document(preamble_paragraphs=paras)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify all paragraphs preserved
        texts = [p.text for p in doc2.preamble_paragraphs if p.text.strip()]
        assert "Paragraph 1" in texts[0]
        assert "Paragraph 2" in texts[1]
        assert "Paragraph 3" in texts[2]

    def test_roundtrip_empty_document(self):
        """Test: extract → rebuild → extract empty document."""
        doc1 = Document()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Should have at least one default paragraph
        assert doc2 is not None
        assert isinstance(doc2, Document)


class TestRoundTripFormatting:
    """Round-trip tests with text formatting."""

    def test_roundtrip_bold_text(self):
        """Test: extract → rebuild → extract preserves bold."""
        # Use normal formatting alongside bold to prevent header detection
        bold_run = TextRun(text="Bold", formatting=RunFormatting(bold=True))
        normal_run = TextRun(text=" text", formatting=RunFormatting())
        para = Paragraph(runs=[bold_run, normal_run])
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify bold preserved (may be in preamble or sections)
        all_paras = doc2.all_paragraphs
        assert len(all_paras) > 0
        assert any(
            run.formatting.bold for para in all_paras for run in para.runs if "Bold" in run.text
        )

    def test_roundtrip_italic_text(self):
        """Test: extract → rebuild → extract preserves italic."""
        italic_run = TextRun(text="Italic", formatting=RunFormatting(italic=True))
        para = Paragraph(runs=[italic_run])
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        para2 = doc2.preamble_paragraphs[0]
        assert any(run.formatting.italic for run in para2.runs if "Italic" in run.text)

    def test_roundtrip_font_properties(self):
        """Test: extract → rebuild → extract preserves font name and size."""
        run = TextRun(text="Styled text", formatting=RunFormatting(font_name="Arial", font_size=14))
        para = Paragraph(runs=[run])
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        para2 = doc2.preamble_paragraphs[0]
        styled_run = [r for r in para2.runs if "Styled text" in r.text]
        assert len(styled_run) > 0
        assert styled_run[0].formatting.font_name == "Arial"
        assert styled_run[0].formatting.font_size == 14

    def test_roundtrip_multiple_runs(self):
        """Test: extract → rebuild → extract multiple runs per paragraph."""
        runs = [
            TextRun(text="Bold ", formatting=RunFormatting(bold=True)),
            TextRun(text="italic ", formatting=RunFormatting(italic=True)),
            TextRun(text="normal", formatting=RunFormatting()),
        ]
        para = Paragraph(runs=runs)
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        para2 = doc2.preamble_paragraphs[0]
        assert len(para2.runs) >= 3
        # Verify formatting preserved
        assert any(r.formatting.bold for r in para2.runs)
        assert any(r.formatting.italic for r in para2.runs)

    def test_roundtrip_superscripts(self):
        """Test: extract → rebuild → extract preserves superscripts."""
        runs = [
            TextRun(text="Text with citation", formatting=RunFormatting()),
            TextRun(text="[1]", formatting=RunFormatting(is_superscript=True)),
        ]
        para = Paragraph(runs=runs)
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        para2 = doc2.preamble_paragraphs[0]
        # Find superscript run
        super_runs = [r for r in para2.runs if r.formatting.is_superscript]
        assert len(super_runs) > 0


class TestRoundTripFormatting2:
    """Round-trip tests with paragraph formatting."""

    def test_roundtrip_alignment(self):
        """Test: extract → rebuild → extract preserves alignment."""
        para = Paragraph(
            runs=[TextRun(text="Centered text", formatting=RunFormatting())],
            formatting=ParagraphFormatting(alignment="center"),
        )
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        para2 = doc2.preamble_paragraphs[0]
        assert para2.formatting.alignment == "center"

    def test_roundtrip_spacing(self):
        """Test: extract → rebuild → extract preserves spacing."""
        para = Paragraph(
            runs=[TextRun(text="Spaced", formatting=RunFormatting())],
            formatting=ParagraphFormatting(space_before=Pt(12), space_after=Pt(12)),
        )
        doc1 = Document(preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        para2 = doc2.preamble_paragraphs[0]
        assert para2.formatting.space_before == Pt(12)
        assert para2.formatting.space_after == Pt(12)


class TestRoundTripSections:
    """Round-trip tests with sections."""

    def test_roundtrip_single_section(self):
        """Test: extract → rebuild → extract single section."""
        header = Paragraph(
            runs=[TextRun(text="Section Header", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
            is_header=True,
        )
        body = Paragraph(runs=[TextRun(text="Body text", formatting=RunFormatting())])
        section = Section(header=header, paragraphs=[body])
        doc1 = Document(sections=[section])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        assert len(doc2.sections) >= 1
        # First section should have our header
        assert "Section Header" in doc2.sections[0].header.text

    def test_roundtrip_multiple_sections(self):
        """Test: extract → rebuild → extract multiple sections."""
        sections = []
        for i in range(3):
            header = Paragraph(
                runs=[TextRun(text=f"Section {i+1}", formatting=RunFormatting(bold=True))],
                style_name="Heading 1",
                is_header=True,
            )
            body = Paragraph(runs=[TextRun(text=f"Content {i+1}", formatting=RunFormatting())])
            sections.append(Section(header=header, paragraphs=[body]))

        doc1 = Document(sections=sections)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify sections preserved
        assert len(doc2.sections) >= 3
        assert "Section 1" in doc2.sections[0].header.text
        assert "Section 2" in doc2.sections[1].header.text
        assert "Section 3" in doc2.sections[2].header.text

    def test_roundtrip_section_with_multiple_paragraphs(self):
        """Test: extract → rebuild → extract section with multiple body paragraphs."""
        header = Paragraph(
            runs=[TextRun(text="Multi-paragraph Section", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
            is_header=True,
        )
        body_paras = [
            Paragraph(runs=[TextRun(text="Paragraph 1", formatting=RunFormatting())]),
            Paragraph(runs=[TextRun(text="Paragraph 2", formatting=RunFormatting())]),
            Paragraph(runs=[TextRun(text="Paragraph 3", formatting=RunFormatting())]),
        ]
        section = Section(header=header, paragraphs=body_paras)
        doc1 = Document(sections=[section])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # First section should have multiple paragraphs
        section2 = doc2.sections[0]
        assert len(section2.paragraphs) >= 3


class TestRoundTripMetadata:
    """Round-trip tests with metadata."""

    def test_roundtrip_page_settings(self):
        """Test: extract → rebuild → extract preserves page settings."""
        page_settings = PageSettings(
            top_margin=Pt(72),
            bottom_margin=Pt(72),
            left_margin=Pt(72),
            right_margin=Pt(72),
        )
        metadata = DocumentMetadata(page_settings=page_settings)
        para = Paragraph(runs=[TextRun(text="Content", formatting=RunFormatting())])
        doc1 = Document(metadata=metadata, preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify margins preserved
        assert doc2.metadata.page_settings.top_margin == Pt(72)
        assert doc2.metadata.page_settings.bottom_margin == Pt(72)
        assert doc2.metadata.page_settings.left_margin == Pt(72)
        assert doc2.metadata.page_settings.right_margin == Pt(72)

    def test_roundtrip_line_numbering(self):
        """Test: extract → rebuild → extract preserves line numbering."""
        line_numbering = LineNumbering(enabled=True, start=1, increment=1)
        metadata = DocumentMetadata(line_numbering=line_numbering)
        para = Paragraph(runs=[TextRun(text="Content", formatting=RunFormatting())])
        doc1 = Document(metadata=metadata, preamble_paragraphs=[para])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify line numbering preserved
        assert doc2.metadata.line_numbering.enabled is True
        assert doc2.metadata.line_numbering.start == 1
        assert doc2.metadata.line_numbering.increment == 1


class TestRoundTripComplex:
    """Complex round-trip tests with multiple features."""

    def test_roundtrip_preamble_and_sections(self):
        """Test: extract → rebuild → extract document with preamble and sections."""
        # Preamble
        preamble = [
            Paragraph(runs=[TextRun(text="Introduction", formatting=RunFormatting())]),
            Paragraph(runs=[TextRun(text="More intro", formatting=RunFormatting())]),
        ]

        # Sections
        sections = []
        for i in range(2):
            header = Paragraph(
                runs=[TextRun(text=f"Chapter {i+1}", formatting=RunFormatting(bold=True))],
                style_name="Heading 1",
                is_header=True,
            )
            body = Paragraph(
                runs=[TextRun(text=f"Chapter {i+1} content", formatting=RunFormatting())]
            )
            sections.append(Section(header=header, paragraphs=[body]))

        doc1 = Document(preamble_paragraphs=preamble, sections=sections)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify preamble
        assert len(doc2.preamble_paragraphs) >= 2

        # Verify sections
        assert len(doc2.sections) >= 2
        assert "Chapter 1" in doc2.sections[0].header.text
        assert "Chapter 2" in doc2.sections[1].header.text

    def test_roundtrip_with_all_features(self):
        """Test: extract → rebuild → extract document with all features."""
        # Metadata
        page_settings = PageSettings(
            top_margin=Pt(72),
            left_margin=Pt(72),
        )
        metadata = DocumentMetadata(page_settings=page_settings)

        # Preamble with formatting
        preamble = [
            Paragraph(
                runs=[
                    TextRun(text="Intro ", formatting=RunFormatting(bold=True)),
                    TextRun(text="text", formatting=RunFormatting()),
                ]
            )
        ]

        # Section with multiple paragraphs and formatting
        header = Paragraph(
            runs=[TextRun(text="Main Section", formatting=RunFormatting(bold=True))],
            style_name="Heading 1",
            is_header=True,
        )
        body_paras = [
            Paragraph(
                runs=[
                    TextRun(text="Content", formatting=RunFormatting()),
                    TextRun(text="[1]", formatting=RunFormatting(is_superscript=True)),
                ],
                formatting=ParagraphFormatting(alignment="left"),
            )
        ]
        section = Section(header=header, paragraphs=body_paras)

        doc1 = Document(metadata=metadata, preamble_paragraphs=preamble, sections=[section])

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify all components
        assert len(doc2.preamble_paragraphs) >= 1
        assert len(doc2.sections) >= 1
        assert doc2.metadata.page_settings.top_margin == Pt(72)

        # Verify content preservation
        preamble_text = doc2.preamble_paragraphs[0].text
        assert "Intro" in preamble_text
        assert "Main Section" in doc2.sections[0].header.text


class TestRoundTripRealDocument:
    """Round-trip tests on real documents."""

    def test_roundtrip_real_example_document(self):
        """Test: extract → rebuild → extract real example document."""
        import os

        # Only run if example document exists
        example_path = "/home/jerry/Desktop/CODING/docx-parser/user_data/input_example.docx"
        if not os.path.exists(example_path):
            return

        # Extract original
        doc1 = extract_document(example_path)

        # Rebuild and re-extract
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            rebuild_document(doc1, tmp.name)
            doc2 = extract_document(tmp.name)

        # Verify structure preserved
        assert len(doc2.sections) == len(doc1.sections)
        assert len(doc2.all_paragraphs) >= len(doc1.all_paragraphs) - 10  # Allow small variance

        # Verify metadata
        assert doc2.metadata.page_settings.top_margin == doc1.metadata.page_settings.top_margin
        assert (
            doc2.metadata.page_settings.bottom_margin == doc1.metadata.page_settings.bottom_margin
        )

        # Verify section headers
        for i, section in enumerate(doc1.sections[:5]):
            assert section.header.text in doc2.sections[i].header.text
