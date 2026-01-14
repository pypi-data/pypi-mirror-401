"""Rebuild DOCX files from structured dataclass objects."""

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from translate_docx.models import (
    Document,
    DocumentMetadata,
    Hyperlink,
    LineNumbering,
    PageSettings,
    Paragraph,
    ParagraphFormatting,
    RunFormatting,
    TextRun,
)
from translate_docx.utils import hex_to_rgb_color


def apply_run_formatting(run, formatting: RunFormatting):
    """Apply formatting attributes to a python-docx Run object.

    Args:
        run: python-docx Run object to format
        formatting: RunFormatting object with desired attributes
    """
    # Apply basic formatting
    if formatting.bold is not None:
        run.bold = formatting.bold

    if formatting.italic is not None:
        run.italic = formatting.italic

    if formatting.underline is not None:
        run.underline = formatting.underline

    # Apply font properties
    font = run.font

    if formatting.font_name:
        font.name = formatting.font_name

    if formatting.font_size:
        font.size = Pt(formatting.font_size)

    # Apply color
    if formatting.color:
        rgb_color = hex_to_rgb_color(formatting.color)
        if rgb_color:
            font.color.rgb = rgb_color

    # Apply superscript and subscript
    if formatting.is_superscript:
        font.superscript = True

    if formatting.is_subscript:
        font.subscript = True


def add_hyperlink_to_paragraph(para, hyperlink_data: Hyperlink):
    """Add a hyperlink with formatting to a python-docx Paragraph.

    This function creates the necessary XML structure and relationship
    for a hyperlink, since python-docx may not fully expose the API.

    Args:
        para: python-docx Paragraph object to add hyperlink to
        hyperlink_data: Hyperlink dataclass object with runs and URL

    Returns:
        The created hyperlink element or None
    """
    from docx.oxml import parse_xml

    url = hyperlink_data.url
    anchor = hyperlink_data.anchor

    # Create external hyperlink relationship if URL exists
    relationship_id = None
    if url:
        try:
            part = para.part
            relationship_id = part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True
            )
        except Exception:
            # Fallback: add as plain text if relationship fails
            for text_run in hyperlink_data.runs:
                run = para.add_run(text_run.text)
                apply_run_formatting(run, text_run.formatting)
            return None

    # Build hyperlink XML element
    xmlns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    hyperlink_xml = f'<w:hyperlink xmlns:w="{xmlns}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'

    if relationship_id:
        hyperlink_xml += f' r:id="{relationship_id}"'

    if anchor:
        hyperlink_xml += f' w:anchor="{anchor}"'

    hyperlink_xml += '></w:hyperlink>'
    hyperlink_element = parse_xml(hyperlink_xml)

    # Add runs to the hyperlink element with formatting
    for text_run in hyperlink_data.runs:
        # Build run XML with text and formatting
        run_xml_parts = [f'<w:r xmlns:w="{xmlns}">']

        # Add run properties (formatting) if present
        if text_run.formatting:
            rpr_parts = []

            # Add bold
            if text_run.formatting.bold is not None:
                rpr_parts.append(f'<w:b w:val="{"1" if text_run.formatting.bold else "0"}"/>')

            # Add italic
            if text_run.formatting.italic is not None:
                rpr_parts.append(f'<w:i w:val="{"1" if text_run.formatting.italic else "0"}"/>')

            # Add underline
            if text_run.formatting.underline is not None:
                underline_val = "single" if text_run.formatting.underline else "none"
                rpr_parts.append(f'<w:u w:val="{underline_val}"/>')

            # Add font name
            if text_run.formatting.font_name:
                rpr_parts.append(
                    f'<w:rFonts w:ascii="{text_run.formatting.font_name}" '
                    f'w:hAnsi="{text_run.formatting.font_name}"/>'
                )

            # Add font size
            if text_run.formatting.font_size:
                half_points = text_run.formatting.font_size * 2
                rpr_parts.append(f'<w:sz w:val="{half_points}"/>')
                rpr_parts.append(f'<w:szCs w:val="{half_points}"/>')

            # Add color
            if text_run.formatting.color:
                # Color is in hex format, remove # if present
                color_val = text_run.formatting.color.lstrip('#')
                rpr_parts.append(f'<w:color w:val="{color_val}"/>')

            # Add superscript
            if text_run.formatting.is_superscript:
                rpr_parts.append('<w:vertAlign w:val="superscript"/>')

            # Add subscript
            if text_run.formatting.is_subscript:
                rpr_parts.append('<w:vertAlign w:val="subscript"/>')

            # Add rPr element if there are any properties
            if rpr_parts:
                run_xml_parts.append('<w:rPr>')
                run_xml_parts.extend(rpr_parts)
                run_xml_parts.append('</w:rPr>')

        # Add text content
        if text_run.text:
            # Escape XML special characters
            text_escaped = (
                text_run.text.replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
            )
            run_xml_parts.append(f'<w:t xml:space="preserve">{text_escaped}</w:t>')

        run_xml_parts.append('</w:r>')

        # Parse and add run to hyperlink
        run_element = parse_xml(''.join(run_xml_parts))
        hyperlink_element.append(run_element)

    # Add hyperlink element to paragraph
    para._element.append(hyperlink_element)

    return hyperlink_element


def apply_paragraph_formatting(para, formatting: ParagraphFormatting):
    """Apply formatting attributes to a python-docx Paragraph object.

    Args:
        para: python-docx Paragraph object to format
        formatting: ParagraphFormatting object with desired attributes
    """
    para_format = para.paragraph_format

    # Apply alignment
    if formatting.alignment:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        }
        alignment_enum = alignment_map.get(formatting.alignment.lower())
        if alignment_enum is not None:
            para.alignment = alignment_enum

    # Apply spacing
    if formatting.line_spacing is not None:
        para_format.line_spacing = formatting.line_spacing

    if formatting.space_before is not None:
        para_format.space_before = formatting.space_before

    if formatting.space_after is not None:
        para_format.space_after = formatting.space_after

    # Apply indentation
    if formatting.first_line_indent is not None:
        para_format.first_line_indent = formatting.first_line_indent

    if formatting.left_indent is not None:
        para_format.left_indent = formatting.left_indent

    if formatting.right_indent is not None:
        para_format.right_indent = formatting.right_indent

    # Apply boolean properties
    if formatting.keep_together:
        para_format.keep_together = True

    if formatting.keep_with_next:
        para_format.keep_with_next = True

    if formatting.page_break_before:
        para_format.page_break_before = True


def set_paragraph_default_font_size(para, font_size: int):
    """Set the default font size for a paragraph (used for empty paragraphs).

    This sets the font size at the paragraph level (pPr > rPr) rather than
    at the run level, which is how Word displays font size for empty paragraphs.

    Args:
        para: python-docx Paragraph object
        font_size: Font size in points
    """
    if font_size is None:
        return

    try:
        from docx.oxml import parse_xml

        p_pr = para._element.get_or_add_pPr()

        # Remove existing rPr if it exists
        rpr = p_pr.find(qn("w:rPr"))
        if rpr is not None:
            p_pr.remove(rpr)

        # Create new rPr with font size (in half-points)
        xmlns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        rpr_xml = (
            f'<w:rPr xmlns:w="{xmlns}">'
            f'<w:sz w:val="{font_size * 2}"/>'
            f'<w:szCs w:val="{font_size * 2}"/>'
            f'</w:rPr>'
        )
        rpr = parse_xml(rpr_xml)
        p_pr.insert(0, rpr)
    except Exception:
        pass


def add_paragraph_to_document(docx_doc, para_data: Paragraph, style: str | None = None, default_font: str | None = None, default_font_size: int | None = None):
    """Add a paragraph to a python-docx Document with all formatting.

    Args:
        docx_doc: python-docx Document object to add paragraph to
        para_data: Paragraph dataclass object to add
        style: Optional style name for the paragraph
        default_font: Optional default font to use for empty paragraphs
        default_font_size: Optional default font size to use for empty paragraphs

    Returns:
        The created python-docx Paragraph object
    """
    # Determine style to use
    if style is None:
        style = para_data.style_name if para_data.style_name else "Normal"

    # Try to use the specified style, fall back to Normal if it doesn't exist
    try:
        para = docx_doc.add_paragraph(style=style)
    except KeyError:
        # Style doesn't exist in this document, use Normal
        para = docx_doc.add_paragraph(style="Normal")

    # Add inline elements (runs and hyperlinks) with formatting
    if para_data.runs:
        for element in para_data.runs:
            if isinstance(element, TextRun):
                # Regular text run
                run = para.add_run(element.text)
                apply_run_formatting(run, element.formatting)
            elif isinstance(element, Hyperlink):
                # Hyperlink with nested runs
                add_hyperlink_to_paragraph(para, element)
    else:
        # Empty paragraph - set font size at paragraph level to preserve formatting
        # Use paragraph's specific font size if available, otherwise use document default
        font_size_to_use = para_data.empty_paragraph_font_size or default_font_size
        if font_size_to_use:
            set_paragraph_default_font_size(para, font_size_to_use)

    # Apply paragraph formatting
    apply_paragraph_formatting(para, para_data.formatting)

    return para


def apply_page_settings(section, page_settings: PageSettings):
    """Apply page layout settings to a document section.

    Args:
        section: python-docx Section object to configure
        page_settings: PageSettings object with desired layout
    """
    if page_settings.width is not None:
        section.page_width = page_settings.width

    if page_settings.height is not None:
        section.page_height = page_settings.height

    if page_settings.top_margin is not None:
        section.top_margin = page_settings.top_margin

    if page_settings.bottom_margin is not None:
        section.bottom_margin = page_settings.bottom_margin

    if page_settings.left_margin is not None:
        section.left_margin = page_settings.left_margin

    if page_settings.right_margin is not None:
        section.right_margin = page_settings.right_margin

    if page_settings.header_distance is not None:
        section.header_distance = page_settings.header_distance

    if page_settings.footer_distance is not None:
        section.footer_distance = page_settings.footer_distance


def apply_line_numbering(section, line_numbering: LineNumbering):
    """Apply line numbering settings to a document section.

    Requires XML manipulation since python-docx doesn't fully expose this API.

    Args:
        section: python-docx Section object to configure
        line_numbering: LineNumbering object with desired settings
    """
    if not line_numbering.enabled:
        return

    try:
        sect_pr = section._sectPr

        # Remove existing lnNumType element if present to avoid state issues
        ln_num_type = sect_pr.find(qn("w:lnNumType"))
        if ln_num_type is not None:
            sect_pr.remove(ln_num_type)

        # Create new lnNumType element
        from docx.oxml import parse_xml

        xmlns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        restart = line_numbering.restart or "continuous"

        # Build XML with all attributes
        # Note: Only include w:start if it's not 1 (Word's default)
        # Explicitly setting start="1" can cause issues with template state
        ln_num_xml = f'<w:lnNumType xmlns:w="{xmlns}" '

        if line_numbering.start != 1:
            ln_num_xml += f'w:start="{line_numbering.start}" '

        ln_num_xml += (
            f'w:countBy="{line_numbering.increment}" '
            f'w:restart="{restart}"'
        )

        if line_numbering.distance is not None:
            ln_num_xml += f' w:distance="{line_numbering.distance}"'

        ln_num_xml += '/>'

        ln_num_type = parse_xml(ln_num_xml)
        sect_pr.append(ln_num_type)

    except Exception:
        # Log but don't fail if line numbering can't be applied
        pass


def apply_metadata(docx_doc, metadata: DocumentMetadata):
    """Apply document-level metadata and settings.

    Args:
        docx_doc: python-docx Document object to configure
        metadata: DocumentMetadata object with desired settings
    """
    # Get the first section
    if not docx_doc.sections:
        return

    section = docx_doc.sections[0]

    # Apply page settings
    if metadata.page_settings:
        apply_page_settings(section, metadata.page_settings)

    # Apply line numbering
    if metadata.line_numbering:
        apply_line_numbering(section, metadata.line_numbering)

    # Apply default font/size to both Normal style and document defaults
    if metadata.default_font or metadata.default_font_size:
        try:
            # Set Normal style font
            normal_style = docx_doc.styles["Normal"]
            if metadata.default_font:
                normal_style.font.name = metadata.default_font
            if metadata.default_font_size:
                normal_style.font.size = Pt(metadata.default_font_size)

            # Also set document defaults (rPrDefault) for empty paragraphs
            if metadata.default_font:
                from docx.oxml import parse_xml
                styles_element = docx_doc.styles.element
                doc_defaults = styles_element.find(qn("w:docDefaults"))

                if doc_defaults is not None:
                    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
                    if rpr_default is None:
                        # Create rPrDefault if it doesn't exist
                        xmlns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        rpr_default_xml = f'<w:rPrDefault xmlns:w="{xmlns}"><w:rPr/></w:rPrDefault>'
                        rpr_default = parse_xml(rpr_default_xml)
                        doc_defaults.insert(0, rpr_default)

                    rpr = rpr_default.find(qn("w:rPr"))
                    if rpr is None:
                        rpr = parse_xml(f'<w:rPr xmlns:w="{xmlns}"/>')
                        rpr_default.append(rpr)

                    # Set or update rFonts
                    r_fonts = rpr.find(qn("w:rFonts"))
                    if r_fonts is None:
                        xmlns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        r_fonts_xml = (
                            f'<w:rFonts xmlns:w="{xmlns}" '
                            f'w:ascii="{metadata.default_font}" '
                            f'w:hAnsi="{metadata.default_font}" '
                            f'w:cs="{metadata.default_font}"/>'
                        )
                        r_fonts = parse_xml(r_fonts_xml)
                        rpr.append(r_fonts)
                    else:
                        # Update existing rFonts
                        r_fonts.set(qn("w:ascii"), metadata.default_font)
                        r_fonts.set(qn("w:hAnsi"), metadata.default_font)
                        r_fonts.set(qn("w:cs"), metadata.default_font)
        except Exception:
            pass


def rebuild_document(doc_data: Document, output_path: str, template_path: str | None = None):
    """Rebuild a DOCX file from a Document dataclass object.

    Creates a new DOCX file with all content and formatting from the
    extracted document structure, including metadata and page settings.

    Args:
        doc_data: Document object with all content and formatting
        output_path: Path where the new DOCX file will be saved
        template_path: Optional path to use as template (preserves styles/defaults)

    Returns:
        None (writes to output_path)
    """
    # Create new python-docx Document (optionally from template)
    if template_path:
        docx_doc = DocxDocument(template_path)
        # Clear all existing paragraphs from template
        # We need to clear the body while preserving section properties
        body = docx_doc.element.body
        # Remove all paragraph and table elements
        for element in list(body):
            if element.tag.endswith(('p', 'tbl')):
                body.remove(element)
    else:
        docx_doc = DocxDocument()

    # Apply document metadata (page settings, line numbering, etc.)
    apply_metadata(docx_doc, doc_data.metadata)

    # Get default font and size for empty paragraphs
    default_font = doc_data.metadata.default_font
    default_font_size = doc_data.metadata.default_font_size

    # Add preamble paragraphs
    for para_data in doc_data.preamble_paragraphs:
        add_paragraph_to_document(docx_doc, para_data, default_font=default_font, default_font_size=default_font_size)

    # Add sections with headers and body paragraphs
    for section in doc_data.sections:
        # Add header paragraph
        add_paragraph_to_document(docx_doc, section.header, default_font=default_font, default_font_size=default_font_size)

        # Add body paragraphs
        for para_data in section.paragraphs:
            add_paragraph_to_document(docx_doc, para_data, default_font=default_font, default_font_size=default_font_size)

    # Save the document
    docx_doc.save(output_path)
