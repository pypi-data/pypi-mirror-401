"""Extract DOCX files into structured dataclass objects."""

from typing import Callable, List, Optional, Tuple

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from translate_docx.models import (
    Document,
    DocumentMetadata,
    Hyperlink,
    LineNumbering,
    PageSettings,
    Paragraph,
    ParagraphFormatting,
    RunFormatting,
    Section,
    TextRun,
)
from translate_docx.utils import color_to_hex, normalize_alignment


def extract_run_formatting(run) -> RunFormatting:
    """Extract formatting attributes from a python-docx Run object.

    Args:
        run: python-docx Run object

    Returns:
        RunFormatting object with extracted attributes
    """
    # Extract basic formatting
    bold = run.bold
    italic = run.italic
    underline = run.underline

    # Extract font properties
    font = run.font
    font_name = font.name
    font_size = int(font.size.pt) if font.size else None

    # Extract color
    color = color_to_hex(font.color)

    # Superscript and subscript
    is_superscript = font.superscript is True
    is_subscript = font.subscript is True

    return RunFormatting(
        bold=bold,
        italic=italic,
        underline=underline,
        font_name=font_name,
        font_size=font_size,
        color=color,
        is_superscript=is_superscript,
        is_subscript=is_subscript,
    )


def extract_text_run(run) -> TextRun:
    """Extract a single text run with its formatting.

    Args:
        run: python-docx Run object

    Returns:
        TextRun object
    """
    formatting = extract_run_formatting(run)
    return TextRun(text=run.text, formatting=formatting)


def extract_hyperlink(hyperlink) -> Hyperlink:
    """Extract a hyperlink with its nested runs and URL.

    Args:
        hyperlink: python-docx Hyperlink object

    Returns:
        Hyperlink object with runs and URL
    """
    # Extract all runs within the hyperlink
    runs = [extract_text_run(run) for run in hyperlink.runs]

    # Extract URL (address for external, fragment for internal anchors)
    url = hyperlink.address  # External URL (e.g., "https://example.com")
    anchor = hyperlink.fragment  # Internal anchor (e.g., "#section1")

    return Hyperlink(runs=runs, url=url, anchor=anchor)


def extract_paragraph_formatting(paragraph) -> ParagraphFormatting:
    """Extract formatting attributes from a python-docx Paragraph object.

    Args:
        paragraph: python-docx Paragraph object

    Returns:
        ParagraphFormatting object with extracted attributes
    """
    para_format = paragraph.paragraph_format

    alignment = normalize_alignment(paragraph.alignment)

    # Extract spacing and indentation
    line_spacing = para_format.line_spacing
    space_before = para_format.space_before
    space_after = para_format.space_after
    first_line_indent = para_format.first_line_indent
    left_indent = para_format.left_indent
    right_indent = para_format.right_indent

    # Extract boolean properties
    keep_together = para_format.keep_together
    keep_with_next = para_format.keep_with_next
    page_break_before = para_format.page_break_before

    return ParagraphFormatting(
        alignment=alignment,
        line_spacing=line_spacing,
        space_before=space_before,
        space_after=space_after,
        first_line_indent=first_line_indent,
        left_indent=left_indent,
        right_indent=right_indent,
        keep_together=keep_together,
        keep_with_next=keep_with_next,
        page_break_before=page_break_before,
    )


def extract_paragraph(paragraph, is_header: bool = False) -> Paragraph:
    """Extract a single paragraph with all formatting.

    Args:
        paragraph: python-docx Paragraph object
        is_header: Whether this paragraph is a section header (default False)

    Returns:
        Paragraph object with extracted content and formatting
    """
    # Extract all inline elements (runs and hyperlinks) in document order
    inline_elements = []

    # Use iter_inner_content() to get runs AND hyperlinks in order
    try:
        for item in paragraph.iter_inner_content():
            # Check type using duck typing - hyperlinks have 'runs' attribute
            if hasattr(item, 'runs') and not isinstance(item, type(paragraph.runs[0])) if paragraph.runs else False:
                # This is a Hyperlink object
                inline_elements.append(extract_hyperlink(item))
            else:
                # This is a Run object
                inline_elements.append(extract_text_run(item))
    except (AttributeError, TypeError):
        # Fallback: if iter_inner_content() not available, use paragraph.runs
        inline_elements = [extract_text_run(run) for run in paragraph.runs]

    # Extract paragraph formatting
    formatting = extract_paragraph_formatting(paragraph)

    # Get style name
    style_name = paragraph.style.name if paragraph.style else None

    # For empty paragraphs, extract font size from paragraph's rPr element
    empty_paragraph_font_size = None
    is_empty = not inline_elements or all(
        (isinstance(el, TextRun) and el.text == "") or
        (isinstance(el, Hyperlink) and el.text == "")
        for el in inline_elements
    )

    if is_empty:
        # Paragraph is empty or contains only empty runs
        try:
            p_pr = paragraph._element.pPr
            if p_pr is not None:
                rpr = p_pr.find(qn("w:rPr"))
                if rpr is not None:
                    sz = rpr.find(qn("w:sz"))
                    if sz is not None:
                        # Font size is in half-points, convert to points
                        half_points = sz.get(qn("w:val"))
                        if half_points:
                            empty_paragraph_font_size = int(half_points) // 2
        except Exception:
            pass

    return Paragraph(
        runs=inline_elements,
        formatting=formatting,
        is_header=is_header,
        style_name=style_name,
        empty_paragraph_font_size=empty_paragraph_font_size,
    )


def extract_line_numbering(section) -> LineNumbering:
    """Extract line numbering settings from a document section.

    Requires direct XML manipulation since python-docx doesn't fully
    expose this API.

    Args:
        section: python-docx Section object

    Returns:
        LineNumbering object with extracted settings
    """
    try:
        sect_pr = section._sectPr
        ln_num_type = sect_pr.find(qn("w:lnNumType"))

        if ln_num_type is None:
            return LineNumbering(enabled=False)

        # Extract attributes
        start_val = ln_num_type.get(qn("w:start"))
        count_by = ln_num_type.get(qn("w:countBy"))
        restart = ln_num_type.get(qn("w:restart"))
        distance = ln_num_type.get(qn("w:distance"))

        return LineNumbering(
            enabled=True,
            start=int(start_val) if start_val else 1,
            increment=int(count_by) if count_by else 1,
            restart=restart,
            distance=int(distance) if distance else None,
        )
    except Exception:
        # If anything fails, return default (no line numbering)
        return LineNumbering(enabled=False)


def extract_page_settings(section) -> PageSettings:
    """Extract page layout settings from a document section.

    Args:
        section: python-docx Section object

    Returns:
        PageSettings object with extracted attributes
    """
    return PageSettings(
        width=section.page_width,
        height=section.page_height,
        top_margin=section.top_margin,
        bottom_margin=section.bottom_margin,
        left_margin=section.left_margin,
        right_margin=section.right_margin,
        header_distance=section.header_distance,
        footer_distance=section.footer_distance,
    )


def extract_metadata(docx_doc: DocxDocument) -> DocumentMetadata: ## type: ignore
    """Extract document-level metadata and settings.

    Args:
        docx_doc: python-docx Document object

    Returns:
        DocumentMetadata object with extracted settings
    """
    # Get the first section (most documents have at least one)
    section = docx_doc.sections[0] if docx_doc.sections else None

    if section is None:
        return DocumentMetadata()

    # Extract page and line numbering settings
    page_settings = extract_page_settings(section)
    line_numbering = extract_line_numbering(section)

    # Try to extract default font and size from styles and document defaults
    default_font = None
    default_font_size = None

    try:
        # First try the Normal style
        styles = docx_doc.styles
        normal_style = styles["Normal"]
        if normal_style.font.name:
            default_font = normal_style.font.name
        if normal_style.font.size:
            default_font_size = int(normal_style.font.size.pt)

        # If not found in Normal style, check document defaults (docDefaults/rPrDefault)
        if not default_font or not default_font_size:
            styles_element = docx_doc.styles.element
            doc_defaults = styles_element.find(qn("w:docDefaults"))

            if doc_defaults is not None:
                rpr_default = doc_defaults.find(qn("w:rPrDefault"))
                if rpr_default is not None:
                    rpr = rpr_default.find(qn("w:rPr"))
                    if rpr is not None:
                        # Extract font family
                        if not default_font:
                            r_fonts = rpr.find(qn("w:rFonts"))
                            if r_fonts is not None:
                                # Try different font attributes (ascii, hAnsi, cs)
                                default_font = (
                                    r_fonts.get(qn("w:ascii"))
                                    or r_fonts.get(qn("w:hAnsi"))
                                    or r_fonts.get(qn("w:cs"))
                                )

                        # Extract font size
                        if not default_font_size:
                            sz = rpr.find(qn("w:sz"))
                            if sz is not None:
                                # Size is in half-points, convert to points
                                half_points = sz.get(qn("w:val"))
                                if half_points:
                                    default_font_size = int(half_points) // 2

        # If still not found, infer from first paragraph's first run (fallback)
        if not default_font:
            for para in docx_doc.paragraphs:
                if para.runs and para.runs[0].font.name:
                    default_font = para.runs[0].font.name
                    break

        if not default_font_size:
            for para in docx_doc.paragraphs:
                if para.runs and para.runs[0].font.size:
                    default_font_size = int(para.runs[0].font.size.pt)
                    break
    except Exception:
        pass

    return DocumentMetadata(
        line_numbering=line_numbering,
        page_settings=page_settings,
        default_font=default_font,
        default_font_size=default_font_size,
    )


def is_section_header(paragraph: Paragraph, header_detector: Optional[Callable] = None) -> bool:
    """Determine if a paragraph is a section header based on formatting.

    Uses multi-criteria heuristic to detect bold headers:
    1. Paragraph style is "Heading X"
    2. All text runs are bold AND paragraph is short (< 200 chars)
    3. Custom header_detector function if provided

    Args:
        paragraph: Paragraph object to check
        header_detector: Optional custom detection function

    Returns:
        True if paragraph is a header, False otherwise
    """
    # If custom detector provided, use it
    if header_detector is not None:
        return header_detector(paragraph)

    # Criterion 1: Check style name (e.g., "Heading 1", "Heading 2")
    if paragraph.style_name and paragraph.style_name.startswith("Heading"):
        return True

    # Criterion 2: All runs are bold AND text is short
    if not paragraph.runs:
        return False

    # Collect all runs (including nested runs in hyperlinks)
    all_runs = []
    for element in paragraph.runs:
        if isinstance(element, TextRun):
            all_runs.append(element)
        elif isinstance(element, Hyperlink):
            all_runs.extend(element.runs)

    # Check if all non-empty runs are bold
    non_empty_runs = [r for r in all_runs if r.text.strip()]
    if not non_empty_runs:
        return False

    all_bold = all(run.formatting.bold for run in non_empty_runs)
    is_short = 0 < len(paragraph.text.strip()) < 200

    if all_bold and is_short:
        return True

    return False


def split_into_sections(
    paragraphs: List[Paragraph],
    header_detector: Optional[Callable] = None,
) -> Tuple[List[Section], List[Paragraph]]:
    """Split paragraphs into sections based on header paragraphs.

    A section consists of a header paragraph followed by body paragraphs.
    Paragraphs before the first header go into preamble.

    Args:
        paragraphs: List of all paragraphs in order
        header_detector: Optional custom header detection function

    Returns:
        Tuple of (sections list, preamble paragraphs list)
    """
    sections = []
    preamble = []
    current_header = None
    current_body = []

    for para in paragraphs:
        # Check if this paragraph is a header
        if is_section_header(para, header_detector):
            # Save previous section if exists
            if current_header is not None:
                sections.append(Section(header=current_header, paragraphs=current_body))

            # Start new section with this header
            current_header = para
            current_body = []
        else:
            # Add to current section or preamble
            if current_header is not None:
                current_body.append(para)
            else:
                preamble.append(para)

    # Don't forget the last section
    if current_header is not None:
        sections.append(Section(header=current_header, paragraphs=current_body))

    return sections, preamble


def extract_document(docx_path: str, header_detector: Optional[Callable] = None) -> Document:
    """Extract a complete DOCX file into a Document object.

    Args:
        docx_path: Path to the DOCX file
        header_detector: Optional custom function to detect section headers

    Returns:
        Document object with all content and formatting extracted

    Raises:
        FileNotFoundError: If the DOCX file doesn't exist
        Exception: If the file is not a valid DOCX
    """
    # Open the DOCX file
    docx_doc = DocxDocument(docx_path)

    # Extract metadata (page settings, line numbering, default fonts)
    metadata = extract_metadata(docx_doc)

    # Extract all paragraphs as Paragraph objects
    paragraphs = [extract_paragraph(para, is_header=False) for para in docx_doc.paragraphs]

    # Split paragraphs into sections based on headers
    sections, preamble = split_into_sections(paragraphs, header_detector)

    # Mark header paragraphs
    for section in sections:
        section.header.is_header = True

    # Create Document object with sections
    document = Document(
        metadata=metadata,
        sections=sections,
        preamble_paragraphs=preamble,
    )

    return document
